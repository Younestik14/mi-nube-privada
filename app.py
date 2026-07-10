"""
================================================================================
 CALCULADORA DE SECCIONES DE CABLE — REBT
================================================================================
Herramienta de apoyo al cálculo de la sección mínima de conductores en
instalaciones de baja tensión, conforme al Reglamento Electrotécnico para
Baja Tensión (REBT, RD 842/2002) y sus Instrucciones Técnicas Complementarias:

    ITC-BT-14   Línea General de Alimentación (LGA)
    ITC-BT-15   Derivación Individual (DI)
    ITC-BT-18   Puestas a tierra (conductor de protección)
    ITC-BT-19   Prescripciones generales de instalaciones interiores
                (intensidades admisibles y caída de tensión) — Tabla 1
    ITC-BT-07   Redes subterráneas (instalación enterrada, resistividad terreno)
    ITC-BT-47   Instalación de receptores. Motores.

Criterios de dimensionado aplicados (art. 19.2 ITC-BT-19):
    1) Criterio térmico              Ib ≤ In ≤ Iz
    2) Criterio de caída de tensión  ΔU% ≤ ΔU%max según tramo
    3) Criterio térmico de cortocircuito (verificación opcional, S ≥ (Icc·√t)/k)

Incluye además:
    - Justificación de fórmulas con los valores sustituidos (pestaña "Fórmulas")
    - Memoria de cálculo exportable en PDF con cajetín tipo plano técnico
    - Presupuesto (mediciones x precios unitarios editables) exportable a Excel

Fuentes numéricas: Guía-BT-19 (Ministerio de Industria, Turismo y Comercio,
Ed. feb-09 Rev.2) Tablas A, C, D, E, F; ITC-BT-14, ITC-BT-15, ITC-BT-47
(textos oficiales); Libro Blanco de la Instalación (Prysmian) para las
constantes k de cortocircuito. Ver pestaña "Metodología" para el detalle de
cada tabla, incluyendo qué valores son directos de la norma y cuáles son
estimaciones ancladas a un valor real verificado.

⚠️  Herramienta de apoyo al diseño, no un sustituto del criterio de un
    técnico competente. Antes de emitir un proyecto o memoria firmada,
    contrasta los valores críticos contra la edición vigente de la
    Guía-BT-19 y la norma UNE-HD 60364-5-52. Los precios del presupuesto
    son orientativos y deben sustituirse por los del proveedor real.

Autor: Younes — IDEA TSG
================================================================================
"""

from __future__ import annotations

import base64
import random
import io
import json
import math
from datetime import date, datetime

import pandas as pd
import streamlit as st

# ==============================================================================
# 1. CONSTANTES Y TABLAS NORMATIVAS
# ==============================================================================

SECCIONES_NORMALIZADAS = [1.5, 2.5, 4, 6, 10, 16, 25, 35, 50, 70,
                          95, 120, 150, 185, 240, 300]

CONDUCTIVIDAD_20C = {"Cobre": 56.0, "Aluminio": 35.0}
COEF_TEMP_RESIST = {"Cobre": 0.00393, "Aluminio": 0.00403}
TEMP_SERVICIO = {"PVC": 70.0, "XLPE/EPR": 90.0}
TEMP_REF_TABLA_AIRE = 40.0
TEMP_REF_TABLA_TERRENO = 25.0
RESISTIVIDAD_REF_TERRENO = 1.5

REACTANCIA_LINEAL_DEFECTO = 0.08

K_CORTOCIRCUITO = {
    ("Cobre", "PVC"): 115,
    ("Cobre", "XLPE/EPR"): 143,
    ("Aluminio", "PVC"): 76,
    ("Aluminio", "XLPE/EPR"): 94,
}

CAIDA_TENSION_MAX = {
    "LGA — centralización única de contadores": 0.5,
    "LGA — centralizaciones parciales": 1.0,
    "Derivación Individual — contadores totalmente centralizados": 1.0,
    "Derivación Individual — contadores en más de un lugar": 0.5,
    "Derivación Individual — usuario único (sin LGA)": 1.5,
    "Instalación interior — Alumbrado": 3.0,
    "Instalación interior — Otros usos / fuerza": 5.0,
    "Circuito de motor": 5.0,
    "Alimentación con transformador propio — Alumbrado": 4.5,
    "Alimentación con transformador propio — Otros usos": 6.5,
    "Instalación generadora / FV (ITC-BT-40)": 1.5,
    "Personalizado": None,
}

SISTEMA_MONO = "Monofásico (230 V)"
SISTEMA_TRI = "Trifásico (400 V)"

# Métodos de instalación disponibles. B1, C/E, F y D salen directamente de la
# Tabla A / Tabla D de la Guía-BT-19. A1 y B2 se derivan de B1 aplicando un
# ratio anclado a un valor real verificado de forma independiente (ver
# comentarios junto a cada factor, y la pestaña "Metodología").
METODO_A1 = "A1 — Tubo empotrado en pared aislante (estimado desde B1)"
METODO_B1 = "B1 — Tubo empotrado en obra o en superficie"
METODO_B2 = "B2 — Cable multiconductor en tubo (estimado desde B1)"
METODO_CE = "C/E — Bandeja no perforada o perforada / directo sobre superficie"
METODO_F = "F — Bandeja perforada, unipolares muy separados (S≥25 mm², XLPE)"
METODO_D = "D — Enterrado bajo tubo (XLPE)"
METODOS_DISPONIBLES = [METODO_A1, METODO_B1, METODO_B2, METODO_CE, METODO_F, METODO_D]

# Ratio A1/B1 anclado en un valor real verificado de forma independiente:
# A1, 2,5 mm², 2 cargados, PVC = 17,5 A (fuente: EleCalculador.com, con cita
# textual de la Guía-BT-19) frente a B1 2,5 mm² 2 cargados PVC = 21 A
# (Guía-BT-19 directa) → ratio 17,5/21 = 0,833.
FACTOR_A1_SOBRE_B1 = 0.833

# Ratio B2/B1 anclado en un ejercicio resuelto independiente (circuitoelectrico.com):
# B2, 25 mm², 2 cargados, PVC = 77 A frente a B1 25 mm² 2 cargados PVC = 84 A
# (Guía-BT-19 directa) → ratio 77/84 = 0,917. Coherente con el rango de
# variación (7-12 %) que cita Prysmian entre columnas B1 y B2.
FACTOR_B2_SOBRE_B1 = 0.917

# ------------------------------------------------------------------------------
# Tabla A — Intensidades admisibles (A), cables de COBRE, NO enterrados,
# temperatura ambiente de referencia 40°C. Fuente: Guía-BT-19, Tabla A.
# Aquí solo se tabula B1 (columnas 5-8) y C/E (columnas 9-12) y F (columna 13),
# que son los tres bloques de columnas de los que hay valores directos
# verificados; A1 y B2 se derivan de B1 mediante los factores anteriores.
# ------------------------------------------------------------------------------
TABLA_A_COBRE = {
    1.5:  {"B1": (13.5, 15.0, 16.0, 16.5), "CE": (19.0, 20.0, 21.0, 24.0), "F": None},
    2.5:  {"B1": (18.5, 21.0, 22.0, 23.0), "CE": (26.0, 26.5, 29.0, 33.0), "F": None},
    4:    {"B1": (24.0, 27.0, 30.0, 31.0), "CE": (34.0, 36.0, 38.0, 45.0), "F": None},
    6:    {"B1": (32.0, 36.0, 37.0, 40.0), "CE": (44.0, 46.0, 49.0, 57.0), "F": None},
    10:   {"B1": (44.0, 50.0, 52.0, 54.0), "CE": (60.0, 65.0, 68.0, 76.0), "F": None},
    16:   {"B1": (59.0, 66.0, 70.0, 73.0), "CE": (81.0, 87.0, 91.0, 105.0), "F": None},
    25:   {"B1": (77.0, 84.0, 88.0, 95.0), "CE": (103.0, 110.0, 116.0, 123.0), "F": 140.0},
    35:   {"B1": (96.0, 104.0, 110.0, 119.0), "CE": (127.0, 137.0, 144.0, 154.0), "F": 174.0},
    50:   {"B1": (117.0, 125.0, 133.0, 145.0), "CE": (155.0, 167.0, 175.0, 188.0), "F": 210.0},
    70:   {"B1": (149.0, 160.0, 171.0, 185.0), "CE": (199.0, 214.0, 224.0, 244.0), "F": 269.0},
    95:   {"B1": (180.0, 194.0, 207.0, 224.0), "CE": (241.0, 259.0, 271.0, 296.0), "F": 327.0},
    120:  {"B1": (208.0, 225.0, 240.0, 260.0), "CE": (280.0, 301.0, 314.0, 348.0), "F": 380.0},
    150:  {"B1": (236.0, 260.0, 278.0, 299.0), "CE": (322.0, 343.0, 363.0, 404.0), "F": 438.0},
    185:  {"B1": (268.0, 297.0, 317.0, 341.0), "CE": (368.0, 391.0, 415.0, 464.0), "F": 500.0},
    240:  {"B1": (315.0, 350.0, 374.0, 401.0), "CE": (435.0, 468.0, 490.0, 552.0), "F": 590.0},
    300:  {"B1": (361.0, 401.0, 430.0, 461.0), "CE": (500.0, 538.0, 563.0, 638.0), "F": 678.0},
}
IDX_3PVC, IDX_2PVC, IDX_3XLPE, IDX_2XLPE = 0, 1, 2, 3

RATIO_ALUMINIO_NO_ENTERRADO = 0.78

TABLA_D_ENTERRADO_XLPE = {
    1.5:  (23.0, None, 27.0, None),
    2.5:  (30.0, 23.0, 36.0, 27.0),
    4:    (39.0, 30.0, 46.0, 36.0),
    6:    (48.0, 37.0, 58.0, 44.0),
    10:   (64.0, 49.0, 77.0, 58.0),
    16:   (82.0, 62.0, 100.0, 77.0),
    25:   (105.0, 82.0, 130.0, 98.0),
    35:   (130.0, 98.0, 155.0, 120.0),
    50:   (155.0, 115.0, 183.0, 139.0),
    70:   (190.0, 145.0, 225.0, 170.0),
    95:   (225.0, 175.0, 265.0, 205.0),
    120:  (260.0, 200.0, 305.0, 230.0),
    150:  (300.0, 230.0, 340.0, 265.0),
    185:  (335.0, 260.0, 385.0, 295.0),
    240:  (400.0, 305.0, 440.0, 340.0),
    300:  (455.0, 350.0, 500.0, 385.0),
}

RESISTIVIDAD_TERRENO = {
    "Inundado (0,40 K·m/W)": 0.40,
    "Muy húmedo (0,50 K·m/W)": 0.50,
    "Húmedo (0,70 K·m/W)": 0.70,
    "Poco húmedo (0,85 K·m/W)": 0.85,
    "Seco (1,00 K·m/W)": 1.00,
    "Arcilloso muy seco (1,20 K·m/W)": 1.20,
    "Arenoso muy seco — referencia Tabla D (1,50 K·m/W)": 1.50,
    "Piedra arenisca (2,00 K·m/W)": 2.00,
    "Piedra caliza (2,50 K·m/W)": 2.50,
    "Piedra granítica (3,00 K·m/W)": 3.00,
}
FACTOR_RESISTIVIDAD_TERRENO = {
    0.40: 1.25, 0.50: 1.21, 0.70: 1.18, 0.85: 1.14, 1.00: 1.10,
    1.20: 1.05, 1.50: 1.00, 2.00: 0.93, 2.50: 0.89, 3.00: 0.85,
}

TABLA_E_AGRUPAMIENTO = {
    "Empotrados o embutidos": {1: 1.00, 2: 0.80, 3: 0.70, 4: 0.70, 6: 0.55, 9: 0.50, 12: 0.45, 16: 0.40, 20: 0.40},
    "Capa única sobre pared/suelo/bandeja no perforada": {1: 1.00, 2: 0.85, 3: 0.80, 4: 0.75, 6: 0.70, 9: 0.70},
    "Capa única fijada bajo techo": {1: 0.95, 2: 0.80, 3: 0.70, 4: 0.70, 6: 0.65, 9: 0.60},
    "Capa única en bandeja perforada (horiz. o vert.)": {1: 1.00, 2: 0.90, 3: 0.80, 4: 0.75, 6: 0.75, 9: 0.70},
    "Capa única en bandeja de escalera / abrazaderas": {1: 1.00, 2: 0.85, 3: 0.80, 4: 0.80, 6: 0.80, 9: 0.80},
}
TABLA_F_CAPAS = {1: 1.00, 2: 0.80, 3: 0.73, 4: 0.70, 5: 0.70, 6: 0.68, 7: 0.68, 8: 0.68, 9: 0.66}

MINIMO_ABSOLUTO_MM2 = 1.5
MINIMO_ALUMINIO_MM2 = 16.0
MINIMO_LGA_COBRE_MM2 = 10.0
MINIMO_LGA_ALUMINIO_MM2 = 16.0

# ------------------------------------------------------------------------------
# Precios orientativos (€) — PUNTO DE PARTIDA EDITABLE, no precios de mercado
# en tiempo real. El precio del cobre fluctúa; sustituye por los de tu
# proveedor antes de presupuestar en firme. Cable unipolar tipo RZ1-K 0,6/1kV
# (Cu), €/m aproximados de referencia.
# ------------------------------------------------------------------------------
PRECIOS_CABLE_COBRE_DEFECTO = {
    1.5: 0.45, 2.5: 0.65, 4: 0.95, 6: 1.35, 10: 2.10, 16: 3.20, 25: 5.50,
    35: 7.50, 50: 10.50, 70: 14.50, 95: 19.50, 120: 24.50, 150: 30.50,
    185: 38.00, 240: 49.00, 300: 62.00,
}
RATIO_PRECIO_ALUMINIO = 0.55  # el cable de aluminio es sensiblemente más barato que el de cobre

PRECIOS_CANALIZACION_DEFECTO = {
    METODO_A1: 1.10, METODO_B1: 1.10, METODO_B2: 1.10,
    METODO_CE: 3.50, METODO_F: 4.50, METODO_D: 6.00,
}
PRECIO_MANO_OBRA_POR_METRO = 2.20
PORCENTAJE_ACCESORIOS = 8.0  # % sobre el subtotal de materiales (cajas, terminales, prensaestopas...)

# Calibres normalizados de interruptor automático (A) y precio orientativo (€)
CALIBRES_MAGNETOTERMICO = [10, 16, 20, 25, 32, 40, 50, 63, 80, 100, 125, 160, 200, 250]

# ITC-BT-47 tabla 1: relación máxima admisible entre intensidad de arranque
# e intensidad nominal, según la potencia del motor (si el motor supera este
# límite en arranque directo, se exige un sistema de arranque que la reduzca,
# como el estrella-triángulo).
RATIO_IA_IN_MAX_MOTOR = [
    (0.0, 1.5, 4.5), (1.5, 5.0, 3.0), (5.0, 15.0, 2.0), (15.0, float("inf"), 1.5),
]

# Rango típico (múltiplos de In) de disparo instantáneo (magnético) según la
# curva del interruptor automático, UNE-EN 60898.
CURVA_MAGNETOTERMICO_RANGOS = {"B": (3, 5), "C": (5, 10), "D": (10, 20)}

# Iluminancias medias recomendadas de referencia, UNE-EN 12464-1 (valores
# orientativos habituales; la norma completa distingue muchas más tareas).
ILUMINANCIA_POR_LOCAL = {
    "Vivienda — estancia/salón": 200, "Vivienda — cocina": 300, "Oficina / despacho": 500,
    "Aula / centro de enseñanza": 300, "Comercio / escaparate": 300, "Almacén (tránsito)": 100,
    "Almacén (manipulación)": 200, "Taller — trabajo basto": 200, "Taller — trabajo fino": 500,
    "Pasillo / circulación": 100, "Garaje / aparcamiento": 75, "Escalera": 100,
}

# Ensayos previos a la puesta en servicio (ITC-BT-05 / UNE-HD 60364-6), usados
# tanto en el Pliego de Condiciones como en la checklist interactiva y el CIE.
ENSAYOS_PUESTA_SERVICIO = [
    ("Continuidad de conductores de protección y equipotenciales", "Continuidad eléctrica verificada"),
    ("Resistencia de aislamiento (circuitos ≤500V, ensayo a 500V c.c.)", "≥ 0,50 MΩ"),
    ("Resistencia de aislamiento (circuitos MBTS/MBTP, ensayo a 250V c.c.)", "≥ 0,25 MΩ"),
    ("Resistencia de aislamiento (circuitos >500V, ensayo a 1000V c.c.)", "≥ 1,00 MΩ"),
    ("Rigidez dieléctrica (1 min a 2U+1000V, mínimo 1.500V)", "Sin perforación del aislamiento"),
    ("Resistencia de puesta a tierra", "Compatible con la sensibilidad del diferencial instalado"),
    ("Disparo de los interruptores diferenciales (botón de test e Idn)", "Disparo correcto"),
    ("Caída de tensión en los circuitos más desfavorables", "Conforme al Anexo de Cálculos"),
    ("Secuencia de fases y tensiones (instalación trifásica)", "Correcta y equilibrada"),
]

PRECIOS_MAGNETOTERMICO_DEFECTO = {
    10: 18, 16: 18, 20: 19, 25: 20, 32: 22, 40: 28, 50: 35, 63: 45,
    80: 120, 100: 145, 125: 180, 160: 260, 200: 320, 250: 410,
}
PRECIOS_DIFERENCIAL_DEFECTO = {
    25: 55, 40: 65, 63: 90, 80: 210, 100: 260, 125: 320, 160: 420, 200: 520, 250: 650,
}

# ------------------------------------------------------------------------------
# Catálogo ampliado de materiales típicos de una instalación eléctrica de BT.
# Precios orientativos (€), punto de partida editable en la pestaña Presupuesto
# — no son precios de mercado en tiempo real. Organizado por categorías para
# que el desplegable de "añadir partida" sea manejable.
# ------------------------------------------------------------------------------
CATALOGO_MATERIALES = {
    "Cables (por tipo, además del cálculo automático)": {
        "Cable RZ1-K 0,6/1kV 1,5mm² (libre halógenos)": ("m", 0.42),
        "Cable RZ1-K 0,6/1kV 2,5mm²": ("m", 0.58),
        "Cable RZ1-K 0,6/1kV 4mm²": ("m", 0.88),
        "Cable RZ1-K 0,6/1kV 6mm²": ("m", 1.25),
        "Cable RZ1-K 0,6/1kV 10mm²": ("m", 1.95),
        "Cable RZ1-K 0,6/1kV 16mm²": ("m", 2.95),
        "Cable H07V-K 1,5mm² (unipolar, interior tubo)": ("m", 0.28),
        "Cable H07V-K 2,5mm²": ("m", 0.40),
        "Cable ES07Z1-K 1,5mm² (libre halógenos, unipolar)": ("m", 0.35),
        "Cable ES07Z1-K 2,5mm²": ("m", 0.48),
        "Manguera redonda RV-K 3x1,5mm²": ("m", 1.10),
        "Manguera redonda RV-K 3x2,5mm²": ("m", 1.55),
        "Manguera redonda RV-K 5x6mm²": ("m", 4.20),
        "Cable coaxial / datos UTP Cat6": ("m", 0.60),
        "Cable de acometida CU 3x35+16mm² (LGA)": ("m", 12.50),
    },
    "Canalizaciones": {
        "Tubo corrugado empotrar Ø16mm": ("m", 0.55),
        "Tubo corrugado empotrar Ø20mm": ("m", 0.65),
        "Tubo corrugado empotrar Ø25mm": ("m", 0.85),
        "Tubo corrugado empotrar Ø32mm": ("m", 1.10),
        "Tubo corrugado empotrar Ø40mm": ("m", 1.55),
        "Tubo corrugado empotrar Ø50mm": ("m", 2.10),
        "Tubo rígido metálico superficie Ø20mm": ("m", 2.80),
        "Tubo rígido PVC superficie Ø20mm": ("m", 1.60),
        "Tubo rígido PVC superficie Ø32mm": ("m", 2.40),
        "Tubo enterrado TPC doble capa Ø63mm": ("m", 3.20),
        "Tubo enterrado TPC doble capa Ø110mm": ("m", 5.60),
        "Bandeja perforada 100mm": ("m", 6.50),
        "Bandeja perforada 200mm": ("m", 9.80),
        "Bandeja perforada 300mm": ("m", 13.50),
        "Bandeja perforada 400mm": ("m", 17.20),
        "Bandeja no perforada (rejiband) 200mm": ("m", 8.20),
        "Bandeja de varilla (rejilla) 100mm": ("m", 5.40),
        "Canaleta PVC 20x12mm": ("m", 2.10),
        "Canaleta PVC 40x25mm": ("m", 4.30),
        "Canaleta PVC 60x40mm con tabique": ("m", 7.80),
        "Columna/canal de suelo técnico 2 vías": ("m", 22.00),
    },
    "Cajas y mecanismos": {
        "Caja de derivación empotrar 100x100": ("ud", 1.80),
        "Caja de derivación empotrar 150x150": ("ud", 3.20),
        "Caja de derivación estanca IP65 150x110": ("ud", 6.50),
        "Caja de derivación estanca IP65 190x140": ("ud", 9.80),
        "Caja de mecanismo universal": ("ud", 0.60),
        "Interruptor simple (mecanismo + tecla)": ("ud", 8.50),
        "Conmutador": ("ud", 9.50),
        "Cruzamiento (conmutador de cruce)": ("ud", 11.50),
        "Pulsador timbre": ("ud", 7.50),
        "Detector de movimiento/presencia empotrar": ("ud", 24.00),
        "Regulador de intensidad (dimmer)": ("ud", 18.50),
        "Base de enchufe Schuko 16A": ("ud", 8.90),
        "Base de enchufe Schuko con USB": ("ud", 16.50),
        "Base de enchufe estanca IP44 16A": ("ud", 12.50),
        "Base industrial CETAC 16A 2P+T": ("ud", 18.00),
        "Base industrial CETAC 16A 3P+N+T": ("ud", 22.00),
        "Base industrial CETAC 32A 3P+N+T": ("ud", 35.00),
        "Base industrial CETAC 63A 3P+N+T": ("ud", 68.00),
        "Base de cocina/horno 25A 2P+T": ("ud", 9.50),
        "Marco embellecedor 1 elemento": ("ud", 3.20),
        "Marco embellecedor 2-3 elementos": ("ud", 5.50),
    },
    "Protección y maniobra": {
        "PIA (magnetotérmico) 6-10A, curva C": ("ud", 18.00),
        "PIA (magnetotérmico) 16-20A, curva C": ("ud", 19.50),
        "PIA (magnetotérmico) 25-32A, curva C": ("ud", 22.00),
        "PIA (magnetotérmico) 40-50A, curva C": ("ud", 32.00),
        "PIA (magnetotérmico) 63A, curva C": ("ud", 48.00),
        "Interruptor general automático (IGA) 25-40A": ("ud", 65.00),
        "Interruptor general automático (IGA) 63A": ("ud", 95.00),
        "Diferencial tipo AC 25A 30mA": ("ud", 32.00),
        "Diferencial tipo A 25A 30mA": ("ud", 45.00),
        "Diferencial tipo A 40A 30mA": ("ud", 58.00),
        "Diferencial tipo A superinmunizado (Si) 40A": ("ud", 85.00),
        "Diferencial tipo B (para inversores/FV) 40A": ("ud", 210.00),
        "Relé diferencial rearmable + toroidal": ("ud", 145.00),
        "Contactor 25A": ("ud", 38.00),
        "Contactor 40A": ("ud", 58.00),
        "Contactor 63A": ("ud", 95.00),
        "Guardamotor 0,4-6,3A (según ajuste)": ("ud", 55.00),
        "Guardamotor 6-40A (según ajuste)": ("ud", 75.00),
        "Relé térmico": ("ud", 42.00),
        "Seccionador bajo carga 40A": ("ud", 48.00),
        "Seccionador bajo carga 100A": ("ud", 95.00),
        "Base + fusible cilíndrico 10x38": ("ud", 6.50),
        "Fusible NH tamaño 00": ("ud", 12.00),
        "Fusible NH tamaño 1": ("ud", 16.00),
        "Protector de sobretensiones DPS tipo 1+2": ("ud", 140.00),
        "Protector de sobretensiones DPS tipo 2": ("ud", 55.00),
        "Contador de energía monofásico": ("ud", 62.00),
        "Contador de energía trifásico": ("ud", 110.00),
    },
    "Cuadros eléctricos": {
        "Cuadro superficie 12 módulos": ("ud", 28.00),
        "Cuadro superficie 24 módulos": ("ud", 42.00),
        "Cuadro superficie 36 módulos": ("ud", 58.00),
        "Cuadro superficie 48 módulos": ("ud", 72.00),
        "Cuadro empotrar 12 módulos": ("ud", 32.00),
        "Cuadro empotrar 24 módulos": ("ud", 48.00),
        "Cuadro empotrar 48 módulos": ("ud", 78.00),
        "Armario metálico IP65 (industrial), pequeño": ("ud", 145.00),
        "Armario metálico IP65 (industrial), grande": ("ud", 320.00),
        "Embarrado de peine 12 módulos": ("ud", 9.50),
        "Caja General de Protección (CGP)": ("ud", 85.00),
        "Caja de Protección y Medida (CPM)": ("ud", 130.00),
    },
    "Puesta a tierra": {
        "Pica de acero cobreado 1m Ø14mm": ("ud", 9.50),
        "Pica de acero cobreado 2m Ø14mm": ("ud", 14.00),
        "Grapa de conexión pica-cable": ("ud", 3.50),
        "Soldadura aluminotérmica": ("ud", 9.00),
        "Placa de tierra de cobre 500x500mm": ("ud", 48.00),
        "Cable desnudo Cu 16mm² (tierra)": ("m", 2.60),
        "Cable desnudo Cu 35mm² (tierra)": ("m", 4.80),
        "Cable desnudo Cu 50mm² (tierra)": ("m", 6.90),
        "Conductor de protección verde/amarillo (según sección)": ("m", 0.55),
        "Arqueta de registro de tierra 30x30": ("ud", 32.00),
        "Punto de puesta a tierra (caja + borne)": ("ud", 15.00),
        "Borne principal de tierra": ("ud", 18.00),
    },
    "Luminarias": {
        "Downlight LED 8W empotrar": ("ud", 8.90),
        "Downlight LED 12W empotrar": ("ud", 11.50),
        "Downlight LED 18W empotrar": ("ud", 15.90),
        "Pantalla LED 60x60 40W": ("ud", 28.00),
        "Regleta LED estanca 1x18W": ("ud", 16.50),
        "Luminaria estanca IP65 1x36W LED": ("ud", 24.00),
        "Luminaria estanca IP65 2x58W LED": ("ud", 38.00),
        "Aparato autónomo de emergencia LED": ("ud", 26.00),
        "Proyector LED exterior 30W": ("ud", 32.00),
        "Proyector LED exterior 50W": ("ud", 45.00),
        "Proyector LED exterior 100W": ("ud", 78.00),
        "Baliza LED exterior baja altura": ("ud", 34.00),
        "Columna/farola LED exterior completa": ("ud", 220.00),
    },
    "Fotovoltaica (además del cálculo automático)": {
        "Panel fotovoltaico 450Wp": ("ud", 110.00),
        "Panel fotovoltaico 550Wp": ("ud", 135.00),
        "Microinversor (por panel)": ("ud", 145.00),
        "Optimizador de potencia (por panel)": ("ud", 65.00),
        "Estructura coplanar cubierta inclinada (por panel)": ("ud", 35.00),
        "Estructura triangular cubierta plana (por panel)": ("ud", 65.00),
        "Estructura sobre suelo/carport (por panel)": ("ud", 95.00),
        "Conector MC4 (par)": ("ud", 3.50),
        "Batería de litio 5 kWh": ("ud", 2400.00),
        "Batería de litio 10 kWh": ("ud", 4500.00),
        "Cargador de vehículo eléctrico 7,4kW": ("ud", 650.00),
        "Monitorización/gestión de energía (kit)": ("ud", 180.00),
    },
    "Domótica y automatización": {
        "Actuador de persiana KNX": ("ud", 85.00),
        "Actuador de encendido KNX 4 canales": ("ud", 120.00),
        "Termostato/sonda de temperatura KNX": ("ud", 95.00),
        "Pasarela KNX-IP": ("ud", 210.00),
        "Detector crepuscular/fotocélula": ("ud", 22.00),
        "Interruptor horario/programador": ("ud", 28.00),
        "Central de alarma básica": ("ud", 145.00),
        "Detector de humo/CO conectado": ("ud", 38.00),
    },
    "Varios / accesorios": {
        "Prensaestopas M16": ("ud", 0.70),
        "Prensaestopas M20": ("ud", 0.90),
        "Prensaestopas M25": ("ud", 1.20),
        "Prensaestopas M32": ("ud", 1.60),
        "Terminal/puntera tubular (según sección)": ("ud", 0.35),
        "Brida de nylon": ("ud", 0.08),
        "Etiqueta de señalización normalizada": ("ud", 1.20),
        "Regleta de conexión": ("ud", 2.50),
        "Regleta de conexión estanca": ("ud", 4.80),
        "Cinta aislante": ("ud", 1.50),
        "Manguito termorretráctil": ("ud", 0.60),
        "Tapa ciega de mecanismo": ("ud", 1.10),
    },
    "Mano de obra": {
        "Encargado / jefe de equipo (hora)": ("h", 26.00),
        "Oficial 1ª electricista (hora)": ("h", 22.00),
        "Oficial 2ª electricista (hora)": ("h", 19.00),
        "Ayudante electricista (hora)": ("h", 16.50),
        "Peón especialista (hora)": ("h", 14.50),
        "Montaje e instalación de punto de luz (unidad)": ("ud", 12.00),
        "Montaje e instalación de toma de corriente (unidad)": ("ud", 10.00),
        "Montaje e instalación de mecanismo de mando (unidad)": ("ud", 9.00),
        "Tendido de cable bajo tubo (por metro)": ("m", 1.10),
        "Montaje de cuadro eléctrico (según nº módulos, unidad)": ("ud", 65.00),
        "Conexionado y etiquetado de cuadro (unidad)": ("ud", 45.00),
        "Ejecución de puesta a tierra (electrodo + conexionado, unidad)": ("ud", 85.00),
        "Puesta en servicio y pruebas reglamentarias (unidad)": ("ud", 120.00),
    },
}

# Tarifas de mano de obra de referencia para el generador de partidas
# compuestas (precios orientativos, no de convenio en tiempo real).
TARIFAS_MANO_OBRA_DEFECTO = {
    "Encargado / jefe de equipo": 26.00,
    "Oficial 1ª electricista": 22.00,
    "Oficial 2ª electricista": 19.00,
    "Ayudante electricista": 16.50,
    "Peón especialista": 14.50,
}

# Tipos de instalación y sus subpartes típicas, para el estimador de
# cantidades del Presupuesto. Cada subparte de tipo "circuito" se calcula
# con el mismo modelo (nº de elementos × longitud media × holgura → tubo,
# cable, cajas); las de tipo "tierra" y "elemento" tienen su propio cálculo.
# Valores por defecto orientativos — siempre editables en la propia UI.
TIPOS_INSTALACION_ESTIMADOR = {
    "🏠 Vivienda — Electrificación básica (ITC-BT-25, ≥5.750 W)": [
        {"nombre": "C1 · Iluminación", "tipo": "circuito", "unidad_elem": "puntos de luz",
         "n_defecto": 10, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "C2 · Tomas de uso general y frigorífico", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 12, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "C3 · Cocina y horno", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 2, "long_defecto": 6.0, "conductores": 3},
        {"nombre": "C4 · Lavadora, lavavajillas y termo eléctrico", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 3, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "C5 · Baño y tomas auxiliares de cocina", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 2, "long_defecto": 6.0, "conductores": 3},
        {"nombre": "Derivación individual", "tipo": "circuito", "unidad_elem": "líneas",
         "n_defecto": 1, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "Cuadro general de mando y protección (básica, ~12 módulos)", "tipo": "elemento",
         "unidad_elem": "cuadros", "n_defecto": 1, "precio_estimado": 42.0},
        {"nombre": "Puesta a tierra", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 1,
         "long_defecto": 20.0},
    ],
    "🏠 Vivienda — Electrificación elevada (ITC-BT-25, ≥9.200 W)": [
        {"nombre": "C1 · Iluminación", "tipo": "circuito", "unidad_elem": "puntos de luz",
         "n_defecto": 14, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "C2 · Tomas de uso general y frigorífico", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 16, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "C3 · Cocina y horno", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 2, "long_defecto": 6.0, "conductores": 3},
        {"nombre": "C4 · Lavadora, lavavajillas y termo eléctrico", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 3, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "C5 · Baño y tomas auxiliares de cocina", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 2, "long_defecto": 6.0, "conductores": 3},
        {"nombre": "C6 · Circuito adicional C1/C2/C3 (exceso de puntos)", "tipo": "circuito",
         "unidad_elem": "tomas/puntos", "n_defecto": 6, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "C7 · Calefacción eléctrica", "tipo": "circuito", "unidad_elem": "unidades",
         "n_defecto": 1, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "C8 · Aire acondicionado", "tipo": "circuito", "unidad_elem": "unidades",
         "n_defecto": 1, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "C9 · Secadora independiente", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 1, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "C10 · Automatización, gestión técnica de energía y seguridad", "tipo": "circuito",
         "unidad_elem": "puntos", "n_defecto": 4, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "C11 · Alimentación circuitos derivados de automatización (C10)", "tipo": "circuito",
         "unidad_elem": "líneas", "n_defecto": 1, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "C12 · Circuito adicional C3/C4/C5 (exceso de puntos o baños)", "tipo": "circuito",
         "unidad_elem": "tomas", "n_defecto": 3, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "Derivación individual", "tipo": "circuito", "unidad_elem": "líneas",
         "n_defecto": 1, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "Cuadro general de mando y protección (elevada, ~24 módulos)", "tipo": "elemento",
         "unidad_elem": "cuadros", "n_defecto": 1, "precio_estimado": 78.0},
        {"nombre": "Puesta a tierra", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 1,
         "long_defecto": 25.0},
    ],
    "🏢 Local comercial / oficina": [
        {"nombre": "Alumbrado general", "tipo": "circuito", "unidad_elem": "puntos de luz",
         "n_defecto": 16, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "Alumbrado de emergencia", "tipo": "circuito", "unidad_elem": "aparatos autónomos",
         "n_defecto": 6, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "Tomas de corriente", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 14, "long_defecto": 12.0, "conductores": 3},
        {"nombre": "Climatización", "tipo": "circuito", "unidad_elem": "unidades",
         "n_defecto": 2, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "Fuerza (maquinaria/electrodomésticos)", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 3, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "Escaparate / rótulo luminoso", "tipo": "circuito", "unidad_elem": "puntos",
         "n_defecto": 2, "long_defecto": 8.0, "conductores": 3},
        {"nombre": "Cuadro general de distribución", "tipo": "elemento", "unidad_elem": "cuadros",
         "n_defecto": 1, "precio_estimado": 78.0},
        {"nombre": "Puesta a tierra", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 1,
         "long_defecto": 20.0},
    ],
    "🏭 Nave industrial": [
        {"nombre": "Acometida y CGP", "tipo": "elemento", "unidad_elem": "conjuntos", "n_defecto": 1,
         "precio_estimado": 85.0},
        {"nombre": "Cuadro general de distribución", "tipo": "elemento", "unidad_elem": "cuadros",
         "n_defecto": 1, "precio_estimado": 145.0},
        {"nombre": "Fuerza motriz (motores)", "tipo": "circuito", "unidad_elem": "motores",
         "n_defecto": 4, "long_defecto": 25.0, "conductores": 5},
        {"nombre": "Alumbrado industrial general", "tipo": "circuito", "unidad_elem": "puntos de luz",
         "n_defecto": 20, "long_defecto": 20.0, "conductores": 3},
        {"nombre": "Alumbrado de emergencia", "tipo": "circuito", "unidad_elem": "aparatos autónomos",
         "n_defecto": 10, "long_defecto": 20.0, "conductores": 3},
        {"nombre": "Tomas de corriente industriales", "tipo": "circuito", "unidad_elem": "tomas",
         "n_defecto": 8, "long_defecto": 18.0, "conductores": 5},
        {"nombre": "Puesta a tierra", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 3,
         "long_defecto": 40.0},
    ],
    "🅿️ Garaje / aparcamiento": [
        {"nombre": "Alumbrado", "tipo": "circuito", "unidad_elem": "puntos de luz",
         "n_defecto": 12, "long_defecto": 20.0, "conductores": 3},
        {"nombre": "Alumbrado de emergencia", "tipo": "circuito", "unidad_elem": "aparatos autónomos",
         "n_defecto": 6, "long_defecto": 20.0, "conductores": 3},
        {"nombre": "Ventilación forzada (extractor)", "tipo": "circuito", "unidad_elem": "motores",
         "n_defecto": 2, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "Detección de CO", "tipo": "circuito", "unidad_elem": "detectores",
         "n_defecto": 3, "long_defecto": 15.0, "conductores": 3},
        {"nombre": "Puerta automática / motor de puerta", "tipo": "circuito", "unidad_elem": "motores",
         "n_defecto": 1, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "Puesta a tierra", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 1,
         "long_defecto": 20.0},
    ],
    "☀️ Instalación fotovoltaica": [
        {"nombre": "Cableado CC (generador → inversor)", "tipo": "circuito", "unidad_elem": "strings",
         "n_defecto": 2, "long_defecto": 15.0, "conductores": 2},
        {"nombre": "Cableado CA (inversor → cuadro)", "tipo": "circuito", "unidad_elem": "líneas",
         "n_defecto": 1, "long_defecto": 10.0, "conductores": 3},
        {"nombre": "Estructura y protecciones DC", "tipo": "elemento", "unidad_elem": "conjuntos",
         "n_defecto": 1, "precio_estimado": 140.0},
        {"nombre": "Cuadro de protecciones CA e interconexión", "tipo": "elemento", "unidad_elem": "cuadros",
         "n_defecto": 1, "precio_estimado": 150.0},
        {"nombre": "Puesta a tierra de estructura", "tipo": "tierra", "unidad_elem": "picas", "n_defecto": 1,
         "long_defecto": 15.0},
    ],
}

# Estructura clásica de presupuesto de instalaciones en España: sobre el
# precio base de cada material/mano de obra se aplican estos dos porcentajes
# para obtener el precio de venta (Precio_venta = Precio_base × (1 + %ben + %amort)).
PORCENTAJE_BENEFICIO_DEFECTO = 15.0       # beneficio industrial
PORCENTAJE_AMORTIZACION_DEFECTO = 5.0     # amortización de medios auxiliares / herramientas


# ==============================================================================
# 2. FUNCIONES DE CALCULO (logica pura, sin dependencias de Streamlit)
# ==============================================================================

def calcular_intensidad_empleo(sistema: str, potencia_w: float, tension: float,
                                cos_phi: float) -> float:
    """Corriente de empleo Ib a partir de la potencia activa (W)."""
    cos_phi = max(cos_phi, 0.01)
    if sistema == SISTEMA_MONO:
        return potencia_w / (tension * cos_phi)
    return potencia_w / (math.sqrt(3) * tension * cos_phi)


def factor_motor_unico(in_motor: float) -> float:
    """ITC-BT-47 3.1: un solo motor -> 125% de la intensidad a plena carga."""
    return 1.25 * in_motor


def factor_varios_motores(corrientes_motores: list) -> float:
    """ITC-BT-47 3.2: 125% del motor de mayor potencia + 100% del resto."""
    if not corrientes_motores:
        return 0.0
    ordenados = sorted(corrientes_motores, reverse=True)
    return 1.25 * ordenados[0] + sum(ordenados[1:])


def factor_correccion_temperatura(aislamiento: str, temp_ambiente: float,
                                   enterrado: bool) -> float:
    """factor = sqrt((Tmax - Ta) / (Tmax - Tref)), Tref=40C aire / 25C terreno."""
    tmax = TEMP_SERVICIO[aislamiento]
    tref = TEMP_REF_TABLA_TERRENO if enterrado else TEMP_REF_TABLA_AIRE
    if temp_ambiente >= tmax:
        return 0.01
    return math.sqrt((tmax - temp_ambiente) / (tmax - tref))


def factor_correccion_agrupamiento(disposicion: str, n_circuitos: int) -> float:
    """Tabla E Guia-BT-19: factor por numero de circuitos agrupados."""
    tabla = TABLA_E_AGRUPAMIENTO[disposicion]
    breakpoints = sorted(tabla.keys())
    factor = tabla[breakpoints[-1]]
    for bp in breakpoints:
        if n_circuitos <= bp:
            factor = tabla[bp]
            break
    return factor


def factor_correccion_capas(n_capas: int) -> float:
    """Tabla F Guia-BT-19: factor adicional por numero de capas."""
    if n_capas <= 1:
        return 1.0
    capas_disponibles = sorted(TABLA_F_CAPAS.keys())
    factor = TABLA_F_CAPAS[capas_disponibles[-1]]
    for c in capas_disponibles:
        if n_capas <= c:
            factor = TABLA_F_CAPAS[c]
            break
    return factor


def factor_correccion_resistividad(resistividad: float) -> float:
    """Factor orientativo por resistividad termica del terreno (ver Tabla C)."""
    claves = sorted(FACTOR_RESISTIVIDAD_TERRENO.keys())
    mas_cercana = min(claves, key=lambda k: abs(k - resistividad))
    return FACTOR_RESISTIVIDAD_TERRENO[mas_cercana]


def iz_tabla(seccion: float, metodo: str, aislamiento: str, conductor: str,
             n_cargados: int):
    """
    Intensidad admisible tabulada (A), antes de factores de correccion.
    A1 y B2 se derivan de B1 aplicando FACTOR_A1_SOBRE_B1 / FACTOR_B2_SOBRE_B1.
    """
    if metodo == METODO_D:
        fila = TABLA_D_ENTERRADO_XLPE.get(seccion)
        if fila is None:
            return None
        cu3, al3, cu2, al2 = fila
        if conductor == "Cobre":
            valor = cu3 if n_cargados == 3 else cu2
        else:
            valor = al3 if n_cargados == 3 else al2
        return valor

    fila = TABLA_A_COBRE.get(seccion)
    if fila is None:
        return None

    if metodo == METODO_F:
        valor_cu = fila["F"]
        if valor_cu is None:
            return None
    else:
        tupla_b1 = fila["B1"]
        tupla_ce = fila["CE"]
        idx = None
        if aislamiento == "PVC":
            idx = IDX_3PVC if n_cargados == 3 else IDX_2PVC
        else:
            idx = IDX_3XLPE if n_cargados == 3 else IDX_2XLPE

        if metodo == METODO_A1:
            valor_cu = round(tupla_b1[idx] * FACTOR_A1_SOBRE_B1, 1)
        elif metodo == METODO_B1:
            valor_cu = tupla_b1[idx]
        elif metodo == METODO_B2:
            valor_cu = round(tupla_b1[idx] * FACTOR_B2_SOBRE_B1, 1)
        else:  # METODO_CE
            valor_cu = tupla_ce[idx]

    if conductor == "Cobre":
        return valor_cu
    return round(valor_cu * RATIO_ALUMINIO_NO_ENTERRADO, 1)


def seccion_por_criterio_termico(ib_calculo: float, metodo: str, aislamiento: str,
                                  conductor: str, n_cargados: int,
                                  factor_correccion: float):
    """Menor seccion normalizada cuya Iz (tabla x factor) >= Ib. Si ninguna basta,
    evalua conductores en paralelo con la mayor seccion disponible."""
    minimo = MINIMO_ALUMINIO_MM2 if conductor == "Aluminio" else MINIMO_ABSOLUTO_MM2
    for s in SECCIONES_NORMALIZADAS:
        if s < minimo:
            continue
        base = iz_tabla(s, metodo, aislamiento, conductor, n_cargados)
        if base is None:
            continue
        iz_real = base * factor_correccion
        if iz_real >= ib_calculo:
            return s, iz_real, False, 1

    s_max = max(SECCIONES_NORMALIZADAS)
    base_max = iz_tabla(s_max, metodo, aislamiento, conductor, n_cargados)
    if base_max is None:
        return None, None, False, 1
    iz_max = base_max * factor_correccion
    n_paralelo = max(2, math.ceil(ib_calculo / iz_max))
    return s_max, iz_max, True, n_paralelo


def kappa_servicio(conductor: str, aislamiento: str, usar_20c: bool = False) -> float:
    """Conductividad (m/ohm*mm2) a la temperatura de servicio del aislamiento."""
    k20 = CONDUCTIVIDAD_20C[conductor]
    if usar_20c:
        return k20
    alpha = COEF_TEMP_RESIST[conductor]
    t_servicio = TEMP_SERVICIO[aislamiento]
    return k20 / (1 + alpha * (t_servicio - 20))


def caida_tension_voltios(sistema: str, ib: float, longitud: float, seccion: float,
                           cos_phi: float, kappa: float,
                           reactancia_ohm_km: float = REACTANCIA_LINEAL_DEFECTO) -> float:
    """e = k_sistema * L * Ib * (R*cosphi + X*sinphi), R = 1/(kappa*S)."""
    sin_phi = math.sqrt(max(0.0, 1 - cos_phi ** 2))
    r_por_metro = 1.0 / (kappa * seccion)
    x_por_metro = reactancia_ohm_km / 1000.0
    k_sistema = 2.0 if sistema == SISTEMA_MONO else math.sqrt(3)
    return k_sistema * longitud * ib * (r_por_metro * cos_phi + x_por_metro * sin_phi)


def seccion_por_caida_tension(sistema: str, ib: float, longitud: float, tension: float,
                               cos_phi: float, kappa: float, delta_u_max_pct: float,
                               conductor: str):
    """Menor seccion normalizada que respeta el limite de caida de tension."""
    minimo = MINIMO_ALUMINIO_MM2 if conductor == "Aluminio" else MINIMO_ABSOLUTO_MM2
    limite_v = tension * delta_u_max_pct / 100.0
    for s in SECCIONES_NORMALIZADAS:
        if s < minimo:
            continue
        e = caida_tension_voltios(sistema, ib, longitud, s, cos_phi, kappa)
        if e <= limite_v:
            pct = e / tension * 100.0
            return s, e, pct
    s_max = max(SECCIONES_NORMALIZADAS)
    e = caida_tension_voltios(sistema, ib, longitud, s_max, cos_phi, kappa)
    return None, e, e / tension * 100.0


def seccion_conductor_neutro(seccion_fase: float, conductor: str,
                              armonicos_significativos: bool) -> float:
    """Sf<=16(Cu)/25(Al) -> Sn=Sf. Si no, Sn>=Sf/2 (redondeada, min 16). Con
    armonicos de 3er orden significativos, Sn=Sf siempre."""
    umbral = 25.0 if conductor == "Aluminio" else 16.0
    if seccion_fase <= umbral or armonicos_significativos:
        return seccion_fase
    objetivo = max(seccion_fase / 2, 16.0)
    for s in SECCIONES_NORMALIZADAS:
        if s >= objetivo:
            return s
    return seccion_fase


def seccion_conductor_proteccion(seccion_fase: float) -> float:
    """Sf<=16 -> Sp=Sf; 16<Sf<=35 -> Sp=16; Sf>35 -> Sp=Sf/2 (redondeada)."""
    if seccion_fase <= 16:
        return seccion_fase
    if seccion_fase <= 35:
        return 16.0
    objetivo = seccion_fase / 2
    for s in SECCIONES_NORMALIZADAS:
        if s >= objetivo:
            return s
    return seccion_fase


def verificar_cortocircuito(seccion: float, icc_ka: float, tiempo_s: float,
                             conductor: str, aislamiento: str):
    """S_min = Icc*sqrt(t) / k (mm2). Devuelve (cumple, seccion_minima)."""
    k = K_CORTOCIRCUITO[(conductor, aislamiento)]
    if seccion > 300:
        if (conductor, aislamiento) == ("Cobre", "PVC"):
            k = 103
        elif (conductor, aislamiento) == ("Aluminio", "PVC"):
            k = 68
    icc_a = icc_ka * 1000.0
    s_min = (icc_a * math.sqrt(tiempo_s)) / k
    return seccion >= s_min, round(s_min, 2)


def calibre_magnetotermico_sugerido(ib_calculo: float) -> int:
    """Menor calibre normalizado de interruptor automático >= Ib de cálculo."""
    for c in CALIBRES_MAGNETOTERMICO:
        if c >= ib_calculo:
            return c
    return CALIBRES_MAGNETOTERMICO[-1]


# ==============================================================================
# 3. ESTILO VISUAL (CSS inyectado)
# ==============================================================================

def generar_css(tema: str = "Oscuro") -> str:
    """Genera el CSS de la app según el tema. Mantiene TODAS las clases ya
    usadas en el resto del código (titleblock, section-label, result-card,
    result-value, badge-ok/fail, regla-chip...) y añade el sistema nuevo
    (sidebar, dashboard, tarjetas de estado, badges, ayuda contextual)."""
    if tema == "Claro":
        v = dict(
            bg_primary="#f4f6fb", bg_panel="#ffffff", bg_panel_alt="#eef1f8", bg_sidebar="#111827",
            border_subtle="rgba(37, 99, 235, 0.14)", border_strong="rgba(37, 99, 235, 0.35)",
            text_primary="#0f172a", text_secondary="#5b6577", text_on_dark="#e5e9f5",
            accent="#2563eb", accent_hover="#1d4ed8", accent_soft="rgba(37, 99, 235, 0.10)",
            copper="#b3711f", success="#16a34a", success_soft="rgba(22,163,74,0.10)",
            warning="#d97706", warning_soft="rgba(217,119,6,0.10)",
            error="#dc2626", error_soft="rgba(220,38,38,0.10)",
            grid_line="rgba(15, 23, 42, 0.035)", shadow_sm="0 1px 2px rgba(15,23,42,0.06)",
            shadow_md="0 6px 16px rgba(15,23,42,0.08)", shadow_lg="0 16px 40px rgba(15,23,42,0.14)",
        )
    else:
        v = dict(
            bg_primary="#0b1220", bg_panel="#121b2e", bg_panel_alt="#16213a", bg_sidebar="#0a0f1a",
            border_subtle="rgba(59, 130, 246, 0.16)", border_strong="rgba(59, 130, 246, 0.40)",
            text_primary="#e8edf4", text_secondary="#8b96a8", text_on_dark="#e8edf4",
            accent="#3b82f6", accent_hover="#60a5fa", accent_soft="rgba(59, 130, 246, 0.14)",
            copper="#e8a33d", success="#22c55e", success_soft="rgba(34,197,94,0.14)",
            warning="#f59e0b", warning_soft="rgba(245,158,11,0.14)",
            error="#ef4444", error_soft="rgba(239,68,68,0.14)",
            grid_line="rgba(232, 237, 244, 0.035)", shadow_sm="0 1px 2px rgba(0,0,0,0.3)",
            shadow_md="0 6px 16px rgba(0,0,0,0.35)", shadow_lg="0 16px 40px rgba(0,0,0,0.45)",
        )
    return f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600;700&family=Inter:wght@400;500;600;700;800&display=swap');

:root {{
    --bg-primary: {v['bg_primary']}; --bg-panel: {v['bg_panel']}; --bg-panel-alt: {v['bg_panel_alt']};
    --bg-sidebar: {v['bg_sidebar']};
    --border-subtle: {v['border_subtle']}; --border-strong: {v['border_strong']};
    --text-primary: {v['text_primary']}; --text-secondary: {v['text_secondary']};
    --text-on-dark: {v['text_on_dark']};
    --accent-copper: {v['copper']};
    --accent-primary: {v['accent']}; --accent-primary-hover: {v['accent_hover']}; --accent-soft: {v['accent_soft']};
    --accent-ok: {v['success']}; --success-soft: {v['success_soft']};
    --accent-warning: {v['warning']}; --warning-soft: {v['warning_soft']};
    --accent-fail: {v['error']}; --error-soft: {v['error_soft']};
    --grid-line: {v['grid_line']};
    --shadow-sm: {v['shadow_sm']}; --shadow-md: {v['shadow_md']}; --shadow-lg: {v['shadow_lg']};
    --radius: 14px; --radius-sm: 8px;
}}

.stApp {{
    background-color: var(--bg-primary);
    background-image:
        linear-gradient(var(--grid-line) 1px, transparent 1px),
        linear-gradient(90deg, var(--grid-line) 1px, transparent 1px);
    background-size: 26px 26px;
}}

html, body, [class*="css"], .stMarkdown, p, span, label, div {{
    font-family: 'Inter', 'IBM Plex Sans', sans-serif;
    color: var(--text-primary);
}}
h1, h2, h3, h4 {{ font-family: 'Inter', sans-serif; letter-spacing: -0.02em; font-weight: 700; }}
code, .stCode, [data-testid="stMetricValue"] {{ font-family: 'JetBrains Mono', monospace; }}

/* ============ Sidebar (navegación SaaS) ============ */
section[data-testid="stSidebar"] {{
    background: var(--bg-sidebar);
    border-right: 1px solid var(--border-subtle);
}}
section[data-testid="stSidebar"] * {{ color: var(--text-on-dark) !important; }}
.sidebar-brand {{
    display: flex; align-items: center; gap: 0.55rem;
    padding: 0.9rem 0.2rem 1.1rem 0.2rem;
    border-bottom: 1px solid rgba(255,255,255,0.08);
    margin-bottom: 0.8rem;
}}
.sidebar-brand .logo {{
    width: 34px; height: 34px; border-radius: 9px;
    background: linear-gradient(135deg, var(--accent-primary), var(--accent-copper));
    display: flex; align-items: center; justify-content: center;
    font-size: 1.15rem; box-shadow: var(--shadow-sm); flex-shrink: 0;
}}
.sidebar-brand .name {{ font-weight: 800; font-size: 1.02rem; line-height: 1.15; }}
.sidebar-brand .sub {{ font-size: 0.68rem; color: #93a0b8 !important; letter-spacing: 0.04em; }}
.nav-group-label {{
    font-size: 0.66rem; font-weight: 700; letter-spacing: 0.12em; text-transform: uppercase;
    color: #6b7690 !important; margin: 1.0rem 0 0.3rem 0.15rem;
}}
section[data-testid="stSidebar"] .stButton button {{
    background: transparent; border: 1px solid transparent; text-align: left; justify-content: flex-start;
    font-weight: 500; border-radius: var(--radius-sm); padding: 0.45rem 0.7rem;
    transition: background 0.15s ease, border-color 0.15s ease; box-shadow: none;
}}
section[data-testid="stSidebar"] .stButton button:hover {{
    background: rgba(255,255,255,0.06); border-color: rgba(255,255,255,0.08);
}}
section[data-testid="stSidebar"] .stButton button[kind="primary"] {{
    background: var(--accent-soft); border: 1px solid var(--accent-primary);
    color: var(--accent-primary-hover) !important; font-weight: 700;
}}
section[data-testid="stSidebar"] .stButton button[kind="primary"] * {{ color: var(--accent-primary-hover) !important; }}

.sidebar-credit {{
    margin-top: 1rem; padding-top: 0.8rem; border-top: 1px solid rgba(255,255,255,0.08);
    text-align: center;
}}
.sidebar-credit-name {{
    font-size: 0.66rem; color: #6b7690 !important; letter-spacing: 0.02em; margin-bottom: 0.35rem;
}}
.sidebar-credit-links {{ display: flex; justify-content: center; gap: 0.7rem; }}
.sidebar-credit-links a {{
    font-size: 1.0rem; text-decoration: none; opacity: 0.7; transition: opacity 0.15s ease, transform 0.15s ease;
}}
.sidebar-credit-links a:hover {{ opacity: 1; transform: translateY(-1px); }}

/* ============ Cabecera tipo cajetín (documentos) ============ */
.titleblock {{
    display: flex; justify-content: space-between; align-items: stretch;
    border: 1px solid var(--border-strong);
    background: linear-gradient(180deg, var(--bg-panel), var(--bg-panel-alt));
    border-radius: var(--radius); margin-bottom: 1.6rem; overflow: hidden; box-shadow: var(--shadow-sm);
}}
.titleblock-main {{ padding: 1.1rem 1.4rem; flex: 1; }}
.titleblock-eyebrow {{
    font-family: 'JetBrains Mono', monospace; color: var(--accent-primary);
    font-size: 0.72rem; letter-spacing: 0.14em; text-transform: uppercase;
}}
.titleblock-main h1 {{ color: var(--text-primary); font-size: 1.55rem; margin: 0.15rem 0 0 0; font-weight: 800; }}
.titleblock-meta {{ display: flex; }}
.titleblock-meta > div {{
    border-left: 1px solid var(--border-subtle); padding: 0.9rem 1.1rem; min-width: 118px;
    display: flex; flex-direction: column; justify-content: center;
}}
.titleblock-meta span {{
    font-family: 'JetBrains Mono', monospace; font-size: 0.62rem; color: var(--text-secondary);
    letter-spacing: 0.12em; text-transform: uppercase;
}}
.titleblock-meta strong {{ font-family: 'JetBrains Mono', monospace; color: var(--text-primary); font-size: 0.82rem; margin-top: 0.15rem; }}

.section-label {{
    font-family: 'JetBrains Mono', monospace; font-size: 0.72rem; letter-spacing: 0.1em; text-transform: uppercase;
    color: var(--accent-primary); border-bottom: 1px solid var(--border-subtle);
    padding-bottom: 0.35rem; margin: 0.4rem 0 0.9rem 0;
}}

/* ============ Tarjetas de resultado (ya usadas en toda la app) ============ */
.result-card {{
    background: var(--bg-panel); border: 1px solid var(--border-subtle); border-radius: var(--radius);
    padding: 0.9rem 1.1rem; margin-bottom: 0.6rem; box-shadow: var(--shadow-sm);
    transition: transform 0.15s ease, box-shadow 0.15s ease;
}}
.result-card:hover {{ box-shadow: var(--shadow-md); transform: translateY(-1px); }}
.result-card.hero {{
    border: 1px solid var(--border-strong);
    background: linear-gradient(135deg, var(--accent-soft), var(--bg-panel));
}}
.result-label {{
    font-family: 'JetBrains Mono', monospace; font-size: 0.68rem; letter-spacing: 0.08em;
    text-transform: uppercase; color: var(--text-secondary);
}}
.result-value {{ font-family: 'JetBrains Mono', monospace; font-size: 1.7rem; font-weight: 700; color: var(--accent-primary); line-height: 1.3; }}
.result-value.small {{ font-size: 1.15rem; color: var(--text-primary); }}
.result-sub {{ font-size: 0.78rem; color: var(--text-secondary); margin-top: 0.15rem; }}
.badge-ok {{ color: var(--accent-ok); font-weight: 600; }}
.badge-fail {{ color: var(--accent-fail); font-weight: 600; }}

.regla-wrap {{ display: flex; gap: 4px; margin: 0.6rem 0 1rem 0; flex-wrap: wrap; }}
.regla-chip {{
    font-family: 'JetBrains Mono', monospace; font-size: 0.72rem; padding: 0.28rem 0.5rem; border-radius: 999px;
    border: 1px solid var(--border-subtle); color: var(--text-secondary); background: var(--bg-panel);
}}
.regla-chip.activa {{ border-color: var(--accent-primary); color: #fff; background: var(--accent-primary); font-weight: 700; }}

/* ============ Dashboard: tarjetas KPI, accesos rápidos, timeline ============ */
.kpi-card {{
    background: var(--bg-panel); border: 1px solid var(--border-subtle); border-radius: var(--radius);
    padding: 1.1rem 1.3rem; box-shadow: var(--shadow-sm); height: 100%;
    transition: transform 0.15s ease, box-shadow 0.15s ease;
}}
.kpi-card:hover {{ transform: translateY(-2px); box-shadow: var(--shadow-md); }}
.kpi-icon {{
    width: 40px; height: 40px; border-radius: 11px; display: flex; align-items: center; justify-content: center;
    font-size: 1.2rem; margin-bottom: 0.6rem;
}}
.kpi-icon.blue {{ background: var(--accent-soft); }}
.kpi-icon.green {{ background: var(--success-soft); }}
.kpi-icon.orange {{ background: var(--warning-soft); }}
.kpi-icon.red {{ background: var(--error-soft); }}
.kpi-value {{ font-size: 1.7rem; font-weight: 800; color: var(--text-primary); line-height: 1.15; }}
.kpi-label {{ font-size: 0.8rem; color: var(--text-secondary); margin-top: 0.15rem; }}

.quick-card {{
    display: block; background: var(--bg-panel); border: 1px solid var(--border-subtle);
    border-radius: var(--radius); padding: 1rem 1.1rem; text-decoration: none !important;
    transition: all 0.15s ease; height: 100%; box-shadow: var(--shadow-sm);
}}
.quick-card:hover {{ border-color: var(--accent-primary); box-shadow: var(--shadow-md); transform: translateY(-2px); }}
.quick-card .qc-icon {{ font-size: 1.4rem; margin-bottom: 0.4rem; }}
.quick-card .qc-title {{ font-weight: 700; color: var(--text-primary) !important; font-size: 0.92rem; }}
.quick-card .qc-sub {{ font-size: 0.76rem; color: var(--text-secondary) !important; margin-top: 0.15rem; }}

.timeline-item {{
    display: flex; gap: 0.7rem; padding: 0.55rem 0; border-bottom: 1px solid var(--border-subtle);
    font-size: 0.85rem;
}}
.timeline-item:last-child {{ border-bottom: none; }}
.timeline-dot {{ width: 7px; height: 7px; border-radius: 50%; background: var(--accent-primary); margin-top: 0.4rem; flex-shrink: 0; }}
.timeline-time {{ color: var(--text-secondary); font-size: 0.72rem; font-family: 'JetBrains Mono', monospace; }}

/* ============ Badges de estado (éxito/aviso/error/info) ============ */
.badge {{
    display: inline-flex; align-items: center; gap: 0.3rem; padding: 0.2rem 0.6rem; border-radius: 999px;
    font-size: 0.72rem; font-weight: 700; letter-spacing: 0.02em;
}}
.badge.success {{ background: var(--success-soft); color: var(--accent-ok); }}
.badge.warning {{ background: var(--warning-soft); color: var(--accent-warning); }}
.badge.error {{ background: var(--error-soft); color: var(--accent-fail); }}
.badge.info {{ background: var(--accent-soft); color: var(--accent-primary); }}

/* ============ Botones y widgets nativos ============ */
.stButton button {{
    border-radius: var(--radius-sm); font-weight: 600; transition: all 0.15s ease; box-shadow: var(--shadow-sm);
}}
.stButton button:hover {{ transform: translateY(-1px); box-shadow: var(--shadow-md); }}
.stButton button[kind="primary"] {{ background: var(--accent-primary); border-color: var(--accent-primary); }}
.stButton button[kind="primary"]:hover {{ background: var(--accent-primary-hover); }}
.stDownloadButton button {{ border-radius: var(--radius-sm); font-weight: 600; box-shadow: var(--shadow-sm); }}

[data-testid="stMetricValue"] {{ color: var(--accent-primary); }}
[data-testid="stExpander"] {{
    border: 1px solid var(--border-subtle); border-radius: var(--radius); box-shadow: var(--shadow-sm);
    overflow: hidden;
}}
[data-testid="stVerticalBlockBorderWrapper"] {{ border-radius: var(--radius) !important; }}
hr {{ border-color: var(--border-subtle); }}
.stTabs [data-baseweb="tab-list"] {{ gap: 4px; }}
.stTabs [data-baseweb="tab"] {{ border-radius: var(--radius-sm) var(--radius-sm) 0 0; }}

/* ============ Ayuda contextual ============ */
.ayuda-texto {{ font-size: 0.82rem; line-height: 1.5; color: var(--text-secondary); }}
.ayuda-texto b {{ color: var(--text-primary); }}

/* ============ Guía de bienvenida ============ */
.welcome-banner {{
    background: linear-gradient(135deg, var(--accent-soft), var(--bg-panel));
    border: 1px solid var(--border-strong);
    border-radius: var(--radius);
    padding: 1.5rem 1.7rem 1rem 1.7rem;
    margin: 0.8rem 0 1.4rem 0;
    box-shadow: var(--shadow-md);
}}
.welcome-banner h4 {{ margin-top: 0; }}

/* ============ Accesibilidad ============ */
/* Foco visible y consistente para navegación por teclado (no solo :hover) */
a:focus-visible, button:focus-visible, [role="button"]:focus-visible,
input:focus-visible, select:focus-visible, textarea:focus-visible,
.stButton button:focus-visible, .stDownloadButton button:focus-visible {{
    outline: 2px solid var(--accent-primary) !important;
    outline-offset: 2px !important;
}}
/* Tamaño mínimo de objetivo táctil/clic recomendado (WCAG 2.5.5, ~44px) */
.stButton button, .stDownloadButton button {{ min-height: 2.5rem; }}
/* Respeta la preferencia del sistema de reducir animaciones */
@media (prefers-reduced-motion: reduce) {{
    *, *::before, *::after {{
        animation-duration: 0.001ms !important;
        transition-duration: 0.001ms !important;
    }}
    .result-card:hover, .kpi-card:hover, .quick-card:hover, .stButton button:hover {{
        transform: none !important;
    }}
}}

/* ============ Responsive ============ */
@media (max-width: 900px) {{
    .titleblock {{ flex-direction: column; }}
    .titleblock-meta {{ border-top: 1px solid var(--border-subtle); }}
    .kpi-value {{ font-size: 1.35rem; }}
    .welcome-banner {{ padding: 1.1rem 1.2rem 0.6rem 1.2rem; }}
}}
</style>
"""


# ==============================================================================
# 3B. INFRAESTRUCTURA DE LA APP — estado de sesión, ayuda contextual,
# actividad reciente y sistema de proyectos (guardar/abrir como archivo).
#
# Nota de arquitectura: esta app no tiene backend propio ni base de datos.
# Desplegada en Streamlit Cloud, el sistema de archivos es efímero (se borra
# en cada reinicio), así que CUALQUIER "base de datos" ahí sería una falsa
# promesa de persistencia. En su lugar: session_state para lo que dura la
# sesión del navegador, y descarga/carga de un .json real para lo que debe
# sobrevivir entre sesiones (eso sí es 100% fiable, lo guarda el usuario).
# ==============================================================================

PAGINAS_HERRAMIENTAS = ["Calculadora", "Fórmulas", "Fotovoltaica", "Presupuesto", "Documentación"]
CLAVES_PROYECTO = ["inputs_cable", "resultado_cable", "inputs_fv", "resultado_fv",
                   "presupuesto_capitulos", "presupuesto_config", "datos_proyecto", "catalogo_precios",
                   "partidas_compuestas", "calculos_guardados", "escenarios_fv_guardados",
                   "checklist_puesta_servicio", "checklist_firma", "fotos_instalacion"]
MAX_HISTORIAL_PROYECTOS = 25  # evita que la sesión acumule memoria sin límite


def _inicializar_estado():
    defaults = {
        "pagina_actual": "Inicio",
        "tema": "Oscuro",
        "config_profesional": {"nombre": "", "empresa": "", "logo_b64": "", "firma": ""},
        "historial_proyectos": [],
        "actividad": [],
        "inputs_cable": {}, "resultado_cable": {}, "inputs_fv": {}, "resultado_fv": {},
        "nombre_proyecto_actual": "Proyecto sin guardar",
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)


def _registrar_actividad(icono: str, texto: str):
    st.session_state.setdefault("actividad", [])
    st.session_state["actividad"].insert(0, {
        "icono": icono, "texto": texto, "hora": datetime.now().strftime("%H:%M"),
    })
    st.session_state["actividad"] = st.session_state["actividad"][:30]


def _ayuda(texto: str, titulo: str = "Ayuda"):
    """Botón de ayuda contextual (?) — usar junto a un campo o sección."""
    with st.popover("❓"):
        st.markdown(f"**{titulo}**")
        st.markdown(f'<div class="ayuda-texto">{texto}</div>', unsafe_allow_html=True)


def _campo_con_ayuda(etiqueta: str, texto_ayuda: str, titulo: str = "Ayuda"):
    """Muestra una etiqueta con un botón (?) al lado, en una fila compacta."""
    c1, c2 = st.columns([8, 1])
    with c1:
        st.markdown(f"**{etiqueta}**")
    with c2:
        _ayuda(texto_ayuda, titulo)


def _estado_vacio(mensaje: str, pagina_destino: str, texto_boton: str, icono: str = "👋"):
    """Estado vacío uniforme: mensaje amistoso + botón que lleva directo a
    donde hay que ir, en vez de un aviso pasivo que hay que interpretar."""
    st.markdown(f'''<div class="result-card" style="text-align:center; padding:2rem 1.5rem;">
        <div style="font-size:2rem; margin-bottom:0.5rem;">{icono}</div>
        <div style="color:var(--text-secondary); margin-bottom:1rem;">{mensaje}</div>
    </div>''', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c2:
        if st.button(texto_boton, type="primary", width='stretch', key=f"estado_vacio_{pagina_destino}"):
            st.session_state["pagina_actual"] = pagina_destino
            st.rerun()


def _serializar_proyecto(nombre: str) -> dict:
    datos = {"__version__": 1, "__nombre__": nombre, "__fecha__": datetime.now().isoformat()}
    for clave in CLAVES_PROYECTO:
        datos[clave] = st.session_state.get(clave)
    return datos


def _cargar_proyecto(datos: dict):
    for clave in CLAVES_PROYECTO:
        if clave in datos:
            st.session_state[clave] = datos[clave]
    st.session_state["nombre_proyecto_actual"] = datos.get("__nombre__", "Proyecto importado")


def _tamano_proyecto_kb(datos: dict) -> float:
    return len(json.dumps(datos, default=str).encode("utf-8")) / 1024


# ==============================================================================
# 4. ORQUESTADOR DE CÁLCULO (puro: dict de entradas -> dict de resultados)
# ==============================================================================

def calcular(inp: dict) -> dict:
    """Ejecuta los 3 criterios de la ITC-BT-19 art. 19.2 a partir del dict de
    entradas producido por _render_inputs(). No toca Streamlit: es testeable
    de forma aislada."""
    avisos = []
    sistema = inp["sistema"]
    n_cargados = 2 if sistema == SISTEMA_MONO else 3

    if inp["modo_entrada"] == "Potencia activa":
        ib = calcular_intensidad_empleo(sistema, inp["potencia_kw"] * 1000.0, inp["tension"], inp["cos_phi"])
    else:
        ib = inp["intensidad_directa"]

    ib_calculo = ib
    ib_motor = None
    if inp["es_motor"] and inp["corrientes_motores"]:
        corrientes = inp["corrientes_motores"]
        if len(corrientes) == 1:
            ib_motor = factor_motor_unico(corrientes[0])
        else:
            ib_motor = factor_varios_motores(corrientes)
        if inp["ascensor_grua"]:
            ib_motor *= 1.3
        ib_calculo = max(ib_calculo, ib_motor)
        avisos.append(f"Corriente de cálculo ajustada según ITC-BT-47 (motores): {ib_motor:.2f} A.")

    if inp["alumbrado_descarga"]:
        ib_calculo *= 1.8
        avisos.append("Corriente de cálculo incrementada ×1,8 por alumbrado de descarga sin corregir (orientativo).")

    f_temp = factor_correccion_temperatura(inp["aislamiento"], inp["temp_ambiente"], inp["enterrado"])
    f_agrup = factor_correccion_agrupamiento(inp["disposicion"], int(inp["n_circuitos"]))
    f_capas = factor_correccion_capas(int(inp["n_capas"]))
    f_resist = factor_correccion_resistividad(inp["resistividad"]) if inp["enterrado"] else 1.0
    factor_total = f_temp * f_agrup * f_capas * f_resist

    if f_temp <= 0.05:
        avisos.append("La temperatura ambiente introducida iguala o supera la temperatura máxima de servicio "
                       "del aislamiento elegido: instalación no viable en estas condiciones.")

    s_termica, iz_termica, necesita_paralelo, n_paralelo = seccion_por_criterio_termico(
        ib_calculo, inp["metodo"], inp["aislamiento"], inp["conductor"], n_cargados, factor_total)

    kappa = kappa_servicio(inp["conductor"], inp["aislamiento"], inp["usar_kappa_20c"])
    ib_paralelo = ib_calculo / n_paralelo if necesita_paralelo else ib_calculo
    s_du, e_voltios, e_pct = seccion_por_caida_tension(
        sistema, ib_paralelo, inp["longitud"], inp["tension"], inp["cos_phi"], kappa,
        inp["delta_u_max"], inp["conductor"])

    candidatos = [s for s in (s_termica, s_du) if s is not None]
    seccion_final = max(candidatos) if candidatos else None
    if s_termica is not None and s_du is None:
        avisos.append("Ni siquiera 300 mm² cumple el criterio de caída de tensión con la longitud indicada: "
                       "valora más conductores en paralelo, reducir la longitud o elevar el nivel de tensión.")

    if seccion_final is not None:
        e_final = caida_tension_voltios(sistema, ib_paralelo, inp["longitud"], seccion_final, inp["cos_phi"], kappa)
        e_final_pct = e_final / inp["tension"] * 100.0
    else:
        e_final, e_final_pct = e_voltios, e_pct

    seccion_neutro = None
    if sistema == SISTEMA_TRI and seccion_final is not None:
        seccion_neutro = seccion_conductor_neutro(seccion_final, inp["conductor"], inp["armonicos"])

    seccion_proteccion = seccion_conductor_proteccion(seccion_final) if seccion_final is not None else None

    if seccion_final is not None:
        if inp["conductor"] == "Aluminio" and seccion_final < MINIMO_ALUMINIO_MM2:
            seccion_final = MINIMO_ALUMINIO_MM2
            avisos.append("Sección elevada a 16 mm² por ser el mínimo normativo del aluminio en instalación "
                           "fija (ITC-BT-07).")
        if inp["tipo_circuito"].startswith("LGA"):
            minimo_lga = MINIMO_LGA_ALUMINIO_MM2 if inp["conductor"] == "Aluminio" else MINIMO_LGA_COBRE_MM2
            if seccion_final < minimo_lga:
                seccion_final = minimo_lga
                avisos.append(f"Sección elevada a {minimo_lga:g} mm² por ser el mínimo normativo de la LGA "
                               "(ITC-BT-14).")

    cumple_cc, s_min_cc = None, None
    if inp["verificar_cc"] and seccion_final is not None:
        cumple_cc, s_min_cc = verificar_cortocircuito(seccion_final, inp["icc_ka"], inp["tiempo_s"],
                                                        inp["conductor"], inp["aislamiento"])
        if not cumple_cc:
            avisos.append(f"La sección adoptada no soporta térmicamente el cortocircuito indicado: se "
                           f"necesitarían ≥ {s_min_cc:g} mm² (o una protección más rápida/limitadora).")

    calibre = calibre_magnetotermico_sugerido(ib_calculo)

    return dict(
        n_cargados=n_cargados, ib=ib, ib_calculo=ib_calculo, ib_motor=ib_motor,
        f_temp=f_temp, f_agrup=f_agrup, f_capas=f_capas, f_resist=f_resist, factor_total=factor_total,
        s_termica=s_termica, iz_termica=iz_termica, necesita_paralelo=necesita_paralelo, n_paralelo=n_paralelo,
        kappa=kappa, ib_paralelo=ib_paralelo, s_du=s_du, e_voltios=e_voltios, e_pct=e_pct,
        seccion_final=seccion_final, e_final=e_final, e_final_pct=e_final_pct,
        seccion_neutro=seccion_neutro, seccion_proteccion=seccion_proteccion,
        cumple_cc=cumple_cc, s_min_cc=s_min_cc, calibre_magnetotermico=calibre, avisos=avisos,
    )


# ==============================================================================
# 5. GENERACIÓN DE PDF (memoria de cálculo con cajetín tipo plano técnico)
# ==============================================================================
# Nota sobre codificación: las fuentes estándar de reportlab (Helvetica) usan
# WinAnsiEncoding, que soporta vocales acentuadas y ñ sin problema, pero NO
# incluye de forma fiable símbolos como Δ, ≤, ≥, √ o ÷. Para evitar glifos
# rotos en el PDF, se sanean esos símbolos a su equivalente ASCII antes de
# escribir cualquier texto; el resto de la app (interfaz Streamlit) sigue
# usando los símbolos Unicode con normalidad, ya que el navegador sí los
# renderiza correctamente.

_PDF_REPLACEMENTS = {
    "Δ": "Delta ", "≤": "<=", "≥": ">=", "√": "raiz", "÷": "/",
    "²": "2", "³": "3", "°": "gr.", "·": "-", "×": "x",
    "—": "-", "–": "-", "✅": "[OK]", "❌": "[NO]", "⚠️": "[AVISO]", "☐": "[ ]",
    "⬇️": "", "🔌": "", "📊": "", "📖": "", "⚙️": "", "🌡️": "", "⚡": "",
}

# Caracteres Unicode de super/subíndice que Helvetica NO tiene como glifo
# real (se comprobó con stringWidth: todos devuelven el mismo ancho de
# "carácter no encontrado"). Se sustituyen por la etiqueta <super>/<sub> de
# reportlab con el dígito normal correspondiente — ESO sí tiene glifo real,
# porque reportlab dibuja el dígito normal más pequeño y desplazado, no
# busca un glifo de superíndice en la fuente. Solo válido dentro de un
# Paragraph (no en celdas de tabla con texto plano).
_PDF_SUPERSCRIPT_SAFE = {
    "⁻": "<super>-</super>", "⁴": "<super>4</super>", "⁵": "<super>5</super>",
    "⁰": "<super>0</super>", "¹": "<super>1</super>", "⁶": "<super>6</super>",
    "⁷": "<super>7</super>", "⁸": "<super>8</super>", "⁹": "<super>9</super>",
}
_PDF_SUBSCRIPT_SAFE = {
    "₀": "<sub>0</sub>", "₁": "<sub>1</sub>", "₂": "<sub>2</sub>", "₃": "<sub>3</sub>",
    "₄": "<sub>4</sub>", "₅": "<sub>5</sub>",
}


def _pdf_safe_markup(texto: str) -> str:
    """Como _pdf_safe, pero para texto que SÍ va dentro de un Paragraph (o
    de una celda de _tabla_pdf, que también usa Paragraph): además de los
    reemplazos habituales, convierte super/subíndices Unicode sin glifo en
    <super>/<sub> con dígitos normales, que sí se ven."""
    texto = str(texto)
    for k, v in _PDF_SUPERSCRIPT_SAFE.items():
        texto = texto.replace(k, v)
    for k, v in _PDF_SUBSCRIPT_SAFE.items():
        texto = texto.replace(k, v)
    for k, v in _PDF_REPLACEMENTS.items():
        if k in ("²", "³", "·", "×"):  # estos SÍ tienen glifo real, no hace falta sustituirlos aquí
            continue
        texto = texto.replace(k, v)
    return texto


def _pdf_safe(texto: str) -> str:
    texto = str(texto)
    for k, v in _PDF_REPLACEMENTS.items():
        texto = texto.replace(k, v)
    return texto


def _tabla_pdf(datos, colw, AZUL, colors, fuente=8.5, fuente_header=8.8, header=True, alinear_num=None,
                extra_estilos=None, negrita_col0=False):
    """Tabla única y robusta para todos los documentos PDF. Cada celda es un
    Paragraph (no texto plano), así que SIEMPRE ajusta su contenido al
    ancho de columna — nunca se sale del recuadro ni se corta a media
    palabra, sea cual sea el número de columnas o lo largo del texto.
    También permite <sub>/<super> en el contenido vía _pdf_safe_markup.
    alinear_num: lista de índices de columna a alinear a la derecha (para
    cifras). extra_estilos: lista adicional de tuplas TableStyle.
    negrita_col0: pinta la primera columna en negrita/azul (nota: al ser
    Paragraph, esto NO se puede hacer con FONTNAME de TableStyle — hace
    falta un ParagraphStyle propio para esa columna)."""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm as _cm

    estilo_celda = ParagraphStyle("celda_tabla", fontName="Helvetica", fontSize=fuente,
                                   leading=fuente * 1.28, textColor=colors.HexColor("#1a2433"))
    estilo_col0_negrita = ParagraphStyle("celda_col0_negrita", parent=estilo_celda, fontName="Helvetica-Bold",
                                          textColor=AZUL)
    estilo_header = ParagraphStyle("celda_header", fontName="Helvetica-Bold", fontSize=fuente_header,
                                    leading=fuente_header * 1.25, textColor=colors.white)
    alinear_num = alinear_num or []

    filas_wrapped = []
    for i_fila, fila in enumerate(datos):
        es_header = header and i_fila == 0
        fila_w = []
        for i_col, valor in enumerate(fila):
            if es_header:
                estilo = estilo_header
            elif negrita_col0 and i_col == 0:
                estilo = estilo_col0_negrita
            else:
                estilo = estilo_celda
            fila_w.append(Paragraph(_pdf_safe_markup(valor), estilo))
        filas_wrapped.append(fila_w)

    t = Table(filas_wrapped, colWidths=[c * _cm for c in colw], repeatRows=1 if header else 0)
    estilo_tabla = [
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9ccd1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if header:
        estilo_tabla.append(("BACKGROUND", (0, 0), (-1, 0), AZUL))
    for idx_col in alinear_num:
        estilo_tabla.append(("ALIGN", (idx_col, 0), (idx_col, -1), "RIGHT"))
    if extra_estilos:
        estilo_tabla.extend(extra_estilos)
    t.setStyle(TableStyle(estilo_tabla))
    return t


def generar_pdf_memoria(inp: dict, res: dict, config_prof: dict = None) -> bytes:
    from reportlab.platypus import Paragraph, Spacer

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "MEMORIA DE CALCULO - SECCION DE CONDUCTORES",
        f"Circuito: {inp['tipo_circuito']}", {"titular": (config_prof or {}).get("empresa", "")}, config_prof)
    from reportlab.lib.styles import ParagraphStyle
    aviso_style = ParagraphStyle("avisoc", parent=normal, textColor=colors.HexColor("#8a4b00"))

    story = _bloque_portada("Memoria de Cálculo", "Sección de conductores — ITC-BT-19",
                             {"titular": (config_prof or {}).get("empresa", "")}, config_prof, AZUL, COBRE,
                             colors, h1, normal)

    story.append(Paragraph("1. Datos del circuito", h2))
    datos = [
        ["Concepto", "Valor"],
        ["Tipo de circuito", inp["tipo_circuito"]],
        ["Sistema", inp["sistema"] + f"  -  {inp['tension']:g} V"],
        ["Conductor / Aislamiento", f"{inp['conductor']} / {inp['aislamiento']}"],
        ["Metodo de instalacion", inp["metodo"]],
        ["Longitud", f"{inp['longitud']:g} m"],
    ]
    story.append(_tabla_pdf(datos, (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("2. Criterio termico (Ib &lt;= In &lt;= Iz) - ITC-BT-19 art. 19", h2))
    termico = [
        ["Magnitud", "Valor"],
        ["Ib (corriente de empleo)", f"{res['ib']:.2f} A"],
        ["Ib de calculo (tras factores)", f"{res['ib_calculo']:.2f} A"],
        ["Factor de correccion total", f"{res['factor_total']:.3f}"],
        ["Seccion por criterio termico", f"{res['s_termica']:g} mm2" if res["s_termica"] else "-"],
        ["Iz obtenida (tabla x factores)", f"{res['iz_termica']:.1f} A" if res["iz_termica"] else "-"],
    ]
    story.append(_tabla_pdf(termico, (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("3. Criterio de caida de tension - ITC-BT-14/15/19/40", h2))
    du = [
        ["Magnitud", "Valor"],
        ["Delta U maxima admisible", f"{res['e_final_pct'] and inp['delta_u_max']:g} %"],
        ["Delta U con la seccion adoptada", f"{res['e_final_pct']:.2f} %"],
        ["Cumple", "SI" if res["e_final_pct"] <= inp["delta_u_max"] else "NO"],
    ]
    story.append(_tabla_pdf(du, (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("4. Seccion final adoptada", h2))
    final_rows = [["Elemento", "Seccion"]]
    if res["seccion_final"] is not None:
        final_rows.append(["Conductor de fase", f"{res['seccion_final']:g} mm2"])
    if res.get("seccion_neutro"):
        final_rows.append(["Conductor de neutro", f"{res['seccion_neutro']:g} mm2"])
    if res.get("seccion_proteccion"):
        final_rows.append(["Conductor de proteccion (PE)", f"{res['seccion_proteccion']:g} mm2"])
    if res.get("necesita_paralelo"):
        final_rows.append(["Conductores en paralelo", f"{res['n_paralelo']} x {res['seccion_final']:g} mm2 por fase"])
    final_rows.append(["Interruptor automatico sugerido", f"{res['calibre_magnetotermico']} A"])
    story.append(_tabla_pdf(final_rows, (6.5, 9.5), AZUL, colors))

    if res.get("cumple_cc") is not None:
        story.append(Paragraph("5. Verificacion termica de cortocircuito", h2))
        estado = "CUMPLE" if res["cumple_cc"] else "NO CUMPLE"
        cc_rows = [
            ["Magnitud", "Valor"],
            ["Resultado", estado],
            ["Seccion minima necesaria", f"{res['s_min_cc']:g} mm2"],
        ]
        story.append(_tabla_pdf(cc_rows, (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("6. Formulas aplicadas", h2))
    for linea in _lineas_formulas_texto(inp, res):
        story.append(Paragraph(_pdf_safe(linea), normal))

    if res.get("avisos"):
        story.append(Paragraph("7. Avisos", h2))
        for a in res["avisos"]:
            story.append(Paragraph("- " + _pdf_safe(a), aviso_style))

    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Herramienta de apoyo al diseno. Los valores de esta memoria deben verificarse contra la edicion "
        "vigente de la Guia-BT-19 y la normativa UNE aplicable antes de incorporarse a un proyecto o "
        "memoria tecnica firmada.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


# ==============================================================================
# 6. JUSTIFICACIÓN DE FÓRMULAS (texto plano, reutilizado por PDF y por la UI)
# ==============================================================================

def _lineas_formulas_texto(inp: dict, res: dict) -> list:
    """Genera la secuencia de fórmulas con los valores del cálculo actual
    sustituidos, en texto plano (sin LaTeX), para incluir en el PDF."""
    L = []
    sistema = inp["sistema"]
    cos_phi = inp["cos_phi"]

    if inp["modo_entrada"] == "Potencia activa":
        p_w = inp["potencia_kw"] * 1000.0
        if sistema == SISTEMA_MONO:
            L.append(f"Ib = P / (V x cos(phi)) = {p_w:.0f} / ({inp['tension']:g} x {cos_phi:.2f}) "
                      f"= {res['ib']:.2f} A")
        else:
            L.append(f"Ib = P / (raiz(3) x V x cos(phi)) = {p_w:.0f} / (1,732 x {inp['tension']:g} x "
                      f"{cos_phi:.2f}) = {res['ib']:.2f} A")
    else:
        L.append(f"Ib = {res['ib']:.2f} A (introducida directamente)")

    if res.get("ib_motor") is not None:
        if len(inp["corrientes_motores"]) == 1:
            L.append(f"Motor unico (ITC-BT-47): Ib_calculo = 1,25 x In = 1,25 x "
                      f"{inp['corrientes_motores'][0]:.2f} = {res['ib_motor']:.2f} A")
        else:
            ordenados = sorted(inp["corrientes_motores"], reverse=True)
            resto = " + ".join(f"{x:.1f}" for x in ordenados[1:])
            L.append(f"Varios motores (ITC-BT-47): Ib_calculo = 1,25 x {ordenados[0]:.2f} + ({resto}) "
                      f"= {res['ib_motor']:.2f} A")
        if inp["ascensor_grua"]:
            L.append("Ascensor/grua: factor adicional x1,3 aplicado sobre la intensidad del motor.")

    if inp["alumbrado_descarga"]:
        L.append("Alumbrado de descarga sin corregir f.d.p.: Ib_calculo se multiplica x1,8 (orientativo).")

    L.append(f"Ib de calculo final = {res['ib_calculo']:.2f} A")
    L.append("")
    L.append(f"Factor de correccion = f_temp x f_agrup x f_capas x f_resist = {res['f_temp']:.3f} x "
              f"{res['f_agrup']:.3f} x {res['f_capas']:.3f} x {res['f_resist']:.3f} = {res['factor_total']:.3f}")
    if res["iz_termica"] is not None:
        L.append(f"Iz = Iz_tabla x factor = {res['iz_termica'] / max(res['factor_total'], 1e-9):.1f} x "
                  f"{res['factor_total']:.3f} = {res['iz_termica']:.2f} A")
        cumple_termico = res["iz_termica"] >= (res["ib_calculo"] / (res["n_paralelo"] if res["necesita_paralelo"] else 1))
        L.append(f"Criterio termico: Iz = {res['iz_termica']:.2f} A >= Ib = {res['ib_calculo']:.2f} A "
                  f"-> {'Cumple' if cumple_termico else 'No cumple'} con S = {res['s_termica']:g} mm2")
    L.append("")

    kappa = res["kappa"]
    L.append(f"Conductividad a temperatura de servicio: kappa = {kappa:.2f} m/(ohm x mm2)")
    k_sist = "2" if sistema == SISTEMA_MONO else "raiz(3)=1,732"
    S = res["seccion_final"] if res["seccion_final"] else res["s_termica"]
    L.append(f"Delta U = {k_sist} x L x Ib x (R x cos(phi) + X x sen(phi)),  R = 1/(kappa x S)")
    if S:
        r = 1.0 / (kappa * S)
        L.append(f"R = 1 / ({kappa:.2f} x {S:g}) = {r:.5f} ohm/m")
        L.append(f"Delta U = {res['e_final']:.2f} V  =  {res['e_final_pct']:.2f} % de {inp['tension']:g} V")
        cumple_du = res["e_final_pct"] <= inp["delta_u_max"]
        L.append(f"Criterio de caida de tension: {res['e_final_pct']:.2f} % <= {inp['delta_u_max']:g} % "
                  f"-> {'Cumple' if cumple_du else 'No cumple'}")
    L.append("")

    if res.get("seccion_neutro"):
        L.append(f"Seccion de neutro: Sf={res['seccion_final']:g} mm2 -> Sn={res['seccion_neutro']:g} mm2 "
                  "(regla REBT / IEC 60364-5-52 punto 524)")
    if res.get("seccion_proteccion"):
        L.append(f"Seccion de proteccion (ITC-BT-18): Sf={res['seccion_final']:g} mm2 -> "
                  f"Sp={res['seccion_proteccion']:g} mm2")

    if res.get("cumple_cc") is not None:
        L.append("")
        k_cc = K_CORTOCIRCUITO[(inp["conductor"], inp["aislamiento"])]
        L.append(f"Cortocircuito: S_min = (Icc x raiz(t)) / k = ({inp['icc_ka']:g}x1000 x "
                  f"raiz({inp['tiempo_s']:g})) / {k_cc} = {res['s_min_cc']:g} mm2")

    return L


def _lineas_formulas_fv_texto(inp: dict, res: dict) -> list:
    """Justificación de cálculo del módulo fotovoltaico en texto plano, para
    el Anexo de cálculos (mismo espíritu que _lineas_formulas_texto)."""
    L = []
    L.append(f"Potencia pico = {res['p_pico_kwp']:.2f} kWp  ({res['n_paneles']} paneles de "
              f"{inp['potencia_panel_wp']:g} Wp segun dimensionado; superficie necesaria = "
              f"{res['superficie_necesaria_m2']:.1f} m2)")
    L.append(f"Perdidas por orientacion/inclinacion (CTE DB-HE5) = 100 x [1,2e-4 x (beta-phi+10)^2 + "
              f"3,5e-5 x alfa^2] , con beta={inp['inclinacion']:g}, phi={inp['latitud']:g}, "
              f"alfa={inp['azimut']:g}  =>  {res['perdidas_orient_pct']:.2f} %")
    L.append(f"PR efectivo = PR_base x (1-perd.orient) x (1-perd.sombras) x eficiencia_inversor = "
              f"{inp['pr']:.2f} x (1-{res['perdidas_orient_pct']:.2f}/100) x "
              f"(1-{inp['perdidas_sombras']:g}/100) x {inp['eficiencia_inversor']/100:.3f} = "
              f"{res['pr_efectivo']:.3f}")
    L.append(f"Produccion anual = Ppico x HSP x 365 x PR_efectivo = {res['p_pico_kwp']:.2f} x "
              f"{inp['hsp']:g} x 365 x {res['pr_efectivo']:.3f} = "
              f"{res['produccion_anual_kwh']:,.0f} kWh/anio".replace(",", "."))
    L.append(f"Produccion a 10 anios (degradacion {inp['degradacion_anual']:g}%/anio) = "
              f"{res['produccion_ano10']:,.0f} kWh/anio".replace(",", "."))
    L.append(f"Produccion a 25 anios = {res['produccion_ano25']:,.0f} kWh/anio".replace(",", "."))
    L.append(f"CO2 evitado = Produccion x factor_red = {res['produccion_anual_kwh']:.0f} x "
              f"{inp.get('factor_co2', FACTOR_CO2_RED_DEFECTO):.2f} = "
              f"{res['co2_evitado_kg_ano']:.0f} kg CO2/anio ({res['co2_evitado_kg_ano']/1000:.2f} t/anio)")
    L.append("")
    L.append(f"Configuracion string: {res['n_serie']} paneles serie x {res['n_paralelo']} strings paralelo "
              f"= {res['n_paneles_configurados']} paneles")
    L.append(f"V string en frio = Nserie x Voc x (1 + coef x (25 - Tmin)) = {res['n_serie']} x {inp['voc']:g} "
              f"x (1 + {inp['coef_temp_voc']/100:.4f} x (25-{inp['temp_min']:g})) = {res['v_string_frio']:.1f} V")
    L.append(f"V string en caliente = Nserie x Vmp x (1 + coef x (Tcel-25)) = {res['v_string_caliente']:.1f} V")
    L.append(f"Verificacion ventana MPPT: {inp['vmin_mppt']:g} V <= V string <= {inp['vmax_mppt']:g} V -> "
              f"{'Cumple' if (res['cumple_vmpp_min'] and res['cumple_vmpp_max']) else 'No cumple'}")
    L.append(f"Verificacion tension max. entrada inversor: V string frio ({res['v_string_frio']:.1f} V) <= "
              f"{inp['vmax_entrada_inversor']:g} V -> {'Cumple' if res['cumple_vmax'] else 'No cumple'}")
    L.append("")
    L.append(f"Cableado CC (ITC-BT-40): I diseno = 1,25 x Isc = 1,25 x {inp['isc']:g} = {res['i_diseno_cc']:.2f} A")
    L.append(f"Seccion CC adoptada = {res['s_cc_final']:g} mm2 Cu XLPE  ->  ΔU parcial CC = {res['du_cc_pct']:.2f} %")
    L.append(f"Cableado CA inversor-cuadro: Ib = {res['ib_ca']:.2f} A ; I diseno (125%) = {res['i_diseno_ca']:.2f} A")
    L.append(f"Seccion CA adoptada = {res['s_ca_final']:g} mm2 Cu XLPE  ->  ΔU parcial CA = {res['e_ca_pct']:.2f} %")
    L.append(f"ΔU combinada CC+CA = {res['du_total_pct']:.2f} %  (limite ITC-BT-40: 1,5% entre generador y "
              "punto de interconexion) -> " + ("Cumple" if res["du_total_pct"] <= 1.5 else "No cumple"))
    L.append("")
    if res.get("calibre_fusible_string"):
        L.append(f"Fusible de string sugerido: {res['calibre_fusible_string']} A / "
                  f"{res.get('tension_fusible_string','-')} V (criterio orientativo ~1,8 x Isc, tension >= "
                  f"V string frio; UNE-EN 62548)")
    L.append(f"Magnetotermico CA sugerido: {res['calibre_magneto_ca']} A")
    if res.get("capacidad_bateria_kwh"):
        L.append("")
        L.append(f"Bateria: Capacidad = (consumo_diario x dias_autonomia) / prof_descarga = "
                  f"({inp.get('consumo_diario_bateria_kwh',0):.2f} x {inp.get('autonomia_dias',0):g}) / "
                  f"({inp.get('profundidad_descarga',80):g}/100) = {res['capacidad_bateria_kwh']:.1f} kWh")
    if res.get("ahorro_anual"):
        L.append("")
        L.append(f"Energia autoconsumida = Produccion x %autoconsumo = {res['produccion_anual_kwh']:.0f} x "
                  f"{inp['pct_autoconsumo']/100:.2f} = {res['energia_autoconsumida']:.0f} kWh/anio")
        L.append(f"Ahorro por autoconsumo = {res['energia_autoconsumida']:.0f} x {inp['precio_kwh']:.2f} = "
                  f"{res['ahorro_autoconsumo']:.2f} EUR/anio")
        if res["ingreso_excedentes"]:
            L.append(f"Ingreso por excedentes compensados = {res['energia_excedente']:.0f} x "
                      f"{inp.get('precio_compensacion', PRECIO_COMPENSACION_DEFECTO):.2f} = "
                      f"{res['ingreso_excedentes']:.2f} EUR/anio")
        L.append(f"Ahorro anual total = {res['ahorro_anual']:.2f} EUR/anio")
        if res.get("payback_anos"):
            L.append(f"Retorno simple = Inversion / Ahorro anual = {inp['inversion_total']:.0f} / "
                      f"{res['ahorro_anual']:.2f} = {res['payback_anos']:.1f} anios")
    return L


def _truncar_ancho(c, texto: str, fuente: str, tamano: float, ancho_max: float) -> str:
    """Trunca 'texto' con puntos suspensivos hasta que quepa en ancho_max,
    midiendo el ancho REAL de la fuente (evita que el texto se salga de su
    hueco y se solape con el bloque contiguo del cajetín)."""
    texto = texto or ""
    if c.stringWidth(texto, fuente, tamano) <= ancho_max:
        return texto
    elipsis = "…"
    lo, hi = 0, len(texto)
    while lo < hi:
        mid = (lo + hi + 1) // 2
        if c.stringWidth(texto[:mid] + elipsis, fuente, tamano) <= ancho_max:
            lo = mid
        else:
            hi = mid - 1
    return texto[:lo].rstrip() + elipsis


def _cajetin_generico(titulo: str, subtitulo: str, datos_proyecto: dict, margin, cajetin_h, AZUL, COBRE, GRIS,
                       config_prof: dict = None):
    """Factoría de función de cajetín reutilizable por los distintos documentos
    (MTD, Anexo, Condiciones Generales), con el mismo estilo que la memoria de
    cálculo de la Calculadora de cables. Incluye el logotipo si hay uno
    configurado en el modo profesional."""
    from reportlab.lib.units import cm
    fecha_hoy = date.today().strftime("%d/%m/%Y")
    config_prof = config_prof or {}

    def _cajetin(c, d):
        from reportlab.lib.pagesizes import A4
        c.saveState()
        width, height = A4
        top = height - margin
        c.setStrokeColor(AZUL)
        c.setLineWidth(1.1)
        c.rect(margin, top - cajetin_h, width - 2 * margin, cajetin_h)
        divisor_x = width - margin - 5.2 * cm
        logo_w = 0
        if config_prof.get("logo_b64"):
            try:
                from reportlab.lib.utils import ImageReader
                img_bytes = io.BytesIO(base64.b64decode(config_prof["logo_b64"]))
                img = ImageReader(img_bytes)
                iw, ih = img.getSize()
                logo_h = cajetin_h - 0.5 * cm
                logo_w = logo_h * iw / ih
                c.drawImage(img, margin + 0.25 * cm, top - cajetin_h + 0.25 * cm, width=logo_w,
                            height=logo_h, preserveAspectRatio=True, mask="auto")
            except Exception:
                logo_w = 0
        text_x = margin + 0.35 * cm + (logo_w + 0.3 * cm if logo_w else 0)
        ancho_texto_disp = divisor_x - text_x - 0.2 * cm  # hueco real antes del panel derecho
        c.line(divisor_x, top - cajetin_h, divisor_x, top)
        c.setFillColor(AZUL)
        c.setFont("Helvetica-Bold", 12.5)
        c.drawString(text_x, top - 0.85 * cm, _pdf_safe(_truncar_ancho(c, titulo, "Helvetica-Bold", 12.5, ancho_texto_disp)))
        c.setFont("Helvetica", 8.3)
        c.setFillColor(GRIS)
        c.drawString(text_x, top - 1.35 * cm, _pdf_safe(_truncar_ancho(c, subtitulo, "Helvetica", 8.3, ancho_texto_disp)))
        linea_titular = f"Titular: {datos_proyecto.get('titular') or '-'}"
        linea_emplazamiento = f"Emplazamiento: {datos_proyecto.get('emplazamiento') or '-'}"
        c.setFont("Helvetica", 7.6)
        c.drawString(text_x, top - 1.80 * cm, _pdf_safe(_truncar_ancho(c, linea_titular, "Helvetica", 7.6, ancho_texto_disp)))
        c.drawString(text_x, top - 2.18 * cm, _pdf_safe(_truncar_ancho(c, linea_emplazamiento, "Helvetica", 7.6, ancho_texto_disp)))
        empresa_txt = config_prof.get("empresa") or config_prof.get("nombre")
        if empresa_txt:
            c.setFont("Helvetica-Oblique", 7.2)
            texto_empresa = _truncar_ancho(c, f"Elaborado por: {empresa_txt}", "Helvetica-Oblique", 7.2,
                                            ancho_texto_disp)
            c.drawString(text_x, top - cajetin_h + 0.28 * cm, _pdf_safe(texto_empresa))
        campos = [("NORMA", "REBT - ITC-BT"), ("FECHA", fecha_hoy), ("REVISION", "1.0")]
        y = top - 0.55 * cm
        for etiqueta, valor in campos:
            c.setFont("Helvetica-Bold", 6.6)
            c.setFillColor(GRIS)
            c.drawString(divisor_x + 0.3 * cm, y, etiqueta)
            c.setFont("Helvetica-Bold", 9)
            c.setFillColor(AZUL)
            c.drawString(divisor_x + 0.3 * cm, y - 0.34 * cm, valor)
            y -= 0.72 * cm
        c.setStrokeColor(COBRE)
        c.setLineWidth(2.2)
        c.line(margin, top - cajetin_h, width - margin, top - cajetin_h)
        c.setFont("Helvetica-Oblique", 6.8)
        c.setFillColor(GRIS)
        c.drawCentredString(width / 2, margin * 0.45,
                              "Documento generado con REBT Suite. Revisar por un tecnico competente antes "
                              "de su presentacion oficial.")
        c.restoreState()

    return _cajetin


def _crear_numbered_canvas(margin, GRIS_HEX="#5a6472"):
    """Canvas que añade 'Página X de Y' al pie, con Y conocido tras construir
    todo el documento (patrón estándar de reportlab de doble pasada)."""
    from reportlab.pdfgen import canvas as canvas_module
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as _colors

    class NumberedCanvas(canvas_module.Canvas):
        def __init__(self, *args, **kwargs):
            canvas_module.Canvas.__init__(self, *args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.setFont("Helvetica-Oblique", 6.8)
                self.setFillColor(_colors.HexColor(GRIS_HEX))
                width, _h = A4
                self.drawRightString(width - margin, margin * 0.45,
                                      f"Página {self._pageNumber} de {total}")
                canvas_module.Canvas.showPage(self)
            canvas_module.Canvas.save(self)

    return NumberedCanvas


def _docx_preparar(titulo: str, subtitulo: str, datos_proyecto: dict, config_prof: dict = None):
    """Crea un documento Word con cabecera tipo cajetín y estilos coherentes
    con los documentos PDF equivalentes. Devuelve (doc, AZUL, COBRE, GRIS)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    config_prof = config_prof or {}
    AZUL = RGBColor(0x12, 0x23, 0x40)
    COBRE = RGBColor(0xb3, 0x71, 0x1f)
    GRIS = RGBColor(0x5a, 0x64, 0x72)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def _sombreado_celda(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color_hex)
        tcPr.append(shd)

    tabla_cajetin = doc.add_table(rows=1, cols=2)
    tabla_cajetin.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla_cajetin.columns[0].width = Cm(12.0)
    tabla_cajetin.columns[1].width = Cm(5.0)
    celda_izq, celda_der = tabla_cajetin.rows[0].cells
    _sombreado_celda(celda_izq, "F4F6FB")
    _sombreado_celda(celda_der, "122340")
    celda_izq.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celda_der.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = celda_izq.paragraphs[0]
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = AZUL
    p2 = celda_izq.add_paragraph()
    r2 = p2.add_run(subtitulo)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRIS
    p3 = celda_izq.add_paragraph()
    r3 = p3.add_run(f"Titular: {datos_proyecto.get('titular') or '-'}  ·  "
                     f"Emplazamiento: {datos_proyecto.get('emplazamiento') or '-'}")
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = GRIS

    for etiqueta, valor in [("NORMA", "REBT · ITC-BT"), ("FECHA", date.today().strftime("%d/%m/%Y")),
                             ("REVISIÓN", "1.0")]:
        pd_ = celda_der.add_paragraph() if celda_der.paragraphs[0].runs else celda_der.paragraphs[0]
        rd1 = pd_.add_run(f"{etiqueta}: ")
        rd1.font.size = Pt(8)
        rd1.font.color.rgb = RGBColor(0xB8, 0xC2, 0xD4)
        rd2 = pd_.add_run(valor)
        rd2.bold = True
        rd2.font.size = Pt(10)
        rd2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if config_prof.get("empresa") or config_prof.get("nombre"):
        pe = celda_izq.add_paragraph()
        re_ = pe.add_run(f"Elaborado por: {' / '.join(filter(None, [config_prof.get('nombre'), config_prof.get('empresa')]))}")
        re_.italic = True
        re_.font.size = Pt(8)
        re_.font.color.rgb = GRIS

    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = pie.add_run("Documento generado con REBT Suite. Revisar por un técnico competente antes de su "
                      "presentación oficial.")
    rp.italic = True
    rp.font.size = Pt(8)
    rp.font.color.rgb = GRIS
    doc.add_paragraph()
    return doc, AZUL, COBRE, GRIS


def _docx_heading(doc, texto: str, nivel: int, color):
    from docx.shared import Pt
    h = doc.add_heading(level=nivel)
    r = h.add_run(texto)
    r.font.color.rgb = color
    r.font.size = Pt(15 if nivel == 1 else (12.5 if nivel == 2 else 11))
    return h


def _docx_parrafo(doc, texto: str, justificar: bool = True, cursiva: bool = False, tamano: float = 10.5):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.italic = cursiva
    r.font.size = Pt(tamano)
    return p


def _docx_lista(doc, items: list):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _docx_tabla(doc, datos: list, azul):
    """Tabla simple con cabecera azul/blanca, ancho de columna automático."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    n_cols = len(datos[0])
    t = doc.add_table(rows=0, cols=n_cols)
    t.style = "Table Grid"
    for i_fila, fila in enumerate(datos):
        celdas = t.add_row().cells
        for i_col, valor in enumerate(fila):
            p = celdas[i_col].paragraphs[0]
            r = p.add_run(str(valor))
            r.font.size = Pt(9.5)
            if i_fila == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tcPr = celdas[i_col]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "122340")
                tcPr.append(shd)
    doc.add_paragraph()
    return t


def _docx_bytes(doc) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _bloque_portada(titulo_doc: str, subtitulo_doc: str, datos_proyecto: dict, config_prof: dict,
                     AZUL, COBRE, colors, h1_style, normal_style):
    """Portada: logo grande, título, datos del proyecto y 'elaborado por'."""
    from reportlab.platypus import Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.units import cm

    config_prof = config_prof or {}
    story = [Spacer(1, 2.5 * cm)]

    if config_prof.get("logo_b64"):
        try:
            img_bytes = io.BytesIO(base64.b64decode(config_prof["logo_b64"]))
            img = Image(img_bytes)
            img._restrictSize(4 * cm, 3 * cm)
            img.hAlign = "CENTER"
            story.append(img)
            story.append(Spacer(1, 0.8 * cm))
        except Exception:
            pass

    story.append(Paragraph(titulo_doc, h1_style))
    story.append(Paragraph(subtitulo_doc, normal_style))
    story.append(Spacer(1, 1.2 * cm))

    filas = [
        ["Titular", datos_proyecto.get("titular") or "—"],
        ["Emplazamiento", datos_proyecto.get("emplazamiento") or "—"],
        ["Fecha", date.today().strftime("%d/%m/%Y")],
    ]
    if config_prof.get("empresa") or config_prof.get("nombre"):
        filas.append(["Elaborado por", " / ".join(filter(None, [config_prof.get("nombre"),
                                                                   config_prof.get("empresa")]))])
    t = _tabla_pdf(filas, (4.5, 9), AZUL, colors, header=False, negrita_col0=True, extra_estilos=[
        ("GRID", (0, 0), (-1, -1), 0, colors.white),  # sin cuadrícula visible en la portada
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9ccd1")),
    ])
    story.append(t)
    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph(
        f'<para textColor="{COBRE.hexval() if hasattr(COBRE, "hexval") else "#b3711f"}">'
        "Documento de apoyo generado con REBT Suite — verificar por un técnico competente antes de su "
        "presentación oficial.</para>", normal_style))
    story.append(PageBreak())
    return story


def _preparar_doc_pdf(titulo: str, subtitulo: str, datos_proyecto: dict, config_prof: dict = None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm as _cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    AZUL = colors.HexColor("#122340")
    COBRE = colors.HexColor("#b3711f")
    GRIS = colors.HexColor("#5a6472")
    buffer = io.BytesIO()
    margin = 1.8 * _cm
    cajetin_h = 3.05 * _cm
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=margin, rightMargin=margin,
                              topMargin=margin + cajetin_h + 0.4 * _cm, bottomMargin=margin + 0.7 * _cm)
    doc._numbered_canvas = _crear_numbered_canvas(margin)
    cajetin = _cajetin_generico(titulo, subtitulo, datos_proyecto, margin, cajetin_h, AZUL, COBRE, GRIS, config_prof)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1doc", parent=styles["Title"], textColor=AZUL, fontSize=20, spaceAfter=10)
    h2 = ParagraphStyle("h2doc", parent=styles["Heading2"], textColor=AZUL, fontSize=12.5, spaceBefore=16,
                         spaceAfter=7, keepWithNext=True, borderWidth=0, leading=15)
    h3 = ParagraphStyle("h3doc", parent=styles["Heading3"], textColor=COBRE, fontSize=10.8, spaceBefore=11,
                         spaceAfter=5, keepWithNext=True, leading=13)
    normal = ParagraphStyle("normaldoc", parent=styles["Normal"], fontSize=9.4, leading=14, spaceAfter=6,
                             alignment=4)  # 4 = TA_JUSTIFY
    lista_item = ParagraphStyle("listaitem", parent=normal, leftIndent=10, spaceAfter=5, bulletIndent=0)
    return buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item


def generar_pdf_mtd(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                     inputs_fv: dict, resultado_fv: dict, total_presupuesto: float,
                     config_prof: dict = None) -> bytes:
    from reportlab.platypus import Paragraph, Spacer, PageBreak

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "MEMORIA TECNICA DE DISENO (MTD)", "REBT - ITC-BT-04", datos_proyecto, config_prof)

    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None
    d = datos_proyecto

    def tabla(datos, colw=(6.5, 9.5)):
        return _tabla_pdf(datos, colw, AZUL, colors)

    story = _bloque_portada("Memoria Técnica de Diseño", "Instalación de baja tensión — REBT / ITC-BT-04",
                             datos_proyecto, config_prof, AZUL, COBRE, colors, h1, normal)

    # ---------------------------------------------------------------- A/B/C
    story.append(Paragraph("A. Datos del titular", h2))
    story.append(tabla([
        ["Campo", "Valor"],
        ["Titular de la instalación (nombre o razón social)", d.get("titular") or "-"],
        ["NIF / CIF", d.get("nif_titular") or "-"],
    ]))

    story.append(Paragraph("B. Instalador autorizado / técnico competente", h2))
    story.append(tabla([
        ["Campo", "Valor"],
        ["Nombre", d.get("instalador") or "-"],
        ["NIF", d.get("nif_instalador") or "-"],
        ["Nº de autorización / colegiado", d.get("n_autorizacion") or "-"],
        ["Categoría", d.get("categoria_instalador") or "-"],
    ]))
    story.append(Paragraph(
        "El instalador autorizado para la categoría de la instalación correspondiente, o el técnico "
        "titulado competente que suscriba esta Memoria, es directamente responsable de que la misma se "
        "adapte a las exigencias reglamentarias vigentes (ITC-BT-04).", normal))

    story.append(Paragraph("C. Datos generales de la instalación", h2))
    story.append(tabla([
        ["Campo", "Valor"],
        ["Emplazamiento", d.get("emplazamiento") or "-"],
        ["Referencia catastral", d.get("referencia_catastral") or "-"],
        ["Uso al que se destina", d.get("uso") or "-"],
        ["Superficie", (d.get("superficie") or "-") + (" m²" if d.get("superficie") else "")],
        ["Tipo de actuación", d.get("tipo_instalacion") or "-"],
    ]))

    # ---------------------------------------------------------------- D. Memoria Descriptiva
    story.append(Paragraph("D. Memoria descriptiva", h2))

    story.append(Paragraph("D.1 Objeto y normativa de aplicación", h3))
    tipo_actuacion = d.get("tipo_instalacion") or "Nueva instalación"
    frases_actuacion = {
        "Nueva instalación": "corresponde a una instalación de nueva planta, sin instalación eléctrica "
                              "previa en el punto descrito",
        "Ampliación": "corresponde a la ampliación de una instalación eléctrica ya existente, "
                      "manteniéndose en servicio las partes no afectadas por esta actuación",
        "Reforma": "corresponde a la reforma de una instalación eléctrica ya existente, sustituyendo o "
                   "adaptando los elementos descritos a las condiciones actuales del REBT",
    }
    frase_actuacion = frases_actuacion.get(tipo_actuacion, frases_actuacion["Nueva instalación"])
    story.append(Paragraph(
        f"La presente Memoria Técnica de Diseño (MTD) tiene por objeto describir y justificar las "
        f"características técnicas de la instalación eléctrica de baja tensión reseñada, así como servir de "
        f"base para su tramitación administrativa y puesta en servicio, de acuerdo con el Reglamento "
        f"Electrotécnico para Baja Tensión (REBT, aprobado por Real Decreto 842/2002) y sus Instrucciones "
        f"Técnicas Complementarias (ITC-BT), en particular la ITC-BT-04 relativa a la documentación y puesta "
        f"en servicio de las instalaciones. La actuación descrita en este documento {frase_actuacion}. Esta "
        f"MTD es válida para instalaciones de las características recogidas en el apartado 3 de la ITC-BT-04 "
        f"que no requieren de Proyecto firmado por técnico titulado competente.", normal))
    if tipo_actuacion in ("Ampliación", "Reforma"):
        story.append(Paragraph(
            "Al tratarse de una actuación sobre una instalación preexistente, deberá verificarse además que "
            "las partes de la instalación no incluidas en esta Memoria mantienen las condiciones de "
            "seguridad exigidas por el REBT, y que la ampliación o reforma descrita no compromete la "
            "protección del conjunto (selectividad de las protecciones, capacidad de la puesta a tierra "
            "existente, etc.).", normal))

    story.append(Paragraph("D.2 Acometida", h3))
    story.append(Paragraph(
        "Es la parte de la instalación de la red de distribución que alimenta la caja general de protección "
        "o unidad funcional equivalente, regulada por la ITC-BT-11. Según su trazado podrá ser aérea posada "
        "sobre fachada, aérea tensada sobre postes, subterránea o aero-subterránea; en cualquier caso los "
        "cables serán aislados, de tensión asignada no inferior a 0,6/1 kV. La acometida forma parte de la "
        "instalación de la empresa distribuidora, por lo que su diseño se ajustará a las normas particulares "
        "de esta.", normal))

    story.append(Paragraph("D.3 Instalaciones de enlace", h3))
    story.append(Paragraph(
        "<b>Caja de Protección y Medida (CPM).</b> Para suministros a un único usuario, al no existir línea "
        "general de alimentación, se instala un único elemento que agrupa la caja general de protección y "
        "el equipo de medida (ITC-BT-13). Se ubicará en fachada, en lugar de libre y permanente acceso, "
        "alojada en nicho con puerta de grado de protección IK10, e incluirá los fusibles de seguridad "
        "correspondientes en todos los conductores de fase.", normal))
    story.append(Paragraph(
        "<b>Derivación individual.</b> Parte de la instalación que, desde la CPM, suministra energía "
        "eléctrica al usuario, comprendiendo fusibles de seguridad, conjunto de medida y dispositivos "
        "generales de mando y protección (ITC-BT-15). Los conductores serán de cobre o aluminio, aislados, "
        "de tensión asignada mínima 450/750 V (0,6/1 kV para multiconductores o tubo enterrado), con sección "
        "mínima de 6 mm² para fases, neutro y protección. La caída de tensión máxima admisible es del 1,5% "
        "cuando no existe línea general de alimentación.", normal))

    if hay_cable:
        story.append(Paragraph("D.4 Instalación de baja tensión calculada", h3))
        story.append(Paragraph(f"Circuito de referencia: <b>{inputs_cable['tipo_circuito']}</b>.", normal))
        story.append(tabla([
            ["Concepto", "Valor"],
            ["Sistema", f"{inputs_cable['sistema']} — {inputs_cable['tension']:g} V"],
            ["Conductor / Aislamiento", f"{inputs_cable['conductor']} / {inputs_cable['aislamiento']}"],
            ["Método de instalación (canalización)", inputs_cable["metodo"]],
            ["Longitud del circuito", f"{inputs_cable['longitud']:g} m"],
            ["Sección de fase adoptada", f"{resultado_cable['seccion_final']:g} mm²"],
            ["Sección de neutro", f"{resultado_cable['seccion_neutro']:g} mm²" if resultado_cable.get("seccion_neutro") else "—"],
            ["Sección de protección (PE)", f"{resultado_cable['seccion_proteccion']:g} mm²" if resultado_cable.get("seccion_proteccion") else "—"],
            ["Protección (interruptor automático)", f"{resultado_cable['calibre_magnetotermico']} A"],
            ["Caída de tensión", f"{resultado_cable['e_final_pct']:.2f} % (máx. {inputs_cable['delta_u_max']:g} %)"],
        ]))
        story.append(Paragraph(
            "Los dispositivos generales e individuales de mando y protección se han seleccionado conforme a "
            "la ITC-BT-17 (protección contra sobreintensidades) y la ITC-BT-24 (protección contra contactos "
            "directos e indirectos), garantizando la selectividad y el poder de corte necesarios frente a la "
            "intensidad de cortocircuito prevista en el punto de instalación.", normal))
    else:
        story.append(Paragraph("D.4 Instalación de baja tensión", h3))
        story.append(Paragraph("No se ha completado ningún cálculo de circuito en la pestaña Calculadora; "
                                "esta sección se completará cuando se disponga de esos datos.", normal))

    story.append(Paragraph("D.5 Cuadro general de mando y protección", h3))
    story.append(Paragraph(
        "El cuadro general de mando y protección aloja, como mínimo: un interruptor general automático "
        "(IGA) de corte omnipolar, con poder de corte suficiente para la intensidad de cortocircuito "
        "prevista (4,5 kA como mínimo en vivienda), que permita su accionamiento manual; un interruptor "
        "diferencial general (o varios, según la previsión de cargas), destinado a la protección contra "
        "contactos indirectos de todos los circuitos; y los interruptores automáticos individuales de cada "
        "circuito derivado, coordinados entre sí para garantizar la selectividad. La altura de los "
        "dispositivos, medida desde el suelo, estará comprendida entre 1 y 2 m. Se recomienda la instalación "
        "de un dispositivo de protección contra sobretensiones (DPS) en cabecera cuando la acometida sea "
        "aérea o exista riesgo de sobretensiones de origen atmosférico (ITC-BT-23).", normal))

    story.append(Paragraph("D.6 Conductores: materiales, dimensionado e identificación", h3))
    story.append(Paragraph(
        "Los conductores serán de cobre o aluminio, siempre aislados, con tensión asignada no inferior a "
        "450/750 V. La sección se determina por el criterio más desfavorable entre intensidad máxima "
        "admisible (ITC-BT-19) y caída de tensión (3% alumbrado / 5% otros usos en instalación interior; "
        "1,5% en derivación individual sin LGA), tal y como se desarrolla en el Anexo de Cálculos. Salvo "
        "justificación por cálculo, la sección del conductor neutro será como mínimo igual a la de las "
        "fases, para tener en cuenta las corrientes armónicas de cargas no lineales. La identificación de "
        "los conductores se realiza por el color de su aislamiento: azul claro para el neutro, "
        "verde-amarillo exclusivamente para el conductor de protección, y marrón, negro o gris para las "
        "fases.", normal))

    story.append(Paragraph("D.7 Subdivisión, equilibrado de cargas y conexiones", h3))
    story.append(Paragraph(
        "La instalación se subdivide en circuitos independientes y adecuadamente protegidos, de forma que "
        "una avería en un punto afecte solo a una parte de la instalación, facilitando además las "
        "verificaciones y el mantenimiento. Se procurará mantener el mayor equilibrio posible de cargas "
        "entre fases. Las conexiones entre conductores se realizan siempre en el interior de cajas de "
        "empalme mediante bornes de conexión, regletas o sistemas equivalentes, sin permitirse la unión por "
        "simple retorcimiento.", normal))

    story.append(Paragraph("D.8 Sistemas de instalación (canalizaciones)", h3))
    metodo_txt = inputs_cable.get("metodo", "") if hay_cable else ""
    story.append(Paragraph(
        f"El sistema de canalización empleado se corresponde con el tipo {metodo_txt or '[a determinar]'} "
        "según la clasificación de la ITC-BT-19 y la norma UNE-HD 60364-5-52, seleccionado en función de "
        "las influencias externas del emplazamiento (humedad, temperatura, riesgo mecánico) conforme a la "
        "clasificación de la norma UNE 20460-3. Las características mínimas de los tubos protectores "
        "(resistencia a la compresión, al impacto y rango de temperatura de servicio) se detallan en el "
        "Pliego de Condiciones Técnicas, apartado 15, en función de si la instalación es superficial, "
        "empotrada o enterrada.", normal))

    story.append(Paragraph("D.9 Protección contra sobreintensidades y sobretensiones", h3))
    story.append(Paragraph(
        "Todo circuito está protegido contra sobrecargas mediante interruptor automático con curva térmica "
        "de corte, y contra cortocircuitos mediante dispositivo cuya capacidad de corte es acorde con la "
        "intensidad de cortocircuito que puede presentarse en su punto de conexión (ITC-BT-22). Cuando la "
        "instalación se alimenta, en todo o en parte, por una línea aérea, o se opta por una protección "
        "adicional, se instala protección contra sobretensiones transitorias de origen atmosférico "
        "(ITC-BT-23), seleccionada según la categoría de sobretensión soportada por los equipos:", normal))
    story.append(tabla([
        ["Categoría", "Tensión soportada a impulsos 230/400V", "Equipos típicos"],
        ["IV", "6 kV", "Contadores, equipos principales de protección, origen de la instalación"],
        ["III", "4 kV", "Cuadros de distribución, aparamenta fija, canalizaciones"],
        ["II", "2,5 kV", "Electrodomésticos, herramientas portátiles"],
        ["I", "1,5 kV", "Equipos electrónicos muy sensibles"],
    ], colw=(2.2, 5.3, 8.5)))

    story.append(Paragraph("D.10 Protección contra contactos directos e indirectos", h3))
    story.append(Paragraph(
        "La protección contra contactos directos se confía al aislamiento de las partes activas y, "
        "complementariamente, a interruptores diferenciales de sensibilidad no superior a 30 mA. La "
        "protección contra contactos indirectos se realiza mediante corte automático de la alimentación: "
        "todas las masas protegidas por un mismo dispositivo se interconectan a un mismo conductor de "
        "protección y toma de tierra, cumpliéndose la condición Ra·Ia ≤ U (siendo U la tensión de contacto "
        "límite convencional: 50 V en locales secos, 24 V en locales húmedos).", normal))

    story.append(Paragraph("D.11 Puesta a tierra", h3))
    story.append(Paragraph(
        "La puesta a tierra de la instalación se ejecutará conforme a la ITC-BT-18, mediante electrodo(s) "
        "(picas y/o conductor enterrado) conectado al borne principal de tierra, del que partirá el "
        "conductor de protección hacia todas las masas metálicas de la instalación. La resistencia de tierra "
        "resultante será compatible con la sensibilidad de las protecciones diferenciales instaladas, de "
        "forma que la tensión de contacto no supere los límites de seguridad de la ITC-BT-24. La instalación "
        "de puesta a tierra se comprobará, como mínimo, en el momento de dar de alta la instalación y "
        "periódicamente por personal técnicamente competente.", normal))

    contador_d = 11
    if hay_fv:
        contador_d += 1
        story.append(Paragraph(f"D.{contador_d} Instalación generadora fotovoltaica", h3))
        story.append(tabla([
            ["Concepto", "Valor"],
            ["Modalidad de autoconsumo (RD 244/2019)", inputs_fv["tipo_autoconsumo"]],
            ["Potencia pico instalada", f"{resultado_fv['p_pico_kwp']:.2f} kWp"],
            ["Nº de paneles / configuración", f"{resultado_fv['n_paneles_configurados']} ({resultado_fv['n_serie']}S{resultado_fv['n_paralelo']}P)"],
            ["Potencia del inversor", f"{inputs_fv['potencia_inversor_kw']:g} kW"],
            ["Sección CC / CA", f"{resultado_fv['s_cc_final']:g} mm² / {resultado_fv['s_ca_final']:g} mm²"],
            ["Producción anual estimada", f"{resultado_fv['produccion_anual_kwh']:,.0f} kWh/año".replace(",", ".")],
            ["CO₂ evitado estimado", f"{resultado_fv['co2_evitado_kg_ano']/1000:.2f} t/año"],
        ]))
        umbral = ("≤10 kW → Certificado eléctrico firmado por instalador autorizado"
                  if resultado_fv["p_pico_kwp"] <= 10 else
                  ">10 kW → Proyecto firmado por técnico competente")
        story.append(Paragraph(f"Régimen documental aplicable según RD 244/2019: {umbral}. La conexión se "
                                "realiza a través de un cuadro de mando y protección específico dotado de "
                                "las protecciones diferenciales necesarias, y de un mecanismo antivertido si "
                                "la modalidad es sin excedentes, conforme al Anexo I de la ITC-BT-40.", normal))

    contador_d += 1
    story.append(Paragraph(f"D.{contador_d} Receptores de alumbrado", h3))
    story.append(Paragraph(
        "Las luminarias empleadas cumplen los requisitos de la serie de normas UNE-EN 60598. Sus partes "
        "metálicas accesibles, cuando no son de Clase II o III, disponen de conexión al conductor de "
        "protección. Si se emplean receptores con lámparas de descarga, se prevé una carga mínima de 1,8 "
        "veces la potencia en vatios de las lámparas (ITC-BT-44) y la compensación del factor de potencia "
        "hasta un valor mínimo de 0,9.", normal))

    if hay_cable and inputs_cable.get("es_motor"):
        contador_d += 1
        story.append(Paragraph(f"D.{contador_d} Receptores a motor", h3))
        story.append(Paragraph(
            "El circuito calculado corresponde a un receptor a motor. Conforme a la ITC-BT-47, el conductor "
            "de conexión se ha dimensionado para el 125% de la intensidad a plena carga cuando se trata de "
            "un motor único, o para el 125% del motor de mayor potencia más el 100% del resto cuando "
            "alimenta a varios motores. El motor dispone de protección contra cortocircuitos y sobrecargas "
            "en todas sus fases, y contra la falta de tensión cuando su arranque espontáneo pudiera provocar "
            "accidentes o perjudicar el equipo.", normal))

    contador_d += 1
    story.append(Paragraph(f"D.{contador_d} Influencias externas", h3))
    story.append(Paragraph(
        "Salvo indicación expresa en contrario, se considera que el emplazamiento presenta condiciones "
        "ambientales normales (AD1, AA4, BA1 según UNE 20460-3), sin riesgo de incendio ni explosión, por lo "
        "que no son de aplicación las prescripciones adicionales de la ITC-BT-29 para locales de "
        "características especiales.", normal))


    # ---------------------------------------------------------------- E. Memoria Justificativa
    story.append(PageBreak())
    story.append(Paragraph("E. Memoria justificativa (resumen de cálculos)", h2))
    story.append(Paragraph(
        "El detalle completo de las fórmulas y valores empleados para justificar las secciones, "
        "protecciones y caídas de tensión se recoge en el Anexo de Cálculos y Mediciones, documento que "
        "forma parte integrante de esta Memoria. A continuación se resume el resultado:", normal))
    filas_resumen = [["Parte de la instalación", "Potencia (kW)", "Longitud (m)", "Sección (mm²)", "ΔU (%)"]]
    if hay_cable:
        filas_resumen.append([
            inputs_cable["tipo_circuito"][:28],
            f"{inputs_cable['potencia_kw']:.2f}" if inputs_cable.get("potencia_kw") else "—",
            f"{inputs_cable['longitud']:g}", f"{resultado_cable['seccion_final']:g}",
            f"{resultado_cable['e_final_pct']:.2f}",
        ])
    if hay_fv:
        filas_resumen.append([
            "Generador FV (tramo CC)", f"{resultado_fv['p_pico_kwp']:.2f}", f"{inputs_fv['longitud_cc']:g}",
            f"{resultado_fv['s_cc_final']:g}", f"{resultado_fv['du_cc_pct']:.2f}",
        ])
        filas_resumen.append([
            "Generador FV (tramo CA)", f"{inputs_fv['potencia_inversor_kw']:g}", f"{inputs_fv['longitud_ca']:g}",
            f"{resultado_fv['s_ca_final']:g}", f"{resultado_fv['e_ca_pct']:.2f}",
        ])
    if len(filas_resumen) > 1:
        story.append(tabla(filas_resumen, colw=(4.5, 3.0, 2.7, 2.9, 2.9)))
    else:
        story.append(Paragraph("No hay cálculos disponibles todavía.", normal))

    # ---------------------------------------------------------------- F/G/H
    story.append(Paragraph("F. Presupuesto", h2))
    if total_presupuesto:
        story.append(Paragraph(f"El presupuesto de la instalación asciende a la cantidad de "
                                f"<b>{_miles(total_presupuesto, 2)} €</b> "
                                f"({numero_a_letras_euros(total_presupuesto)}), "
                                "IVA incluido, según el desglose por capítulos que se adjunta en el documento "
                                "de Presupuesto.", normal))
    else:
        story.append(Paragraph("No se ha generado presupuesto en la pestaña correspondiente.", normal))

    story.append(Paragraph("G. Documentos que integran el proyecto", h2))
    for linea in ["Memoria Técnica de Diseño (este documento)", "Anexo de Cálculos y Mediciones",
                  "Pliego de Condiciones Generales", "Presupuesto",
                  "Esquema unifilar y croquis de emplazamiento (a incorporar por el técnico)",
                  "Certificado de Instalación Eléctrica (CIE), tras la puesta en servicio"]:
        story.append(Paragraph("• " + linea, normal))

    story.append(Paragraph("H. Declaración y firma", h2))
    story.append(Paragraph(
        f"D./Dña. {d.get('instalador') or '_______________________'}, en calidad de "
        f"{(d.get('categoria_instalador') or 'instalador autorizado').lower()}, declara que la instalación "
        "descrita en la presente Memoria Técnica de Diseño se ajusta a las prescripciones del Reglamento "
        "Electrotécnico para Baja Tensión y sus Instrucciones Técnicas Complementarias.", normal))
    story.append(Spacer(1, 24))
    story.append(Paragraph("En ______________________, a ____ de ______________ de 20____", normal))
    story.append(Spacer(1, 24))
    story.append(Paragraph("Firma: _______________________________________", normal))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Documento de apoyo generado automáticamente a partir de los datos introducidos en la aplicación. "
        "Antes de su presentación ante el organismo competente, debe ser revisado, completado y firmado por "
        "el instalador autorizado o técnico competente responsable.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


def _parrafos_calculo_cable_pdf(inp: dict, res: dict, normal, formula) -> list:
    """Redacta la justificación de cálculo del circuito de baja tensión como
    una memoria de cálculo al uso (fórmula tipográfica + texto explicativo +
    sustitución numérica), en vez de líneas sueltas tipo calculadora."""
    from reportlab.platypus import Paragraph
    P = []
    sistema = inp["sistema"]
    cosphi = inp["cos_phi"]

    if inp["modo_entrada"] == "Potencia activa":
        p_w = inp["potencia_kw"] * 1000.0
        p_w_txt = _miles(p_w)
        if sistema == SISTEMA_MONO:
            P.append(Paragraph("La intensidad de empleo del circuito, para un sistema monofásico, se "
                                "determina mediante la expresión:", normal))
            P.append(Paragraph("I<sub>b</sub> = P / (V · cos φ)", formula))
            P.append(Paragraph(f"Sustituyendo los valores de la instalación (P = {p_w_txt} W, "
                                f"V = {inp['tension']:g} V, cos φ = {cosphi:.2f}), resulta:", normal))
            P.append(Paragraph(f"I<sub>b</sub> = {p_w_txt} / ({inp['tension']:g} × {cosphi:.2f}) = "
                                f"<b>{res['ib']:.2f} A</b>", formula))
        else:
            P.append(Paragraph("La intensidad de empleo del circuito, para un sistema trifásico "
                                "equilibrado, se determina mediante la expresión:", normal))
            P.append(Paragraph("I<sub>b</sub> = P / (√3 · V · cos φ)", formula))
            P.append(Paragraph(f"Sustituyendo los valores de la instalación (P = {p_w_txt} W, "
                                f"V = {inp['tension']:g} V, cos φ = {cosphi:.2f}), resulta:", normal))
            P.append(Paragraph(f"I<sub>b</sub> = {p_w_txt} / (√3 × {inp['tension']:g} × {cosphi:.2f}) = "
                                f"<b>{res['ib']:.2f} A</b>", formula))
    else:
        P.append(Paragraph(f"La intensidad de empleo se ha introducido de forma directa: "
                            f"I<sub>b</sub> = <b>{res['ib']:.2f} A</b>.", normal))

    if res.get("ib_motor") is not None:
        if len(inp["corrientes_motores"]) == 1:
            P.append(Paragraph("Al tratarse de un circuito de motor único, la ITC-BT-47 exige dimensionar "
                                "el conductor para el 125% de la intensidad a plena carga del motor:", normal))
            P.append(Paragraph(f"I<sub>b,motor</sub> = 1,25 · I<sub>n</sub> = 1,25 × "
                                f"{inp['corrientes_motores'][0]:.2f} = <b>{res['ib_motor']:.2f} A</b>", formula))
        else:
            ordenados = sorted(inp["corrientes_motores"], reverse=True)
            resto = " + ".join(f"{x:.1f}" for x in ordenados[1:])
            P.append(Paragraph("Al existir varios motores en el mismo circuito, la ITC-BT-47 exige "
                                "dimensionar el conductor para el 125% de la intensidad del motor de mayor "
                                "potencia, más el 100% de la intensidad del resto:", normal))
            P.append(Paragraph(f"I<sub>b,motor</sub> = 1,25 × {ordenados[0]:.2f} + ({resto}) = "
                                f"<b>{res['ib_motor']:.2f} A</b>", formula))
        if inp["ascensor_grua"]:
            P.append(Paragraph("Por tratarse de un ascensor o grúa, se aplica adicionalmente el factor "
                                "×1,3 previsto en la ITC-BT-47.", normal))
    if inp["alumbrado_descarga"]:
        P.append(Paragraph("Al tratarse de alumbrado con lámparas de descarga sin corregir el factor de "
                            "potencia, se incrementa la intensidad de cálculo con un factor orientativo de "
                            "1,8, según la ITC-BT-44:", normal))
        P.append(Paragraph(f"I<sub>b,cálculo</sub> = I<sub>b</sub> × 1,8 = <b>{res['ib_calculo']:.2f} A</b>",
                            formula))

    P.append(Paragraph("El criterio térmico del artículo 19.2 de la ITC-BT-19 exige que la intensidad "
                        "admisible del cable, corregida por las condiciones reales de la instalación, no "
                        "sea inferior a la intensidad de empleo:", normal))
    P.append(Paragraph("I<sub>b</sub> ≤ I<sub>n</sub> ≤ I<sub>z</sub> ,  con  I<sub>z</sub> = "
                        "I<sub>z,tabla</sub> · f<sub>temp</sub> · f<sub>agrup</sub> · f<sub>capas</sub> · "
                        "f<sub>resist</sub>", formula))
    factor_base = res["iz_termica"] / max(res["factor_total"], 1e-9)
    P.append(Paragraph(f"Para la sección adoptada, la Guía-BT-19 establece una intensidad admisible de "
                        f"base de {factor_base:.1f} A que, corregida por los factores de temperatura "
                        f"({res['f_temp']:.3f}), agrupamiento ({res['f_agrup']:.3f}), capas "
                        f"({res['f_capas']:.3f}) y resistividad del terreno ({res['f_resist']:.3f}), "
                        "resulta:", normal))
    P.append(Paragraph(f"I<sub>z</sub> = {factor_base:.1f} × {res['f_temp']:.3f} × {res['f_agrup']:.3f} × "
                        f"{res['f_capas']:.3f} × {res['f_resist']:.3f} = "
                        f"<b>{res['iz_termica']:.2f} A</b>", formula))
    i_comparar = res["ib_calculo"] / (res["n_paralelo"] if res["necesita_paralelo"] else 1)
    cumple_termico = res["iz_termica"] >= i_comparar
    P.append(Paragraph(f"Al ser I<sub>z</sub> = {res['iz_termica']:.2f} A "
                        f"{'superior' if cumple_termico else 'inferior'} a la intensidad de cálculo "
                        f"({i_comparar:.2f} A), la sección de <b>{res['s_termica']:g} mm²</b> "
                        f"{'satisface' if cumple_termico else 'NO satisface'} el criterio térmico.", normal))

    P.append(Paragraph("Para el criterio de caída de tensión, la conductividad del conductor se toma a su "
                        "temperatura de servicio (más conservadora que a 20°C):", normal))
    P.append(Paragraph("κ(T) = κ<sub>20°C</sub> / [1 + α · (T<sub>servicio</sub> − 20)]", formula))
    kappa = res["kappa"]
    P.append(Paragraph(f"κ = {CONDUCTIVIDAD_20C[inp['conductor']]:g} / [1 + "
                        f"{COEF_TEMP_RESIST[inp['conductor']]:g} × "
                        f"({TEMP_SERVICIO[inp['aislamiento']]:g} − 20)] = <b>{kappa:.2f} m/(Ω·mm²)</b>", formula))
    S = res["seccion_final"]
    r_metro = 1.0 / (kappa * S)
    k_sist_txt = "2" if sistema == SISTEMA_MONO else "√3"
    P.append(Paragraph("La caída de tensión se calcula, con resistencia R = 1/(κ·S), mediante:", normal))
    P.append(Paragraph(f"ΔU = {k_sist_txt} · L · I<sub>b</sub> · (R·cos φ + X·sen φ)", formula))
    P.append(Paragraph(f"Para L = {inp['longitud']:g} m y S = {S:g} mm², R = 1/({kappa:.2f}×{S:g}) = "
                        f"{r_metro:.5f} Ω/m, resultando:", normal))
    P.append(Paragraph(f"ΔU = {res['e_final']:.2f} V = <b>{res['e_final_pct']:.2f} %</b> de "
                        f"{inp['tension']:g} V", formula))
    cumple_du = res["e_final_pct"] <= inp["delta_u_max"]
    P.append(Paragraph(f"Como {res['e_final_pct']:.2f} % es {'inferior' if cumple_du else 'superior'} al "
                        f"máximo admisible del {inp['delta_u_max']:g} % para este tramo, la sección "
                        f"{'satisface' if cumple_du else 'NO satisface'} el criterio de caída de tensión.",
                        normal))

    if res.get("seccion_neutro"):
        P.append(Paragraph("La sección del conductor neutro se obtiene según el criterio general del REBT "
                            "(igual a la de fase hasta 16 mm², o la mitad redondeada para secciones "
                            "mayores sin armónicos significativos):", normal))
        P.append(Paragraph(f"S<sub>f</sub> = {S:g} mm² → <b>S<sub>n</sub> = "
                            f"{res['seccion_neutro']:g} mm²</b>", formula))
    if res.get("seccion_proteccion"):
        P.append(Paragraph("La sección del conductor de protección se obtiene según la tabla de la "
                            "ITC-BT-18:", normal))
        P.append(Paragraph(f"S<sub>f</sub> = {S:g} mm² → <b>S<sub>p</sub> = "
                            f"{res['seccion_proteccion']:g} mm²</b>", formula))

    if res.get("cumple_cc") is not None:
        k_cc = K_CORTOCIRCUITO[(inp["conductor"], inp["aislamiento"])]
        P.append(Paragraph("Por último, se verifica el criterio térmico de cortocircuito (IEC 60364-5-54):",
                            normal))
        P.append(Paragraph("S<sub>mín</sub> = I<sub>cc</sub> · √t / k", formula))
        P.append(Paragraph(f"S<sub>mín</sub> = ({inp['icc_ka']:g}×1000 × √{inp['tiempo_s']:g}) / {k_cc} = "
                            f"<b>{res['s_min_cc']:g} mm²</b>, frente a los {S:g} mm² adoptados: "
                            f"{'cumple' if res['cumple_cc'] else 'NO cumple'}.", normal))
    return P


def _parrafos_calculo_fv_pdf(inp: dict, res: dict, normal, formula) -> list:
    """Redacta la justificación de cálculo de la instalación fotovoltaica
    como memoria de cálculo (mismo espíritu que la función de cable)."""
    from reportlab.platypus import Paragraph
    P = []

    P.append(Paragraph(f"El generador fotovoltaico se dimensiona para una potencia pico de "
                        f"<b>{res['p_pico_kwp']:.2f} kWp</b>, equivalente a {res['n_paneles']} paneles de "
                        f"{inp['potencia_panel_wp']:g} Wp y una superficie de captación de "
                        f"{res['superficie_necesaria_m2']:.1f} m².", normal))

    P.append(Paragraph("Las pérdidas por orientación e inclinación no óptimas se calculan según la "
                        "expresión del Documento Básico HE5 del Código Técnico de la Edificación:", normal))
    P.append(Paragraph("Pérdidas (%) = 100 · [1,2·10<super>-4</super>·(β − φ + 10)² + "
                        "3,5·10<super>-5</super>·α²]", formula))
    P.append(Paragraph(f"con inclinación β = {inp['inclinacion']:g}°, latitud φ = {inp['latitud']:g}° y "
                        f"azimut α = {inp['azimut']:g}°, lo que da unas pérdidas de "
                        f"<b>{res['perdidas_orient_pct']:.2f} %</b>.", normal))
    P.append(Paragraph("El rendimiento efectivo (Performance Ratio) incorpora además las pérdidas por "
                        "sombras/suciedad y la eficiencia del inversor:", normal))
    P.append(Paragraph(f"PR<sub>efectivo</sub> = PR<sub>base</sub> · (1 − p<sub>orient</sub>) · "
                        f"(1 − p<sub>sombras</sub>) · η<sub>inversor</sub> = {inp['pr']:.2f} × "
                        f"(1 − {res['perdidas_orient_pct']/100:.3f}) × "
                        f"(1 − {inp['perdidas_sombras']/100:.2f}) × {inp['eficiencia_inversor']/100:.3f} = "
                        f"<b>{res['pr_efectivo']:.3f}</b>", formula))
    P.append(Paragraph("La producción anual estimada resulta de aplicar la fórmula del IDAE:", normal))
    P.append(Paragraph("E<sub>anual</sub> = P<sub>pico</sub> · HSP · 365 · PR<sub>efectivo</sub>", formula))
    P.append(Paragraph(f"E<sub>anual</sub> = {res['p_pico_kwp']:.2f} × {inp['hsp']:g} × 365 × "
                        f"{res['pr_efectivo']:.3f} = <b>{_miles(res['produccion_anual_kwh'])} kWh/año</b>",
                        formula))
    P.append(Paragraph(f"Considerando una degradación anual del panel del {inp['degradacion_anual']:g}%, la "
                        f"producción estimada será de {_miles(res['produccion_ano10'])} kWh/año en el año "
                        f"10 y de {_miles(res['produccion_ano25'])} kWh/año en el año 25. El ahorro de "
                        f"emisiones asociado, con un factor de emisión de la red de "
                        f"{inp.get('factor_co2', FACTOR_CO2_RED_DEFECTO):.2f} kg CO<sub>2</sub>/kWh, es de "
                        f"<b>{res['co2_evitado_kg_ano']/1000:.2f} toneladas de CO<sub>2</sub> al año</b>.", normal))

    P.append(Paragraph(f"El generador se configura con {res['n_serie']} paneles en serie por string y "
                        f"{res['n_paralelo']} strings en paralelo. La tensión de circuito abierto del "
                        "string en la condición más desfavorable (frío) se calcula como:", normal))
    P.append(Paragraph("V<sub>string,frío</sub> = N<sub>serie</sub> · V<sub>oc</sub> · "
                        "[1 + α<sub>V</sub> · (25 − T<sub>mín</sub>)]", formula))
    P.append(Paragraph(f"V<sub>string,frío</sub> = {res['n_serie']} × {inp['voc']:g} × "
                        f"[1 + {inp['coef_temp_voc']/100:.4f} × (25 − {inp['temp_min']:g})] = "
                        f"<b>{res['v_string_frio']:.1f} V</b>", formula))
    cumple_v = res["cumple_vmax"] and res["cumple_vmpp_min"] and res["cumple_vmpp_max"]
    P.append(Paragraph(f"Este valor {'se mantiene dentro' if cumple_v else 'NO se mantiene dentro'} de la "
                        f"ventana de tensión admisible del inversor "
                        f"({inp['vmin_mppt']:g}-{inp['vmax_mppt']:g} V de rango MPPT, "
                        f"{inp['vmax_entrada_inversor']:g} V de tensión máxima de entrada).", normal))

    P.append(Paragraph("De acuerdo con la ITC-BT-40, los cables de conexión de instalaciones generadoras "
                        "se dimensionan para una intensidad no inferior al 125% de la intensidad máxima del "
                        "generador, con una caída de tensión conjunta (continua + alterna) no superior al "
                        "1,5% entre el generador y el punto de interconexión:", normal))
    P.append(Paragraph(f"I<sub>diseño,CC</sub> = 1,25 · I<sub>sc</sub> = 1,25 × {inp['isc']:g} = "
                        f"<b>{res['i_diseno_cc']:.2f} A</b>  →  sección adoptada "
                        f"<b>{res['s_cc_final']:g} mm²</b> (ΔU = {res['du_cc_pct']:.2f} %)", formula))
    P.append(Paragraph(f"I<sub>diseño,CA</sub> = 1,25 · I<sub>b</sub> = <b>{res['i_diseno_ca']:.2f} A</b>  →  "
                        f"sección adoptada <b>{res['s_ca_final']:g} mm²</b> "
                        f"(ΔU = {res['e_ca_pct']:.2f} %)", formula))
    cumple_total = res["du_total_pct"] <= 1.5
    P.append(Paragraph(f"La caída de tensión conjunta resulta ΔU = {res['du_total_pct']:.2f} %, que "
                        f"{'cumple' if cumple_total else 'NO cumple'} el límite del 1,5% establecido por "
                        "la ITC-BT-40.", normal))

    if res.get("calibre_fusible_string"):
        P.append(Paragraph(f"Al haber más de un string en paralelo, se dispone fusible de protección de "
                            f"cada string, de calibre <b>{res['calibre_fusible_string']} A</b> y tensión "
                            f"nominal <b>{res.get('tension_fusible_string','—')} V</b> (criterio orientativo "
                            "1,5-2,4 veces la Isc del string, UNE-EN 62548).", normal))
    if res.get("capacidad_bateria_kwh"):
        P.append(Paragraph("La capacidad de la batería de acumulación se determina a partir del consumo "
                            "diario a cubrir, la autonomía deseada y la profundidad de descarga admisible:",
                            normal))
        P.append(Paragraph("C<sub>batería</sub> = (Consumo<sub>diario</sub> · Días<sub>autonomía</sub>) / "
                            "Profundidad<sub>descarga</sub>", formula))
        P.append(Paragraph(f"C<sub>batería</sub> = ({inp.get('consumo_diario_bateria_kwh',0):.2f} × "
                            f"{inp.get('autonomia_dias',0):g}) / {inp.get('profundidad_descarga',80)/100:.2f} "
                            f"= <b>{res['capacidad_bateria_kwh']:.1f} kWh</b>", formula))

    if res.get("ahorro_anual"):
        texto_ahorro = (f"Del total producido, se estima que un {inp['pct_autoconsumo']:g}% se "
                        f"autoconsume directamente ({_miles(res['energia_autoconsumida'])} kWh/año), "
                        f"generando un ahorro de <b>{_fmt_eur(res['ahorro_autoconsumo'])}/año</b>")
        if res["ingreso_excedentes"]:
            texto_ahorro += (f", más un ingreso por compensación de excedentes de "
                              f"<b>{_fmt_eur(res['ingreso_excedentes'])}/año</b>")
        texto_ahorro += "."
        P.append(Paragraph(texto_ahorro, normal))
        if res.get("payback_anos"):
            P.append(Paragraph(f"Con una inversión estimada de {_fmt_eur(inp['inversion_total'])} y un "
                                f"ahorro anual de {_fmt_eur(res['ahorro_anual'])}, el retorno simple de la "
                                f"inversión es de <b>{res['payback_anos']:.1f} años</b>.", normal))
    return P


def generar_docx_mtd(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                      inputs_fv: dict, resultado_fv: dict, total_presupuesto: float,
                      config_prof: dict = None) -> bytes:
    """Versión en Word (.docx) de la Memoria Técnica de Diseño, con el mismo
    contenido y estructura que la versión en PDF, para quien necesite
    entregarla editable."""
    doc, AZUL, COBRE, GRIS = _docx_preparar("MEMORIA TÉCNICA DE DISEÑO (MTD)", "REBT · ITC-BT-04",
                                             datos_proyecto, config_prof)
    d = datos_proyecto
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    _docx_heading(doc, "A. Datos del titular", 2, AZUL)
    _docx_tabla(doc, [
        ["Campo", "Valor"],
        ["Titular de la instalación", d.get("titular") or "-"],
        ["NIF / CIF", d.get("nif_titular") or "-"],
    ], AZUL)

    _docx_heading(doc, "B. Instalador autorizado / técnico competente", 2, AZUL)
    _docx_tabla(doc, [
        ["Campo", "Valor"],
        ["Nombre", d.get("instalador") or "-"],
        ["NIF", d.get("nif_instalador") or "-"],
        ["Nº de autorización / colegiado", d.get("n_autorizacion") or "-"],
        ["Categoría", d.get("categoria_instalador") or "-"],
    ], AZUL)
    _docx_parrafo(doc, "El instalador autorizado para la categoría de la instalación correspondiente, o el "
                  "técnico titulado competente que suscriba esta Memoria, es directamente responsable de "
                  "que la misma se adapte a las exigencias reglamentarias vigentes (ITC-BT-04).")

    _docx_heading(doc, "C. Datos generales de la instalación", 2, AZUL)
    _docx_tabla(doc, [
        ["Campo", "Valor"],
        ["Emplazamiento", d.get("emplazamiento") or "-"],
        ["Referencia catastral", d.get("referencia_catastral") or "-"],
        ["Uso al que se destina", d.get("uso") or "-"],
        ["Superficie", (d.get("superficie") or "-") + (" m²" if d.get("superficie") else "")],
        ["Tipo de actuación", d.get("tipo_instalacion") or "-"],
    ], AZUL)

    _docx_heading(doc, "D. Memoria descriptiva", 2, AZUL)

    _docx_heading(doc, "D.1 Objeto y normativa de aplicación", 3, COBRE)
    tipo_actuacion = d.get("tipo_instalacion") or "Nueva instalación"
    frases_actuacion = {
        "Nueva instalación": "corresponde a una instalación de nueva planta, sin instalación eléctrica "
                              "previa en el punto descrito",
        "Ampliación": "corresponde a la ampliación de una instalación eléctrica ya existente, "
                      "manteniéndose en servicio las partes no afectadas por esta actuación",
        "Reforma": "corresponde a la reforma de una instalación eléctrica ya existente, sustituyendo o "
                   "adaptando los elementos descritos a las condiciones actuales del REBT",
    }
    frase_actuacion = frases_actuacion.get(tipo_actuacion, frases_actuacion["Nueva instalación"])
    _docx_parrafo(doc, "La presente Memoria Técnica de Diseño (MTD) tiene por objeto describir y justificar "
                  "las características técnicas de la instalación eléctrica de baja tensión reseñada, así "
                  "como servir de base para su tramitación administrativa y puesta en servicio, de acuerdo "
                  "con el Reglamento Electrotécnico para Baja Tensión (REBT, Real Decreto 842/2002) y sus "
                  "Instrucciones Técnicas Complementarias (ITC-BT), en particular la ITC-BT-04. La "
                  f"actuación descrita en este documento {frase_actuacion}. Esta MTD es válida para "
                  "instalaciones de las características recogidas en el apartado 3 de la ITC-BT-04 que no "
                  "requieren de Proyecto firmado por técnico titulado competente.")
    if tipo_actuacion in ("Ampliación", "Reforma"):
        _docx_parrafo(doc, "Al tratarse de una actuación sobre una instalación preexistente, deberá "
                      "verificarse además que las partes no incluidas en esta Memoria mantienen las "
                      "condiciones de seguridad exigidas por el REBT, y que la actuación descrita no "
                      "compromete la protección del conjunto.")

    _docx_heading(doc, "D.2 Acometida", 3, COBRE)
    _docx_parrafo(doc, "Parte de la instalación de la red de distribución que alimenta la caja general de "
                  "protección o unidad funcional equivalente (ITC-BT-11). Según su trazado podrá ser aérea "
                  "posada sobre fachada, aérea tensada sobre postes, subterránea o aero-subterránea; los "
                  "cables serán aislados, de tensión asignada no inferior a 0,6/1 kV. La acometida forma "
                  "parte de la instalación de la empresa distribuidora, ajustándose a sus normas "
                  "particulares.")

    _docx_heading(doc, "D.3 Instalaciones de enlace", 3, COBRE)
    _docx_parrafo(doc, "Caja de Protección y Medida (CPM): para suministros a un único usuario se instala "
                  "un único elemento que agrupa la caja general de protección y el equipo de medida "
                  "(ITC-BT-13), en fachada, en nicho con puerta IK10, con los fusibles de seguridad "
                  "correspondientes.")
    _docx_parrafo(doc, "Derivación individual: desde la CPM suministra energía al usuario, comprendiendo "
                  "fusibles de seguridad, conjunto de medida y dispositivos generales de mando y protección "
                  "(ITC-BT-15). Conductores de cobre o aluminio, sección mínima 6 mm² para fases/neutro/"
                  "protección, caída de tensión máxima 1,5% sin línea general de alimentación.")

    _docx_heading(doc, "D.4 Instalación de baja tensión calculada", 3, COBRE)
    if hay_cable:
        _docx_parrafo(doc, f"Circuito de referencia: {inputs_cable['tipo_circuito']}.")
        _docx_tabla(doc, [
            ["Concepto", "Valor"],
            ["Sistema", f"{inputs_cable['sistema']} — {inputs_cable['tension']:g} V"],
            ["Conductor / Aislamiento", f"{inputs_cable['conductor']} / {inputs_cable['aislamiento']}"],
            ["Método de instalación", inputs_cable["metodo"]],
            ["Longitud del circuito", f"{inputs_cable['longitud']:g} m"],
            ["Sección de fase adoptada", f"{resultado_cable['seccion_final']:g} mm²"],
            ["Protección (interruptor automático)", f"{resultado_cable['calibre_magnetotermico']} A"],
            ["Caída de tensión", f"{resultado_cable['e_final_pct']:.2f} % (máx. {inputs_cable['delta_u_max']:g} %)"],
        ], AZUL)
    else:
        _docx_parrafo(doc, "No se ha completado ningún cálculo de circuito en la pestaña Calculadora; esta "
                      "sección se completará cuando se disponga de esos datos.")

    _docx_heading(doc, "D.5 Cuadro general de mando y protección", 3, COBRE)
    _docx_parrafo(doc, "Aloja, como mínimo: interruptor general automático (IGA) de corte omnipolar con "
                  "poder de corte suficiente (4,5 kA mínimo en vivienda); interruptor diferencial general "
                  "(o varios, según la previsión de cargas); e interruptores automáticos individuales de "
                  "cada circuito derivado, coordinados para garantizar la selectividad. Altura de los "
                  "dispositivos entre 1 y 2 m. Se recomienda DPS en cabecera si la acometida es aérea o "
                  "existe riesgo de sobretensión atmosférica (ITC-BT-23).")

    _docx_heading(doc, "D.6 Conductores: materiales, dimensionado e identificación", 3, COBRE)
    _docx_parrafo(doc, "Conductores de cobre o aluminio, siempre aislados, tensión asignada no inferior a "
                  "450/750 V. Sección determinada por el criterio más desfavorable entre intensidad máxima "
                  "admisible (ITC-BT-19) y caída de tensión (3% alumbrado / 5% otros usos en instalación "
                  "interior; 1,5% en derivación individual sin LGA). Identificación por color: azul claro "
                  "el neutro, verde-amarillo la protección, marrón/negro/gris las fases.")

    _docx_heading(doc, "D.7 Subdivisión, equilibrado de cargas y conexiones", 3, COBRE)
    _docx_parrafo(doc, "La instalación se subdivide en circuitos independientes y adecuadamente protegidos. "
                  "Se procura el mayor equilibrio posible de cargas entre fases. Las conexiones se realizan "
                  "siempre en cajas de empalme mediante bornes de conexión, sin permitirse la unión por "
                  "simple retorcimiento.")

    _docx_heading(doc, "D.8 Sistemas de instalación (canalizaciones)", 3, COBRE)
    metodo_txt = inputs_cable.get("metodo", "") if hay_cable else "[a determinar]"
    _docx_parrafo(doc, f"El sistema de canalización empleado se corresponde con el tipo {metodo_txt} según "
                  "la clasificación de la ITC-BT-19 y la norma UNE-HD 60364-5-52, según las influencias "
                  "externas del emplazamiento. Las características mínimas de los tubos protectores se "
                  "detallan en el Pliego de Condiciones Técnicas.")

    _docx_heading(doc, "D.9 Protección contra sobreintensidades y sobretensiones", 3, COBRE)
    _docx_parrafo(doc, "Todo circuito está protegido contra sobrecargas (interruptor automático, curva "
                  "térmica) y contra cortocircuitos (capacidad de corte acorde a la Icc prevista, "
                  "ITC-BT-22). Si la instalación se alimenta por línea aérea, o se opta por protección "
                  "adicional, se instala protección contra sobretensiones transitorias (ITC-BT-23).")
    _docx_tabla(doc, [
        ["Categoría", "Tensión soportada (230/400V)", "Equipos típicos"],
        ["IV", "6 kV", "Contadores, equipos principales de protección"],
        ["III", "4 kV", "Cuadros de distribución, aparamenta fija"],
        ["II", "2,5 kV", "Electrodomésticos, herramientas portátiles"],
        ["I", "1,5 kV", "Equipos electrónicos muy sensibles"],
    ], AZUL)

    _docx_heading(doc, "D.10 Protección contra contactos directos e indirectos", 3, COBRE)
    _docx_parrafo(doc, "Contactos directos: aislamiento de las partes activas, complementado con "
                  "diferenciales de sensibilidad ≤30 mA. Contactos indirectos: corte automático de la "
                  "alimentación, cumpliendo Ra·Ia ≤ U (50 V en locales secos, 24 V en húmedos).")

    _docx_heading(doc, "D.11 Puesta a tierra", 3, COBRE)
    _docx_parrafo(doc, "Ejecutada conforme a la ITC-BT-18, mediante electrodo(s) conectado al borne "
                  "principal de tierra, del que parte el conductor de protección hacia todas las masas "
                  "metálicas. La resistencia de tierra será compatible con la sensibilidad de las "
                  "protecciones diferenciales instaladas.")

    contador_d = 11
    if hay_fv:
        contador_d += 1
        _docx_heading(doc, f"D.{contador_d} Instalación generadora fotovoltaica", 3, COBRE)
        _docx_tabla(doc, [
            ["Concepto", "Valor"],
            ["Modalidad de autoconsumo (RD 244/2019)", inputs_fv["tipo_autoconsumo"]],
            ["Potencia pico instalada", f"{resultado_fv['p_pico_kwp']:.2f} kWp"],
            ["Nº de paneles / configuración",
             f"{resultado_fv['n_paneles_configurados']} ({resultado_fv['n_serie']}S{resultado_fv['n_paralelo']}P)"],
            ["Potencia del inversor", f"{inputs_fv['potencia_inversor_kw']:g} kW"],
            ["Producción anual estimada", f"{_miles(resultado_fv['produccion_anual_kwh'])} kWh/año"],
        ], AZUL)
        umbral = ("≤10 kW → Certificado eléctrico firmado por instalador autorizado" if resultado_fv["p_pico_kwp"] <= 10
                  else ">10 kW → Proyecto firmado por técnico competente")
        _docx_parrafo(doc, f"Régimen documental aplicable según RD 244/2019: {umbral}.")

    contador_d += 1
    _docx_heading(doc, f"D.{contador_d} Receptores de alumbrado", 3, COBRE)
    _docx_parrafo(doc, "Luminarias conformes a la serie UNE-EN 60598. Partes metálicas accesibles (Clase I) "
                  "conectadas al conductor de protección. Con lámparas de descarga, carga mínima prevista "
                  "1,8 veces la potencia en vatios (ITC-BT-44), factor de potencia compensado a ≥0,9.")

    if hay_cable and inputs_cable.get("es_motor"):
        contador_d += 1
        _docx_heading(doc, f"D.{contador_d} Receptores a motor", 3, COBRE)
        _docx_parrafo(doc, "Conforme a la ITC-BT-47, el conductor se dimensiona para el 125% de la "
                      "intensidad a plena carga (motor único) o el 125% del mayor más el 100% del resto "
                      "(varios motores). Protección contra cortocircuitos y sobrecargas en todas las fases.")

    contador_d += 1
    _docx_heading(doc, f"D.{contador_d} Influencias externas", 3, COBRE)
    _docx_parrafo(doc, "Salvo indicación contraria, se considera que el emplazamiento presenta condiciones "
                  "ambientales normales (AD1, AA4, BA1 según UNE 20460-3), sin riesgo de incendio ni "
                  "explosión.")

    _docx_heading(doc, "E. Memoria justificativa (resumen de cálculos)", 2, AZUL)
    _docx_parrafo(doc, "El detalle completo de fórmulas y valores se recoge en el Anexo de Cálculos y "
                  "Mediciones, documento que forma parte integrante de esta Memoria.")
    if not hay_cable and not hay_fv:
        _docx_parrafo(doc, "No hay cálculos disponibles todavía.", cursiva=True)

    _docx_heading(doc, "F. Presupuesto", 2, AZUL)
    _docx_parrafo(doc, f"El presupuesto de la instalación asciende a la cantidad de "
                  f"{_fmt_eur(total_presupuesto)} ({numero_a_letras_euros(total_presupuesto).capitalize()}), "
                  "IVA incluido, según el desglose por capítulos que se adjunta en el documento de "
                  "Presupuesto.")

    _docx_heading(doc, "G. Documentos que integran el proyecto", 2, AZUL)
    _docx_lista(doc, ["Memoria Técnica de Diseño (este documento)", "Anexo de Cálculos y Mediciones",
                       "Pliego de Condiciones Generales", "Presupuesto",
                       "Esquema unifilar y croquis de emplazamiento (a incorporar por el técnico)",
                       "Certificado de Instalación Eléctrica (CIE), tras la puesta en servicio"])

    _docx_heading(doc, "H. Declaración y firma", 2, AZUL)
    _docx_parrafo(doc, f"D./Dña. _______________________, en calidad de {d.get('categoria_instalador', 'básica').lower()}, "
                  "declara que la instalación descrita en la presente Memoria Técnica de Diseño se ajusta "
                  "a las prescripciones del Reglamento Electrotécnico para Baja Tensión y sus Instrucciones "
                  "Técnicas Complementarias.")
    _docx_parrafo(doc, "En ______________________, a ____ de ______________ de 20____")
    _docx_parrafo(doc, "Firma: _______________________________________")

    return _docx_bytes(doc)


def generar_pdf_anexo_calculos(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                                 inputs_fv: dict, resultado_fv: dict, capitulos_presupuesto: list,
                                 pct_beneficio: float, pct_amortizacion: float,
                                 config_prof: dict = None) -> bytes:
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "ANEXO DE CALCULOS Y MEDICIONES", "Justificacion tecnica y mediciones", datos_proyecto, config_prof)
    formula = ParagraphStyle("formula", parent=normal, alignment=TA_CENTER, textColor=COBRE,
                              fontSize=10, leading=15, spaceBefore=4, spaceAfter=8)

    story = _bloque_portada("Anexo de Cálculos y Mediciones", "Justificación técnica y mediciones",
                             datos_proyecto, config_prof, AZUL, COBRE, colors, h1, normal)
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    story.append(Paragraph("0. Introducción y criterios generales", h2))
    story.append(Paragraph(
        "Este Anexo justifica, mediante las fórmulas y valores empleados, las secciones de conductor, "
        "protecciones y caídas de tensión adoptadas en la instalación descrita en la Memoria Técnica de "
        "Diseño, así como las mediciones que dan origen al Presupuesto. Los criterios generales aplicados "
        "son los del artículo 19.2 de la ITC-BT-19 (criterio térmico, criterio de caída de tensión y, "
        "cuando se aporta, criterio térmico de cortocircuito) y, para la instalación fotovoltaica en su "
        "caso, la ITC-BT-40 y el Real Decreto 244/2019.", normal))

    story.append(Paragraph("1. Cálculos justificativos", h2))
    if hay_cable:
        story.append(Paragraph(f"1.1. Instalación de baja tensión — {inputs_cable['tipo_circuito']}", h3))
        story.extend(_parrafos_calculo_cable_pdf(inputs_cable, resultado_cable, normal, formula))
    if hay_fv:
        story.append(Paragraph("1.2. Instalación fotovoltaica", h3))
        story.extend(_parrafos_calculo_fv_pdf(inputs_fv, resultado_fv, normal, formula))
    if not hay_cable and not hay_fv:
        story.append(Paragraph("No hay cálculos disponibles: completa la Calculadora de cables y/o el "
                                "módulo Fotovoltaico antes de generar este anexo.", normal))

    story.append(Paragraph("2. Mediciones", h2))
    story.append(Paragraph(
        "Relación de unidades de obra y materiales, agrupadas por capítulo, que dan origen al Presupuesto "
        "de la instalación. El precio unitario (Pu) incluye ya el beneficio industrial y la amortización de "
        "medios auxiliares aplicados en el documento de Presupuesto.", normal))
    if capitulos_presupuesto:
        for cap in capitulos_presupuesto:
            if not cap["items"]:
                continue
            story.append(Paragraph(cap["nombre"], h3))
            filas = [["Partida", "Designación", "Ud.", "Cantidad", "Pu (€)", "Importe (€)"]]
            total_cap = 0.0
            for it in cap["items"]:
                pu = calcular_precio_venta(it["precio_base"], pct_beneficio, pct_amortizacion)
                importe = round(it["cantidad"] * pu, 2)
                total_cap += importe
                filas.append([it.get("partida", "-"), it["designacion"], it["unidades"],
                              f"{it['cantidad']:g}", _miles(pu, 2), _miles(importe, 2)])
            filas.append(["", "", "", "", "<b>TOTAL</b>", f"<b>{_miles(total_cap, 2)} €</b>"])
            story.append(_tabla_pdf(filas, (1.6, 6.6, 1.4, 1.9, 2.1, 2.4), AZUL, colors, fuente=7.8,
                                     alinear_num=[3, 4, 5]))
            story.append(Spacer(1, 4))
        story.append(Paragraph(f"Porcentajes aplicados: Beneficio industrial {pct_beneficio:g}%, "
                                f"Amortización de medios auxiliares {pct_amortizacion:g}%. El desglose "
                                "económico completo, con IVA y capítulos agrupados, se encuentra en el "
                                "documento de Presupuesto.", normal))
    else:
        story.append(Paragraph("No se han definido capítulos en la pestaña Presupuesto.", normal))

    story.append(Paragraph("3. Conclusión", h2))
    story.append(Paragraph(
        "Con los cálculos y mediciones anteriores queda justificada, a juicio del técnico que suscribe, la "
        "idoneidad de las secciones, protecciones y demás elementos descritos en la Memoria Técnica de "
        "Diseño para el uso previsto de la instalación.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


def generar_docx_anexo_calculos(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                                 inputs_fv: dict, resultado_fv: dict, capitulos_presupuesto: list,
                                 pct_beneficio: float, pct_amortizacion: float,
                                 config_prof: dict = None) -> bytes:
    """Versión en Word del Anexo de Cálculos y Mediciones."""
    doc, AZUL, COBRE, GRIS = _docx_preparar("ANEXO DE CÁLCULOS Y MEDICIONES", "Justificación técnica y mediciones",
                                             datos_proyecto, config_prof)
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    _docx_heading(doc, "0. Introducción y criterios generales", 2, AZUL)
    _docx_parrafo(doc, "Este Anexo justifica, mediante las fórmulas y valores empleados, las secciones de "
                  "conductor, protecciones y caídas de tensión adoptadas en la instalación descrita en la "
                  "Memoria Técnica de Diseño, así como las mediciones que dan origen al Presupuesto. "
                  "Criterios generales: artículo 19.2 de la ITC-BT-19 (criterio térmico, criterio de caída "
                  "de tensión y, cuando se aporta, criterio térmico de cortocircuito) y, para la instalación "
                  "fotovoltaica, la ITC-BT-40 y el Real Decreto 244/2019.")

    _docx_heading(doc, "1. Cálculos justificativos", 2, AZUL)
    if hay_cable:
        _docx_heading(doc, f"1.1. Instalación de baja tensión — {inputs_cable['tipo_circuito']}", 3, COBRE)
        _docx_tabla(doc, [
            ["Magnitud", "Valor"],
            ["Intensidad de empleo Ib", f"{resultado_cable['ib']:.2f} A"],
            ["Intensidad de cálculo (tras factores)", f"{resultado_cable['ib_calculo']:.2f} A"],
            ["Iz obtenida (tabla × factores)", f"{resultado_cable['iz_termica']:.1f} A" if resultado_cable.get("iz_termica") else "-"],
            ["Sección adoptada", f"{resultado_cable['seccion_final']:g} mm²"],
            ["Caída de tensión", f"{resultado_cable['e_final_pct']:.2f} % (máx. {inputs_cable['delta_u_max']:g} %)"],
            ["Protección", f"{resultado_cable['calibre_magnetotermico']} A"],
        ], AZUL)
        _docx_parrafo(doc, "La sección se ha determinado por el criterio más desfavorable entre intensidad "
                      "máxima admisible (Ib ≤ In ≤ Iz, ITC-BT-19) y caída de tensión máxima admisible, "
                      "conforme se detalla en la tabla anterior.")
    else:
        _docx_parrafo(doc, "No se ha completado ningún cálculo de circuito en la pestaña Calculadora.",
                      cursiva=True)

    if hay_fv:
        _docx_heading(doc, "1.2. Instalación fotovoltaica", 3, COBRE)
        _docx_tabla(doc, [
            ["Magnitud", "Valor"],
            ["Potencia pico", f"{resultado_fv['p_pico_kwp']:.2f} kWp"],
            ["Nº de paneles / configuración",
             f"{resultado_fv['n_paneles_configurados']} ({resultado_fv['n_serie']}S{resultado_fv['n_paralelo']}P)"],
            ["Producción anual estimada", f"{_miles(resultado_fv['produccion_anual_kwh'])} kWh/año"],
            ["Producción año 10", f"{_miles(resultado_fv.get('produccion_ano10', 0))} kWh/año"],
            ["Producción año 25", f"{_miles(resultado_fv.get('produccion_ano25', 0))} kWh/año"],
            ["Sección CC / CA", f"{resultado_fv.get('s_cc_final', 0):g} mm² / {resultado_fv.get('s_ca_final', 0):g} mm²"],
            ["CO₂ evitado", f"{resultado_fv.get('co2_evitado_kg_ano', 0)/1000:.2f} t/año"],
        ], AZUL)
        _docx_parrafo(doc, "Dimensionado conforme a la ITC-BT-40: cables de conexión dimensionados al 125% "
                      "de la intensidad máxima del generador, con caída de tensión conjunta (CC+CA) no "
                      "superior al 1,5% entre el generador y el punto de interconexión.")

    _docx_heading(doc, "2. Mediciones", 2, AZUL)
    _docx_parrafo(doc, "Relación de unidades de obra y materiales, agrupadas por capítulo, que dan origen "
                  "al Presupuesto. El precio unitario (Pu) incluye ya el beneficio industrial y la "
                  "amortización de medios auxiliares.")
    if capitulos_presupuesto:
        for cap in capitulos_presupuesto:
            if not cap["items"]:
                continue
            _docx_heading(doc, cap["nombre"], 3, COBRE)
            filas = [["Partida", "Designación", "Ud.", "Cantidad", "Pu (€)", "Importe (€)"]]
            total_cap = 0.0
            for it in cap["items"]:
                pu = calcular_precio_venta(it["precio_base"], pct_beneficio, pct_amortizacion)
                importe = round(it["cantidad"] * pu, 2)
                total_cap += importe
                filas.append([it.get("partida", "-"), it["designacion"], it["unidades"],
                              f"{it['cantidad']:g}", _miles(pu, 2), _miles(importe, 2)])
            filas.append(["", "", "", "", "TOTAL", f"{_miles(total_cap, 2)} €"])
            _docx_tabla(doc, filas, AZUL)
        _docx_parrafo(doc, f"Porcentajes aplicados: Beneficio industrial {pct_beneficio:g}%, Amortización de "
                      f"medios auxiliares {pct_amortizacion:g}%. El desglose económico completo, con IVA, "
                      "se encuentra en el documento de Presupuesto.")
    else:
        _docx_parrafo(doc, "No se han definido capítulos en la pestaña Presupuesto.", cursiva=True)

    _docx_heading(doc, "3. Conclusión", 2, AZUL)
    _docx_parrafo(doc, "Con los cálculos y mediciones anteriores queda justificada, a juicio del técnico "
                  "que suscribe, la idoneidad de las secciones, protecciones y demás elementos descritos en "
                  "la Memoria Técnica de Diseño para el uso previsto de la instalación.")

    return _docx_bytes(doc)


def generar_docx_condiciones_generales(datos_proyecto: dict, hay_fv: bool, config_prof: dict = None) -> bytes:
    """Versión en Word del Pliego de Condiciones (facultativas, económicas y
    técnicas particulares), con el mismo contenido que la versión en PDF."""
    doc, AZUL, COBRE, GRIS = _docx_preparar("PLIEGO DE CONDICIONES", "Facultativas, económicas y técnicas particulares",
                                             datos_proyecto, config_prof)

    _docx_heading(doc, "0. Objeto y alcance", 2, AZUL)
    _docx_parrafo(doc, "El presente Pliego establece las prescripciones facultativas, económicas y técnicas "
                  "particulares que rigen la ejecución de la instalación eléctrica descrita en la Memoria "
                  "Técnica de Diseño y su Anexo de Cálculos, sin perjuicio de condiciones particulares más "
                  "restrictivas que pueda fijar la Dirección Facultativa, la empresa distribuidora o las "
                  "ordenanzas municipales y autonómicas de aplicación.")

    _docx_heading(doc, "1. Normativa de aplicación", 2, AZUL)
    _docx_lista(doc, [
        "Reglamento Electrotécnico para Baja Tensión (REBT), RD 842/2002, e Instrucciones Técnicas "
        "Complementarias (ITC-BT-01 a ITC-BT-53).",
        "Normas UNE armonizadas de materiales, cables y aparamenta (UNE-EN 50525, UNE-HD 60364-5-52, "
        "UNE-EN 60898, UNE-EN 61008/61009, UNE-EN 61439).",
        "Real Decreto 244/2019, autoconsumo de energía eléctrica (si la instalación incluye FV).",
        "Ley 31/1995 de Prevención de Riesgos Laborales y RD 1627/1997.",
        "Código Técnico de la Edificación y ordenanzas municipales o autonómicas aplicables.",
    ])

    _docx_heading(doc, "PARTE I — Condiciones facultativas", 1, AZUL)
    secciones_facultativas = [
        ("2. Técnico director / instalador autorizado", "Redacta complementos y rectificaciones, asiste a "
         "la obra, supervisa el replanteo, comprueba instalaciones provisionales y condiciones de "
         "seguridad, dirige la ejecución material, realiza las pruebas de calidad y suscribe el CIE."),
        ("3. Constructor o instalador", "Organiza los trabajos, elabora el Plan de Seguridad y Salud, "
         "asegura la idoneidad de los materiales, y suscribe las actas de recepción."),
        ("4. Verificación de la documentación y replanteo", "El instalador consignará por escrito que la "
         "documentación aportada resulta suficiente antes de iniciar los trabajos, e iniciará éstos con el "
         "replanteo de la instalación."),
        ("5. Orden y ritmo de ejecución", "Facultad del instalador, salvo variación técnica que disponga la "
         "Dirección Facultativa, dentro del plazo acordado con el titular."),
        ("6. Trabajos no estipulados y modificaciones", "El instalador ejecutará cuanto sea necesario para "
         "la buena ejecución, reflejando por escrito cualquier modificación sobre lo proyectado."),
        ("7. Obras y conexiones ocultas", "Se levantará croquis acotado de los tramos que queden ocultos, "
         "si el titular lo requiere, antes de su cierre."),
        ("8. Trabajos defectuosos y vicios ocultos", "El instalador es responsable de los defectos por mala "
         "gestión o deficiente calidad de los materiales, pudiendo la Dirección Facultativa exigir la "
         "demolición y reconstrucción a su costa."),
        ("9. Recepción y plazo de garantía", "Doce meses de garantía frente a defectos de ejecución, sin "
         "perjuicio de las garantías comerciales de los fabricantes de los equipos."),
    ]
    for titulo_s, texto_s in secciones_facultativas:
        _docx_heading(doc, titulo_s, 3, COBRE)
        _docx_parrafo(doc, texto_s)

    _docx_heading(doc, "PARTE II — Condiciones económicas", 1, AZUL)
    secciones_economicas = [
        ("10. Composición de precios y presupuesto", "El precio de cada partida resulta de sumar el precio "
         "base (coste directo) más los porcentajes de beneficio industrial y amortización de medios "
         "auxiliares. El IVA se aplica sobre dicha suma y no está incluido en los precios unitarios."),
        ("11. Certificaciones y forma de pago", "El instalador podrá presentar relaciones valoradas de la "
         "instalación ejecutada, tomando como base las mediciones practicadas y los precios del "
         "Presupuesto."),
        ("12. Precios contradictorios y revisión de precios", "Las unidades no previstas se fijarán por "
         "precio contradictorio antes de su ejecución. Salvo pacto en contrario, no se admite revisión de "
         "precios sobre unidades ya contratadas."),
        ("13. Seguro y conservación de la instalación", "Durante el plazo de garantía, la conservación "
         "corresponde al instalador salvo asunción expresa del titular. Se recomienda seguro de "
         "responsabilidad civil y daños a terceros durante la ejecución."),
    ]
    for titulo_s, texto_s in secciones_economicas:
        _docx_heading(doc, titulo_s, 3, COBRE)
        _docx_parrafo(doc, texto_s)

    _docx_heading(doc, "PARTE III — Condiciones técnicas particulares", 1, AZUL)
    _docx_parrafo(doc, "Ejecución y montaje de instalaciones eléctricas en baja tensión.", cursiva=True)
    secciones_tecnicas = [
        ("14. Condiciones generales de los materiales", "Todos los materiales serán de primera calidad, "
         "con marcado CE cuando sea de aplicación, y podrán someterse a análisis o pruebas por cuenta del "
         "instalador."),
        ("15. Canalizaciones eléctricas", "Cables bajo tubo, fijados sobre paredes, enterrados, empotrados, "
         "en huecos, bajo canales o en bandeja, según la Memoria y el Anexo de Cálculos. Diámetro de tubos "
         "según ITC-BT-21."),
        ("16. Conductores", "Cobre o aluminio, siempre aislados, tensión asignada ≥450/750 V (≥0,6/1 kV en "
         "acometidas, derivaciones individuales, enterrados o FV). Identificación por color según norma."),
        ("17. Cajas de empalme y derivación", "Material aislante no propagador de la llama o metálicas "
         "protegidas contra la corrosión; nunca unión de conductores por simple retorcimiento."),
        ("18. Mecanismos y tomas de corriente", "Corte de la corriente máxima sin arco permanente, mínimo "
         "10.000 maniobras; tomas de corriente con puesta a tierra."),
        ("19. Aparamenta de mando y protección", "Cuadros nuevos, ensamblados en fábrica, IP 30 mínimo. "
         "Interruptores automáticos UNE-EN 60898, diferenciales UNE-EN 61008/61009 (tipo A mínimo, tipo B "
         "en FV si procede), fusibles de alta capacidad de ruptura."),
        ("20. Receptores de alumbrado", "Conformes a UNE-EN 60598, con conexión a tierra en Clase I; carga "
         "mínima 1,8× la potencia en descarga; factor de potencia compensado ≥0,9."),
        ("21. Receptores a motor", "Conductores dimensionados al 125% de la intensidad a plena carga "
         "(ITC-BT-47); protección contra cortocircuitos y sobrecargas en todas las fases."),
        ("22. Puesta a tierra", "Conforme a la ITC-BT-18; conductores de tierra enterrados sin protección "
         "mínimo 25 mm² de cobre; tensión de contacto ≤24V en locales húmedos, ≤50V en los demás."),
    ]
    for titulo_s, texto_s in secciones_tecnicas:
        _docx_heading(doc, titulo_s, 3, COBRE)
        _docx_parrafo(doc, texto_s)

    if hay_fv:
        _docx_heading(doc, "23. Condiciones específicas de la instalación fotovoltaica", 3, COBRE)
        _docx_lista(doc, [
            "Cumplimiento de la ITC-BT-40 y el RD 244/2019.",
            "Cables CC/CA dimensionados al 125% de la intensidad máxima del generador, ΔU conjunta ≤1,5%.",
            "Mecanismo antivertido en autoconsumo sin excedentes (Anexo I, ITC-BT-40).",
            "Conectores CC normalizados, estancos e inconfundibles con otros sistemas.",
            "Estructura de soporte conectada equipotencialmente a tierra.",
        ])

    _docx_heading(doc, "24. Inspecciones y pruebas antes de la puesta en servicio", 3, COBRE)
    _docx_parrafo(doc, "Antes de la puesta en servicio se realizarán, como mínimo, las siguientes "
                  "comprobaciones (ITC-BT-05 y UNE-HD 60364-6):")
    _docx_tabla(doc, [["Ensayo", "Criterio de aceptación"]] + [list(e) for e in ENSAYOS_PUESTA_SERVICIO], AZUL)

    for titulo_s, texto_s in [
        ("25. Seguridad en los trabajos eléctricos", "Los trabajos se realizarán sin tensión, verificando "
         "su ausencia (las 'cinco reglas de oro'). Guantes y herramientas aislantes; aparatos portátiles de "
         "clase II o a tensión de seguridad."),
        ("26. Limpieza y mantenimiento", "Antes de la recepción, los cuadros se limpiarán de polvo y "
         "residuos de la ejecución. El titular revisará periódicamente conductores, protecciones y puesta "
         "a tierra (test del diferencial mensual)."),
        ("27. Criterios de medición", "Cables, bandejas y tubos por metro lineal; cuadros y receptores por "
         "unidad montada y conexionada, con los accesorios de montaje incluidos en el precio."),
    ]:
        _docx_heading(doc, titulo_s, 3, COBRE)
        _docx_parrafo(doc, texto_s)

    _docx_parrafo(doc, "Este pliego es un documento de apoyo estándar; adáptalo a las particularidades de "
                  "cada proyecto y a las ordenanzas municipales o autonómicas que puedan resultar de "
                  "aplicación.", cursiva=True)

    return _docx_bytes(doc)


def generar_pdf_condiciones_generales(datos_proyecto: dict, hay_fv: bool, config_prof: dict = None) -> bytes:
    from reportlab.platypus import Paragraph, Spacer, PageBreak

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "PLIEGO DE CONDICIONES", "Facultativas, económicas y técnicas particulares", datos_proyecto,
        config_prof)

    def lista(items, style=None):
        return [Paragraph("• " + t, style or lista_item) for t in items]

    def tabla(datos, colw=(6.0, 5.0, 5.0), fuente=8.3):
        return _tabla_pdf(datos, colw, AZUL, colors, fuente=fuente)

    story = _bloque_portada("Pliego de Condiciones", "Facultativas, económicas y técnicas particulares",
                             datos_proyecto, config_prof, AZUL, COBRE, colors, h1, normal)

    story.append(Paragraph("0. Objeto y alcance", h2))
    story.append(Paragraph(
        "El presente Pliego de Condiciones establece las prescripciones facultativas, económicas y técnicas "
        "particulares que rigen la ejecución de la instalación eléctrica descrita en la Memoria Técnica de "
        "Diseño y su Anexo de Cálculos, sin perjuicio de condiciones particulares más restrictivas que pueda "
        "fijar la Dirección Facultativa, la empresa distribuidora o las ordenanzas municipales y autonómicas "
        "de aplicación. Se estructura en tres partes: condiciones facultativas (relaciones entre los agentes "
        "que intervienen en la obra), condiciones económicas (valoración y abono de los trabajos) y "
        "condiciones técnicas particulares de la instalación eléctrica de baja tensión.", normal))

    story.append(Paragraph("1. Normativa de aplicación", h2))
    story.extend(lista([
        "Reglamento Electrotécnico para Baja Tensión (REBT), Real Decreto 842/2002, y sus Instrucciones "
        "Técnicas Complementarias (ITC-BT-01 a ITC-BT-53).",
        "Normas UNE armonizadas de materiales, cables (UNE-EN 50525, UNE-HD 60364-5-52) y aparamenta "
        "(UNE-EN 60898, UNE-EN 61008/61009, UNE-EN 61439).",
        "Real Decreto 244/2019, condiciones administrativas, técnicas y económicas del autoconsumo de "
        "energía eléctrica (cuando la instalación incluya generación fotovoltaica).",
        "Ley 31/1995 de Prevención de Riesgos Laborales y Real Decreto 1627/1997 sobre disposiciones "
        "mínimas de seguridad y salud en las obras de construcción.",
        "Código Técnico de la Edificación y ordenanzas municipales o autonómicas que resulten de aplicación.",
    ]))

    # ============================================================ PARTE I
    story.append(PageBreak())
    story.append(Paragraph("PARTE I — CONDICIONES FACULTATIVAS", h1))

    story.append(Paragraph("2. Técnico director / instalador autorizado", h2))
    story.append(Paragraph("Corresponde al técnico director o instalador autorizado que suscribe la "
                            "documentación técnica:", normal))
    story.extend(lista([
        "Redactar los complementos o rectificaciones del proyecto o memoria que se precisen.",
        "Asistir a la obra cuantas veces lo requiera su naturaleza, resolviendo las incidencias que se "
        "produzcan e impartiendo las órdenes complementarias necesarias.",
        "Efectuar o supervisar el replanteo de la instalación.",
        "Comprobar las instalaciones provisionales, medios auxiliares y condiciones de seguridad y salud.",
        "Ordenar y dirigir la ejecución material con arreglo al proyecto, a las normas técnicas y a las "
        "reglas de la buena práctica constructiva.",
        "Realizar o disponer las pruebas y verificaciones de materiales e instalaciones necesarias para "
        "asegurar la calidad de la ejecución.",
        "Suscribir el Certificado de Instalación Eléctrica (CIE) o el certificado final que corresponda.",
    ]))

    story.append(Paragraph("3. Constructor o instalador", h2))
    story.append(Paragraph("Corresponde al constructor o instalador:", normal))
    story.extend(lista([
        "Organizar los trabajos y disponer las instalaciones provisionales y medios auxiliares necesarios.",
        "Elaborar, cuando se requiera, el Plan de Seguridad y Salud en aplicación del estudio "
        "correspondiente, velando por su cumplimiento.",
        "Ostentar la jefatura del personal que intervenga en la obra y coordinar a los subcontratistas.",
        "Asegurar la idoneidad de todos los materiales y elementos que se utilicen, rechazando los "
        "suministros que no cuenten con las garantías o documentos de idoneidad exigidos.",
        "Facilitar al técnico director, con antelación suficiente, los datos precisos para el cumplimiento "
        "de su cometido.",
        "Suscribir con el titular las actas de recepción provisional y definitiva de la instalación.",
    ]))

    story.append(Paragraph("4. Verificación de la documentación y replanteo", h2))
    story.append(Paragraph(
        "Antes de dar comienzo a los trabajos, el instalador consignará por escrito que la documentación "
        "aportada (Memoria, Anexo de Cálculos, planos si los hubiera) le resulta suficiente para la "
        "comprensión de la totalidad de la instalación contratada o, en caso contrario, solicitará las "
        "aclaraciones pertinentes al técnico director. El instalador iniciará los trabajos con el replanteo "
        "de la instalación, sometiéndolo a la aprobación del técnico director cuando este exista.", normal))

    story.append(Paragraph("5. Orden y ritmo de ejecución de los trabajos", h2))
    story.append(Paragraph(
        "La determinación del orden de los trabajos es, con carácter general, facultad del instalador, "
        "salvo que por circunstancias de orden técnico convenga su variación a juicio de la Dirección "
        "Facultativa. Los trabajos se desarrollarán de forma que la ejecución total se lleve a efecto dentro "
        "del plazo acordado con el titular, comunicando el instalador el inicio de los trabajos con la "
        "antelación que se acuerde.", normal))

    story.append(Paragraph("6. Trabajos no estipulados expresamente y modificaciones", h2))
    story.append(Paragraph(
        "Es obligación del instalador ejecutar cuanto sea necesario para la buena ejecución de la "
        "instalación, aun cuando no se halle expresamente detallado en la Memoria o el Anexo de Cálculos, "
        "siempre que, sin apartarse de su espíritu, lo disponga el técnico director. Cualquier modificación "
        "sobre lo proyectado deberá quedar reflejada por escrito y, en su caso, en la documentación "
        "'as built' entregada al titular a la finalización de los trabajos.", normal))

    story.append(Paragraph("7. Obras y conexiones ocultas", h2))
    story.append(Paragraph(
        "De todos los tramos de canalización y conexiones que hayan de quedar ocultos (empotrados, "
        "enterrados o bajo falso techo) se levantará, si el titular lo requiere, croquis suficientemente "
        "acotado antes de proceder a su cierre, de forma que quede constancia de su trazado para futuras "
        "intervenciones o ampliaciones.", normal))

    story.append(Paragraph("8. Trabajos defectuosos y vicios ocultos", h2))
    story.append(Paragraph(
        "El instalador es responsable de la ejecución de los trabajos contratados y de los defectos que en "
        "ellos puedan existir por su mala gestión o por deficiente calidad de los materiales empleados, sin "
        "que le exima de responsabilidad el hecho de que los trabajos hayan sido facturados o certificados a "
        "buena cuenta. Si el técnico director advirtiese vicios o defectos, podrá disponer que las partes "
        "defectuosas sean demolidas y reconstruidas de acuerdo con lo contratado, a costa del instalador.",
        normal))

    story.append(Paragraph("9. Recepción y plazo de garantía", h2))
    story.append(Paragraph(
        "Concluidos los trabajos y realizadas con resultado favorable las verificaciones descritas en el "
        "apartado 24, se procederá a la recepción de la instalación por el titular. El instalador garantizará "
        "la instalación frente a defectos de ejecución durante el plazo legal aplicable (con carácter "
        "general, doce meses), sin perjuicio de las garantías comerciales de cada fabricante sobre los "
        "equipos suministrados (aparamenta, luminarias, inversor, paneles, baterías). Durante dicho plazo, "
        "el instalador corregirá a su cargo los defectos observados y reparará las averías que por causa "
        "imputable a la ejecución se produjeran.", normal))

    # ============================================================ PARTE II
    story.append(PageBreak())
    story.append(Paragraph("PARTE II — CONDICIONES ECONÓMICAS", h1))

    story.append(Paragraph("10. Composición de precios y presupuesto", h2))
    story.append(Paragraph(
        "El precio de cada partida del Presupuesto resulta de sumar el precio base del material o unidad de "
        "obra (coste directo) más los porcentajes de beneficio industrial y amortización de medios "
        "auxiliares que se detallan en el propio documento de Presupuesto. El IVA se aplica sobre dicha "
        "suma y no está incluido en los precios unitarios.", normal))

    story.append(Paragraph("11. Certificaciones y forma de pago", h2))
    story.append(Paragraph(
        "Cuando así se acuerde entre las partes, el instalador podrá presentar relaciones valoradas de la "
        "instalación ejecutada en los plazos que se establezcan, tomando como base las mediciones "
        "practicadas y los precios del Presupuesto. Los pagos se efectuarán por el titular en los plazos "
        "acordados, correspondiendo su importe al de las certificaciones o facturas conformadas.", normal))

    story.append(Paragraph("12. Precios contradictorios y revisión de precios", h2))
    story.append(Paragraph(
        "Si durante la ejecución fuese necesario introducir unidades no previstas en el Presupuesto, se "
        "fijará un precio contradictorio antes de su ejecución, tomando como referencia los precios más "
        "análogos del propio Presupuesto o, en su defecto, precios de mercado de uso habitual en la zona. "
        "Salvo pacto en contrario, no se admitirá revisión de precios sobre las unidades ya contratadas.",
        normal))

    story.append(Paragraph("13. Seguro y conservación de la instalación", h2))
    story.append(Paragraph(
        "Durante el plazo de garantía, la conservación de la instalación corresponderá al instalador salvo "
        "que el titular la hubiese asumido expresamente. Se recomienda que la instalación quede cubierta por "
        "un seguro de responsabilidad civil y daños a terceros durante la ejecución de los trabajos, a "
        "suscribir por el instalador conforme a la normativa vigente.", normal))

    # ============================================================ PARTE III
    story.append(PageBreak())
    story.append(Paragraph("PARTE III — CONDICIONES TÉCNICAS PARTICULARES", h1))
    story.append(Paragraph("Ejecución y montaje de instalaciones eléctricas en baja tensión", normal))

    story.append(Paragraph("14. Condiciones generales de los materiales", h2))
    story.extend(lista([
        "Todos los materiales serán de primera calidad y reunirán las condiciones exigidas por el REBT y "
        "demás disposiciones vigentes, disponiendo del marcado CE cuando sea de aplicación.",
        "Podrán ser sometidos a los análisis o pruebas que se crean necesarios para acreditar su calidad, "
        "por cuenta del instalador.",
        "Los materiales no consignados expresamente que den lugar a precios contradictorios reunirán las "
        "condiciones de idoneidad necesarias a juicio de la Dirección Facultativa.",
        "Todos los trabajos se ejecutarán esmeradamente, con arreglo a las buenas prácticas de las "
        "instalaciones eléctricas, cumpliendo estrictamente el REBT y las instrucciones de la Dirección "
        "Facultativa.",
    ]))

    story.append(Paragraph("15. Canalizaciones eléctricas", h2))
    story.append(Paragraph(
        "Los cables se colocarán bajo tubo, fijados directamente sobre paredes, enterrados, empotrados en "
        "estructuras, en huecos de la construcción, bajo canales o molduras, o en bandeja, según se haya "
        "definido en la Memoria y el Anexo de Cálculos. Antes de iniciar el tendido deberán estar ejecutados "
        "los elementos estructurales que hayan de soportar o alojar la canalización.", normal))
    story.append(Paragraph("15.1 Tubos protectores — características mínimas según norma UNE-EN 61386",
                            h3))
    story.append(tabla([
        ["Tipo de instalación", "Resist. compresión", "Resist. impacto", "Temp. mín / máx"],
        ["Superficie", "Fuerte (4)", "Media (3)", "-5 ºC / +60 ºC"],
        ["Empotrado en obra de fábrica", "Ligera (2)", "Ligera (2)", "-5 ºC / +60 ºC"],
        ["Empotrado embebido en hormigón", "Media (3)", "Media (3)", "-5 ºC / +90 ºC"],
        ["Aéreo / flexible", "Fuerte (4)", "Media (3)", "-5 ºC / +60 ºC"],
        ["Enterrado", "250-750 N según terreno", "Ligero-Normal", "No aplicable"],
    ], colw=(5.5, 4.2, 3.8, 4.5)))
    story.append(Paragraph(
        "El diámetro exterior mínimo de los tubos, en función del número y sección de los conductores, se "
        "obtendrá de las tablas de la ITC-BT-21. En tramos rectos los registros no distarán más de 15 m "
        "entre sí, ni habrá más de 3 curvas entre dos registros consecutivos. Los tubos metálicos accesibles "
        "se pondrán a tierra y no se emplearán como conductor de protección o de neutro.", normal))
    story.append(Paragraph("15.2 Otros sistemas de canalización", h3))
    story.extend(lista([
        "Conductores fijados directamente sobre paredes: cables de tensión asignada no inferior a 0,6/1 kV, "
        "fijados con bridas o abrazaderas a una distancia máxima de 0,40 m entre puntos de fijación.",
        "Conductores enterrados: bajo tubo salvo que dispongan de cubierta y tensión asignada 0,6/1 kV, "
        "según ITC-BT-07 e ITC-BT-21.",
        "Conductores en bandeja: unipolares o multipolares con cubierta, bandeja de acero galvanizado con "
        "ancho mínimo de 100 mm, sujeta de forma que no se produzcan flechas superiores a 10 mm.",
        "Conductores bajo canal protectora: canal con grado de protección IP4X mínimo, tapa desmontable "
        "solo con herramienta, conectada a la red de tierra si es metálica.",
    ]))

    story.append(Paragraph("16. Conductores", h2))
    story.append(Paragraph(
        "Los conductores serán de cobre (o aluminio cuando lo justifique el cálculo) y siempre aislados, con "
        "tensión asignada no inferior a 450/750 V en instalación interior y 0,6/1 kV en acometidas, "
        "derivaciones individuales, enterrados o instalación fotovoltaica. Los conductores de sección igual "
        "o superior a 6 mm² estarán constituidos por cable trenzado de hilo de cobre.", normal))
    story.append(Paragraph(
        "La sección se determinará por el criterio más desfavorable entre intensidad máxima admisible "
        "(ITC-BT-19, con los factores de corrección que procedan) y caída de tensión (3% alumbrado / 5% "
        "otros usos en instalación interior; 1,5% en derivación individual sin LGA), según se desarrolla en "
        "el Anexo de Cálculos. La sección del conductor neutro y del conductor de protección seguirán, "
        "respectivamente, el criterio general del REBT y la tabla de la ITC-BT-18.", normal))
    story.append(Paragraph(
        "Identificación: neutro en azul claro, protección en verde-amarillo, fases en marrón, negro o gris. "
        "Todo conductor de fase, o aquel para el que no se prevea su paso posterior a neutro, se identificará "
        "con estos últimos colores.", normal))

    story.append(Paragraph("17. Cajas de empalme y derivación", h2))
    story.append(Paragraph(
        "Las conexiones entre conductores se realizarán en el interior de cajas de material aislante no "
        "propagador de la llama o metálicas protegidas contra la corrosión, de dimensiones suficientes para "
        "alojar holgadamente los conductores que deban contener (profundidad mínima 40 mm, lado o diámetro "
        "mínimo 60-80 mm según el tipo de instalación). No se permitirá la unión de conductores por simple "
        "retorcimiento; se utilizarán siempre bornes de conexión, regletas o sistemas equivalentes.", normal))

    story.append(Paragraph("18. Mecanismos y tomas de corriente", h2))
    story.append(Paragraph(
        "Los interruptores y conmutadores cortarán la corriente máxima del circuito sin formación de arco "
        "permanente, serán de tipo cerrado y material aislante, y soportarán un mínimo de 10.000 maniobras "
        "con su carga nominal. Las tomas de corriente dispondrán, como norma general, de puesta a tierra e "
        "irán instaladas en caja empotrada, quedando al exterior únicamente el mando y la tapa "
        "embellecedora.", normal))

    story.append(Paragraph("19. Aparamenta de mando y protección", h2))
    story.append(Paragraph("19.1 Cuadros eléctricos", h3))
    story.append(Paragraph(
        "Serán nuevos, diseñados para servicio interior, ensamblados y cableados en fábrica, con estructura "
        "adecuada para montaje mural o sobre el suelo y grado de protección acorde a su ubicación (mínimo "
        "IP 30 según UNE-EN 60529). Cada circuito de salida estará protegido contra sobrecargas y "
        "cortocircuitos, y contra corrientes de defecto mediante diferencial de sensibilidad adecuada "
        "(ITC-BT-24). Dispondrán de un espacio de reserva no inferior al 20% de los módulos instalados y "
        "todos sus componentes serán accesibles desde el frente.", normal))
    story.append(Paragraph("19.2 Interruptores automáticos, diferenciales y fusibles", h3))
    story.extend(lista([
        "Los interruptores automáticos (PIA) cumplirán la norma UNE-EN 60898, con poder de corte no "
        "inferior a la intensidad de cortocircuito prevista en su punto de instalación, y curva térmica y "
        "electromagnética adecuadas a la carga protegida.",
        "Los interruptores diferenciales cumplirán UNE-EN 61008/61009, siendo de tipo A o superinmunizado "
        "en circuitos con receptores electrónicos, y de tipo B en instalación fotovoltaica cuando el "
        "inversor no garantice por diseño la ausencia de componente continua de defecto.",
        "Los fusibles serán de alta capacidad de ruptura; de acción lenta en protección de motores y de "
        "acción rápida en protección de circuitos de consumidores óhmicos. Se montarán de forma que no "
        "pueda proyectarse metal fundido al fundirse.",
        "Los guardamotores dispondrán de protección térmica de rearme manual, con relé de característica "
        "retardada en arranques de larga duración.",
    ]))
    story.append(Paragraph("19.3 Seccionadores y embarrados", h3))
    story.append(Paragraph(
        "Los seccionadores serán de conexión y desconexión brusca, independiente de la acción del operador, "
        "adecuados para abrir y cerrar la corriente nominal a tensión nominal. El embarrado principal "
        "constará de barras de cobre electrolítico de alta conductividad para las fases y una barra "
        "(seccionable) para el neutro, más una barra independiente de puesta a tierra.", normal))

    story.append(Paragraph("20. Receptores de alumbrado", h2))
    story.append(Paragraph(
        "Las luminarias cumplirán la serie de normas UNE-EN 60598. Sus partes metálicas accesibles, si no "
        "son de Clase II o III, dispondrán de conexión al conductor de protección. Para receptores con "
        "lámparas de descarga, la carga mínima prevista será de 1,8 veces la potencia en vatios de las "
        "lámparas (ITC-BT-44), siendo obligatoria la compensación del factor de potencia hasta un valor "
        "mínimo de 0,9.", normal))

    story.append(Paragraph("21. Receptores a motor", h2))
    story.append(Paragraph(
        "Los conductores de conexión de un motor único se dimensionarán para el 125% de su intensidad a "
        "plena carga; si alimentan varios motores, para el 125% del de mayor potencia más el 100% del resto "
        "(ITC-BT-47). Los motores estarán protegidos contra cortocircuitos y sobrecargas en todas sus fases "
        "y, cuando proceda, contra la falta de tensión. Los motores de potencia superior a 0,75 kW estarán "
        "provistos de dispositivo de arranque que limite la relación entre la corriente de arranque y la "
        "nominal, conforme a la tabla de la ITC-BT-47 según su potencia.", normal))

    story.append(Paragraph("22. Puesta a tierra", h2))
    story.append(Paragraph(
        "La puesta a tierra se ejecutará conforme a la ITC-BT-18, mediante electrodo o grupo de electrodos "
        "(picas, placas o conductor enterrado) conectados al borne principal de tierra, del que partirá el "
        "conductor de protección hacia todas las masas metálicas de la instalación. Los conductores de "
        "tierra enterrados sin protección mecánica ni contra la corrosión serán de cobre de 25 mm² mínimo. "
        "El valor de la resistencia de tierra será tal que ninguna masa pueda dar lugar a tensiones de "
        "contacto superiores a 24 V en locales húmedos o conductores y 50 V en los demás casos, condición "
        "que se verificará mediante la relación Ra·Ia ≤ U.", normal))

    if hay_fv:
        story.append(Paragraph("23. Condiciones específicas de la instalación fotovoltaica", h2))
        story.extend(lista([
            "La instalación generadora cumplirá la ITC-BT-40 y el Real Decreto 244/2019 sobre condiciones "
            "administrativas, técnicas y económicas del autoconsumo de energía eléctrica.",
            "Los cables de conexión (tramos CC y CA) se dimensionarán para una intensidad no inferior al "
            "125% de la intensidad máxima del generador, con una caída de tensión conjunta no superior al "
            "1,5% entre el generador y el punto de interconexión.",
            "Las instalaciones de autoconsumo sin excedentes dispondrán de un mecanismo antivertido conforme "
            "al Anexo I de la ITC-BT-40.",
            "Los conectores del lado de corriente continua serán del tipo normalizado para aplicaciones "
            "fotovoltaicas, estancos e inconfundibles con los de otros sistemas.",
            "Las estructuras metálicas de soporte de los módulos se conectarán equipotencialmente a tierra.",
        ]))
        n = 24
    else:
        n = 23

    story.append(Paragraph(f"{n}. Inspecciones y pruebas antes de la puesta en servicio", h2))
    story.append(Paragraph(
        "Antes de la puesta en servicio se realizarán, como mínimo, las siguientes comprobaciones "
        "(ITC-BT-05 y UNE-HD 60364-6):", normal))
    story.append(tabla([
        ["Ensayo", "Criterio de aceptación"],
        ["Continuidad de conductores de protección y equipotenciales", "Continuidad eléctrica verificada"],
        ["Resistencia de aislamiento (circuitos ≤500V, ensayo a 500V c.c.)", "≥ 0,50 MΩ"],
        ["Resistencia de aislamiento (circuitos MBTS/MBTP, ensayo a 250V c.c.)", "≥ 0,25 MΩ"],
        ["Resistencia de aislamiento (circuitos >500V, ensayo a 1000V c.c.)", "≥ 1,00 MΩ"],
        ["Rigidez dieléctrica (1 min a 2U+1000V, mínimo 1.500V)", "Sin perforación del aislamiento"],
        ["Resistencia de puesta a tierra", "Compatible con la sensibilidad del diferencial instalado"],
        ["Disparo de los interruptores diferenciales (botón de test e Idn)", "Disparo correcto"],
        ["Caída de tensión en los circuitos más desfavorables", "Conforme al Anexo de Cálculos"],
        ["Secuencia de fases y tensiones (instalación trifásica)", "Correcta y equilibrada"],
    ], colw=(9.5, 6.5)))
    story.append(Paragraph(
        "En fábrica, la aparamenta habrá sido sometida a idénticos ensayos de aislamiento y rigidez "
        "dieléctrica, además de comprobación visual y funcional de todas sus partes móviles; cuando se "
        "exijan, el instalador aportará los certificados de ensayo del fabricante. Realizadas las "
        "comprobaciones con resultado favorable, se emitirá el correspondiente Certificado de Instalación "
        "Eléctrica (CIE), que se presentará ante el órgano competente de la Comunidad Autónoma para su "
        "registro.", normal))

    story.append(Paragraph(f"{n+1}. Seguridad en los trabajos eléctricos", h2))
    story.append(Paragraph(
        "Siempre que se intervenga en una instalación eléctrica, en su ejecución o mantenimiento, los "
        "trabajos se realizarán sin tensión, verificando su ausencia mediante los aparatos de medición "
        "correspondientes (las 'cinco reglas de oro'). Se utilizarán guantes y herramientas aislantes; las "
        "herramientas y aparatos eléctricos portátiles dispondrán de aislamiento de clase II o se "
        "alimentarán a tensión de seguridad. Los aparatos de protección, seccionamiento y maniobra se "
        "bloquearán en posición de apertura durante la intervención, y no se restablecerá el servicio sin "
        "haber comprobado que no existe peligro alguno. Se cumplirá, en lo que corresponda, la Ley 31/1995 "
        "de Prevención de Riesgos Laborales y su normativa de desarrollo.", normal))

    story.append(Paragraph(f"{n+2}. Limpieza y mantenimiento", h2))
    story.append(Paragraph(
        "Antes de la recepción, los cuadros y demás elementos se limpiarán de polvo, pintura y cualquier "
        "material acumulado durante la ejecución de los trabajos. El titular velará por el correcto "
        "mantenimiento de la instalación, revisando periódicamente el estado de conductores, protecciones y "
        "puesta a tierra (accionamiento del botón de test del diferencial al menos una vez al mes), y "
        "encargando las inspecciones periódicas que, en su caso, sean obligatorias según la potencia y el "
        "uso de la instalación (ITC-BT-05). Cuando sea necesario intervenir nuevamente en la instalación, "
        "por avería o modificación, se aplicarán las mismas condiciones de ejecución, control y seguridad "
        "que si se tratase de una instalación nueva.", normal))

    story.append(Paragraph(f"{n+3}. Criterios de medición", h2))
    story.append(Paragraph(
        "Los cables, bandejas y tubos se medirán por unidad de longitud (metro), según tipo y sección o "
        "diámetro; en la medición se entienden incluidos los accesorios necesarios para el montaje (grapas, "
        "terminales, bornes, prensaestopas, cajas de derivación) y la mano de obra de transporte interior, "
        "montaje y pruebas de recepción. Los cuadros y receptores eléctricos se medirán por unidad montada y "
        "conexionada. A las unidades medidas se aplicarán los precios del Presupuesto, en los que se "
        "consideran incluidos los gastos generales del instalador; si fuese necesaria alguna unidad no "
        "contemplada, se formalizará el correspondiente precio contradictorio conforme al apartado 12.",
        normal))

    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Este pliego es un documento de apoyo estándar, redactado a partir de la práctica habitual en "
        "instalaciones de baja tensión; adáptalo a las particularidades de cada proyecto, al contrato "
        "suscrito entre las partes y a las ordenanzas municipales o autonómicas que puedan resultar de "
        "aplicación.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()



# ==============================================================================
# 7. PRESUPUESTO POR CAPÍTULOS (mismo formato que el excel de referencia:
# Partida/Designación/Unidades/Cantidad/Precio/Importe + panel de Precio
# base/Beneficio/Amortización, y Resumen con IVA e importe en letra) — SIN la
# columna de "Código" que traía el excel original.
# ==============================================================================

def generar_pdf_cie(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                     inputs_fv: dict, resultado_fv: dict, checklist: list, firma: dict,
                     config_prof: dict = None) -> bytes:
    """Certificado de Instalación Eléctrica (CIE): documento de apoyo
    rellenado a partir de los datos ya introducidos en el proyecto y de la
    checklist de puesta en servicio. NO sustituye al CIE oficial, que emite
    el instalador autorizado a través de la aplicación de su Comunidad
    Autónoma — pero deja todo listo para trasladarlo."""
    from reportlab.platypus import Paragraph, Spacer, PageBreak

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "CERTIFICADO DE INSTALACIÓN ELÉCTRICA (CIE)", "Documento de apoyo — ITC-BT-04/05", datos_proyecto,
        config_prof)
    d = datos_proyecto
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    story = _bloque_portada("Certificado de Instalación Eléctrica", "Documento de apoyo para su tramitación",
                             datos_proyecto, config_prof, AZUL, COBRE, colors, h1, normal)

    story.append(Paragraph("1. Datos del titular de la instalación", h2))
    story.append(_tabla_pdf([
        ["Campo", "Valor"],
        ["Titular", d.get("titular") or "-"],
        ["NIF / CIF", d.get("nif_titular") or "-"],
        ["Emplazamiento", d.get("emplazamiento") or "-"],
    ], (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("2. Datos de la instalación", h2))
    story.append(_tabla_pdf([
        ["Campo", "Valor"],
        ["Uso", d.get("uso") or "-"],
        ["Superficie", (d.get("superficie") or "-") + (" m²" if d.get("superficie") else "")],
        ["Tipo de actuación", d.get("tipo_instalacion") or "-"],
        ["Potencia instalada (circuito calculado)",
         f"{inputs_cable.get('potencia_kw', 0):g} kW" if hay_cable else "-"],
        ["Potencia fotovoltaica generadora",
         f"{resultado_fv.get('p_pico_kwp', 0):.2f} kWp" if hay_fv else "No aplica"],
    ], (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("3. Instalador autorizado / técnico competente", h2))
    story.append(_tabla_pdf([
        ["Campo", "Valor"],
        ["Nombre", d.get("instalador") or "-"],
        ["NIF", d.get("nif_instalador") or "-"],
        ["Nº de autorización / colegiado", d.get("n_autorizacion") or "-"],
        ["Categoría", d.get("categoria_instalador") or "-"],
    ], (6.5, 9.5), AZUL, colors))

    story.append(Paragraph("4. Resumen de las características técnicas", h2))
    if hay_cable:
        story.append(Paragraph(
            f"Circuito de referencia: <b>{inputs_cable['tipo_circuito']}</b>. Sección de fase adoptada: "
            f"<b>{resultado_cable['seccion_final']:g} mm²</b>. Protección: <b>{resultado_cable['calibre_magnetotermico']} A</b>. "
            f"Caída de tensión: <b>{resultado_cable['e_final_pct']:.2f} %</b> (máx. {inputs_cable['delta_u_max']:g} %).",
            normal))
    if hay_fv:
        story.append(Paragraph(
            f"Instalación generadora fotovoltaica: <b>{resultado_fv['p_pico_kwp']:.2f} kWp</b>, modalidad "
            f"{inputs_fv.get('tipo_autoconsumo', '-')} (RD 244/2019).", normal))
    if not hay_cable and not hay_fv:
        story.append(Paragraph("No hay cálculos disponibles todavía en este proyecto.", normal))

    story.append(PageBreak())
    story.append(Paragraph("5. Verificaciones previas a la puesta en servicio (ITC-BT-05)", h2))
    story.append(Paragraph(
        "Resultado de las comprobaciones reglamentarias, según la checklist de puesta en servicio "
        "cumplimentada en la aplicación:", normal))
    filas_check = [["Ensayo", "Criterio", "¿Realizado?", "Valor medido"]]
    for (nombre_ens, criterio_ens), resultado_ens in zip(ENSAYOS_PUESTA_SERVICIO, checklist):
        filas_check.append([nombre_ens, criterio_ens, "✅ Sí" if resultado_ens["realizado"] else "☐ No",
                             resultado_ens["valor_medido"] or "-"])
    story.append(_tabla_pdf(filas_check, (6.0, 4.5, 2.2, 3.3), AZUL, colors, fuente=7.8))
    n_hechos = sum(1 for e in checklist if e["realizado"])
    if n_hechos < len(checklist):
        story.append(Paragraph(_pdf_safe_markup(
            f"⚠️ {len(checklist) - n_hechos} de {len(checklist)} verificaciones quedan pendientes de "
            "completar antes de poder emitir el certificado definitivo."), normal))
    else:
        story.append(Paragraph(_pdf_safe_markup("✅ Todas las verificaciones están completadas."), normal))

    story.append(Paragraph("6. Declaración", h2))
    fecha_pruebas = firma.get("fecha")
    fecha_pruebas_txt = fecha_pruebas.strftime("%d/%m/%Y") if hasattr(fecha_pruebas, "strftime") else str(fecha_pruebas or "-")
    story.append(Paragraph(
        f"D./Dña. <b>{firma.get('instalador') or '_______________________'}</b>, en calidad de "
        f"{(d.get('categoria_instalador') or 'básica').lower()}, certifica que la instalación descrita "
        f"reúne las condiciones y garantías reglamentarias exigidas por el REBT, habiéndose realizado las "
        f"comprobaciones anteriores el día <b>{fecha_pruebas_txt}</b>.", normal))
    story.append(Spacer(1, 40))
    story.append(Paragraph("Firma: _______________________________________", normal))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Este documento es un <b>apoyo</b> para preparar la información necesaria; el Certificado de "
        "Instalación Eléctrica oficial se emite y registra por el instalador autorizado a través de la "
        "aplicación telemática de la Comunidad Autónoma correspondiente.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


def generar_pdf_presentacion_cliente(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                                      inputs_fv: dict, resultado_fv: dict, total_presupuesto: float,
                                      config_prof: dict = None) -> bytes:
    """Ficha para el cliente final: mismo contenido que la vista de
    Presentación cliente, en lenguaje llano y sin siglas técnicas."""
    from reportlab.platypus import Paragraph, Spacer

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "RESUMEN DE TU PROYECTO", "Preparado para ti", datos_proyecto, config_prof)
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None
    nombre_cliente = datos_proyecto.get("titular") or "tu proyecto"

    story = [Paragraph(f"Esto es lo que hemos preparado para {nombre_cliente}", h1), Spacer(1, 10)]

    if hay_cable:
        story.append(Paragraph("La instalación eléctrica", h2))
        grosor_mm = resultado_cable["seccion_final"]
        comparacion = ("un cable fino, como el de un cargador de móvil" if grosor_mm <= 2.5 else
                      "un cable de grosor medio, como el de un electrodoméstico grande" if grosor_mm <= 10 else
                      "un cable considerablemente grueso, pensado para mover mucha potencia con seguridad")
        story.append(Paragraph(
            f"Para esta parte de la instalación hace falta un cable de <b>{grosor_mm:g} mm²</b> de grosor "
            f"— a modo de referencia, es {comparacion}.", normal))
        story.append(Paragraph(
            f"Se instala además una protección automática de <b>{resultado_cable['calibre_magnetotermico']} A</b> "
            "que corta la luz sola si algo va mal, antes de que pueda ser peligroso.", normal))

    if hay_fv:
        story.append(Paragraph("Los paneles solares", h2))
        p_pico = resultado_fv["p_pico_kwp"]
        produccion = resultado_fv["produccion_anual_kwh"]
        n_paneles = resultado_fv["n_paneles_configurados"]
        story.append(Paragraph(
            f"Se instalan <b>{n_paneles} paneles solares</b>, con una potencia conjunta de "
            f"<b>{p_pico:.1f} kW</b> — producen aproximadamente <b>{_miles(produccion)} kWh al año</b>.",
            normal))
        if resultado_fv.get("ahorro_anual"):
            story.append(Paragraph(
                f"Esto supone un ahorro estimado de <b>{_fmt_eur(resultado_fv['ahorro_anual'])} al año</b> "
                "en la factura de la luz.", normal))
        if resultado_fv.get("payback_anos"):
            story.append(Paragraph(
                f"Con la inversión indicada, los paneles se pagan solos en unos "
                f"<b>{resultado_fv['payback_anos']:.0f} años</b>.", normal))
        co2 = resultado_fv.get("co2_evitado_kg_ano", 0) / 1000
        if co2:
            story.append(Paragraph(f"De paso, se evitan unas <b>{co2:.1f} toneladas de CO₂ al año</b>.",
                                    normal))

    if total_presupuesto:
        story.append(Paragraph("El coste total", h2))
        story.append(Paragraph(f"<font size=22 color='#b3711f'>"
                                f"<b>{_fmt_eur(total_presupuesto)}</b></font>", normal))
        story.append(Paragraph("Impuestos incluidos.", normal))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "Los números técnicos completos (secciones exactas, normativa aplicada, desglose de precios) "
        "están disponibles en la documentación técnica del proyecto.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


def generar_pdf_resumen_una_pagina(datos_proyecto: dict, inputs_cable: dict, resultado_cable: dict,
                                    inputs_fv: dict, resultado_fv: dict, total_presupuesto: float,
                                    config_prof: dict = None) -> bytes:
    """Ficha resumen del proyecto en una sola página, para imprimir o
    compartir rápido sin abrir los documentos completos."""
    from reportlab.platypus import Paragraph, Spacer

    buffer, doc, cajetin, AZUL, colors, h2, h3, normal, h1, COBRE, lista_item = _preparar_doc_pdf(
        "FICHA RESUMEN DEL PROYECTO", "Resumen de una página", datos_proyecto, config_prof)
    d = datos_proyecto
    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    story = []
    story.append(Paragraph("Datos generales", h2))
    story.append(_tabla_pdf([
        ["Titular", d.get("titular") or "-"],
        ["Emplazamiento", d.get("emplazamiento") or "-"],
        ["Instalador", d.get("instalador") or "-"],
        ["Tipo de actuación", d.get("tipo_instalacion") or "-"],
    ], (6.5, 9.5), AZUL, colors, header=False, negrita_col0=True))

    if hay_cable:
        story.append(Paragraph("Instalación eléctrica", h2))
        story.append(_tabla_pdf([
            ["Circuito", inputs_cable["tipo_circuito"]],
            ["Sistema", f"{inputs_cable['sistema']} — {inputs_cable['tension']:g} V"],
            ["Sección adoptada", f"{resultado_cable['seccion_final']:g} mm²"],
            ["Protección", f"{resultado_cable['calibre_magnetotermico']} A"],
            ["Caída de tensión", f"{resultado_cable['e_final_pct']:.2f} %"],
        ], (6.5, 9.5), AZUL, colors, header=False, negrita_col0=True))

    if hay_fv:
        story.append(Paragraph("Instalación fotovoltaica", h2))
        story.append(_tabla_pdf([
            ["Potencia pico", f"{resultado_fv['p_pico_kwp']:.2f} kWp"],
            ["Nº de paneles", f"{resultado_fv['n_paneles_configurados']}"],
            ["Producción anual estimada", f"{_miles(resultado_fv['produccion_anual_kwh'])} kWh/año"],
            ["Modalidad", inputs_fv.get("tipo_autoconsumo", "-")],
        ], (6.5, 9.5), AZUL, colors, header=False, negrita_col0=True))

    story.append(Paragraph("Presupuesto", h2))
    story.append(Paragraph(
        f"<b>{_fmt_eur(total_presupuesto)}</b> IVA incluido." if total_presupuesto else
        "Presupuesto pendiente de completar.", normal))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "Ficha resumen orientativa; el detalle completo (cálculos, mediciones, condiciones) se encuentra "
        "en la Memoria Técnica de Diseño, el Anexo de Cálculos y el Pliego de Condiciones.", normal))

    doc.build(story, onFirstPage=cajetin, onLaterPages=cajetin, canvasmaker=doc._numbered_canvas)
    return buffer.getvalue()


def calcular_precio_venta(precio_base: float, pct_beneficio: float, pct_amortizacion: float) -> float:
    """Precio_venta = Precio_base x (1 + %beneficio + %amortizacion), tal y como
    lo calcula la 'Justificación de Precio' del excel de referencia."""
    return round(precio_base * (1 + (pct_beneficio + pct_amortizacion) / 100.0), 4)


def calcular_totales_capitulo(items: list, pct_beneficio: float, pct_amortizacion: float) -> float:
    total = 0.0
    for it in items:
        precio_venta = calcular_precio_venta(it["precio_base"], pct_beneficio, pct_amortizacion)
        total += it["cantidad"] * precio_venta
    return round(total, 2)


# --- Conversión de importes a letras (estilo "Cuatro mil ochocientos treinta
# y nueve euros y setenta céntimos", igual que en el excel de referencia) ---

_UNIDADES_L = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
_DIECI_L = ["diez", "once", "doce", "trece", "catorce", "quince", "dieciséis", "diecisiete",
            "dieciocho", "diecinueve"]
_VEINTI_L = ["veinte", "veintiuno", "veintidós", "veintitrés", "veinticuatro", "veinticinco",
             "veintiséis", "veintisiete", "veintiocho", "veintinueve"]
_DECENAS_L = ["", "", "", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
_CENTENAS_L = ["", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos",
               "seiscientos", "setecientos", "ochocientos", "novecientos"]


def _dos_digitos_letras(n: int) -> str:
    if n < 10:
        return _UNIDADES_L[n]
    if n < 20:
        return _DIECI_L[n - 10]
    if n < 30:
        return _VEINTI_L[n - 20]
    d, u = divmod(n, 10)
    return _DECENAS_L[d] + (f" y {_UNIDADES_L[u]}" if u else "")


def _tres_digitos_letras(n: int) -> str:
    if n == 0:
        return ""
    if n == 100:
        return "cien"
    c, resto = divmod(n, 100)
    partes = []
    if c:
        partes.append(_CENTENAS_L[c])
    if resto:
        partes.append(_dos_digitos_letras(resto))
    return " ".join(partes)


def _apocope_uno(texto: str) -> str:
    if texto.endswith("veintiuno"):
        return texto[:-3] + "ún"
    if texto.endswith("uno"):
        return texto[:-3] + "un"
    return texto


def numero_entero_a_letras(n: int) -> str:
    if n == 0:
        return "cero"
    partes = []
    millones, resto = divmod(n, 1_000_000)
    miles, unidades = divmod(resto, 1000)
    if millones:
        texto_millones = "un millón" if millones == 1 else _tres_digitos_letras(millones) + " millones"
        if miles == 0 and unidades == 0:
            texto_millones += " de"
        partes.append(texto_millones)
    if miles:
        partes.append("mil" if miles == 1 else _tres_digitos_letras(miles) + " mil")
    if unidades:
        partes.append(_tres_digitos_letras(unidades))
    return " ".join(partes)


def numero_a_letras_euros(importe: float) -> str:
    importe = round(importe, 2)
    euros = int(importe)
    centimos = round((importe - euros) * 100)
    texto_euros = _apocope_uno(numero_entero_a_letras(euros))
    resultado = f"{texto_euros.capitalize()} {'euro' if euros == 1 else 'euros'}"
    if centimos:
        texto_cent = _apocope_uno(numero_entero_a_letras(centimos))
        resultado += f" y {texto_cent} {'céntimo' if centimos == 1 else 'céntimos'}"
    return resultado


IVA_DEFECTO_PCT = 21.0


def item_desde_calculo_cable(inp: dict, res: dict) -> list:
    """Traduce el resultado de la Calculadora de cables en líneas de
    presupuesto (acción explícita del usuario: 'importar', nunca automática)."""
    if res.get("seccion_final") is None:
        return []
    seccion = res["seccion_final"]
    n_paralelo = res["n_paralelo"] if res["necesita_paralelo"] else 1
    n_fases = 3 if inp["sistema"] == SISTEMA_TRI else 1
    precio_cable = PRECIOS_CABLE_COBRE_DEFECTO.get(seccion, 5.0) if inp["conductor"] == "Cobre" else \
        round(PRECIOS_CABLE_COBRE_DEFECTO.get(seccion, 5.0) * RATIO_PRECIO_ALUMINIO, 2)
    items = [{
        "designacion": f"Cable {inp['conductor']} {seccion:g} mm² ({inp['aislamiento']}) — fase",
        "unidades": "m", "cantidad": round(inp["longitud"] * n_fases * n_paralelo, 2),
        "precio_base": precio_cable,
    }]
    if res.get("seccion_neutro"):
        items.append({
            "designacion": f"Cable {inp['conductor']} {res['seccion_neutro']:g} mm² — neutro",
            "unidades": "m", "cantidad": round(inp["longitud"] * n_paralelo, 2),
            "precio_base": PRECIOS_CABLE_COBRE_DEFECTO.get(res["seccion_neutro"], 5.0),
        })
    if res.get("seccion_proteccion"):
        items.append({
            "designacion": f"Cable {inp['conductor']} {res['seccion_proteccion']:g} mm² — protección (PE)",
            "unidades": "m", "cantidad": round(inp["longitud"] * n_paralelo, 2),
            "precio_base": PRECIOS_CABLE_COBRE_DEFECTO.get(res["seccion_proteccion"], 5.0),
        })
    items.append({
        "designacion": f"Interruptor automático {res['calibre_magnetotermico']} A",
        "unidades": "ud", "cantidad": 1,
        "precio_base": PRECIOS_MAGNETOTERMICO_DEFECTO.get(res["calibre_magnetotermico"], 30.0),
    })
    items.append({
        "designacion": "Mano de obra de instalación (oficial 1ª electricista)",
        "unidades": "horas", "cantidad": round(max(inp["longitud"] / 10.0, 1.0), 1),
        "precio_base": 33.0,
    })
    return items


def items_desde_calculo_fv(inp: dict, res: dict) -> list:
    """Traduce el resultado del módulo Fotovoltaico en líneas de presupuesto
    (también una acción explícita del usuario, no automática)."""
    items = [
        {"designacion": f"Panel fotovoltaico {inp['potencia_panel_wp']:g} Wp", "unidades": "ud",
         "cantidad": res["n_paneles_configurados"], "precio_base": PRECIOS_FV_DEFECTO["Panel fotovoltaico (según Wp introducido)"][1]},
        {"designacion": f"Inversor {inp['potencia_inversor_kw']:g} kW", "unidades": "ud", "cantidad": 1,
         "precio_base": PRECIOS_FV_DEFECTO["Inversor (según kW introducido)"][1]},
        {"designacion": "Estructura soporte por panel", "unidades": "ud",
         "cantidad": res["n_paneles_configurados"],
         "precio_base": PRECIOS_FV_DEFECTO["Estructura soporte por panel (cubierta inclinada)"][1]},
        {"designacion": f"Cable solar H1Z2Z2-K — tramo CC ({res['s_cc_final']:g} mm²)", "unidades": "m",
         "cantidad": round(inp["longitud_cc"] * 2 * inp["n_strings_paralelo"], 1),
         "precio_base": PRECIOS_FV_DEFECTO["Cable solar H1Z2Z2-K (CC)"][1]},
        {"designacion": f"Cable CA inversor-cuadro ({res['s_ca_final']:g} mm²)", "unidades": "m",
         "cantidad": round(inp["longitud_ca"] * (3 if inp["sistema_ca"] == SISTEMA_TRI else 2), 1),
         "precio_base": PRECIOS_FV_DEFECTO["Cable CA inversor-cuadro (según sección calculada)"][1]},
        {"designacion": "Caja de conexión / string box con fusibles", "unidades": "ud", "cantidad": 1,
         "precio_base": PRECIOS_FV_DEFECTO["Caja de conexión / string box con fusibles"][1]},
        {"designacion": "Protector de sobretensiones DC (DPS tipo 2)", "unidades": "ud", "cantidad": 1,
         "precio_base": PRECIOS_FV_DEFECTO["Protector de sobretensiones DC (DPS tipo 2)"][1]},
        {"designacion": "Interruptor-seccionador CC", "unidades": "ud", "cantidad": 1,
         "precio_base": PRECIOS_FV_DEFECTO["Interruptor-seccionador CC"][1]},
        {"designacion": f"Magnetotérmico CA {res['calibre_magneto_ca']} A (salida inversor)", "unidades": "ud",
         "cantidad": 1, "precio_base": PRECIOS_FV_DEFECTO["Magnetotérmico CA salida inversor"][1]},
        {"designacion": "Diferencial tipo A (salida inversor)", "unidades": "ud", "cantidad": 1,
         "precio_base": PRECIOS_FV_DEFECTO["Diferencial tipo A (salida inversor)"][1]},
        {"designacion": "Mano de obra instalador FV", "unidades": "kWp",
         "cantidad": round(res["p_pico_kwp"], 2),
         "precio_base": PRECIOS_FV_DEFECTO["Mano de obra instalador FV (por kWp)"][1]},
        {"designacion": "Legalización y tramitación (notificación/certificado)", "unidades": "ud",
         "cantidad": 1, "precio_base": PRECIOS_FV_DEFECTO["Legalización y tramitación (notificación/certificado)"][1]},
    ]
    if res.get("calibre_fusible_string"):
        items.append({"designacion": f"Fusible de string {res['calibre_fusible_string']} A / "
                                      f"{res.get('tension_fusible_string', 1000)} V (par)", "unidades": "ud",
                       "cantidad": inp["n_strings_paralelo"], "precio_base": 8.0})
    if res.get("capacidad_bateria_kwh"):
        items.append({"designacion": f"Batería de litio {res['capacidad_bateria_kwh']:.1f} kWh",
                       "unidades": "kWh", "cantidad": round(res["capacidad_bateria_kwh"], 1),
                       "precio_base": PRECIOS_FV_DEFECTO["Batería de litio (por kWh de capacidad)"][1]})
    return items


def _sanear_nombre_hoja_excel(nombre: str, existentes: set = None) -> str:
    """Los nombres de hoja de Excel no admiten \\ / ? * [ ] : y tienen un
    límite de 31 caracteres; además deben ser únicos dentro del libro."""
    caracteres_invalidos = '\\/?*[]:'
    limpio = "".join(c for c in nombre if c not in caracteres_invalidos).strip()
    limpio = limpio[:31] if limpio else "Capitulo"
    if existentes is not None:
        base = limpio[:28]
        sufijo = 1
        while limpio in existentes:
            sufijo += 1
            limpio = f"{base} ({sufijo})"
        existentes.add(limpio)
    return limpio


def generar_excel_presupuesto_capitulos(capitulos: list, pct_beneficio: float, pct_amortizacion: float,
                                          pct_iva: float, nombre_proyecto: str = "") -> bytes:
    """Genera un Excel con fórmulas REALES (no valores estáticos): Precio,
    Importe, totales de capítulo, subtotal, IVA y TOTAL son celdas con
    fórmula que recalculan solas si cambias una cantidad o un precio base.
    Estructura: hoja 'Presupuesto' (todos los capítulos), una hoja por
    capítulo con el desglose Precio base/Beneficio/Amortización, y
    'Resumen del presupuesto' con el IVA y el importe final en letra."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    azul = "122340"
    cobre = "E8A33D"
    gris_claro = "F4F6FB"
    thin = Side(style="thin", color="C9CCD1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    cols_principales = ["Partida", "Designación", "Unidades", "Cantidad", "Precio (€)", "Importe (€)"]

    def _estilo_input(cell):
        cell.fill = PatternFill("solid", fgColor=gris_claro)
        cell.font = Font(color="8A5A1F", italic=True)

    wb = Workbook()

    # ============================================================ Hoja Presupuesto (combinada)
    ws_presu = wb.active
    ws_presu.title = "Presupuesto"
    ws_presu.append(["PRESUPUESTO" + (f" — {nombre_proyecto}" if nombre_proyecto else "")])
    ws_presu["A1"].font = Font(bold=True, size=14, color=azul)
    ws_presu.append([])
    ws_presu.cell(row=3, column=1, value="Beneficio industrial (%)").font = Font(bold=True)
    celda_benef = ws_presu.cell(row=3, column=2, value=pct_beneficio)
    _estilo_input(celda_benef)
    ws_presu.cell(row=3, column=4, value="Amortización medios aux. (%)").font = Font(bold=True)
    celda_amort = ws_presu.cell(row=3, column=5, value=pct_amortizacion)
    _estilo_input(celda_amort)
    ws_presu.cell(row=4, column=1, value="(Puedes editar estos dos porcentajes: el precio y el total de "
                  "todo el presupuesto se recalculan solos.)").font = Font(italic=True, size=8, color="5A6472")
    ref_benef, ref_amort = "$B$3", "$E$3"

    fila = 6
    for cap in capitulos:
        ws_presu.cell(row=fila, column=1, value=cap["nombre"]).font = Font(bold=True, color=azul)
        fila += 1
        for j, col in enumerate(cols_principales, start=1):
            c = ws_presu.cell(row=fila, column=j, value=col)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor=azul)
        ws_presu.cell(row=fila, column=8, value="Precio base (€)").font = Font(bold=True, color="FFFFFF")
        ws_presu.cell(row=fila, column=8).fill = PatternFill("solid", fgColor=azul)
        fila += 1
        fila_inicio_items = fila
        for it in cap["items"]:
            precio_base = it["precio_base"]
            ws_presu.cell(row=fila, column=1, value=it.get("partida", "-")).border = border
            ws_presu.cell(row=fila, column=2, value=it["designacion"]).border = border
            ws_presu.cell(row=fila, column=3, value=it["unidades"]).border = border
            ws_presu.cell(row=fila, column=4, value=it["cantidad"]).border = border
            c_precio = ws_presu.cell(row=fila, column=5, value=f"=H{fila}*(1+({ref_benef}+{ref_amort})/100)")
            c_precio.number_format = '#,##0.00 €'
            c_precio.border = border
            c_importe = ws_presu.cell(row=fila, column=6, value=f"=D{fila}*E{fila}")
            c_importe.number_format = '#,##0.00 €'
            c_importe.border = border
            c_base = ws_presu.cell(row=fila, column=8, value=precio_base)
            c_base.number_format = '#,##0.0000 €'
            _estilo_input(c_base)
            fila += 1
        fila_fin_items = fila - 1
        ws_presu.cell(row=fila, column=4, value="TOTAL").font = Font(bold=True)
        tc = ws_presu.cell(row=fila, column=6,
                            value=f"=SUM(F{fila_inicio_items}:F{fila_fin_items})" if cap["items"] else 0)
        tc.font = Font(bold=True, color=azul)
        tc.number_format = '#,##0.00 €'
        fila += 2
    widths = [10, 50, 9, 10, 12, 12, 3, 13]
    for j, w in enumerate(widths, start=1):
        ws_presu.column_dimensions[get_column_letter(j)].width = w

    # ============================================================ Una hoja por capítulo
    nombres_hojas_usados = {"Presupuesto", "Resumen del presupuesto"}
    hojas_capitulo = {}  # nombre_original -> nombre_hoja_saneado
    filas_total_capitulo = []  # (nombre_capitulo, fila_del_total_EN_SU_PROPIA_HOJA) para el Resumen
    for cap in capitulos:
        nombre_hoja = _sanear_nombre_hoja_excel(cap["nombre"], nombres_hojas_usados)
        hojas_capitulo[cap["nombre"]] = nombre_hoja
        ws = wb.create_sheet(nombre_hoja)
        ws.append([cap["nombre"]])
        ws["A1"].font = Font(bold=True, size=12, color=azul)
        ws.cell(row=2, column=8, value="% Beneficio").font = Font(bold=True, size=8)
        c_b = ws.cell(row=2, column=9, value=pct_beneficio)
        _estilo_input(c_b)
        ws.cell(row=2, column=10, value="% Amortización").font = Font(bold=True, size=8)
        c_a = ws.cell(row=2, column=11, value=pct_amortizacion)
        _estilo_input(c_a)
        cabecera = cols_principales + ["", "Precio base (€)", "Beneficio (€)", "Amortización (€)"]
        for j, col in enumerate(cabecera, start=1):
            if not col:
                continue
            c = ws.cell(row=3, column=j, value=col)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor=azul)
        r = 4
        r_inicio = r
        for it in cap["items"]:
            ws.cell(row=r, column=1, value=it.get("partida", "-")).border = border
            ws.cell(row=r, column=2, value=it["designacion"]).border = border
            ws.cell(row=r, column=3, value=it["unidades"]).border = border
            ws.cell(row=r, column=4, value=it["cantidad"]).border = border
            c_precio = ws.cell(row=r, column=5, value=f"=H{r}+I{r}+J{r}")
            c_precio.number_format = '#,##0.0000 €'
            c_precio.border = border
            c_importe = ws.cell(row=r, column=6, value=f"=D{r}*E{r}")
            c_importe.number_format = '#,##0.00 €'
            c_importe.border = border
            c_base = ws.cell(row=r, column=8, value=it["precio_base"])
            c_base.number_format = '#,##0.0000 €'
            _estilo_input(c_base)
            ws.cell(row=r, column=9, value=f"=H{r}*$I$2/100").number_format = '#,##0.0000 €'
            ws.cell(row=r, column=10, value=f"=H{r}*$K$2/100").number_format = '#,##0.0000 €'
            r += 1
        r_fin = r - 1
        ws.cell(row=r, column=4, value="TOTAL").font = Font(bold=True)
        ws.cell(row=r, column=6, value=f"=SUM(F{r_inicio}:F{r_fin})" if cap["items"] else 0).number_format = '#,##0.00 €'
        filas_total_capitulo.append((cap["nombre"], r))
        widths_cap = [10, 50, 9, 10, 12, 12, 3, 13, 12, 14]
        for j, w in enumerate(widths_cap, start=1):
            ws.column_dimensions[get_column_letter(j)].width = w

    # ============================================================ Resumen del presupuesto
    ws_r = wb.create_sheet("Resumen del presupuesto")
    ws_r.merge_cells("A1:F1")
    ws_r["A1"] = "RESUMEN DEL PRESUPUESTO"
    ws_r["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws_r["A1"].fill = PatternFill("solid", fgColor=azul)
    ws_r.append([])
    ws_r.append(["Capítulo", "", "", "", "", "Coste parcial (€)"])
    for c in ws_r[3]:
        c.font = Font(bold=True)
    fila_r = 4
    fila_inicio_capitulos = fila_r
    for nombre_cap, fila_total_cap in filas_total_capitulo:
        nombre_hoja = hojas_capitulo[nombre_cap]
        ws_r.cell(row=fila_r, column=1, value=nombre_cap)
        c = ws_r.cell(row=fila_r, column=6, value=f"='{nombre_hoja}'!F{fila_total_cap}")
        c.number_format = '#,##0.00 €'
        fila_r += 1
    fila_fin_capitulos = fila_r - 1
    fila_r += 1
    ws_r.cell(row=fila_r, column=5, value="Subtotal").font = Font(bold=True)
    c_subtotal = ws_r.cell(row=fila_r, column=6,
                            value=f"=SUM(F{fila_inicio_capitulos}:F{fila_fin_capitulos})" if filas_total_capitulo else 0)
    c_subtotal.number_format = '#,##0.00 €'
    fila_subtotal = fila_r
    fila_r += 1
    ws_r.cell(row=fila_r, column=5, value="IVA (%)").font = Font(italic=True)
    c_iva_pct = ws_r.cell(row=fila_r, column=6, value=pct_iva)
    _estilo_input(c_iva_pct)
    c_iva_pct.number_format = '0.0"%"'
    fila_iva_pct = fila_r
    fila_r += 1
    ws_r.cell(row=fila_r, column=5, value="Importe IVA").font = Font(bold=True)
    c_iva = ws_r.cell(row=fila_r, column=6, value=f"=F{fila_subtotal}*F{fila_iva_pct}/100")
    c_iva.number_format = '#,##0.00 €'
    fila_r += 1
    ws_r.cell(row=fila_r, column=5, value="TOTAL").font = Font(bold=True)
    c_total = ws_r.cell(row=fila_r, column=6, value=f"=F{fila_subtotal}+F{fila_r - 1}")
    fila_total_final = fila_r
    for col in (5, 6):
        ws_r.cell(row=fila_r, column=col).font = Font(bold=True, color=azul)
        ws_r.cell(row=fila_r, column=col).fill = PatternFill("solid", fgColor=cobre)
    c_total.number_format = '#,##0.00 €'

    # El importe en letra no puede ser fórmula fiable en Excel sin macros; se
    # calcula aquí con el total YA resuelto en Python (mismo valor que dará
    # la fórmula anterior) y se anota que es informativo.
    subtotal_calc = sum(calcular_totales_capitulo(cap["items"], pct_beneficio, pct_amortizacion)
                         for cap in capitulos)
    total_calc = round(subtotal_calc * (1 + pct_iva / 100.0), 2)
    fila_r += 1
    ws_r.cell(row=fila_r, column=1,
              value=f"{numero_a_letras_euros(total_calc).capitalize()} "
                    "(referencia con los porcentajes actuales; si los editas, recalcula la celda F"
                    f"{fila_total_final} de arriba).").font = Font(italic=True, size=9)
    ws_r.column_dimensions["A"].width = 50
    for col in "BCDE":
        ws_r.column_dimensions[col].width = 13
    ws_r.column_dimensions["F"].width = 16

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ==============================================================================
# 7B. CÁLCULOS BT — calculadoras sueltas de referencia rápida (independientes
# entre sí), inspiradas en el catálogo de calculadoras de circuitoelectrico.com:
# cables, cortocircuito, tierras, fotovoltaica rápida y eléctricas generales.
# ==============================================================================

# --- Conversión AWG <-> mm² (fórmula estándar: d(mm) = 0,127 x 92^((36-AWG)/39)) ---
AWG_DISPONIBLES = [0000, 000, 00, 0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15,
                   16, 17, 18, 19, 20, 21, 22, 23, 24]


def awg_a_mm2(awg) -> float:
    n = {"0000": -3, "000": -2, "00": -1}.get(str(awg), awg if isinstance(awg, (int, float)) else int(awg))
    d_mm = 0.127 * (92 ** ((36 - n) / 39))
    return round(math.pi / 4 * d_mm ** 2, 4)


def mm2_a_awg_mas_cercano(mm2: float) -> tuple:
    mejor = min(AWG_DISPONIBLES, key=lambda a: abs(awg_a_mm2(a) - mm2))
    return mejor, awg_a_mm2(mejor)


# --- Resistencia de electrodos de tierra (ITC-BT-18, Guía técnica, Tabla 5) ---
RESISTIVIDAD_TERRENOS_REF = {
    "Terrenos pantanosos (algunos)": 30,
    "Limo": 100,
    "Humus": 150,
    "Turba húmeda": 5,
    "Arcilla plástica": 50,
    "Terrenos cultivables y fértiles, terraplenes compactos y húmedos": 50,
    "Terraplenes cultivables poco fértiles y otros terraplenes": 500,
    "Suelos pedregosos desnudos, arenas secas permeables": 3000,
    "Arena arcillosa": 100,
    "Calizas blandas": 300,
    "Calizas compactas": 1000,
    "Pizarras": 150,
    "Rocas eruptivas (granito, basalto...)": 5000,
    "Personalizado (introducir ρ manualmente)": None,
}


def resistencia_electrodo_pica(rho: float, longitud: float) -> float:
    """R = ρ / L (ITC-BT-18, Tabla 5)."""
    return rho / max(longitud, 1e-6)


def resistencia_electrodo_placa(rho: float, perimetro: float) -> float:
    """R = 0,8 · ρ / P (ITC-BT-18, Tabla 5)."""
    return 0.8 * rho / max(perimetro, 1e-6)


def resistencia_electrodo_conductor(rho: float, longitud: float) -> float:
    """R = 2 · ρ / L (ITC-BT-18, Tabla 5), conductor enterrado horizontalmente."""
    return 2 * rho / max(longitud, 1e-6)


def resistencia_picas_paralelo(r_una_pica: float, n_picas: int) -> float:
    """Aproximación habitual para picas bien separadas (distancia >= 2xL):
    R_total ≈ R_una_pica / n. Orientativa; en terrenos heterogéneos o picas
    próximas la reducción real puede ser menor."""
    return r_una_pica / max(n_picas, 1)


# --- Cortocircuito simplificado (con/sin datos de red aguas arriba) ---
def impedancia_linea(conductor: str, aislamiento: str, seccion: float, longitud: float,
                      n_cargados: int, reactancia_ohm_km: float = REACTANCIA_LINEAL_DEFECTO) -> tuple:
    """Devuelve (R, X) en ohmios de un tramo de línea, ida simple (no ida+vuelta)."""
    kappa = kappa_servicio(conductor, aislamiento)
    r = longitud / (kappa * seccion)
    x = longitud * reactancia_ohm_km / 1000.0
    return r, x


def icc_trifasico(u: float, r_total: float, x_total: float) -> float:
    """Icc = U / (√3 · Z), Z = sqrt(R²+X²). U en V, resultado en A."""
    z = math.sqrt(r_total ** 2 + x_total ** 2)
    return u / (math.sqrt(3) * max(z, 1e-9))


# --- Colores de resistencias (código de 4 bandas) ---
COLORES_DIGITO = {
    "Negro": 0, "Marrón": 1, "Rojo": 2, "Naranja": 3, "Amarillo": 4,
    "Verde": 5, "Azul": 6, "Violeta": 7, "Gris": 8, "Blanco": 9,
}
COLORES_MULTIPLICADOR = {
    "Negro": 1, "Marrón": 10, "Rojo": 100, "Naranja": 1_000, "Amarillo": 10_000,
    "Verde": 100_000, "Azul": 1_000_000, "Oro": 0.1, "Plata": 0.01,
}
COLORES_TOLERANCIA = {
    "Marrón": "±1%", "Rojo": "±2%", "Verde": "±0,5%", "Azul": "±0,25%",
    "Violeta": "±0,1%", "Gris": "±0,05%", "Oro": "±5%", "Plata": "±10%", "Ninguno": "±20%",
}


def valor_resistencia_4_bandas(c1: str, c2: str, c3: str, c4: str) -> tuple:
    valor = (COLORES_DIGITO[c1] * 10 + COLORES_DIGITO[c2]) * COLORES_MULTIPLICADOR[c3]
    return valor, COLORES_TOLERANCIA.get(c4, "±20%")


# ==============================================================================
# 8. INSTALACIONES FOTOVOLTAICAS — módulo de cálculo independiente
# ==============================================================================
# Independiente del todo de la Calculadora de cables: entradas, sesión y motor
# de cálculo propios. Solo comparten las funciones puras de sección de cable
# (iz_tabla, seccion_por_criterio_termico, caida_tension_voltios...) porque son
# la misma física, no porque haya un acoplamiento de datos entre apartados.

ZONAS_CLIMATICAS_HSP = {
    "Norte (Galicia, Asturias, Cantabria, País Vasco)": (3.7, 43.0),
    "Centro / Meseta (Madrid, Castilla)": (4.4, 40.4),
    "Mediterráneo / Levante (Cataluña, C. Valenciana, Baleares)": (4.8, 39.5),
    "Sur (Andalucía, Murcia, Extremadura)": (5.3, 37.2),
    "Canarias": (5.4, 28.3),
    "Personalizado (introducir HSP y latitud manualmente)": (None, None),
}

TIPO_AUTOCONSUMO_FV = [
    "Sin excedentes",
    "Con excedentes acogido a compensación",
    "Con excedentes no acogido a compensación",
    "Instalación aislada (con batería)",
]

PR_DEFECTO_FV = 0.80
EFICIENCIA_INVERSOR_DEFECTO = 97.0        # % — típico 95-98% (SunFields, fabricantes)
DEGRADACION_ANUAL_DEFECTO = 0.5           # %/año — típico 0,3-0,8%/año
FACTOR_CO2_RED_DEFECTO = 0.20             # kg CO2/kWh — orientativo, mix peninsular reciente (REE, 130-200 gCO2/kWh)
AREA_PANEL_DEFECTO = 1.95                 # m² — panel estándar ~450Wp (aprox. 1,75x1,13 m)
PRECIO_COMPENSACION_DEFECTO = 0.08        # €/kWh excedente compensado (orientativo, varía por comercializadora)

# Módulo y catálogo de referencia orientativos (editables en la propia app).
PANEL_DEFECTO = dict(potencia_wp=450, voc=41.8, isc=13.9, vmp=34.6, imp=13.0, coef_temp_voc=-0.27)
INVERSOR_DEFECTO = dict(potencia_kw=5.0, vmin_mppt=80, vmax_mppt=550, vmax_entrada=600, n_mppt=2)

# Calibres de fusible de string normalizados (A) — criterio orientativo del
# fabricante: In entre 1,5 y 2,4 veces Isc del string (UNE-EN 62548 / IEC 60269-6).
CALIBRES_FUSIBLE_STRING = [2, 4, 6, 8, 10, 12, 15, 16, 20, 25, 32]

# Tensiones nominales normalizadas de fusible cilíndrico fotovoltaico (V CC).
TENSIONES_FUSIBLE_FV = [500, 600, 1000, 1500]

PRECIOS_FV_DEFECTO = {
    "Panel fotovoltaico (según Wp introducido)": ("ud", 110.0),
    "Inversor (según kW introducido)": ("ud", 850.0),
    "Estructura soporte por panel (cubierta inclinada)": ("ud", 35.0),
    "Cable solar H1Z2Z2-K (CC)": ("m", 0.95),
    "Cable CA inversor-cuadro (según sección calculada)": ("m", 2.5),
    "Caja de conexión / string box con fusibles": ("ud", 85.0),
    "Protector de sobretensiones DC (DPS tipo 2)": ("ud", 55.0),
    "Interruptor-seccionador CC": ("ud", 45.0),
    "Magnetotérmico CA salida inversor": ("ud", 28.0),
    "Diferencial tipo A (salida inversor)": ("ud", 65.0),
    "Conectores MC4 (par)": ("ud", 3.5),
    "Batería de litio (por kWh de capacidad)": ("kWh", 450.0),
    "Mano de obra instalador FV (por kWp)": ("kWp", 180.0),
    "Legalización y tramitación (notificación/certificado)": ("ud", 250.0),
}


def perdidas_orientacion_inclinacion_cte(inclinacion: float, azimut: float, latitud: float) -> float:
    """Pérdidas (%) por orientación e inclinación no óptimas, según la
    fórmula oficial del CTE DB-HE5: Pérdidas(%) = 100 x [1,2e-4·(β-φ+10)² +
    3,5e-5·α²], con β=inclinación, φ=latitud, α=azimut (0°=sur, ±180°=norte).
    Válida para |azimut| < 90° y ángulos de inclinación habituales; fuera de
    ese rango se satura al valor en el límite para evitar resultados
    absurdos."""
    beta = max(0.0, min(90.0, inclinacion))
    alfa = max(-180.0, min(180.0, azimut))
    perdidas = 100 * (1.2e-4 * (beta - latitud + 10) ** 2 + 3.5e-5 * alfa ** 2)
    return max(0.0, min(perdidas, 40.0))


@st.cache_data(ttl=3600, show_spinner=False)
def _consultar_pvgis(lat: float, lon: float, inclinacion: float, azimut: float, loss_pct: float = 14.0) -> dict:
    """Consulta la API pública y gratuita PVGIS (Comisión Europea, JRC) para
    obtener la irradiación real sobre el plano inclinado en la ubicación
    exacta del proyecto, en vez de una aproximación por zona climática.
    Devuelve un diccionario con los datos relevantes. Lanza una excepción
    con un mensaje claro y en español si la consulta falla (sin conexión,
    fuera de cobertura, parámetros inválidos...); quien llame debe capturarla
    y ofrecer el método por zona climática como alternativa.
    Resultado cacheado 1 hora por combinación de parámetros, para no repetir
    la misma consulta en cada rerun de Streamlit."""
    import requests
    url = "https://re.jrc.ec.europa.eu/api/v5_2/PVcalc"
    params = {"lat": lat, "lon": lon, "peakpower": 1, "loss": loss_pct,
              "angle": inclinacion, "aspect": azimut, "outputformat": "json"}
    try:
        resp = requests.get(url, params=params, timeout=15)
    except requests.exceptions.ConnectionError as e:
        raise RuntimeError("No se ha podido conectar con PVGIS (sin conexión a internet o el servicio no "
                           "responde). Puedes seguir usando el HSP por zona climática mientras tanto.") from e
    except requests.exceptions.Timeout as e:
        raise RuntimeError("PVGIS ha tardado demasiado en responder (más de 15 s). Inténtalo de nuevo en "
                           "un momento.") from e
    if resp.status_code != 200:
        try:
            detalle = resp.json().get("message", resp.text[:200])
        except Exception:
            detalle = resp.text[:200]
        raise RuntimeError(f"PVGIS ha devuelto un error (HTTP {resp.status_code}): {detalle}. Revisa que "
                           "la latitud/longitud caigan sobre tierra firme (PVGIS no cubre océanos ni los "
                           "polos).")
    try:
        datos = resp.json()
        totales = datos["outputs"]["totals"]["fixed"]
        mensuales = datos["outputs"]["monthly"]["fixed"]
    except (ValueError, KeyError) as e:
        raise RuntimeError("PVGIS ha devuelto una respuesta con un formato inesperado. Puede que las "
                           "coordenadas no tengan cobertura de datos solares.") from e
    return {
        "hi_y": totales["H(i)_y"],
        "e_y_especifica": totales["E_y"],
        "perdidas_totales_pct": totales.get("l_total"),
        "mensual": [{"mes": m["month"], "e_m": m["E_m"], "hi_m": m["H(i)_m"]} for m in mensuales],
        "elevacion": (datos.get("inputs", {}).get("location", {}) or {}).get("elevation"),
    }


def calcular_fv(inp: dict) -> dict:
    """Motor de cálculo puro de la instalación fotovoltaica. No depende de la
    Calculadora de cables ni comparte estado con ella."""
    avisos = []

    hsp = inp["hsp"]
    pr_base = inp["pr"]

    if inp["modo_dimensionado"] == "Por consumo anual (kWh)":
        p_pico_kwp = inp["consumo_anual_kwh"] / max(hsp * 365 * pr_base, 1e-6)
        n_paneles = math.ceil(p_pico_kwp * 1000 / inp["potencia_panel_wp"])
    elif inp["modo_dimensionado"] == "Por potencia pico deseada (kWp)":
        p_pico_kwp = inp["potencia_pico_deseada"]
        n_paneles = math.ceil(p_pico_kwp * 1000 / inp["potencia_panel_wp"])
    else:  # Por número de paneles
        n_paneles = inp["n_paneles_manual"]
        p_pico_kwp = n_paneles * inp["potencia_panel_wp"] / 1000

    # --- Pérdidas por orientación/inclinación (CTE DB-HE5) + sombras/suciedad
    # + eficiencia del inversor, todas aplicadas sobre la producción bruta.
    # Si el HSP viene de PVGIS (irradiación real ya calculada para esa
    # inclinación/azimut exactos), esa pérdida ya está incluida en el propio
    # HSP y NO debe volver a aplicarse aquí — se pasa perdidas_orientacion_pvgis=0. ---
    if inp.get("perdidas_orientacion_pvgis") is not None:
        perdidas_orient_pct = inp["perdidas_orientacion_pvgis"]
    else:
        perdidas_orient_pct = perdidas_orientacion_inclinacion_cte(
            inp["inclinacion"], inp["azimut"], inp["latitud"])
    perdidas_sombra_pct = inp.get("perdidas_sombras", 0.0)
    eficiencia_inversor_pct = inp.get("eficiencia_inversor", EFICIENCIA_INVERSOR_DEFECTO)

    pr_efectivo = pr_base * (1 - perdidas_orient_pct / 100.0) * (1 - perdidas_sombra_pct / 100.0) * \
        (eficiencia_inversor_pct / 100.0)
    produccion_anual_kwh = p_pico_kwp * hsp * 365 * pr_efectivo

    if perdidas_orient_pct > 15:
        avisos.append(f"Pérdidas por orientación/inclinación estimadas en {perdidas_orient_pct:.1f}% "
                      "(CTE DB-HE5): revisa si es posible acercar el campo solar a la orientación/"
                      "inclinación óptimas para esta latitud.")

    degradacion_pct = inp.get("degradacion_anual", DEGRADACION_ANUAL_DEFECTO)
    produccion_ano10 = produccion_anual_kwh * (1 - degradacion_pct / 100.0) ** 9
    produccion_ano25 = produccion_anual_kwh * (1 - degradacion_pct / 100.0) ** 24

    superficie_necesaria_m2 = n_paneles * inp.get("area_panel", AREA_PANEL_DEFECTO)

    # --- Configuración del generador (strings) ---
    n_serie = inp["n_paneles_serie"]
    n_paralelo = inp["n_strings_paralelo"]
    n_paneles_configurados = n_serie * n_paralelo

    voc, isc, vmp = inp["voc"], inp["isc"], inp["vmp"]
    coef_v = inp["coef_temp_voc"] / 100.0

    v_string_frio = n_serie * voc * (1 + coef_v * (25 - inp["temp_min"]))
    v_string_caliente = n_serie * vmp * (1 + coef_v * (inp["temp_max_celula"] - 25))
    i_generador = n_paralelo * isc

    cumple_vmax = v_string_frio <= inp["vmax_entrada_inversor"]
    cumple_vmpp_min = v_string_caliente >= inp["vmin_mppt"]
    cumple_vmpp_max = v_string_frio <= inp["vmax_mppt"]
    if not cumple_vmax:
        avisos.append(f"La tensión de string en frío ({v_string_frio:.0f} V) supera la tensión máxima de "
                       f"entrada del inversor ({inp['vmax_entrada_inversor']:.0f} V): reduce paneles en serie.")
    if not cumple_vmpp_min:
        avisos.append(f"La tensión de string en caliente ({v_string_caliente:.0f} V) queda por debajo del "
                       f"mínimo MPPT ({inp['vmin_mppt']:.0f} V): añade paneles en serie o revisa el inversor.")
    if not cumple_vmpp_max:
        avisos.append(f"La tensión de string en frío ({v_string_frio:.0f} V) supera el máximo de la ventana "
                       f"MPPT ({inp['vmax_mppt']:.0f} V).")

    ratio_dc_ac = (p_pico_kwp / inp["potencia_inversor_kw"]) if inp["potencia_inversor_kw"] else None
    if ratio_dc_ac and (ratio_dc_ac < 0.9 or ratio_dc_ac > 1.3):
        avisos.append(f"Ratio DC/AC = {ratio_dc_ac:.2f} fuera del rango habitual (0,9-1,3); revisa la potencia "
                       "del inversor frente a la del generador.")

    # --- Cableado CC (ITC-BT-40: 125% Imax; DeltaU <= 1,5% ENTRE GENERADOR Y
    # PUNTO DE INTERCONEXIÓN, es decir CC+CA combinados, no 1,5% en cada tramo).
    # Se dimensiona cada tramo con un objetivo individual del 1% como margen de
    # trabajo habitual, y se verifica el conjunto contra el 1,5% real. ---
    i_diseno_cc = 1.25 * isc
    kappa_cc = kappa_servicio("Cobre", "XLPE/EPR", usar_20c=False)
    objetivo_parcial_pct = 0.6
    s_cc_termica = None
    for s in SECCIONES_NORMALIZADAS:
        base = iz_tabla(s, METODO_B1, "XLPE/EPR", "Cobre", 2)
        if base is not None and base >= i_diseno_cc:
            s_cc_termica = s
            break
    s_cc_termica = s_cc_termica or max(SECCIONES_NORMALIZADAS)
    s_cc_du = None
    for s in SECCIONES_NORMALIZADAS:
        e = 2 * inp["longitud_cc"] * i_diseno_cc / (kappa_cc * s)
        if e / inp["tension_cc_ref"] * 100 <= objetivo_parcial_pct:
            s_cc_du = s
            break
    s_cc_du = s_cc_du or max(SECCIONES_NORMALIZADAS)
    s_cc_final = max(s_cc_termica, s_cc_du)
    du_cc_pct = (2 * inp["longitud_cc"] * i_diseno_cc / (kappa_cc * s_cc_final)) / inp["tension_cc_ref"] * 100

    # --- Cableado CA (inversor -> cuadro), mismas funciones que la Calculadora
    # de cables pero con datos 100% propios de este módulo. ---
    ib_ca = calcular_intensidad_empleo(inp["sistema_ca"], inp["potencia_inversor_kw"] * 1000.0,
                                        inp["tension_ca"], 1.0)
    i_diseno_ca = 1.25 * ib_ca
    n_cargados_ca = 2 if inp["sistema_ca"] == SISTEMA_MONO else 3
    s_ca_termica, iz_ca, _, _ = seccion_por_criterio_termico(
        i_diseno_ca, METODO_B1, "XLPE/EPR", "Cobre", n_cargados_ca, 1.0)
    kappa_ca = kappa_servicio("Cobre", "XLPE/EPR")
    s_ca_du, e_ca, e_ca_pct = seccion_por_caida_tension(
        inp["sistema_ca"], i_diseno_ca, inp["longitud_ca"], inp["tension_ca"], 1.0, kappa_ca,
        objetivo_parcial_pct, "Cobre")
    candidatos_ca = [s for s in (s_ca_termica, s_ca_du) if s is not None]
    s_ca_final = max(candidatos_ca) if candidatos_ca else None
    if s_ca_final is not None:
        e_ca_real = caida_tension_voltios(inp["sistema_ca"], i_diseno_ca, inp["longitud_ca"], s_ca_final,
                                            1.0, kappa_ca)
        e_ca_pct = e_ca_real / inp["tension_ca"] * 100

    du_total_pct = du_cc_pct + (e_ca_pct if s_ca_final is not None else 0.0)
    if du_total_pct > 1.5:
        avisos.append(f"Caída de tensión combinada CC+CA = {du_total_pct:.2f}% > 1,5% (ITC-BT-40, entre el "
                       "generador y el punto de interconexión): aumenta la sección de algún tramo o reduce "
                       "la longitud.")

    # --- Protecciones ---
    calibre_fusible_string = None
    tension_fusible_string = None
    if n_paralelo >= 2:
        objetivo = 1.8 * isc
        for c in CALIBRES_FUSIBLE_STRING:
            if c >= objetivo:
                calibre_fusible_string = c
                break
        calibre_fusible_string = calibre_fusible_string or CALIBRES_FUSIBLE_STRING[-1]
        for v in TENSIONES_FUSIBLE_FV:
            if v >= v_string_frio:
                tension_fusible_string = v
                break
        tension_fusible_string = tension_fusible_string or TENSIONES_FUSIBLE_FV[-1]
        if tension_fusible_string < v_string_frio:
            avisos.append(f"Ningún calibre estándar de tensión de fusible ({TENSIONES_FUSIBLE_FV} V) cubre "
                          f"la tensión de string en frío ({v_string_frio:.0f} V): usa fusibles de mayor "
                          "tensión nominal.")
    calibre_magneto_ca = calibre_magnetotermico_sugerido(i_diseno_ca)

    # --- Batería (opcional: aislada o autoconsumo con almacenamiento) ---
    capacidad_bateria_kwh = None
    if inp.get("con_bateria"):
        consumo_diario_kwh = inp.get("consumo_diario_bateria_kwh") or (produccion_anual_kwh / 365.0)
        autonomia_dias = inp.get("autonomia_dias", 1.0)
        prof_descarga_pct = inp.get("profundidad_descarga", 80.0)
        capacidad_bateria_kwh = (consumo_diario_kwh * autonomia_dias) / max(prof_descarga_pct / 100.0, 0.01)

    # --- Económico: autoconsumo + compensación de excedentes ---
    energia_autoconsumida = produccion_anual_kwh * (inp["pct_autoconsumo"] / 100.0)
    energia_excedente = produccion_anual_kwh - energia_autoconsumida
    ahorro_autoconsumo = energia_autoconsumida * inp["precio_kwh"]
    ingreso_excedentes = 0.0
    if inp["tipo_autoconsumo"] == "Con excedentes acogido a compensación":
        ingreso_excedentes = energia_excedente * inp.get("precio_compensacion", PRECIO_COMPENSACION_DEFECTO)
    ahorro_anual = ahorro_autoconsumo + ingreso_excedentes
    payback_anos = (inp["inversion_total"] / ahorro_anual) if ahorro_anual > 0 and inp["inversion_total"] else None

    # --- CO2 evitado ---
    factor_co2 = inp.get("factor_co2", FACTOR_CO2_RED_DEFECTO)
    co2_evitado_kg_ano = produccion_anual_kwh * factor_co2

    return dict(
        p_pico_kwp=p_pico_kwp, n_paneles=n_paneles, produccion_anual_kwh=produccion_anual_kwh,
        produccion_ano10=produccion_ano10, produccion_ano25=produccion_ano25,
        perdidas_orient_pct=perdidas_orient_pct, pr_efectivo=pr_efectivo,
        superficie_necesaria_m2=superficie_necesaria_m2,
        n_serie=n_serie, n_paralelo=n_paralelo, n_paneles_configurados=n_paneles_configurados,
        v_string_frio=v_string_frio, v_string_caliente=v_string_caliente, i_generador=i_generador,
        cumple_vmax=cumple_vmax, cumple_vmpp_min=cumple_vmpp_min, cumple_vmpp_max=cumple_vmpp_max,
        ratio_dc_ac=ratio_dc_ac, i_diseno_cc=i_diseno_cc, s_cc_final=s_cc_final, du_cc_pct=du_cc_pct,
        ib_ca=ib_ca, i_diseno_ca=i_diseno_ca, s_ca_final=s_ca_final, iz_ca=iz_ca, e_ca_pct=e_ca_pct,
        du_total_pct=du_total_pct,
        calibre_fusible_string=calibre_fusible_string, tension_fusible_string=tension_fusible_string,
        calibre_magneto_ca=calibre_magneto_ca, capacidad_bateria_kwh=capacidad_bateria_kwh,
        energia_autoconsumida=energia_autoconsumida, energia_excedente=energia_excedente,
        ahorro_autoconsumo=ahorro_autoconsumo, ingreso_excedentes=ingreso_excedentes,
        ahorro_anual=ahorro_anual, payback_anos=payback_anos,
        co2_evitado_kg_ano=co2_evitado_kg_ano, avisos=avisos,
    )


# ==============================================================================
# 9. INTERFAZ STREAMLIT
# ==============================================================================

def _fmt_eur(valor: float) -> str:
    """Formatea un importe en estilo español: 1.234,56 €."""
    texto = f"{valor:,.2f}"
    texto = texto.replace(",", "TMP").replace(".", ",").replace("TMP", ".")
    return f"{texto} €"


def _miles(valor: float, decimales: int = 0) -> str:
    """Formatea un número en estilo español (1.234,56), SOLO el número — a
    diferencia de un .replace(',', '.') sobre toda la frase, esto no toca las
    comas literales del texto que lo rodea."""
    texto = f"{valor:,.{decimales}f}"
    return texto.replace(",", "TMP").replace(".", ",").replace("TMP", ".")


def _render_inputs() -> dict:
    plantilla = st.session_state.pop("plantilla_activa", None) or {}
    if plantilla:
        st.success("📋 Plantilla aplicada — ajusta lo que necesites, el resto de valores parten de aquí.")

    st.markdown('<p class="section-label">1 · Datos eléctricos del circuito</p>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        cc1, cc2 = st.columns([6, 1])
        opciones_circuito = list(CAIDA_TENSION_MAX.keys())
        idx_circuito = opciones_circuito.index(plantilla["tipo_circuito"]) if plantilla.get("tipo_circuito") in opciones_circuito else 6
        with cc1:
            tipo_circuito = st.selectbox("Tipo de circuito / tramo", opciones_circuito, index=idx_circuito)
        with cc2:
            st.markdown("<div style='height:1.85rem;'></div>", unsafe_allow_html=True)
            _ayuda("Determina la caída de tensión máxima admisible del tramo (ITC-BT-14/15/19/40). "
                   "Ej.: una Derivación Individual admite solo 1-1,5%, mientras que un circuito interior "
                   "de fuerza admite hasta el 5%. Si dudas, usa 'Instalación interior — Otros usos'.",
                   "Tipo de circuito")
        sistema = st.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI],
                               index=0 if plantilla.get("sistema", SISTEMA_MONO) == SISTEMA_MONO else 1)
    with c2:
        tension_defecto = 230.0 if sistema == SISTEMA_MONO else 400.0
        tension = st.number_input("Tensión de servicio (V)", min_value=100.0, max_value=1000.0,
                                   value=tension_defecto, step=1.0)
        modo_entrada = st.radio("Datos de partida", ["Potencia activa", "Intensidad directa"], horizontal=True)
    with c3:
        if modo_entrada == "Potencia activa":
            potencia_kw = st.number_input("Potencia activa (kW)", min_value=0.01,
                                           value=float(plantilla.get("potencia_kw", 5.0)), step=0.1)
            cc1, cc2 = st.columns([6, 1])
            with cc1:
                cos_phi = st.number_input("cos φ", min_value=0.10, max_value=1.00,
                                           value=float(plantilla.get("cos_phi", 0.90)), step=0.01)
            with cc2:
                st.markdown("<div style='height:1.85rem;'></div>", unsafe_allow_html=True)
                _ayuda("Factor de potencia de la carga. Orientativo: resistivo puro (calefacción, "
                       "incandescencia) ≈ 1,0 · uso general/electrónica ≈ 0,90-0,95 · motores ≈ 0,80-0,85.",
                       "cos φ")
            intensidad_directa = None
        else:
            potencia_kw = None
            intensidad_directa = st.number_input("Intensidad (A)", min_value=0.1, value=20.0, step=0.5)
            cos_phi = st.number_input("cos φ (para ΔU)", min_value=0.10, max_value=1.00, value=0.90, step=0.01)

    st.markdown('<p class="section-label">2 · Datos de la instalación</p>', unsafe_allow_html=True)
    c4, c5, c6 = st.columns(3)
    with c4:
        conductor = st.selectbox("Material conductor", ["Cobre", "Aluminio"])
        metodo = st.selectbox("Método de instalación", METODOS_DISPONIBLES,
                               index=METODOS_DISPONIBLES.index(plantilla["metodo"]) if plantilla.get("metodo") in METODOS_DISPONIBLES else 1,
                               help="A1/B1/B2: tubo (empotrado o superficie). C/E: bandeja o superficie. "
                                    "F: bandeja perforada de grandes secciones, típica en industria. "
                                    "D: enterrado. A1 y B2 son estimaciones ancladas a B1 — ver pestaña "
                                    "Metodología.")
    with c5:
        if metodo in (METODO_D, METODO_F):
            aislamiento = "XLPE/EPR"
            st.selectbox("Aislamiento", ["XLPE/EPR"], disabled=True,
                         help="Este método se modela sólo para XLPE/EPR, el estándar actual en enterrado y "
                              "bandejas de gran sección.")
        else:
            opciones_aisl = ["PVC", "XLPE/EPR"]
            idx_aisl = opciones_aisl.index(plantilla["aislamiento"]) if plantilla.get("aislamiento") in opciones_aisl else 0
            aislamiento = st.selectbox("Aislamiento", opciones_aisl, index=idx_aisl)
        longitud = st.number_input("Longitud del circuito (m)", min_value=0.1,
                                    value=float(plantilla.get("longitud", 20.0)), step=1.0)
    with c6:
        delta_u_defecto = CAIDA_TENSION_MAX[tipo_circuito]
        valor_defecto = float(delta_u_defecto) if delta_u_defecto is not None else 3.0
        delta_u_max = st.number_input("ΔU máx. admisible (%)", min_value=0.1, max_value=10.0,
                                       value=valor_defecto, step=0.1,
                                       help="Editable: útil si tu proyecto aplica un criterio de compensación "
                                            "de caídas entre tramos distinto del habitual.")

    with st.expander("⚙️ Casuística especial — motores, alumbrado, armónicos"):
        cc1, cc2 = st.columns(2)
        with cc1:
            es_motor = st.checkbox("Circuito de motor (ITC-BT-47)")
            corrientes_motores = []
            ascensor_grua = False
            if es_motor:
                tipo_motor = st.radio("Motores en el circuito", ["Motor único", "Varios motores"], horizontal=True)
                if tipo_motor == "Motor único":
                    in_motor = st.number_input("In del motor (A)", min_value=0.1, value=20.0, step=0.5)
                    corrientes_motores = [in_motor]
                else:
                    texto_motores = st.text_input("In de cada motor, separadas por coma (A)", "20, 10, 8")
                    try:
                        corrientes_motores = [float(x.strip()) for x in texto_motores.split(",") if x.strip()]
                    except ValueError:
                        corrientes_motores = []
                ascensor_grua = st.checkbox("Ascensor / grúa (factor adicional ITC-BT-47, ×1,3)")
        with cc2:
            alumbrado_descarga = st.checkbox(
                "Alumbrado de descarga sin corregir f.d.p.",
                help="Aplica un factor orientativo ×1,8 a la intensidad de cálculo (ITC-BT-44) para lámparas "
                     "con reactancia sin condensador de compensación.")
            armonicos = st.checkbox(
                "Cargas no lineales / armónicos de 3er orden significativos",
                help="Informática, variadores de frecuencia, alumbrado electrónico... impide reducir la "
                     "sección del neutro por debajo de la de fase.")

    enterrado = metodo == METODO_D
    with st.expander("🌡️ Condiciones ambientales y agrupamiento"):
        ce1, ce2, ce3 = st.columns(3)
        with ce1:
            temp_defecto = 25.0 if enterrado else 40.0
            temp_ambiente = st.number_input(
                "Temperatura ambiente (°C)", min_value=-10.0, max_value=80.0,
                value=temp_defecto, step=1.0,
                help="La Tabla A/D de la Guía-BT-19 está referida a 40°C (aire) / 25°C (terreno); se corrige "
                     "analíticamente si la temperatura real es distinta.")
            usar_kappa_20c = st.checkbox("Usar conductividad a 20°C para ΔU (menos conservador)", value=False)
        with ce2:
            disposicion = st.selectbox("Disposición de agrupamiento", list(TABLA_E_AGRUPAMIENTO.keys()))
            n_circuitos = st.number_input("Nº de circuitos/ternas agrupados", min_value=1, max_value=50,
                                           value=1, step=1)
        with ce3:
            n_capas = st.number_input("Nº de capas de bandejas/tubos", min_value=1, max_value=9, value=1, step=1)
            if enterrado:
                claves_resist = list(RESISTIVIDAD_TERRENO.keys())
                resistividad_label = st.selectbox("Resistividad térmica del terreno", claves_resist,
                                                   index=len(claves_resist) - 4)
                resistividad = RESISTIVIDAD_TERRENO[resistividad_label]
            else:
                resistividad = RESISTIVIDAD_REF_TERRENO

    with st.expander("⚡ Verificación de cortocircuito (opcional)"):
        verificar_cc = st.checkbox("Verificar criterio térmico de cortocircuito")
        icc_ka, tiempo_s = None, None
        if verificar_cc:
            ccc1, ccc2 = st.columns(2)
            with ccc1:
                icc_ka = st.number_input("Icc en el origen del circuito (kA)", min_value=0.01, value=6.0, step=0.1)
            with ccc2:
                tiempo_s = st.number_input("Tiempo de actuación de la protección (s)", min_value=0.001,
                                            value=0.1, step=0.01, format="%.3f")

    return dict(
        tipo_circuito=tipo_circuito, sistema=sistema, tension=tension, modo_entrada=modo_entrada,
        potencia_kw=potencia_kw, cos_phi=cos_phi, intensidad_directa=intensidad_directa,
        conductor=conductor, metodo=metodo, aislamiento=aislamiento, longitud=longitud,
        delta_u_max=delta_u_max, es_motor=es_motor, corrientes_motores=corrientes_motores,
        ascensor_grua=ascensor_grua, alumbrado_descarga=alumbrado_descarga, armonicos=armonicos,
        temp_ambiente=temp_ambiente, usar_kappa_20c=usar_kappa_20c, disposicion=disposicion,
        n_circuitos=n_circuitos, n_capas=n_capas, enterrado=enterrado, resistividad=resistividad,
        verificar_cc=verificar_cc, icc_ka=icc_ka, tiempo_s=tiempo_s,
    )


def _render_resultados(inp: dict, res: dict):
    st.markdown('<p class="section-label">Resultado</p>', unsafe_allow_html=True)

    seccion_final = res["seccion_final"]
    if seccion_final is None:
        st.error("No se ha podido determinar una sección con los parámetros indicados. Revisa los valores.")
        return

    r1, r2, r3, r4 = st.columns(4)
    with r1:
        st.markdown(f'''<div class="result-card hero">
            <div class="result-label">Sección de fase adoptada</div>
            <div class="result-value">{seccion_final:g} mm²</div>
            <div class="result-sub">{inp['conductor']} · {inp['aislamiento']}</div>
        </div>''', unsafe_allow_html=True)
    with r2:
        st.markdown(f'''<div class="result-card">
            <div class="result-label">Ib de cálculo</div>
            <div class="result-value small">{res['ib_calculo']:.2f} A</div>
            <div class="result-sub">Iz tabla: {res['iz_termica']:.1f} A (S={res['s_termica']:g} mm²)</div>
        </div>''', unsafe_allow_html=True)
    with r3:
        cumple_du = res["e_final_pct"] <= inp["delta_u_max"]
        badge = "badge-ok" if cumple_du else "badge-fail"
        texto_badge = "Cumple" if cumple_du else "No cumple"
        st.markdown(f'''<div class="result-card">
            <div class="result-label">Caída de tensión</div>
            <div class="result-value small">{res['e_final_pct']:.2f} %</div>
            <div class="result-sub"><span class="{badge}">{texto_badge}</span> · máx {inp['delta_u_max']:g}%</div>
        </div>''', unsafe_allow_html=True)
    with r4:
        neutro_txt = f"{res['seccion_neutro']:g}" if res.get("seccion_neutro") else "—"
        st.markdown(f'''<div class="result-card">
            <div class="result-label">Neutro / Protección (mm²)</div>
            <div class="result-value small">{neutro_txt} / {res['seccion_proteccion']:g}</div>
            <div class="result-sub">Interruptor sugerido: {res['calibre_magnetotermico']} A</div>
        </div>''', unsafe_allow_html=True)

    if res["necesita_paralelo"]:
        st.warning(f"La corriente de cálculo supera la capacidad de un único conductor de "
                   f"{max(SECCIONES_NORMALIZADAS):g} mm². Se necesitan **{res['n_paralelo']} conductores en "
                   f"paralelo** de {seccion_final:g} mm² por fase (mismo material, longitud y aislamiento, "
                   "condición del REBT para la puesta en paralelo de conductores).")

    if inp["verificar_cc"]:
        if res["cumple_cc"]:
            st.success(f"✅ Verificación de cortocircuito: {seccion_final:g} mm² soporta térmicamente "
                       f"Icc={inp['icc_ka']:g} kA durante {inp['tiempo_s']:g} s (mínimo requerido: "
                       f"{res['s_min_cc']:g} mm²).")
        else:
            st.error(f"⚠️ Verificación de cortocircuito: {seccion_final:g} mm² NO soporta térmicamente "
                     f"Icc={inp['icc_ka']:g} kA durante {inp['tiempo_s']:g} s. Se requieren ≥ "
                     f"{res['s_min_cc']:g} mm², o una protección con actuación más rápida.")

    st.markdown('<p class="section-label">Secciones normalizadas evaluadas</p>', unsafe_allow_html=True)
    chips = "".join(
        f'<span class="regla-chip{" activa" if s == seccion_final else ""}">{s:g}</span>'
        for s in SECCIONES_NORMALIZADAS
    )
    st.markdown(f'<div class="regla-wrap">{chips}</div>', unsafe_allow_html=True)

    for aviso in res["avisos"]:
        st.info(aviso)

    with st.expander("📋 Ver detalle del cálculo, sección por sección"):
        filas = []
        for s in SECCIONES_NORMALIZADAS:
            base = iz_tabla(s, inp["metodo"], inp["aislamiento"], inp["conductor"], res["n_cargados"])
            if base is None:
                continue
            iz_real = base * res["factor_total"]
            e = caida_tension_voltios(inp["sistema"], res["ib_calculo"], inp["longitud"], s,
                                       inp["cos_phi"], res["kappa"])
            pct = e / inp["tension"] * 100
            filas.append({
                "Sección (mm²)": s,
                "Iz tabla (A)": base,
                "Iz corregida (A)": round(iz_real, 1),
                "Cumple térmico": "✅" if iz_real >= res["ib_calculo"] else "❌",
                "ΔU (%)": round(pct, 2),
                "Cumple ΔU": "✅" if pct <= inp["delta_u_max"] else "❌",
            })
        st.dataframe(pd.DataFrame(filas), width='stretch', hide_index=True)
        st.caption("ΔU aquí calculada con la intensidad total (sin repartir en paralelo), para mostrar por "
                   "qué una sección concreta no bastaría por sí sola.")

    with st.spinner("Generando memoria en PDF..."):
        pdf_bytes = generar_pdf_memoria(inp, res, st.session_state.get("config_profesional"))
    dl1, dl2 = st.columns([1.4, 1])
    with dl1:
        if st.download_button("⬇️ Descargar memoria de cálculo (PDF)", data=pdf_bytes,
                            file_name="memoria_calculo_cable.pdf", mime="application/pdf"):
            _registrar_actividad("📄", "Memoria de cálculo descargada")

    st.markdown('<p class="section-label">Guardar este cálculo</p>', unsafe_allow_html=True)
    st.caption("Le pones un nombre y queda disponible para importarlo en cualquier capítulo del Presupuesto "
               "— útil cuando el proyecto tiene varios circuitos distintos (alumbrado, tomas, motor...) y "
               "quieres tenerlos todos calculados antes de presupuestar.")
    st.session_state.setdefault("calculos_guardados", [])
    sg1, sg2 = st.columns([3, 1])
    with sg1:
        nombre_calculo = st.text_input(
            "Nombre para este cálculo", key="nombre_calculo_guardar",
            placeholder=f"{inp['tipo_circuito'][:35]} — {seccion_final:g} mm²")
    with sg2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("💾 Guardar cálculo", width='stretch'):
            nombre_final = nombre_calculo.strip() or f"{inp['tipo_circuito'][:35]} — {seccion_final:g} mm²"
            st.session_state["calculos_guardados"].append({
                "nombre": nombre_final, "inputs_cable": dict(inp), "resultado_cable": dict(res),
            })
            _registrar_actividad("💾", f"Cálculo de cable guardado: {nombre_final}")
            st.success(f"Guardado como '{nombre_final}'. Ya está disponible en Presupuesto → Importar cable.")

    guardados = st.session_state["calculos_guardados"]
    if guardados:
        with st.expander(f"📋 Cálculos guardados en esta sesión ({len(guardados)})"):
            for i_g, g in enumerate(guardados):
                gc1, gc2 = st.columns([5, 1])
                sc = g["resultado_cable"]
                gc1.markdown(f"**{g['nombre']}** — {sc.get('seccion_final','—')} mm², "
                            f"{g['inputs_cable'].get('tipo_circuito','')}")
                if gc2.button("🗑️", key=f"del_calc_guardado_{i_g}"):
                    guardados.pop(i_g)
                    st.rerun()


def _render_inputs_fv() -> dict:
    plantilla = st.session_state.pop("plantilla_activa_fv", None) or {}
    if plantilla:
        st.success("📋 Plantilla aplicada — ajusta lo que necesites, el resto de valores parten de aquí.")

    st.markdown('<p class="section-label">1 · Dimensionado</p>', unsafe_allow_html=True)
    d1, d2, d3 = st.columns(3)
    opciones_modo = ["Por consumo anual (kWh)", "Por potencia pico deseada (kWp)", "Por número de paneles"]
    with d1:
        modo_dimensionado = st.selectbox(
            "Punto de partida", opciones_modo,
            index=opciones_modo.index(plantilla["modo_dimensionado"]) if plantilla.get("modo_dimensionado") in opciones_modo else 0)
    consumo_anual_kwh = potencia_pico_deseada = n_paneles_manual = None
    with d2:
        if modo_dimensionado == "Por consumo anual (kWh)":
            consumo_anual_kwh = st.number_input("Consumo anual estimado (kWh)", min_value=1.0, value=4000.0, step=100.0)
        elif modo_dimensionado == "Por potencia pico deseada (kWp)":
            potencia_pico_deseada = st.number_input("Potencia pico deseada (kWp)", min_value=0.1,
                                                     value=float(plantilla.get("potencia_pico_deseada", 5.0)), step=0.5)
        else:
            n_paneles_manual = st.number_input("Número de paneles", min_value=1, value=12, step=1)
    with d3:
        potencia_panel_wp = st.number_input("Potencia unitaria del panel (Wp)", min_value=50.0, value=450.0, step=5.0)
        area_panel = st.number_input("Superficie unitaria del panel (m²)", min_value=0.5, value=AREA_PANEL_DEFECTO, step=0.05)

    st.markdown('<p class="section-label">2 · Ubicación, orientación y producción</p>', unsafe_allow_html=True)
    u1, u2, u3, u4 = st.columns(4)
    with u1:
        zona = st.selectbox("Zona climática", list(ZONAS_CLIMATICAS_HSP.keys()), index=2,
                             help="HSP y latitud orientativos por zona (IDAE). Para el cálculo con datos "
                                  "reales de tu ubicación exacta, usa el botón «Consultar PVGIS» de abajo.")
        hsp_defecto, lat_defecto = ZONAS_CLIMATICAS_HSP[zona]
        hsp = st.number_input("HSP (horas de sol pico, h/día)", min_value=1.0, max_value=8.0,
                               value=float(hsp_defecto or 4.5), step=0.1, key="fv_hsp")
        latitud = st.number_input("Latitud (°)", min_value=0.0, max_value=90.0,
                                   value=float(lat_defecto or 40.0), step=0.5, key="fv_latitud")
    with u2:
        inclinacion = st.number_input("Inclinación de los paneles (°)", min_value=0.0, max_value=90.0,
                                       value=max(latitud - 10, 0.0), step=1.0,
                                       help="Óptima orientativa en España: entre la latitud y la latitud-10°.")
        azimut = st.number_input("Azimut / orientación (°, 0=sur, ±90=este/oeste, ±180=norte)",
                                  min_value=-180.0, max_value=180.0, value=0.0, step=5.0, key="fv_azimut")
        perdidas_orient_preview = perdidas_orientacion_inclinacion_cte(inclinacion, azimut, latitud)
        st.caption(f"Pérdida por orientación/inclinación (CTE DB-HE5): **{perdidas_orient_preview:.1f}%**"
                   if not st.session_state.get("pvgis_activo") else
                   "Pérdida por orientación ya incluida en el dato de PVGIS (ver abajo).")
    with u3:
        pr = st.number_input("Performance Ratio base (PR)", min_value=0.5, max_value=0.95, value=PR_DEFECTO_FV,
                              step=0.01, help="Típico 0,75-0,85. Aquí NO incluyas orientación/inclinación ni "
                                   "eficiencia del inversor: se aplican aparte para desglosar cada pérdida.")
        perdidas_sombras = st.number_input("Pérdidas adicionales por sombras/suciedad (%)", min_value=0.0,
                                            max_value=40.0, value=3.0, step=0.5)
        eficiencia_inversor = st.number_input("Eficiencia del inversor (%)", min_value=80.0, max_value=99.5,
                                               value=EFICIENCIA_INVERSOR_DEFECTO, step=0.5)
    with u4:
        degradacion_anual = st.number_input("Degradación anual del panel (%/año)", min_value=0.0, max_value=2.0,
                                             value=DEGRADACION_ANUAL_DEFECTO, step=0.05,
                                             help="Típico 0,3-0,8%/año según el fabricante.")
        tipo_autoconsumo = st.selectbox(
            "Modalidad (RD 244/2019)", TIPO_AUTOCONSUMO_FV,
            index=TIPO_AUTOCONSUMO_FV.index(plantilla["tipo_autoconsumo"]) if plantilla.get("tipo_autoconsumo") in TIPO_AUTOCONSUMO_FV else 0)
        precio_kwh = st.number_input("Precio kWh evitado (€/kWh)", min_value=0.01, value=0.18, step=0.01)

    with st.container(border=True):
        st.markdown("**🌍 PVGIS — irradiación real para esta ubicación exacta**")
        st.caption("Consulta gratuita a la API pública de la Comisión Europea (JRC): sustituye el HSP "
                   "orientativo por zona climática por la irradiación real, calculada a partir de datos "
                   "satelitales para tus coordenadas, inclinación y azimut exactos. Necesita conexión a "
                   "internet.")
        pv1, pv2, pv3 = st.columns([1, 1, 1.2])
        with pv1:
            longitud_geo = st.number_input("Longitud (°, negativo = oeste)", min_value=-180.0, max_value=180.0,
                                            value=st.session_state.get("fv_longitud_geo", -3.7), step=0.1,
                                            key="fv_longitud_geo")
        with pv2:
            st.markdown("<br>", unsafe_allow_html=True)
            consultar_pvgis_click = st.button("🌍 Consultar PVGIS", width='stretch')
        with pv3:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.session_state.get("pvgis_activo"):
                if st.button("↩️ Volver al HSP por zona", width='stretch'):
                    st.session_state["pvgis_activo"] = False
                    st.session_state.pop("pvgis_resultado", None)
                    st.rerun()

        if consultar_pvgis_click:
            loss_equivalente = 100 * (1 - (1 - perdidas_sombras / 100) * (eficiencia_inversor / 100))
            with st.spinner("Consultando PVGIS..."):
                try:
                    resultado_pvgis = _consultar_pvgis(latitud, longitud_geo, inclinacion, azimut,
                                                        loss_equivalente)
                    st.session_state["pvgis_resultado"] = resultado_pvgis
                    st.session_state["pvgis_activo"] = True
                    st.session_state["fv_hsp"] = round(resultado_pvgis["hi_y"] / 365, 3)
                    _registrar_actividad("🌍", f"PVGIS consultado: HSP real = {resultado_pvgis['hi_y']/365:.2f} h/día")
                    st.rerun()
                except RuntimeError as e:
                    st.error(f"❌ {e}")

        if st.session_state.get("pvgis_activo") and st.session_state.get("pvgis_resultado"):
            r_pv = st.session_state["pvgis_resultado"]
            st.success(f"✅ Usando datos reales de PVGIS: irradiación anual en el plano inclinado "
                       f"**{r_pv['hi_y']:.0f} kWh/m²/año** → HSP equivalente **{r_pv['hi_y']/365:.2f} h/día** "
                       f"(elevación del terreno: {r_pv['elevacion']:.0f} m)." if r_pv.get('elevacion') is not None
                       else f"✅ Usando datos reales de PVGIS: irradiación anual en el plano inclinado "
                       f"**{r_pv['hi_y']:.0f} kWh/m²/año** → HSP equivalente **{r_pv['hi_y']/365:.2f} h/día**.")
            import plotly.graph_objects as go
            meses_nom = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
            fig_pvgis = go.Figure(go.Bar(x=meses_nom, y=[m["e_m"] for m in r_pv["mensual"]],
                                         marker=dict(color="#3b82f6")))
            fig_pvgis.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=220,
                                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                                    font=dict(color="#8b96a8"), yaxis_title="kWh/kWp",
                                    title=dict(text="Producción específica mensual (PVGIS, por kWp instalado)",
                                              font=dict(size=11)))
            st.plotly_chart(fig_pvgis, width='stretch')

    st.markdown('<p class="section-label">3 · Módulo fotovoltaico (valores de ficha técnica STC)</p>',
                unsafe_allow_html=True)
    m1, m2, m3, m4 = st.columns(4)
    with m1:
        voc = st.number_input("Voc (V)", min_value=1.0, value=PANEL_DEFECTO["voc"], step=0.1)
        isc = st.number_input("Isc (A)", min_value=0.1, value=PANEL_DEFECTO["isc"], step=0.1)
    with m2:
        vmp = st.number_input("Vmp (V)", min_value=1.0, value=PANEL_DEFECTO["vmp"], step=0.1)
        st.number_input("Imp (A)", min_value=0.1, value=PANEL_DEFECTO["imp"], step=0.1, key="imp_fv_display",
                         disabled=True)
    with m3:
        coef_temp_voc = st.number_input("Coef. temperatura Voc (%/°C)", value=PANEL_DEFECTO["coef_temp_voc"],
                                         step=0.01, format="%.2f",
                                         help="Negativo: la tensión SUBE al bajar la temperatura.")
    with m4:
        n_paneles_serie = st.number_input("Paneles en serie por string", min_value=1, value=11, step=1)
        n_strings_paralelo = st.number_input("Strings en paralelo", min_value=1, value=1, step=1)

    st.markdown('<p class="section-label">4 · Inversor y condiciones ambientales</p>', unsafe_allow_html=True)
    i1, i2, i3, i4 = st.columns(4)
    with i1:
        potencia_inversor_kw = st.number_input("Potencia nominal inversor (kW)", min_value=0.1,
                                                value=float(plantilla.get("potencia_inversor_kw", INVERSOR_DEFECTO["potencia_kw"])), step=0.5)
        sistema_ca = st.selectbox("Sistema CA de salida", [SISTEMA_MONO, SISTEMA_TRI],
                                  index=0 if plantilla.get("sistema_ca", SISTEMA_MONO) == SISTEMA_MONO else 1)
    with i2:
        vmin_mppt = st.number_input("Tensión mínima MPPT (V)", min_value=1.0,
                                     value=float(INVERSOR_DEFECTO["vmin_mppt"]), step=5.0)
        vmax_mppt = st.number_input("Tensión máxima MPPT (V)", min_value=1.0,
                                     value=float(INVERSOR_DEFECTO["vmax_mppt"]), step=5.0)
    with i3:
        vmax_entrada_inversor = st.number_input("Tensión máxima de entrada (V)", min_value=1.0,
                                                 value=float(INVERSOR_DEFECTO["vmax_entrada"]), step=5.0)
        tension_ca = st.number_input("Tensión de salida CA (V)", min_value=100.0,
                                      value=230.0 if sistema_ca == SISTEMA_MONO else 400.0, step=1.0)
    with i4:
        temp_min = st.number_input("Temperatura mínima histórica (°C)", value=-5.0, step=1.0,
                                    help="Para verificar la tensión de string en frío.")
        temp_max_celula = st.number_input("Temperatura máx. célula en servicio (°C)", value=70.0, step=1.0,
                                           help="Orientativa: temperatura ambiente máxima + 25-30°C.")

    st.markdown('<p class="section-label">5 · Cableado</p>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        longitud_cc = st.number_input("Longitud de cada string CC (m, ida)", min_value=0.1, value=15.0, step=1.0)
    with c2:
        tension_cc_ref = st.number_input("Tensión CC de referencia para ΔU (V)", min_value=1.0, value=400.0,
                                          step=10.0, help="Orientativa: tensión de string en el punto de máxima potencia.")
    with c3:
        longitud_ca = st.number_input("Longitud CA inversor→cuadro (m)", min_value=0.1, value=10.0, step=1.0)

    with st.expander("🔋 Batería / almacenamiento (opcional)"):
        con_bateria = st.checkbox("Instalación con batería (autoconsumo con almacenamiento o aislada)",
                                   value=(tipo_autoconsumo == "Instalación aislada (con batería)"))
        consumo_diario_bateria_kwh = autonomia_dias = profundidad_descarga = None
        if con_bateria:
            b1, b2, b3 = st.columns(3)
            with b1:
                consumo_diario_bateria_kwh = st.number_input("Consumo diario a cubrir (kWh)", min_value=0.1,
                                                              value=8.0, step=0.5)
            with b2:
                autonomia_dias = st.number_input("Días de autonomía deseados", min_value=0.5, value=1.0, step=0.5)
            with b3:
                profundidad_descarga = st.number_input("Profundidad de descarga admisible (%)", min_value=20.0,
                                                        max_value=100.0, value=80.0, step=5.0,
                                                        help="Típico 80% para baterías de litio (LiFePO4).")

    with st.expander("💶 Estimación económica y ambiental (opcional)"):
        e1, e2, e3 = st.columns(3)
        with e1:
            pct_autoconsumo = st.number_input("% de la producción autoconsumida", min_value=0.0, max_value=100.0,
                                               value=65.0, step=5.0)
            inversion_total = st.number_input("Inversión total estimada (€)", min_value=0.0, value=6000.0, step=100.0)
        with e2:
            precio_compensacion = st.number_input(
                "Precio de compensación de excedentes (€/kWh)", min_value=0.0,
                value=PRECIO_COMPENSACION_DEFECTO, step=0.01,
                help="Solo aplica si la modalidad es 'con excedentes acogido a compensación'; varía por comercializadora.",
                disabled=(tipo_autoconsumo != "Con excedentes acogido a compensación"))
        with e3:
            factor_co2 = st.number_input("Factor de emisión de la red (kg CO₂/kWh)", min_value=0.0,
                                          value=FACTOR_CO2_RED_DEFECTO, step=0.01,
                                          help="Orientativo (REE); el mix eléctrico varía año a año.")

    return dict(
        modo_dimensionado=modo_dimensionado, consumo_anual_kwh=consumo_anual_kwh,
        potencia_pico_deseada=potencia_pico_deseada, n_paneles_manual=n_paneles_manual,
        potencia_panel_wp=potencia_panel_wp, area_panel=area_panel, zona=zona, hsp=hsp, latitud=latitud,
        longitud_geo=longitud_geo, inclinacion=inclinacion, azimut=azimut, pr=pr,
        perdidas_sombras=perdidas_sombras,
        perdidas_orientacion_pvgis=(0.0 if st.session_state.get("pvgis_activo") else None),
        eficiencia_inversor=eficiencia_inversor, degradacion_anual=degradacion_anual,
        tipo_autoconsumo=tipo_autoconsumo, precio_kwh=precio_kwh, voc=voc, isc=isc, vmp=vmp,
        coef_temp_voc=coef_temp_voc, n_paneles_serie=n_paneles_serie, n_strings_paralelo=n_strings_paralelo,
        potencia_inversor_kw=potencia_inversor_kw, sistema_ca=sistema_ca, vmin_mppt=vmin_mppt,
        vmax_mppt=vmax_mppt, vmax_entrada_inversor=vmax_entrada_inversor, tension_ca=tension_ca,
        temp_min=temp_min, temp_max_celula=temp_max_celula, longitud_cc=longitud_cc,
        tension_cc_ref=tension_cc_ref, longitud_ca=longitud_ca, pct_autoconsumo=pct_autoconsumo,
        inversion_total=inversion_total, con_bateria=con_bateria,
        consumo_diario_bateria_kwh=consumo_diario_bateria_kwh, autonomia_dias=autonomia_dias,
        profundidad_descarga=profundidad_descarga, precio_compensacion=precio_compensacion,
        factor_co2=factor_co2,
    )


def _render_comparador_fv(inp_actual: dict, res_actual: dict):
    """Guarda escenarios FV con nombre (p.ej. 'con batería' vs 'sin batería',
    o dos potencias distintas) y los compara lado a lado."""
    st.session_state.setdefault("escenarios_fv_guardados", [])
    guardados = st.session_state["escenarios_fv_guardados"]

    st.markdown('<p class="section-label">Comparador de escenarios</p>', unsafe_allow_html=True)
    with st.expander("🔄 Guardar y comparar escenarios (con batería / sin batería, distintas potencias...)"):
        st.caption("Guarda el cálculo actual con un nombre, cambia parámetros arriba (p. ej. activa la "
                   "batería o cambia la potencia) y guarda otra vez con otro nombre — luego compáralos "
                   "uno junto al otro.")
        sg1, sg2 = st.columns([3, 1])
        with sg1:
            nombre_escenario = st.text_input(
                "Nombre para el escenario actual", key="nombre_escenario_fv",
                placeholder=f"{'Con batería' if inp_actual.get('con_bateria') else 'Sin batería'} — "
                           f"{res_actual.get('p_pico_kwp', 0):.1f} kWp")
        with sg2:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("💾 Guardar escenario", width='stretch', key="btn_guardar_escenario_fv"):
                nombre_final = nombre_escenario.strip() or (
                    f"{'Con batería' if inp_actual.get('con_bateria') else 'Sin batería'} — "
                    f"{res_actual.get('p_pico_kwp', 0):.1f} kWp")
                guardados.append({"nombre": nombre_final, "inputs_fv": dict(inp_actual), "resultado_fv": dict(res_actual)})
                _registrar_actividad("🔄", f"Escenario FV guardado: {nombre_final}")
                st.success(f"Guardado como «{nombre_final}».")

        if guardados:
            with st.container(border=True):
                st.markdown(f"**Escenarios guardados ({len(guardados)})**")
                for i_esc, esc in enumerate(guardados):
                    ec1, ec2 = st.columns([5, 1])
                    ec1.markdown(f"{esc['nombre']} — {esc['resultado_fv'].get('p_pico_kwp', 0):.1f} kWp, "
                                f"{_miles(esc['resultado_fv'].get('produccion_anual_kwh', 0))} kWh/año")
                    if ec2.button("🗑️", key=f"del_escenario_fv_{i_esc}"):
                        guardados.pop(i_esc)
                        st.rerun()

        if len(guardados) >= 1:
            st.divider()
            st.markdown("**Comparar dos escenarios**")
            opciones_comp = {"— Cálculo actual —": ("actual", None)}
            for i_e, esc in enumerate(guardados):
                opciones_comp[esc["nombre"]] = ("guardado", i_e)
            cc1, cc2 = st.columns(2)
            with cc1:
                sel_a = st.selectbox("Escenario A", list(opciones_comp.keys()), index=0, key="comp_fv_a")
            with cc2:
                idx_b_defecto = 1 if len(opciones_comp) > 1 else 0
                sel_b = st.selectbox("Escenario B", list(opciones_comp.keys()), index=idx_b_defecto, key="comp_fv_b")

            def _resolver_escenario(sel):
                tipo, idx = opciones_comp[sel]
                if tipo == "actual":
                    return inp_actual, res_actual
                return guardados[idx]["inputs_fv"], guardados[idx]["resultado_fv"]

            inp_a, res_a = _resolver_escenario(sel_a)
            inp_b, res_b = _resolver_escenario(sel_b)

            filas_comp = [
                ("Potencia pico", f"{res_a.get('p_pico_kwp', 0):.2f} kWp", f"{res_b.get('p_pico_kwp', 0):.2f} kWp"),
                ("Nº de paneles", f"{res_a.get('n_paneles_configurados', '—')}", f"{res_b.get('n_paneles_configurados', '—')}"),
                ("Con batería", "Sí" if inp_a.get("con_bateria") else "No", "Sí" if inp_b.get("con_bateria") else "No"),
                ("Producción año 1", f"{_miles(res_a.get('produccion_anual_kwh', 0))} kWh",
                 f"{_miles(res_b.get('produccion_anual_kwh', 0))} kWh"),
                ("Producción año 25", f"{_miles(res_a.get('produccion_ano25', 0))} kWh",
                 f"{_miles(res_b.get('produccion_ano25', 0))} kWh"),
                ("Ahorro anual estimado", _fmt_eur(res_a.get("ahorro_anual", 0)), _fmt_eur(res_b.get("ahorro_anual", 0))),
                ("Inversión estimada", _fmt_eur(inp_a.get("inversion_total", 0)), _fmt_eur(inp_b.get("inversion_total", 0))),
                ("Retorno de la inversión",
                 f"{res_a['payback_anos']:.1f} años" if res_a.get("payback_anos") else "—",
                 f"{res_b['payback_anos']:.1f} años" if res_b.get("payback_anos") else "—"),
                ("CO₂ evitado/año", f"{res_a.get('co2_evitado_kg_ano', 0)/1000:.2f} t", f"{res_b.get('co2_evitado_kg_ano', 0)/1000:.2f} t"),
            ]
            etiqueta_a = f"A: {sel_a}"
            etiqueta_b = f"B: {sel_b}" if sel_b != sel_a else f"B: {sel_b} (2)"
            df_comp = pd.DataFrame(filas_comp, columns=["Magnitud", etiqueta_a, etiqueta_b])
            if sel_a == sel_b:
                st.caption("💡 Has elegido el mismo escenario en A y B — cambia uno de los dos para "
                           "comparar de verdad.")
            st.dataframe(df_comp, hide_index=True, width='stretch')

            import plotly.graph_objects as go
            metricas_grafico = ["Potencia pico (kWp)", "Producción año 1 (MWh)", "Ahorro anual (k€)"]
            valores_a = [res_a.get("p_pico_kwp", 0), res_a.get("produccion_anual_kwh", 0) / 1000,
                        res_a.get("ahorro_anual", 0) / 1000]
            valores_b = [res_b.get("p_pico_kwp", 0), res_b.get("produccion_anual_kwh", 0) / 1000,
                        res_b.get("ahorro_anual", 0) / 1000]
            fig_comp = go.Figure(data=[
                go.Bar(name=etiqueta_a, x=metricas_grafico, y=valores_a, marker_color="#3b82f6"),
                go.Bar(name=etiqueta_b, x=metricas_grafico, y=valores_b, marker_color="#e8a33d"),
            ])
            fig_comp.update_layout(barmode="group", margin=dict(t=10, b=10, l=10, r=10), height=300,
                                   paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                                   font=dict(color="#8b96a8"), legend=dict(orientation="h", y=-0.15))
            st.plotly_chart(fig_comp, width='stretch')


def _render_resultados_fv(inp: dict, res: dict):
    st.markdown('<p class="section-label">Resultado</p>', unsafe_allow_html=True)

    r1, r2, r3, r4 = st.columns(4)
    with r1:
        st.markdown(f'''<div class="result-card hero">
            <div class="result-label">Potencia pico</div>
            <div class="result-value">{res['p_pico_kwp']:.2f} kWp</div>
            <div class="result-sub">{res['n_paneles']} paneles · {res['superficie_necesaria_m2']:.1f} m² necesarios</div>
        </div>''', unsafe_allow_html=True)
    with r2:
        st.markdown(f'''<div class="result-card">
            <div class="result-label">Producción anual estimada</div>
            <div class="result-value small">{res['produccion_anual_kwh']:,.0f} kWh/año</div>
            <div class="result-sub">PR efectivo={res['pr_efectivo']:.3f} (incl. orient./sombra/inversor)</div>
        </div>'''.replace(",", "."), unsafe_allow_html=True)
    with r3:
        badge_v = "badge-ok" if (res["cumple_vmax"] and res["cumple_vmpp_min"] and res["cumple_vmpp_max"]) else "badge-fail"
        texto_v = "Compatible" if badge_v == "badge-ok" else "Revisar"
        st.markdown(f'''<div class="result-card">
            <div class="result-label">String {res['n_serie']}S{res['n_paralelo']}P</div>
            <div class="result-value small">{res['v_string_frio']:.0f} V (frío)</div>
            <div class="result-sub"><span class="{badge_v}">{texto_v}</span> con el inversor</div>
        </div>''', unsafe_allow_html=True)
    with r4:
        badge_du = "badge-ok" if res["du_total_pct"] <= 1.5 else "badge-fail"
        st.markdown(f'''<div class="result-card">
            <div class="result-label">ΔU CC+CA combinada</div>
            <div class="result-value small">{res['du_total_pct']:.2f} %</div>
            <div class="result-sub"><span class="{badge_du}">{"Cumple" if res["du_total_pct"]<=1.5 else "No cumple"}</span> · máx 1,5% (ITC-BT-40)</div>
        </div>''', unsafe_allow_html=True)

    for aviso in res["avisos"]:
        (st.error if "supera" in aviso or "combinada" in aviso else st.warning)(aviso)

    st.markdown('<p class="section-label">Pérdidas y producción a largo plazo</p>', unsafe_allow_html=True)
    pl1, pl2, pl3, pl4 = st.columns(4)
    pl1.metric("Pérdidas orientación/inclinación", f"{res['perdidas_orient_pct']:.1f} %",
               help="CTE DB-HE5, según azimut/inclinación/latitud introducidos.")
    pl2.metric("Producción año 1", f"{res['produccion_anual_kwh']:,.0f} kWh".replace(",", "."))
    pl3.metric("Producción año 10", f"{res['produccion_ano10']:,.0f} kWh".replace(",", "."),
               help=f"Con degradación de {inp['degradacion_anual']:g}%/año.")
    pl4.metric("Producción año 25", f"{res['produccion_ano25']:,.0f} kWh".replace(",", "."))

    st.markdown('<p class="section-label">Generador y string</p>', unsafe_allow_html=True)
    g1, g2, g3, g4 = st.columns(4)
    g1.metric("Paneles configurados", f"{res['n_paneles_configurados']}")
    g2.metric("V string en caliente", f"{res['v_string_caliente']:.0f} V")
    g3.metric("I generador (Icc total)", f"{res['i_generador']:.1f} A")
    g4.metric("Ratio DC/AC", f"{res['ratio_dc_ac']:.2f}" if res["ratio_dc_ac"] else "—")

    st.markdown('<p class="section-label">Cableado y protecciones</p>', unsafe_allow_html=True)
    cbl = pd.DataFrame([
        {"Tramo": "CC (string → caja de conexión)", "I diseño (125%)": f"{res['i_diseno_cc']:.2f} A",
         "Sección": f"{res['s_cc_final']:g} mm² Cu XLPE", "ΔU parcial": f"{res['du_cc_pct']:.2f} %"},
        {"Tramo": "CA (inversor → cuadro)", "I diseño (125%)": f"{res['i_diseno_ca']:.2f} A",
         "Sección": f"{res['s_ca_final']:g} mm² Cu XLPE" if res["s_ca_final"] else "—",
         "ΔU parcial": f"{res['e_ca_pct']:.2f} %"},
    ])
    st.dataframe(cbl, width='stretch', hide_index=True)

    p1, p2, p3 = st.columns(3)
    p1.metric("Fusible de string sugerido",
              f"{res['calibre_fusible_string']} A / {res['tension_fusible_string']} V"
              if res["calibre_fusible_string"] else "No requerido (1 string)")
    p2.metric("Magnetotérmico CA sugerido", f"{res['calibre_magneto_ca']} A")
    p3.metric("CO₂ evitado", f"{res['co2_evitado_kg_ano']/1000:.2f} t/año")

    if res.get("capacidad_bateria_kwh"):
        st.markdown('<p class="section-label">Batería / almacenamiento</p>', unsafe_allow_html=True)
        b1, b2 = st.columns(2)
        b1.metric("Capacidad de batería necesaria", f"{res['capacidad_bateria_kwh']:.1f} kWh")
        b2.metric("Autonomía configurada", f"{inp['autonomia_dias']:g} días · "
                  f"{inp['profundidad_descarga']:g}% descarga admisible")

    if res["ahorro_anual"]:
        st.markdown('<p class="section-label">Estimación económica</p>', unsafe_allow_html=True)
        e1, e2, e3, e4 = st.columns(4)
        e1.metric("Energía autoconsumida", f"{res['energia_autoconsumida']:,.0f} kWh".replace(",", "."))
        e2.metric("Ahorro por autoconsumo", _fmt_eur(res["ahorro_autoconsumo"]))
        e3.metric("Ingreso por excedentes", _fmt_eur(res["ingreso_excedentes"]))
        e4.metric("Retorno simple (payback)", f"{res['payback_anos']:.1f} años" if res["payback_anos"] else "—")
        st.caption("Estimación simplificada (no considera IPC de la electricidad, degradación del panel ni "
                   "financiación). Para un estudio de viabilidad usa PVGIS y un análisis de flujo de caja.")


def _fila_formula(latex_str: str, calculo_md: str):
    """Muestra una fórmula en LaTeX junto a su cálculo con valores sustituidos."""
    col_formula, col_calculo = st.columns([1, 1])
    with col_formula:
        st.latex(latex_str)
    with col_calculo:
        st.markdown(calculo_md)
    st.markdown("<hr style='margin:0.3rem 0;'>", unsafe_allow_html=True)


def _render_formulas(inp: dict, res: dict):
    st.markdown('<p class="section-label">Justificación del cálculo</p>', unsafe_allow_html=True)
    st.caption("Cada fórmula, con el cálculo de este circuito justo al lado. Mismo contenido que incluye "
               "la memoria en PDF.")

    if res["seccion_final"] is None:
        st.warning("Ajusta los datos en la pestaña Calculadora para poder mostrar la justificación.")
        return

    sistema = inp["sistema"]
    cos_phi = inp["cos_phi"]

    st.markdown("##### 1 · Corriente de empleo")
    if inp["modo_entrada"] == "Potencia activa":
        p_w = inp["potencia_kw"] * 1000.0
        if sistema == SISTEMA_MONO:
            _fila_formula(
                r"I_b = \dfrac{P}{V \cdot \cos\varphi}",
                f"`Ib = {p_w:.0f} / ({inp['tension']:g} × {cos_phi:.2f})`  \n**Ib = {res['ib']:.2f} A**")
        else:
            _fila_formula(
                r"I_b = \dfrac{P}{\sqrt{3} \cdot V \cdot \cos\varphi}",
                f"`Ib = {p_w:.0f} / (1,732 × {inp['tension']:g} × {cos_phi:.2f})`  \n**Ib = {res['ib']:.2f} A**")
    else:
        _fila_formula(r"I_b = \text{dato directo}", f"**Ib = {res['ib']:.2f} A** (introducida)")

    if res.get("ib_motor") is not None:
        if len(inp["corrientes_motores"]) == 1:
            _fila_formula(
                r"I_{b,\,motor} = 1{,}25 \cdot I_n",
                f"`Ib = 1,25 × {inp['corrientes_motores'][0]:.2f}`  \n**Ib,motor = {res['ib_motor']:.2f} A**")
        else:
            ordenados = sorted(inp["corrientes_motores"], reverse=True)
            resto = " + ".join(f"{x:.1f}" for x in ordenados[1:])
            _fila_formula(
                r"I_{b,\,motor} = 1{,}25 \cdot I_{n,\,mayor} + \textstyle\sum I_{n,\,resto}",
                f"`Ib = 1,25 × {ordenados[0]:.2f} + ({resto})`  \n**Ib,motor = {res['ib_motor']:.2f} A**")
    if inp["alumbrado_descarga"]:
        _fila_formula(r"I_{b,\,cálculo} = I_b \times 1{,}8",
                       f"Alumbrado de descarga sin corregir f.d.p.  \n**Ib,cálculo = {res['ib_calculo']:.2f} A**")

    st.markdown("##### 2 · Criterio térmico (ITC-BT-19 art. 19)")
    _fila_formula(
        r"I_z = I_{z,\,tabla} \cdot f_{temp} \cdot f_{agrup} \cdot f_{capas} \cdot f_{resist}",
        f"`Iz = {res['iz_termica']/max(res['factor_total'],1e-9):.1f} × {res['f_temp']:.3f} × "
        f"{res['f_agrup']:.3f} × {res['f_capas']:.3f} × {res['f_resist']:.3f}`  \n"
        f"**Iz = {res['iz_termica']:.2f} A**")
    cumple_termico = res["iz_termica"] >= (res["ib_calculo"] / (res["n_paralelo"] if res["necesita_paralelo"] else 1))
    _fila_formula(
        r"I_b \leq I_n \leq I_z",
        f"`{res['ib_calculo']:.2f} A <= Iz={res['iz_termica']:.2f} A`  \n"
        f"**{'✅ Cumple' if cumple_termico else '❌ No cumple'} con S = {res['s_termica']:g} mm²**")

    st.markdown("##### 3 · Criterio de caída de tensión")
    kappa = res["kappa"]
    _fila_formula(
        r"\kappa(T) = \dfrac{\kappa_{20°C}}{1 + \alpha \cdot (T_{servicio} - 20)}",
        f"`kappa = {CONDUCTIVIDAD_20C[inp['conductor']]:g} / (1 + {COEF_TEMP_RESIST[inp['conductor']]:g} × "
        f"({TEMP_SERVICIO[inp['aislamiento']]:g} - 20))`  \n**κ = {kappa:.2f} m/(Ω·mm²)**")
    S = res["seccion_final"]
    r_metro = 1.0 / (kappa * S)
    k_sist_txt = "2" if sistema == SISTEMA_MONO else r"\sqrt{3}"
    _fila_formula(
        rf"\Delta U = {k_sist_txt} \cdot L \cdot I_b \cdot (R\cos\varphi + X\sin\varphi), \; R = \dfrac{{1}}{{\kappa S}}",
        f"`R = 1/({kappa:.2f}×{S:g}) = {r_metro:.5f} Ω/m`  \n`ΔU = {res['e_final']:.2f} V`  \n"
        f"**ΔU = {res['e_final_pct']:.2f} %** (máx {inp['delta_u_max']:g} %)")

    if res.get("seccion_neutro"):
        st.markdown("##### 4 · Sección de neutro y conductor de protección")
        _fila_formula(
            r"S_n = S_f \;(S_f\leq16) \qquad S_n \geq S_f/2 \;(S_f>16,\text{ sin armónicos})",
            f"`Sf = {S:g} mm²` → **Sn = {res['seccion_neutro']:g} mm²**")
    if res.get("seccion_proteccion"):
        _fila_formula(
            r"S_p{=}S_f (S_f{\leq}16) \;\; S_p{=}16 (16{<}S_f{\leq}35) \;\; S_p{=}S_f/2 (S_f{>}35)",
            f"`Sf = {S:g} mm²` → **Sp = {res['seccion_proteccion']:g} mm²**")

    if res.get("cumple_cc") is not None:
        st.markdown("##### 5 · Verificación de cortocircuito")
        k_cc = K_CORTOCIRCUITO[(inp["conductor"], inp["aislamiento"])]
        _fila_formula(
            r"S_{min} = \dfrac{I_{cc} \cdot \sqrt{t}}{k}",
            f"`Smin = ({inp['icc_ka']:g}×1000 × √{inp['tiempo_s']:g}) / {k_cc}`  \n"
            f"**Smin = {res['s_min_cc']:g} mm²** → "
            f"{'✅ Cumple' if res['cumple_cc'] else '❌ No cumple'} con S={S:g} mm²")



def _numero_romano(n: int) -> str:
    valores = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    resultado = ""
    for valor, letra in valores:
        while n >= valor:
            resultado += letra
            n -= valor
    return resultado


def _render_presupuesto(inputs_cable: dict, resultado_cable: dict, inputs_fv: dict, resultado_fv: dict):
    st.markdown('<p class="section-label">Presupuesto</p>', unsafe_allow_html=True)
    st.caption("Formato por capítulos (Partida / Designación / Unidades / Cantidad / Precio / Importe), con "
               "desglose de Precio base + Beneficio + Amortización, igual que un presupuesto de instalación "
               "al uso. Los precios son un punto de partida editable, no precios de mercado en tiempo real.")

    st.session_state.setdefault("presupuesto_capitulos", [])
    st.session_state.setdefault("presupuesto_config", {
        "nombre_proyecto": "", "pct_beneficio": PORCENTAJE_BENEFICIO_DEFECTO,
        "pct_amortizacion": PORCENTAJE_AMORTIZACION_DEFECTO, "pct_iva": IVA_DEFECTO_PCT,
    })
    st.session_state.setdefault("catalogo_precios", {
        cat: {item: precio for item, (unidad, precio) in items.items()}
        for cat, items in CATALOGO_MATERIALES.items()
    })
    st.session_state.setdefault("partidas_compuestas", {})
    st.session_state.setdefault("partida_compuesta_en_construccion", [])
    st.session_state.setdefault("calculos_guardados", [])
    capitulos = st.session_state["presupuesto_capitulos"]
    cfg = st.session_state["presupuesto_config"]
    catalogo = st.session_state["catalogo_precios"]
    partidas_compuestas = st.session_state["partidas_compuestas"]

    with st.expander("⚙️ Configuración general", expanded=not capitulos):
        cf1, cf2, cf3, cf4 = st.columns(4)
        with cf1:
            cfg["nombre_proyecto"] = st.text_input("Nombre del proyecto/instalación", cfg["nombre_proyecto"])
        with cf2:
            cfg["pct_beneficio"] = st.number_input("Beneficio industrial (%)", min_value=0.0, max_value=50.0,
                                                    value=cfg["pct_beneficio"], step=0.5)
        with cf3:
            cfg["pct_amortizacion"] = st.number_input("Amortización medios auxiliares (%)", min_value=0.0,
                                                       max_value=20.0, value=cfg["pct_amortizacion"], step=0.5)
        with cf4:
            cfg["pct_iva"] = st.number_input("IVA (%)", min_value=0.0, max_value=25.0, value=cfg["pct_iva"], step=1.0)

    with st.expander("📚 Base de precios — catálogo editable de materiales típicos", expanded=False):
        st.caption(f"{sum(len(v) for v in catalogo.values())} materiales en {len(catalogo)} categorías, "
                   "con todas sus variantes habituales (secciones, calibres, tipos, mano de obra...). Edita "
                   "el precio base de cualquier fila; se guarda para toda la sesión y se usa al añadir "
                   "partidas a un capítulo.")
        cat_editar = st.selectbox("Categoría a revisar/editar", list(catalogo.keys()), key="cat_editar_catalogo")
        df_cat = pd.DataFrame([{"Concepto": item, "Precio base (€)": precio}
                                for item, precio in catalogo[cat_editar].items()])
        df_cat_editada = st.data_editor(df_cat, key=f"editor_catalogo_{cat_editar}", width='stretch',
                                         hide_index=True, num_rows="fixed")
        for _, fila in df_cat_editada.iterrows():
            catalogo[cat_editar][fila["Concepto"]] = float(fila["Precio base (€)"])

        st.divider()
        st.markdown("**📥 Importar precios desde un Excel/CSV de proveedor**")
        st.caption("Sube una lista de precios (una columna con el nombre del artículo, otra con el precio). "
                   "Los que coincidan exactamente con un material ya existente actualizan su precio; el "
                   "resto se añade como materiales nuevos en la categoría que elijas.")
        archivo_precios = st.file_uploader("Archivo del proveedor", type=["xlsx", "xls", "csv"],
                                            key="upload_precios_proveedor")
        if archivo_precios is not None:
            try:
                if archivo_precios.name.lower().endswith(".csv"):
                    df_import = pd.read_csv(archivo_precios)
                else:
                    df_import = pd.read_excel(archivo_precios)
            except Exception as e:
                st.error(f"No se ha podido leer el archivo: {e}")
                df_import = None

            if df_import is not None and not df_import.empty:
                st.dataframe(df_import.head(8), width='stretch', hide_index=True)
                st.caption(f"{len(df_import)} filas leídas. Muestra de las 8 primeras.")
                cols_disponibles = list(df_import.columns)
                imp1, imp2, imp3 = st.columns(3)
                with imp1:
                    col_nombre = st.selectbox("Columna con el nombre del artículo", cols_disponibles,
                                               key="imp_col_nombre")
                with imp2:
                    cols_numericas = [c for c in cols_disponibles
                                       if pd.api.types.is_numeric_dtype(df_import[c])] or cols_disponibles
                    idx_precio_defecto = cols_numericas.index(cols_numericas[-1]) if cols_numericas else 0
                    col_precio = st.selectbox("Columna con el precio", cols_disponibles,
                                               index=cols_disponibles.index(cols_numericas[idx_precio_defecto])
                                               if cols_numericas else 0, key="imp_col_precio")
                with imp3:
                    categoria_nuevos = st.selectbox(
                        "Categoría para los artículos que NO coincidan con ninguno existente",
                        list(catalogo.keys()) + ["📦 Importado de proveedor (nueva categoría)"],
                        key="imp_categoria_nuevos")

                todas_designaciones = {desig: cat for cat, items in catalogo.items() for desig in items}
                filas_preview, n_actualizar, n_nuevos = [], 0, 0
                for _, fila in df_import.iterrows():
                    nombre_raw = str(fila[col_nombre]).strip()
                    try:
                        precio_raw = float(str(fila[col_precio]).replace(",", ".").replace("€", "").strip())
                    except (ValueError, TypeError):
                        continue
                    if not nombre_raw or nombre_raw.lower() == "nan":
                        continue
                    if nombre_raw in todas_designaciones:
                        filas_preview.append({"Artículo": nombre_raw, "Precio nuevo (€)": precio_raw,
                                              "Acción": f"Actualiza precio en «{todas_designaciones[nombre_raw]}»"})
                        n_actualizar += 1
                    else:
                        filas_preview.append({"Artículo": nombre_raw, "Precio nuevo (€)": precio_raw,
                                              "Acción": "Nuevo artículo"})
                        n_nuevos += 1

                if filas_preview:
                    st.dataframe(pd.DataFrame(filas_preview), width='stretch', hide_index=True, height=200)
                    st.caption(f"**{n_actualizar}** artículos actualizarán su precio · **{n_nuevos}** se "
                               "añadirán como nuevos.")
                    if st.button("📥 Importar al catálogo", type="primary", key="btn_importar_precios"):
                        cat_destino_nuevos = categoria_nuevos
                        if cat_destino_nuevos == "📦 Importado de proveedor (nueva categoría)":
                            catalogo.setdefault("📦 Importado de proveedor", {})
                            cat_destino_nuevos = "📦 Importado de proveedor"
                        for fp in filas_preview:
                            nombre_a = fp["Artículo"]
                            precio_a = fp["Precio nuevo (€)"]
                            if nombre_a in todas_designaciones:
                                catalogo[todas_designaciones[nombre_a]][nombre_a] = precio_a
                            else:
                                catalogo[cat_destino_nuevos][nombre_a] = precio_a
                        _registrar_actividad("📥", f"Catálogo actualizado desde {archivo_precios.name}: "
                                             f"{n_actualizar} precios actualizados, {n_nuevos} artículos nuevos")
                        st.success(f"Catálogo actualizado: {n_actualizar} precios actualizados, {n_nuevos} "
                                   "artículos nuevos añadidos.")
                        st.rerun()
                else:
                    st.warning("No se ha podido interpretar ninguna fila válida con las columnas "
                               "seleccionadas — revisa que la columna de precio contenga números.")

    with st.expander("🧩 Partidas compuestas — presupuesto por unidades de obra completas", expanded=False):
        st.caption(
            "Un presupuesto **detallado** lista cada material por separado (una fila = un material). Un "
            "presupuesto **compuesto** agrupa materiales + mano de obra en una sola unidad de obra ya "
            "'montada y puesta en servicio' (p. ej. 'Circuito de alumbrado completo, instalado y probado' "
            "en una sola línea). Aquí construyes esas partidas compuestas; luego están disponibles como una "
            "categoría más al añadir líneas a cualquier capítulo — puedes usar ambos estilos, o mezclarlos, "
            "en el mismo presupuesto.")

        if partidas_compuestas:
            st.markdown("**Partidas compuestas guardadas**")
            for nombre_pc, datos_pc in list(partidas_compuestas.items()):
                precio_pc = sum(c["cantidad"] * c["precio_unitario"] for c in datos_pc["componentes"])
                with st.container(border=True):
                    pc1, pc2 = st.columns([5, 1])
                    pc1.markdown(f"**{nombre_pc}** — {_fmt_eur(precio_pc)} / {datos_pc['unidad']} "
                                 f"({len(datos_pc['componentes'])} componentes)")
                    if pc2.button("🗑️", key=f"del_pc_{nombre_pc}"):
                        del partidas_compuestas[nombre_pc]
                        st.rerun()
                    st.dataframe(pd.DataFrame(datos_pc["componentes"]), hide_index=True, width='stretch')

        st.markdown("**Construir una partida compuesta nueva**")
        en_construccion = st.session_state["partida_compuesta_en_construccion"]
        cc1, cc2, cc3, cc4 = st.columns([2, 2.3, 1, 1])
        with cc1:
            cat_pc = st.selectbox("Categoría", list(catalogo.keys()), key="pc_cat")
        with cc2:
            item_pc = st.selectbox("Material / mano de obra", list(catalogo[cat_pc].keys()), key="pc_item")
        with cc3:
            cant_pc = st.number_input("Cantidad", min_value=0.01, value=1.0, step=0.5, key="pc_cant")
        with cc4:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("➕ Añadir", key="pc_add_comp"):
                en_construccion.append({"designacion": item_pc, "cantidad": cant_pc,
                                        "precio_unitario": catalogo[cat_pc][item_pc]})
                st.rerun()

        if en_construccion:
            st.dataframe(pd.DataFrame(en_construccion), hide_index=True, width='stretch')
            precio_total_construccion = sum(c["cantidad"] * c["precio_unitario"] for c in en_construccion)
            st.markdown(f"Precio unitario resultante: **{_fmt_eur(precio_total_construccion)}**")
            fc1, fc2, fc3, fc4 = st.columns([3, 1, 1, 1])
            nombre_pc_final = fc1.text_input("Nombre de la partida", key="pc_nombre_final",
                                              placeholder="Circuito de alumbrado completo, instalado y probado")
            unidad_pc_final = fc2.selectbox("Unidad", ["ud", "m", "kWp", "punto", "circuito"],
                                             key="pc_unidad_final")
            with fc3:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("💾 Guardar", key="pc_guardar"):
                    if nombre_pc_final:
                        partidas_compuestas[nombre_pc_final] = {"unidad": unidad_pc_final,
                                                                 "componentes": list(en_construccion)}
                        st.session_state["partida_compuesta_en_construccion"] = []
                        st.session_state.pop("pc_nombre_final", None)
                        _registrar_actividad("🧩", f"Partida compuesta creada: {nombre_pc_final}")
                        st.rerun()
            with fc4:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🗑️ Vaciar", key="pc_vaciar"):
                    st.session_state["partida_compuesta_en_construccion"] = []
                    st.session_state.pop("pc_nombre_final", None)
                    st.rerun()

    with st.expander("📏 Estimador de cantidades — por tipo de instalación", expanded=False):
        st.caption(
            "⚠️ Estimación **paramétrica** a partir del tipo de instalación y sus subpartes típicas — no es "
            "una medición real sobre plano. Desmarca las subpartes que no apliquen y ajusta las cantidades; "
            "luego exporta el resultado a un capítulo del presupuesto.")

        tipo_inst_sel = st.selectbox("Tipo de instalación", list(TIPOS_INSTALACION_ESTIMADOR.keys()),
                                      key="est_tipo_inst")
        if "Electrificación básica" in tipo_inst_sel:
            st.info("ITC-BT-25: grado básico, potencia mínima **5.750 W**. Obligatorio en toda vivienda; "
                     "cubre los circuitos C1 a C5 (iluminación, tomas generales, cocina, "
                     "lavadora/lavavajillas/termo, y baño).")
        elif "Electrificación elevada" in tipo_inst_sel:
            st.info("ITC-BT-25: grado elevado, potencia mínima **9.200 W**. Obligatorio si la superficie útil "
                     "supera 160 m², o hay previsión de calefacción/aire acondicionado eléctricos, secadora, "
                     "automatización, o más de 30 puntos de utilización. Incluye C1-C5 de la básica más "
                     "C6 a C12.")
        est_holgura = st.slider("Holgura por curvas/subidas en tubo, cable y conductor de tierra (%)",
                                 0, 40, 15, key="est_holgura_global")

        subpartes = TIPOS_INSTALACION_ESTIMADOR[tipo_inst_sel]
        st.markdown(f"**Subpartes de «{tipo_inst_sel}»**")
        datos_subpartes = []
        for i_sp, sp in enumerate(subpartes):
            with st.container(border=True):
                sc1, sc2, sc3, sc4 = st.columns([0.4, 2.6, 1.3, 1.3])
                with sc1:
                    incluido = st.checkbox("Incluir", value=True, label_visibility="collapsed",
                                            key=f"est_incl_{tipo_inst_sel}_{i_sp}")
                with sc2:
                    st.markdown(f"**{sp['nombre']}**  \n"
                                f"<span style='font-size:0.72rem; color:var(--text-secondary);'>"
                                f"{sp['tipo']}</span>", unsafe_allow_html=True)
                with sc3:
                    n_val = st.number_input(f"Nº {sp['unidad_elem']}", min_value=0, value=sp["n_defecto"],
                                             step=1, key=f"est_n_{tipo_inst_sel}_{i_sp}")
                with sc4:
                    if sp["tipo"] in ("circuito", "tierra"):
                        long_val = st.number_input("Long. media (m)", min_value=0.0,
                                                    value=float(sp.get("long_defecto", 10.0)), step=1.0,
                                                    key=f"est_long_{tipo_inst_sel}_{i_sp}")
                    else:
                        long_val = None
                        st.caption("(elemento único, sin tubo/cable)")
                datos_subpartes.append((sp, incluido, n_val, long_val))

        filas_resumen, items_por_subparte = [], {}
        total_tubo = total_cable = total_cajas = 0.0
        for sp, incluido, n_val, long_val in datos_subpartes:
            if not incluido or n_val <= 0:
                continue
            items_sp = items_por_subparte.setdefault(sp["nombre"], [])
            if sp["tipo"] == "circuito":
                tubo = n_val * long_val * (1 + est_holgura / 100)
                cable = tubo * sp["conductores"]
                cajas = n_val
                total_tubo += tubo
                total_cable += cable
                total_cajas += cajas
                filas_resumen.append({"Subparte": sp["nombre"], "Tubo (m)": f"{tubo:.1f}",
                                      "Cable (m)": f"{cable:.1f}", "Cajas/uds": str(cajas)})
                items_sp.append({"designacion": f"Tubo corrugado Ø20mm — {sp['nombre']} (estimado)",
                                 "unidades": "m", "cantidad": round(tubo, 1),
                                 "precio_base": catalogo["Canalizaciones"]["Tubo corrugado empotrar Ø20mm"]})
                items_sp.append({"designacion": f"Cable H07V-K 2,5mm² — {sp['nombre']} (estimado)",
                                 "unidades": "m", "cantidad": round(cable, 1),
                                 "precio_base": catalogo["Cables (por tipo, además del cálculo automático)"]["Cable H07V-K 2,5mm²"]})
                items_sp.append({"designacion": f"Caja de derivación — {sp['nombre']} (estimado)",
                                 "unidades": "ud", "cantidad": cajas,
                                 "precio_base": catalogo["Cajas y mecanismos"]["Caja de derivación empotrar 100x100"]})
            elif sp["tipo"] == "tierra":
                conductor = long_val * (1 + est_holgura / 100)
                filas_resumen.append({"Subparte": sp["nombre"], "Tubo (m)": "—", "Cable (m)": "—",
                                      "Cajas/uds": f"{n_val} picas + {conductor:.0f} m cond."})
                items_sp.append({"designacion": f"Pica de tierra — {sp['nombre']} (estimado)",
                                 "unidades": "ud", "cantidad": n_val,
                                 "precio_base": catalogo["Puesta a tierra"]["Pica de acero cobreado 2m Ø14mm"]})
                items_sp.append({"designacion": f"Cable desnudo Cu 16mm² — {sp['nombre']} (estimado)",
                                 "unidades": "m", "cantidad": round(conductor, 1),
                                 "precio_base": catalogo["Puesta a tierra"]["Cable desnudo Cu 16mm² (tierra)"]})
            else:  # elemento
                filas_resumen.append({"Subparte": sp["nombre"], "Tubo (m)": "—", "Cable (m)": "—",
                                      "Cajas/uds": f"{n_val} ud"})
                items_sp.append({"designacion": f"{sp['nombre']} (estimado)", "unidades": "ud",
                                 "cantidad": n_val, "precio_base": sp["precio_estimado"]})

        if filas_resumen:
            st.markdown("**Resumen por subparte**")
            st.dataframe(pd.DataFrame(filas_resumen), hide_index=True, width='stretch')
            rt1, rt2, rt3 = st.columns(3)
            rt1.metric("Total tubo", f"{total_tubo:.0f} m")
            rt2.metric("Total cable", f"{total_cable:.0f} m")
            rt3.metric("Total cajas de derivación", f"{total_cajas:.0f}")

            st.caption(f"Al exportar, cada subparte crea (o usa, si ya existe) **su propio capítulo** — "
                       f"se crearán/rellenarán {len(items_por_subparte)} capítulos, uno por subparte, en vez "
                       "de mezclarlo todo en uno solo.")
            if st.button("📤 Exportar — cada subparte a su propio capítulo", key="est_exportar_btn"):
                nombres_actuales = {c["nombre"]: c for c in capitulos}
                n_nuevos, n_reutilizados, total_items = 0, 0, 0
                for nombre_sp, items_sp in items_por_subparte.items():
                    if nombre_sp in nombres_actuales:
                        cap_dest = nombres_actuales[nombre_sp]
                        n_reutilizados += 1
                    else:
                        cap_dest = {"nombre": nombre_sp, "items": []}
                        capitulos.append(cap_dest)
                        nombres_actuales[nombre_sp] = cap_dest
                        n_nuevos += 1
                    idx_cap = capitulos.index(cap_dest)
                    for j, it in enumerate(items_sp):
                        it["partida"] = f"{idx_cap + 1}.{len(cap_dest['items']) + j + 1}"
                    cap_dest["items"].extend(items_sp)
                    total_items += len(items_sp)
                _registrar_actividad("📏", f"Estimación de «{tipo_inst_sel}» exportada: {n_nuevos} capítulos "
                                     f"nuevos, {n_reutilizados} reutilizados, {total_items} partidas")
                st.success(f"{total_items} partidas repartidas en {len(items_por_subparte)} capítulos "
                           f"({n_nuevos} nuevos, {n_reutilizados} ya existentes y reutilizados).")
                st.rerun()
        else:
            st.info("Marca al menos una subparte con cantidad mayor que 0 para ver el resumen.")

    nc1, nc2 = st.columns([3, 1])
    with nc1:
        nombre_defecto = f"CAPÍTULO {_numero_romano(len(capitulos) + 1)}: "
        nuevo_nombre = st.text_input("Nombre del nuevo capítulo", nombre_defecto, key="nuevo_cap_nombre")
    with nc2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("➕ Nuevo capítulo"):
            capitulos.append({"nombre": nuevo_nombre, "items": []})
            st.rerun()

    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None

    for idx, cap in enumerate(capitulos):
        with st.container(border=True):
            tc1, tc2 = st.columns([5, 0.6])
            tc1.markdown(f"**{cap['nombre']}**")
            if tc2.button("🗑️ Eliminar capítulo", key=f"del_cap_{idx}"):
                capitulos.pop(idx)
                st.rerun()

            calculos_disp = {"— Cálculo actual de la Calculadora —": ("actual_cable", None)}
            for i_calc, calc in enumerate(st.session_state["calculos_guardados"]):
                calculos_disp[f"💾 {calc['nombre']}"] = ("guardado", i_calc)

            ib1, ib2, ib3 = st.columns([2.2, 1.3, 1])
            with ib1:
                origen_cable = st.selectbox("Cálculo de cable a importar", list(calculos_disp.keys()),
                                            key=f"origen_cable_{idx}")
            with ib2:
                st.markdown("<br>", unsafe_allow_html=True)
                tipo_origen, i_origen = calculos_disp[origen_cable]
                puede_importar = (tipo_origen == "actual_cable" and hay_cable) or tipo_origen == "guardado"
                if st.button("📥 Importar cable", key=f"imp_cable_{idx}", disabled=not puede_importar):
                    if tipo_origen == "actual_cable":
                        ic_origen, rc_origen = inputs_cable, resultado_cable
                    else:
                        guardado = st.session_state["calculos_guardados"][i_origen]
                        ic_origen, rc_origen = guardado["inputs_cable"], guardado["resultado_cable"]
                    nuevos = item_desde_calculo_cable(ic_origen, rc_origen)
                    for j, it in enumerate(nuevos):
                        it["partida"] = f"{idx + 1}.{len(cap['items']) + j + 1}"
                    cap["items"].extend(nuevos)
                    st.rerun()
            with ib3:
                st.markdown("<br>", unsafe_allow_html=True)
                if hay_fv and st.button("📥 Importar FV", key=f"imp_fv_{idx}"):
                    nuevos = items_desde_calculo_fv(inputs_fv, resultado_fv)
                    for j, it in enumerate(nuevos):
                        it["partida"] = f"{idx + 1}.{len(cap['items']) + j + 1}"
                    cap["items"].extend(nuevos)
                    st.rerun()

            fuentes_disponibles = list(catalogo.keys())
            if partidas_compuestas:
                fuentes_disponibles = ["🧩 Partidas compuestas"] + fuentes_disponibles
            ac1, ac2, ac3, ac4 = st.columns([2, 2.5, 1, 1])
            with ac1:
                cat_sel = st.selectbox("Categoría / origen", fuentes_disponibles, key=f"cat_sel_{idx}")
            es_compuesta = cat_sel == "🧩 Partidas compuestas"
            opciones_item = list(partidas_compuestas.keys()) if es_compuesta else list(catalogo[cat_sel].keys())
            with ac2:
                item_sel = st.selectbox("Material / partida", opciones_item, key=f"item_sel_{idx}")
            with ac3:
                cant_sel = st.number_input("Cantidad", min_value=0.1, value=1.0, step=1.0, key=f"cant_sel_{idx}")
            with ac4:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("➕ Añadir", key=f"add_cat_{idx}", disabled=not opciones_item):
                    if es_compuesta:
                        datos_pc = partidas_compuestas[item_sel]
                        precio_pc = sum(c["cantidad"] * c["precio_unitario"] for c in datos_pc["componentes"])
                        unidad, precio = datos_pc["unidad"], precio_pc
                    else:
                        unidad = CATALOGO_MATERIALES[cat_sel][item_sel][0]
                        precio = catalogo[cat_sel][item_sel]
                    cap["items"].append({
                        "partida": f"{idx + 1}.{len(cap['items']) + 1}", "designacion": item_sel,
                        "unidades": unidad, "cantidad": cant_sel, "precio_base": precio,
                    })
                    st.rerun()

            df_items = pd.DataFrame([
                {"Partida": it.get("partida", "-"), "Designación": it["designacion"], "Unidades": it["unidades"],
                 "Cantidad": it["cantidad"], "Precio base": it["precio_base"]}
                for it in cap["items"]
            ]) if cap["items"] else pd.DataFrame(columns=["Partida", "Designación", "Unidades", "Cantidad", "Precio base"])

            st.caption("También puedes editar, añadir o borrar filas directamente en la tabla (con texto libre).")
            df_editada = st.data_editor(
                df_items, key=f"items_editor_{idx}", num_rows="dynamic", width='stretch', hide_index=True,
                column_config={
                    "Cantidad": st.column_config.NumberColumn(min_value=0.0, step=0.1),
                    "Precio base": st.column_config.NumberColumn(min_value=0.0, step=0.01, format="%.4f €"),
                },
            )

            nuevos_items = []
            for i, row in df_editada.iterrows():
                designacion = row.get("Designación")
                if not designacion or pd.isna(designacion):
                    continue
                cantidad = row.get("Cantidad")
                precio_base = row.get("Precio base")
                nuevos_items.append({
                    "partida": row.get("Partida") or f"{idx + 1}.{i + 1}",
                    "designacion": designacion,
                    "unidades": row.get("Unidades") or "ud",
                    "cantidad": float(cantidad) if pd.notna(cantidad) else 0.0,
                    "precio_base": float(precio_base) if pd.notna(precio_base) else 0.0,
                })
            cap["items"] = nuevos_items

            if cap["items"]:
                filas_venta = []
                for it in cap["items"]:
                    pv = calcular_precio_venta(it["precio_base"], cfg["pct_beneficio"], cfg["pct_amortizacion"])
                    filas_venta.append({
                        "Partida": it.get("partida", "-"), "Designación": it["designacion"], "Unidades": it["unidades"],
                        "Cantidad": it["cantidad"], "Precio": pv, "Importe": round(it["cantidad"] * pv, 2),
                    })
                st.dataframe(pd.DataFrame(filas_venta), width='stretch', hide_index=True)
                total_cap = calcular_totales_capitulo(cap["items"], cfg["pct_beneficio"], cfg["pct_amortizacion"])
                st.markdown(f"**TOTAL {cap['nombre']}: {_fmt_eur(total_cap)}**")

    if not capitulos:
        st.info("Añade al menos un capítulo para empezar a presupuestar (por ejemplo, uno por cada circuito "
                "o instalación: derivación individual, cuadro de protección, iluminación, fotovoltaica...). "
                "💡 Atajo: en **Inicio → Plantillas → Presupuesto** puedes crear de golpe la estructura típica "
                "de una vivienda, nave industrial o instalación FV.")
        return

    st.markdown('<p class="section-label">Resumen del presupuesto</p>', unsafe_allow_html=True)
    filas_resumen = []
    subtotal = 0.0
    for cap in capitulos:
        parcial = calcular_totales_capitulo(cap["items"], cfg["pct_beneficio"], cfg["pct_amortizacion"])
        subtotal += parcial
        filas_resumen.append({"Capítulo": cap["nombre"], "Coste parcial": _fmt_eur(parcial)})
    st.dataframe(pd.DataFrame(filas_resumen), width='stretch', hide_index=True)

    importe_iva = round(subtotal * cfg["pct_iva"] / 100.0, 2)
    total = round(subtotal + importe_iva, 2)

    s1, s2, s3 = st.columns(3)
    s1.markdown(f'''<div class="result-card"><div class="result-label">Subtotal</div>
        <div class="result-value small">{_fmt_eur(subtotal)}</div></div>''', unsafe_allow_html=True)
    s2.markdown(f'''<div class="result-card"><div class="result-label">IVA ({cfg["pct_iva"]:g}%)</div>
        <div class="result-value small">{_fmt_eur(importe_iva)}</div></div>''', unsafe_allow_html=True)
    s3.markdown(f'''<div class="result-card hero"><div class="result-label">TOTAL</div>
        <div class="result-value">{_fmt_eur(total)}</div></div>''', unsafe_allow_html=True)
    st.caption(numero_a_letras_euros(total).capitalize() + ".")

    with st.spinner("Generando Excel..."):
        excel_bytes = generar_excel_presupuesto_capitulos(
            capitulos, cfg["pct_beneficio"], cfg["pct_amortizacion"], cfg["pct_iva"], cfg["nombre_proyecto"])
    if st.download_button("⬇️ Descargar presupuesto (Excel, por capítulos)", data=excel_bytes,
                        file_name="presupuesto.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"):
        _registrar_actividad("💰", "Presupuesto descargado (Excel)")


def _render_documentacion(inputs_cable: dict, resultado_cable: dict, inputs_fv: dict, resultado_fv: dict):
    st.markdown('<p class="section-label">Documentación</p>', unsafe_allow_html=True)
    st.caption("Genera la Memoria Técnica de Diseño (MTD, ITC-BT-04), el Anexo de Cálculos y Mediciones, y "
               "las Condiciones Generales de ejecución, reuniendo lo que ya tengas calculado en las demás "
               "pestañas. A diferencia de la Calculadora y el módulo Fotovoltaico (independientes entre sí), "
               "estos tres documentos SÍ están pensados para juntar toda la información del proyecto.")

    st.session_state.setdefault("datos_proyecto", {
        "titular": "", "nif_titular": "", "emplazamiento": "", "referencia_catastral": "", "uso": "",
        "superficie": "", "tipo_instalacion": "Nueva instalación",
        "instalador": "", "nif_instalador": "", "n_autorizacion": "", "categoria_instalador": "Básica",
    })
    datos = st.session_state["datos_proyecto"]

    st.markdown("**A · Titular**")
    d1, d2, d3 = st.columns(3)
    with d1:
        datos["titular"] = st.text_input("Titular de la instalación (nombre o razón social)", datos["titular"])
    with d2:
        datos["nif_titular"] = st.text_input("NIF/CIF del titular", datos["nif_titular"])
    with d3:
        datos["tipo_instalacion"] = st.selectbox(
            "Tipo de actuación", ["Nueva instalación", "Ampliación", "Reforma"],
            index=["Nueva instalación", "Ampliación", "Reforma"].index(datos["tipo_instalacion"]))

    st.markdown("**B · Emplazamiento**")
    e1, e2, e3, e4 = st.columns(4)
    with e1:
        datos["emplazamiento"] = st.text_input("Emplazamiento (dirección)", datos["emplazamiento"])
    with e2:
        datos["referencia_catastral"] = st.text_input("Referencia catastral (opcional)",
                                                        datos["referencia_catastral"])
    with e3:
        datos["uso"] = st.text_input("Uso de la instalación", datos["uso"],
                                      placeholder="Vivienda unifamiliar, local comercial...")
    with e4:
        datos["superficie"] = st.text_input("Superficie (m²)", datos["superficie"])

    st.markdown("**C · Instalador autorizado / técnico competente**")
    i1, i2, i3, i4 = st.columns(4)
    with i1:
        datos["instalador"] = st.text_input("Nombre del instalador/técnico", datos["instalador"])
    with i2:
        datos["nif_instalador"] = st.text_input("NIF del instalador/técnico", datos["nif_instalador"])
    with i3:
        datos["n_autorizacion"] = st.text_input("Nº de autorización / colegiado", datos["n_autorizacion"])
    with i4:
        datos["categoria_instalador"] = st.selectbox(
            "Categoría", ["Básica", "Especialista", "Técnico titulado competente"],
            index=["Básica", "Especialista", "Técnico titulado competente"].index(datos["categoria_instalador"]))

    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None
    capitulos = st.session_state.get("presupuesto_capitulos", [])
    cfg_presu = st.session_state.get("presupuesto_config", {
        "pct_beneficio": PORCENTAJE_BENEFICIO_DEFECTO, "pct_amortizacion": PORCENTAJE_AMORTIZACION_DEFECTO,
        "pct_iva": IVA_DEFECTO_PCT,
    })

    st.markdown('<p class="section-label">Disponible para incluir</p>', unsafe_allow_html=True)
    e1, e2, e3 = st.columns(3)
    e1.metric("Cálculo de cable", "Sí" if hay_cable else "No calculado")
    e2.metric("Cálculo fotovoltaico", "Sí" if hay_fv else "No calculado")
    e3.metric("Capítulos de presupuesto", f"{len(capitulos)}")

    subtotal_presupuesto = sum(
        calcular_totales_capitulo(cap["items"], cfg_presu["pct_beneficio"], cfg_presu["pct_amortizacion"])
        for cap in capitulos
    )
    total_presupuesto = subtotal_presupuesto * (1 + cfg_presu["pct_iva"] / 100.0) if capitulos else 0.0

    st.divider()
    cfg_prof = st.session_state.get("config_profesional", {})
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**Memoria Técnica de Diseño**")
        st.caption("Memoria descriptiva + justificativa, ITC-BT-04. Documento extenso multi-página.")
        with st.spinner("Generando MTD..."):
            pdf_mtd = generar_pdf_mtd(datos, inputs_cable, resultado_cable, inputs_fv, resultado_fv,
                                       total_presupuesto, cfg_prof)
        if st.download_button("⬇️ Descargar MTD (PDF)", data=pdf_mtd, file_name="MTD.pdf",
                               mime="application/pdf"):
            _registrar_actividad("📄", "MTD descargada")
        with st.spinner("Generando MTD en Word..."):
            docx_mtd = generar_docx_mtd(datos, inputs_cable, resultado_cable, inputs_fv, resultado_fv,
                                         total_presupuesto, cfg_prof)
        if st.download_button("⬇️ Descargar MTD (Word)", data=docx_mtd, file_name="MTD.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            _registrar_actividad("📝", "MTD descargada en Word")
    with c2:
        st.markdown("**Anexo de Cálculos y Mediciones**")
        st.caption("Justificación técnica completa + mediciones por capítulo.")
        with st.spinner("Generando Anexo..."):
            pdf_anexo = generar_pdf_anexo_calculos(datos, inputs_cable, resultado_cable, inputs_fv, resultado_fv,
                                                    capitulos, cfg_presu["pct_beneficio"], cfg_presu["pct_amortizacion"],
                                                    cfg_prof)
        if st.download_button("⬇️ Descargar Anexo (PDF)", data=pdf_anexo, file_name="anexo_calculos.pdf",
                               mime="application/pdf"):
            _registrar_actividad("📄", "Anexo de cálculos descargado")
        with st.spinner("Generando Anexo en Word..."):
            docx_anexo = generar_docx_anexo_calculos(datos, inputs_cable, resultado_cable, inputs_fv, resultado_fv,
                                                       capitulos, cfg_presu["pct_beneficio"], cfg_presu["pct_amortizacion"],
                                                       cfg_prof)
        if st.download_button("⬇️ Descargar Anexo (Word)", data=docx_anexo, file_name="anexo_calculos.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            _registrar_actividad("📝", "Anexo descargado en Word")
    with c3:
        st.markdown("**Pliego de Condiciones**")
        st.caption("Condiciones generales de materiales, ejecución y pruebas.")
        with st.spinner("Generando Pliego..."):
            pdf_cond = generar_pdf_condiciones_generales(datos, hay_fv, cfg_prof)
        if st.download_button("⬇️ Descargar Condiciones (PDF)", data=pdf_cond, file_name="condiciones_generales.pdf",
                               mime="application/pdf"):
            _registrar_actividad("📄", "Pliego de condiciones descargado")
        with st.spinner("Generando Pliego en Word..."):
            docx_cond = generar_docx_condiciones_generales(datos, hay_fv, cfg_prof)
        if st.download_button("⬇️ Descargar Condiciones (Word)", data=docx_cond, file_name="condiciones_generales.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            _registrar_actividad("📝", "Pliego descargado en Word")

    if not hay_cable and not hay_fv:
        st.info("Todavía no hay ningún cálculo hecho: los documentos se generarán igualmente, pero con las "
                "secciones de cálculo vacías. Completa la Calculadora y/o la Fotovoltaica para un contenido "
                "completo.")

    # ---------------------------------------------------------------- CHECKLIST DE PUESTA EN SERVICIO
    st.markdown('<p class="section-label">Puesta en servicio</p>', unsafe_allow_html=True)
    st.session_state.setdefault("checklist_puesta_servicio", [
        {"realizado": False, "valor_medido": ""} for _ in ENSAYOS_PUESTA_SERVICIO
    ])
    st.session_state.setdefault("checklist_firma", {"instalador": "", "fecha": date.today()})
    checklist = st.session_state["checklist_puesta_servicio"]
    if len(checklist) != len(ENSAYOS_PUESTA_SERVICIO):  # por si la lista de ensayos cambia de tamaño
        checklist = [{"realizado": False, "valor_medido": ""} for _ in ENSAYOS_PUESTA_SERVICIO]
        st.session_state["checklist_puesta_servicio"] = checklist

    with st.expander("✅ Checklist de puesta en servicio (anexo del CIE)", expanded=False):
        st.caption("Los mismos 9 ensayos del Pliego de Condiciones (ITC-BT-05), pero aquí los marcas "
                   "conforme los vas haciendo en la instalación real, con el valor medido. Se incluye como "
                   "anexo al generar el CIE.")
        for i_ens, (nombre_ens, criterio_ens) in enumerate(ENSAYOS_PUESTA_SERVICIO):
            ch1, ch2, ch3 = st.columns([0.5, 3, 2])
            with ch1:
                checklist[i_ens]["realizado"] = st.checkbox("Hecho", value=checklist[i_ens]["realizado"],
                                                             key=f"chk_ens_{i_ens}", label_visibility="collapsed")
            with ch2:
                st.markdown(f"**{nombre_ens}**  \n<span style='font-size:0.75rem; color:var(--text-secondary);'>"
                            f"Criterio: {criterio_ens}</span>", unsafe_allow_html=True)
            with ch3:
                checklist[i_ens]["valor_medido"] = st.text_input("Valor medido", checklist[i_ens]["valor_medido"],
                                                                  key=f"txt_ens_{i_ens}", label_visibility="collapsed",
                                                                  placeholder="Valor medido")
        n_hechos = sum(1 for e in checklist if e["realizado"])
        st.progress(n_hechos / len(checklist), text=f"{n_hechos}/{len(checklist)} ensayos completados")
        st.divider()
        fi1, fi2 = st.columns(2)
        with fi1:
            st.session_state["checklist_firma"]["instalador"] = st.text_input(
                "Nombre de quien realiza las pruebas", st.session_state["checklist_firma"]["instalador"])
        with fi2:
            st.session_state["checklist_firma"]["fecha"] = st.date_input(
                "Fecha de las pruebas", st.session_state["checklist_firma"]["fecha"])

    # ---------------------------------------------------------------- REGISTRO FOTOGRAFICO
    st.session_state.setdefault("fotos_instalacion", [])
    fotos = st.session_state["fotos_instalacion"]
    with st.expander(f"📷 Registro fotográfico de la instalación ({len(fotos)})", expanded=False):
        st.caption("Sube fotos del antes/durante/después de la instalación — se incluyen como anexo en la "
                   "MTD y en el CIE. Máximo recomendado: 12 fotos, para no disparar el tamaño del PDF.")
        fo1, fo2 = st.columns([1, 2])
        with fo1:
            etapa_foto = st.selectbox("Etapa", ["Antes", "Durante", "Después"], key="etapa_foto_nueva")
        with fo2:
            archivos_foto = st.file_uploader("Sube una o varias fotos", type=["jpg", "jpeg", "png"],
                                              accept_multiple_files=True, key="uploader_fotos")
        if archivos_foto:
            if st.button("➕ Añadir al registro", key="btn_add_fotos"):
                for archivo in archivos_foto:
                    fotos.append({"etapa": etapa_foto, "nombre": archivo.name,
                                  "b64": base64.b64encode(archivo.read()).decode(), "descripcion": ""})
                _registrar_actividad("📷", f"{len(archivos_foto)} foto(s) añadida(s) al registro ({etapa_foto})")
                st.rerun()

        if fotos:
            for i_f, foto in enumerate(fotos):
                fc1, fc2, fc3 = st.columns([1, 3, 0.5])
                with fc1:
                    st.image(base64.b64decode(foto["b64"]), width=100)
                with fc2:
                    st.caption(f"{foto['etapa']} — {foto['nombre']}")
                    foto["descripcion"] = st.text_input("Descripción (opcional)", foto["descripcion"],
                                                         key=f"desc_foto_{i_f}", label_visibility="collapsed",
                                                         placeholder="Descripción (opcional)")
                with fc3:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("🗑️", key=f"del_foto_{i_f}"):
                        fotos.pop(i_f)
                        st.rerun()

    # ---------------------------------------------------------------- CIE Y FICHA RESUMEN
    st.markdown('<p class="section-label">Otros documentos</p>', unsafe_allow_html=True)
    oc1, oc2 = st.columns(2)
    with oc1:
        st.markdown("**Certificado de Instalación Eléctrica (CIE)**")
        st.caption("Rellenado a partir de los datos del proyecto y la checklist de puesta en servicio "
                   "de arriba. Documento de apoyo — el CIE oficial lo emite el instalador autorizado a "
                   "través de la aplicación de su Comunidad Autónoma.")
        with st.spinner("Generando CIE..."):
            pdf_cie = generar_pdf_cie(datos, inputs_cable, resultado_cable, inputs_fv, resultado_fv,
                                       checklist, st.session_state["checklist_firma"], cfg_prof)
        if st.download_button("⬇️ Descargar CIE (PDF)", data=pdf_cie, file_name="CIE.pdf",
                               mime="application/pdf"):
            _registrar_actividad("📄", "CIE descargado")
    with oc2:
        st.markdown("**Ficha resumen (una página)**")
        st.caption("Lo esencial del proyecto en una sola hoja: para imprimir o enviar rápido sin abrir los "
                   "documentos completos.")
        with st.spinner("Generando ficha resumen..."):
            pdf_resumen = generar_pdf_resumen_una_pagina(datos, inputs_cable, resultado_cable, inputs_fv,
                                                          resultado_fv, total_presupuesto, cfg_prof)
        if st.download_button("⬇️ Descargar ficha resumen (PDF)", data=pdf_resumen, file_name="ficha_resumen.pdf",
                               mime="application/pdf"):
            _registrar_actividad("📄", "Ficha resumen descargada")


def _render_calculos_bt():
    st.markdown('<p class="section-label">Cálculos BT</p>', unsafe_allow_html=True)
    st.caption("Calculadoras sueltas de referencia rápida, independientes entre sí y de las demás pestañas "
               "— inspiradas en el catálogo de calculadoras de circuitoelectrico.com (cables, cortocircuito, "
               "tierras, fotovoltaica rápida y eléctricas generales). Para el dimensionado completo de un "
               "circuito, usa la pestaña Calculadora.")

    # ---------------------------------------------------------------- CABLES
    with st.expander("🔌 Cables eléctricos", expanded=True):
        st.markdown("**Caída de tensión** (dados sección, longitud e intensidad)")
        c1, c2, c3, c4, c5 = st.columns(5)
        sistema_ct = c1.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI], key="bt_ct_sistema")
        s_ct = c2.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=4, key="bt_ct_seccion")
        l_ct = c3.number_input("Longitud (m)", min_value=0.1, value=20.0, key="bt_ct_l")
        i_ct = c4.number_input("Intensidad (A)", min_value=0.1, value=20.0, key="bt_ct_i")
        cond_ct = c5.selectbox("Conductor", ["Cobre", "Aluminio"], key="bt_ct_cond")
        cosphi_ct = st.slider("cos φ", 0.1, 1.0, 0.9, key="bt_ct_cosphi")
        kappa_ct = kappa_servicio(cond_ct, "PVC")
        e_ct = caida_tension_voltios(sistema_ct, i_ct, l_ct, s_ct, cosphi_ct, kappa_ct)
        tension_ref = 230.0 if sistema_ct == SISTEMA_MONO else 400.0
        st.markdown(f"→ ΔU = **{e_ct:.2f} V** = **{e_ct/tension_ref*100:.2f} %** (sobre {tension_ref:g} V)")

        st.divider()
        st.markdown("**Longitud máxima** (para no superar una ΔU objetivo)")
        c1, c2, c3, c4 = st.columns(4)
        sistema_lm = c1.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI], key="bt_lm_sistema")
        s_lm = c2.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=4, key="bt_lm_seccion")
        i_lm = c3.number_input("Intensidad (A)", min_value=0.1, value=20.0, key="bt_lm_i")
        du_lm = c4.number_input("ΔU máx. (%)", min_value=0.1, value=3.0, key="bt_lm_du")
        tension_lm = 230.0 if sistema_lm == SISTEMA_MONO else 400.0
        k_sist_lm = 2.0 if sistema_lm == SISTEMA_MONO else math.sqrt(3)
        kappa_lm = kappa_servicio("Cobre", "PVC")
        l_max = (du_lm / 100 * tension_lm) / (k_sist_lm * i_lm * (0.9 / (kappa_lm * s_lm)))
        st.markdown(f"→ Longitud máxima ≈ **{l_max:.1f} m** (cobre, PVC, cos φ=0,9 orientativo)")

        st.divider()
        st.markdown("**Sección por caída de tensión** / **Sección por corriente (Iz)**")
        c1, c2, c3, c4, c5 = st.columns(5)
        sistema_s = c1.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI], key="bt_s_sistema")
        i_s = c2.number_input("Intensidad (A)", min_value=0.1, value=20.0, key="bt_s_i")
        l_s = c3.number_input("Longitud (m)", min_value=0.1, value=20.0, key="bt_s_l")
        du_s = c4.number_input("ΔU máx. (%)", min_value=0.1, value=3.0, key="bt_s_du")
        metodo_s = c5.selectbox("Método", METODOS_DISPONIBLES, index=1, key="bt_s_metodo")
        aisl_s = "XLPE/EPR" if metodo_s in (METODO_D, METODO_F) else st.selectbox(
            "Aislamiento", ["PVC", "XLPE/EPR"], key="bt_s_aisl")
        tension_s = 230.0 if sistema_s == SISTEMA_MONO else 400.0
        kappa_s = kappa_servicio("Cobre", aisl_s)
        n_carg_s = 2 if sistema_s == SISTEMA_MONO else 3
        s_por_du, _, _ = seccion_por_caida_tension(sistema_s, i_s, l_s, tension_s, 0.9, kappa_s, du_s, "Cobre")
        s_por_iz, iz_val, _, _ = seccion_por_criterio_termico(i_s, metodo_s, aisl_s, "Cobre", n_carg_s, 1.0)
        cc1, cc2 = st.columns(2)
        cc1.metric("Sección mínima por ΔU", f"{s_por_du:g} mm²" if s_por_du else "—")
        cc2.metric("Sección mínima por Iz", f"{s_por_iz:g} mm²" if s_por_iz else "—",
                   help=f"Iz tabla = {iz_val:.1f} A" if iz_val else None)
        if s_por_du and s_por_iz:
            st.markdown(f"→ Sección final adoptada (la mayor de ambas): **{max(s_por_du, s_por_iz):g} mm²**")

        st.divider()
        st.markdown("**Resistencia de un conductor**")
        c1, c2, c3, c4 = st.columns(4)
        cond_r = c1.selectbox("Conductor", ["Cobre", "Aluminio"], key="bt_r_cond")
        s_r = c2.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=4, key="bt_r_seccion")
        l_r = c3.number_input("Longitud (m)", min_value=0.1, value=100.0, key="bt_r_l")
        temp_r = c4.number_input("Temperatura (°C)", value=20.0, key="bt_r_temp")
        kappa_20 = CONDUCTIVIDAD_20C[cond_r]
        alpha_r = COEF_TEMP_RESIST[cond_r]
        kappa_t = kappa_20 / (1 + alpha_r * (temp_r - 20))
        r_cond = l_r / (kappa_t * s_r)
        st.markdown(f"→ R = L/(κ·S) = {l_r:g}/({kappa_t:.2f}×{s_r:g}) = **{r_cond:.4f} Ω** (ida simple)")

        st.divider()
        st.markdown("**Conversión AWG ↔ mm²**")
        c1, c2, c3 = st.columns(3)
        modo_awg = c1.radio("Convertir", ["AWG → mm²", "mm² → AWG"], key="bt_awg_modo", horizontal=True)
        if modo_awg == "AWG → mm²":
            awg_in = c2.selectbox("Calibre AWG", AWG_DISPONIBLES, index=10, key="bt_awg_in")
            c3.metric("Sección equivalente", f"{awg_a_mm2(awg_in):.3f} mm²")
        else:
            mm2_in = c2.number_input("Sección (mm²)", min_value=0.01, value=2.5, key="bt_mm2_in")
            awg_cercano, mm2_exacto = mm2_a_awg_mas_cercano(mm2_in)
            c3.metric("AWG más cercano", f"{awg_cercano} AWG ({mm2_exacto:.3f} mm²)")

        st.divider()
        st.markdown("**Comparador de métodos de instalación** (misma sección, ¿cuánto cambia la Iz?)")
        cm1, cm2, cm3, cm4 = st.columns(4)
        s_cmp = cm1.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=4, key="bt_cmp_s")
        cond_cmp = cm2.selectbox("Conductor", ["Cobre", "Aluminio"], key="bt_cmp_cond")
        aisl_cmp = cm3.selectbox("Aislamiento", ["PVC", "XLPE/EPR"], key="bt_cmp_aisl")
        n_carg_cmp = cm4.selectbox("Conductores cargados", [2, 3], index=1, key="bt_cmp_ncarg")
        filas_metodos = []
        for metodo_cmp in METODOS_DISPONIBLES:
            iz_m = iz_tabla(s_cmp, metodo_cmp, aisl_cmp, cond_cmp, n_carg_cmp)
            filas_metodos.append({"Método": metodo_cmp.split(" — ")[0], "Descripción": metodo_cmp.split(" — ")[1],
                                  "Iz (A)": iz_m if iz_m is not None else None})
        df_metodos = pd.DataFrame(filas_metodos)
        import plotly.express as px
        df_validos = df_metodos.dropna(subset=["Iz (A)"])
        if not df_validos.empty:
            fig_metodos = px.bar(df_validos, x="Método", y="Iz (A)", color="Método", text="Iz (A)",
                                 color_discrete_sequence=["#3b82f6", "#e8a33d", "#22c55e", "#ef4444", "#a78bfa", "#06b6d4"])
            fig_metodos.update_traces(textposition="outside")
            fig_metodos.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=280, showlegend=False,
                                      paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                                      font=dict(color="#8b96a8"))
            st.plotly_chart(fig_metodos, width='stretch')
            iz_min_row = df_validos.loc[df_validos["Iz (A)"].idxmin()]
            iz_max_row = df_validos.loc[df_validos["Iz (A)"].idxmax()]
            diferencia_pct = (iz_max_row["Iz (A)"] / iz_min_row["Iz (A)"] - 1) * 100
            st.caption(f"Entre el método más restrictivo ({iz_min_row['Método']}, {iz_min_row['Iz (A)']:g} A) y "
                       f"el más favorable ({iz_max_row['Método']}, {iz_max_row['Iz (A)']:g} A) hay un "
                       f"**{diferencia_pct:.0f}%** de diferencia en la intensidad admisible para la misma "
                       f"sección de {s_cmp:g} mm² — el método de instalación importa tanto como la sección "
                       "a la hora de dimensionar.")
        st.dataframe(df_metodos, width='stretch', hide_index=True)

        st.divider()
        st.markdown("**Protección contra sobrecarga** (Ib ≤ In ≤ Iz)")
        c1, c2, c3 = st.columns(3)
        ib_p = c1.number_input("Ib, intensidad de empleo (A)", min_value=0.1, value=18.0, key="bt_p_ib")
        in_p = c2.selectbox("In, calibre del interruptor (A)", CALIBRES_MAGNETOTERMICO, index=2, key="bt_p_in")
        iz_p = c3.number_input("Iz, admisible del cable (A)", min_value=0.1, value=21.0, key="bt_p_iz")
        cumple1 = ib_p <= in_p
        cumple2 = in_p <= iz_p
        st.markdown(f"→ Ib≤In: {'✅' if cumple1 else '❌'}  ·  In≤Iz: {'✅' if cumple2 else '❌'}  ·  "
                    f"conjunto: {'✅ Cumple ITC-BT-22' if (cumple1 and cumple2) else '❌ No cumple'}")

        st.markdown("**Protección contra sobrecarga y cortocircuito** (añade verificación térmica)")
        c1, c2, c3, c4 = st.columns(4)
        s_pc = c1.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=4, key="bt_pc_s")
        icc_pc = c2.number_input("Icc en el punto (kA)", min_value=0.01, value=4.0, key="bt_pc_icc")
        t_pc = c3.number_input("Tiempo de actuación (s)", min_value=0.001, value=0.1, format="%.3f", key="bt_pc_t")
        aisl_pc = c4.selectbox("Aislamiento", ["PVC", "XLPE/EPR"], key="bt_pc_aisl")
        cumple_pc, smin_pc = verificar_cortocircuito(s_pc, icc_pc, t_pc, "Cobre", aisl_pc)
        st.markdown(f"→ S_mín térmica = **{smin_pc:g} mm²** frente a {s_pc:g} mm² adoptados: "
                    f"{'✅ Cumple' if cumple_pc else '❌ No cumple'}")

    # ---------------------------------------------------------------- CORTOCIRCUITO
    with st.expander("⚡ Corriente de cortocircuito"):
        st.markdown("**Con datos de la red de BT** (se conoce la Icc o la Scc de origen)")
        c1, c2, c3, c4, c5 = st.columns(5)
        u_cc1 = c1.number_input("Tensión (V)", min_value=100.0, value=400.0, key="bt_cc1_u")
        icc_origen = c2.number_input("Icc en el origen (kA)", min_value=0.1, value=10.0, key="bt_cc1_icc0")
        s_cc1 = c3.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=6, key="bt_cc1_s")
        l_cc1 = c4.number_input("Longitud (m)", min_value=0.1, value=30.0, key="bt_cc1_l")
        n_carg_cc1 = c5.selectbox("Conductores cargados", [2, 3], index=1, key="bt_cc1_n")
        z_red = u_cc1 / (math.sqrt(3) * icc_origen * 1000.0)
        r_lin1, x_lin1 = impedancia_linea("Cobre", "PVC", s_cc1, l_cc1, n_carg_cc1)
        icc_final1 = icc_trifasico(u_cc1, z_red + r_lin1, x_lin1)
        st.markdown(f"→ Z_red = {z_red*1000:.2f} mΩ · Z_línea ≈ {r_lin1*1000:.2f} mΩ  →  "
                    f"**Icc en el punto final = {icc_final1/1000:.2f} kA**")

        st.divider()
        st.markdown("**Sin datos de la red de BT** (se asume red de potencia infinita en origen)")
        c1, c2, c3, c4 = st.columns(4)
        u_cc2 = c1.number_input("Tensión (V)", min_value=100.0, value=400.0, key="bt_cc2_u")
        s_cc2 = c2.selectbox("Sección (mm²)", SECCIONES_NORMALIZADAS, index=6, key="bt_cc2_s")
        l_cc2 = c3.number_input("Longitud (m)", min_value=0.1, value=30.0, key="bt_cc2_l")
        n_carg_cc2 = c4.selectbox("Conductores cargados", [2, 3], index=1, key="bt_cc2_n")
        r_lin2, x_lin2 = impedancia_linea("Cobre", "PVC", s_cc2, l_cc2, n_carg_cc2)
        icc_final2 = icc_trifasico(u_cc2, r_lin2, x_lin2)
        st.markdown(f"→ **Icc estimada = {icc_final2/1000:.2f} kA** (valor máximo teórico, orientativo al "
                    "no considerar la impedancia de la red aguas arriba)")

    # ---------------------------------------------------------------- TIERRAS
    with st.expander("🌍 Resistencia de tierras"):
        st.markdown("**Electrodos de tierra** (ITC-BT-18, Tabla 5)")
        c1, c2 = st.columns(2)
        with c1:
            claves_rho = list(RESISTIVIDAD_TERRENOS_REF.keys())
            terreno_sel = st.selectbox("Naturaleza del terreno", claves_rho,
                                        index=claves_rho.index("Terrenos cultivables y fértiles, "
                                                                "terraplenes compactos y húmedos"))
            rho_defecto = RESISTIVIDAD_TERRENOS_REF[terreno_sel] or 100.0
            rho_t = st.number_input("Resistividad ρ (Ω·m)", min_value=1.0, value=float(rho_defecto))
        with c2:
            tipo_electrodo = st.radio("Tipo de electrodo", ["Pica vertical", "Placa enterrada",
                                                             "Conductor enterrado horizontal"])
        if tipo_electrodo == "Pica vertical":
            long_pica = st.number_input("Longitud de la pica (m)", min_value=0.5, value=2.0, step=0.5)
            r_electrodo = resistencia_electrodo_pica(rho_t, long_pica)
            st.markdown(f"→ R = ρ/L = {rho_t:g}/{long_pica:g} = **{r_electrodo:.2f} Ω**")
        elif tipo_electrodo == "Placa enterrada":
            perimetro_placa = st.number_input("Perímetro de la placa (m)", min_value=0.5, value=2.0, step=0.1)
            r_electrodo = resistencia_electrodo_placa(rho_t, perimetro_placa)
            st.markdown(f"→ R = 0,8·ρ/P = 0,8×{rho_t:g}/{perimetro_placa:g} = **{r_electrodo:.2f} Ω**")
        else:
            long_cond = st.number_input("Longitud del conductor enterrado (m)", min_value=1.0, value=20.0)
            r_electrodo = resistencia_electrodo_conductor(rho_t, long_cond)
            st.markdown(f"→ R = 2·ρ/L = 2×{rho_t:g}/{long_cond:g} = **{r_electrodo:.2f} Ω**")

        st.divider()
        st.markdown("**Picas de tierra en paralelo**")
        c1, c2 = st.columns(2)
        r_una = c1.number_input("Resistencia de una pica (Ω)", min_value=0.1, value=75.0, key="bt_picas_r1")
        n_picas = c2.number_input("Número de picas (bien separadas, ≥2×L)", min_value=1, value=3, key="bt_picas_n")
        r_total_picas = resistencia_picas_paralelo(r_una, int(n_picas))
        st.markdown(f"→ R_total ≈ R/n = {r_una:g}/{n_picas:g} = **{r_total_picas:.2f} Ω** "
                    "(aproximación para picas bien separadas)")

    # ---------------------------------------------------------------- FOTOVOLTAICA RAPIDA
    with st.expander("☀️ Fotovoltaicas — cálculo rápido"):
        st.caption("Para el dimensionado completo (strings, inversor, cableado CC/CA, batería...) usa la "
                   "pestaña Fotovoltaica. Aquí solo estimaciones puntuales rápidas.")
        st.markdown("**Energía máxima diaria** (consumo)")
        c1, c2, c3 = st.columns(3)
        pot_ap = c1.number_input("Potencia del aparato (W)", min_value=1.0, value=100.0, key="bt_fv_pot")
        horas_ap = c2.number_input("Horas de uso al día", min_value=0.1, value=4.0, key="bt_fv_horas")
        n_ap = c3.number_input("Número de aparatos iguales", min_value=1, value=1, key="bt_fv_n")
        energia_dia = pot_ap * horas_ap * n_ap
        st.markdown(f"→ Energía diaria = **{energia_dia:.0f} Wh/día** ({energia_dia/1000:.2f} kWh/día)")

        st.divider()
        st.markdown("**Número de paneles solares**")
        c1, c2, c3, c4 = st.columns(4)
        energia_nec = c1.number_input("Energía necesaria (Wh/día)", min_value=1.0, value=3000.0, key="bt_fv_en")
        pot_panel = c2.number_input("Potencia del panel (Wp)", min_value=10.0, value=450.0, key="bt_fv_pp")
        hsp_fv = c3.number_input("HSP (h/día)", min_value=1.0, value=4.5, key="bt_fv_hsp")
        pr_fv = c4.number_input("PR", min_value=0.5, max_value=0.95, value=0.80, key="bt_fv_pr")
        n_paneles_bt = math.ceil(energia_nec / (pot_panel * hsp_fv * pr_fv))
        st.markdown(f"→ Nº de paneles ≈ **{n_paneles_bt}** paneles de {pot_panel:g} Wp")

        st.divider()
        st.markdown("**Baterías solares**")
        c1, c2, c3, c4 = st.columns(4)
        energia_bat = c1.number_input("Energía a acumular (Wh/día)", min_value=1.0, value=3000.0, key="bt_bat_e")
        dias_aut = c2.number_input("Días de autonomía", min_value=0.5, value=1.0, key="bt_bat_dias")
        v_bat = c3.number_input("Tensión del banco (V)", min_value=1.0, value=48.0, key="bt_bat_v")
        pd_bat = c4.number_input("Profundidad de descarga (%)", min_value=10.0, max_value=100.0, value=80.0,
                                  key="bt_bat_pd")
        capacidad_ah = (energia_bat * dias_aut) / (v_bat * (pd_bat / 100.0))
        st.markdown(f"→ Capacidad necesaria ≈ **{capacidad_ah:.0f} Ah** a {v_bat:g} V "
                    f"({capacidad_ah*v_bat/1000:.2f} kWh)")

    # ---------------------------------------------------------------- ELECTRICAS GENERALES
    with st.expander("📐 Eléctricas generales"):
        st.markdown("**Consumo eléctrico**")
        c1, c2, c3 = st.columns(3)
        pot_ce = c1.number_input("Potencia (W)", min_value=1.0, value=1500.0, key="bt_ce_p")
        horas_ce = c2.number_input("Horas de uso al día", min_value=0.1, value=2.0, key="bt_ce_h")
        precio_ce = c3.number_input("Precio kWh (€)", min_value=0.01, value=0.18, key="bt_ce_precio")
        kwh_dia = pot_ce / 1000 * horas_ce
        st.markdown(f"→ {kwh_dia:.3f} kWh/día · {kwh_dia*30:.1f} kWh/mes · "
                    f"**{_fmt_eur(kwh_dia*30*precio_ce)}/mes**")

        st.divider()
        st.markdown("**Ley de Ohm**")
        c1, c2 = st.columns(2)
        v_ohm = c1.number_input("Tensión V (V)", min_value=0.0, value=230.0, key="bt_ohm_v")
        i_ohm = c2.number_input("Intensidad I (A)", min_value=0.001, value=2.0, key="bt_ohm_i")
        st.markdown(f"→ Con V={v_ohm:g} V e I={i_ohm:g} A: R = V/I = **{v_ohm/max(i_ohm,1e-9):.2f} Ω**")

        st.divider()
        st.markdown("**Potencia eléctrica**")
        c1, c2, c3 = st.columns(3)
        v_pot = c1.number_input("Tensión (V)", min_value=0.0, value=230.0, key="bt_pot_v")
        i_pot = c2.number_input("Intensidad (A)", min_value=0.0, value=5.0, key="bt_pot_i")
        cosphi_pot = c3.slider("cos φ (1 si es CC o carga resistiva)", 0.1, 1.0, 1.0, key="bt_pot_cosphi")
        p_pot = v_pot * i_pot * cosphi_pot
        st.markdown(f"→ P = V·I·cos φ = **{p_pot:.1f} W**")

        st.divider()
        st.markdown("**Código de colores de resistencias** (4 bandas)")
        c1, c2, c3, c4 = st.columns(4)
        col1 = c1.selectbox("1ª banda", list(COLORES_DIGITO.keys()), index=1, key="bt_col1")
        col2 = c2.selectbox("2ª banda", list(COLORES_DIGITO.keys()), index=0, key="bt_col2")
        col3 = c3.selectbox("Multiplicador", list(COLORES_MULTIPLICADOR.keys()), index=2, key="bt_col3")
        col4 = c4.selectbox("Tolerancia", list(COLORES_TOLERANCIA.keys()), index=6, key="bt_col4")
        valor_r, tol_r = valor_resistencia_4_bandas(col1, col2, col3, col4)
        st.markdown(f"→ Valor = **{valor_r:g} Ω** ({tol_r})")

        st.divider()
        st.markdown("**Resistencias en paralelo**")
        texto_r = st.text_input("Valores separados por coma (Ω)", "100, 220, 330", key="bt_rp_texto")
        try:
            valores_r = [float(x.strip()) for x in texto_r.split(",") if x.strip()]
            req = 1.0 / sum(1.0 / v for v in valores_r) if valores_r else 0.0
            st.markdown(f"→ R_eq = **{req:.2f} Ω**")
        except (ValueError, ZeroDivisionError):
            st.warning("Introduce valores numéricos separados por comas.")

        st.divider()
        st.markdown("**Factor de potencia y batería de condensadores**")
        c1, c2, c3 = st.columns(3)
        p_fp = c1.number_input("Potencia activa P (kW)", min_value=0.1, value=10.0, key="bt_fp_p")
        cosphi_actual = c2.number_input("cos φ actual", min_value=0.1, max_value=1.0, value=0.75, key="bt_fp_actual")
        cosphi_obj = c3.number_input("cos φ objetivo", min_value=0.1, max_value=1.0, value=0.95, key="bt_fp_obj")
        tan_actual = math.tan(math.acos(cosphi_actual))
        tan_obj = math.tan(math.acos(cosphi_obj))
        q_condensador = max(p_fp * (tan_actual - tan_obj), 0)
        escalones_kvar = [2.5, 5, 7.5, 10, 12.5, 15, 20, 25, 30, 40, 50, 60, 75, 100, 125, 150]
        equipo_sugerido = next((e for e in escalones_kvar if e >= q_condensador), escalones_kvar[-1])
        st.markdown(f"→ Potencia reactiva necesaria: **{q_condensador:.2f} kVAr** → equipo comercial "
                    f"recomendado: batería de **{equipo_sugerido:g} kVAr** (siguiente escalón normalizado "
                    "por encima del cálculo).")
        if cosphi_actual < 0.95:
            st.caption("⚠️ Con cos φ < 0,95 la compañía eléctrica puede aplicar penalización por energía "
                       "reactiva en la factura (según la normativa de peajes vigente); compensar hasta "
                       "0,95-0,98 suele amortizarse en poco tiempo.")

        st.divider()
        st.markdown("**Divisor de tensión**")
        c1, c2, c3 = st.columns(3)
        vin_dt = c1.number_input("Tensión de entrada Vin (V)", min_value=0.0, value=12.0, key="bt_dt_vin")
        r1_dt = c2.number_input("R1 (Ω)", min_value=0.01, value=1000.0, key="bt_dt_r1")
        r2_dt = c3.number_input("R2 (Ω)", min_value=0.01, value=1000.0, key="bt_dt_r2")
        vout_dt = vin_dt * r2_dt / (r1_dt + r2_dt)
        st.markdown(f"→ Vout = Vin·R2/(R1+R2) = **{vout_dt:.3f} V**")

    # ---------------------------------------------------------------- SELECTIVIDAD
    with st.expander("🔀 Selectividad de protecciones"):
        st.markdown("**Selectividad entre magnetotérmicos en serie**")
        st.caption("Comprobación orientativa (amperimétrica + rango del disparo magnético según curva). "
                   "La selectividad garantizada al 100% en toda la gama de sobrecorriente solo la certifican "
                   "las tablas de selectividad del fabricante para cada pareja concreta de aparatos.")
        c1, c2, c3, c4 = st.columns(4)
        in_arriba = c1.selectbox("In aguas arriba (A)", CALIBRES_MAGNETOTERMICO, index=6, key="bt_sel_in_arriba")
        curva_arriba = c2.selectbox("Curva aguas arriba", ["B", "C", "D"], index=1, key="bt_sel_curva_arriba")
        in_abajo = c3.selectbox("In aguas abajo (A)", CALIBRES_MAGNETOTERMICO, index=2, key="bt_sel_in_abajo")
        curva_abajo = c4.selectbox("Curva aguas abajo", ["B", "C", "D"], index=1, key="bt_sel_curva_abajo")
        icc_sel = st.number_input("Icc prevista en el punto de instalación del de aguas abajo (kA) — opcional",
                                   min_value=0.0, value=0.0, step=0.1, key="bt_sel_icc")

        ratio_sel = in_arriba / in_abajo
        mult_min_arriba, _ = CURVA_MAGNETOTERMICO_RANGOS[curva_arriba]
        _, mult_max_abajo = CURVA_MAGNETOTERMICO_RANGOS[curva_abajo]
        im_min_arriba_a = in_arriba * mult_min_arriba
        im_max_abajo_a = in_abajo * mult_max_abajo
        st.markdown(f"→ Relación de calibres In arriba/In abajo = **{ratio_sel:.2f}**")
        if ratio_sel >= 2:
            st.success("✅ Ratio ≥ 2: selectividad amperimétrica probable en la zona de sobrecarga.")
        elif ratio_sel >= 1.6:
            st.warning("⚠️ Ratio entre 1,6 y 2: selectividad parcial únicamente (regla práctica orientativa).")
        else:
            st.error("❌ Ratio < 1,6: selectividad no garantizada ni en sobrecarga — revisa los calibres.")
        st.markdown(f"→ Umbral de disparo magnético mínimo de aguas arriba: **{im_min_arriba_a:.0f} A** "
                    f"({mult_min_arriba}×{in_arriba} A, curva {curva_arriba}) · disparo magnético máximo de "
                    f"aguas abajo: **{im_max_abajo_a:.0f} A** ({mult_max_abajo}×{in_abajo} A, curva "
                    f"{curva_abajo}). Si el máximo de abajo no alcanza el mínimo de arriba, el de abajo "
                    "siempre dispara primero también en la zona magnética.")
        if icc_sel > 0:
            icc_sel_a = icc_sel * 1000
            if icc_sel_a <= im_min_arriba_a:
                st.success(f"✅ Con Icc={icc_sel:g} kA en ese punto, no se alcanza el umbral magnético del de "
                           "aguas arriba: selectividad total también frente a cortocircuito.")
            else:
                st.error(f"❌ Con Icc={icc_sel:g} kA se supera el umbral magnético del de aguas arriba "
                         f"({im_min_arriba_a/1000:.2f} kA): éste podría dispararse también — selectividad "
                         f"limitada hasta {im_min_arriba_a/1000:.2f} kA.")

        st.divider()
        st.markdown("**Selectividad entre interruptores diferenciales en cascada**")
        c1, c2, c3 = st.columns(3)
        idn_arriba = c1.selectbox("I∆n aguas arriba (mA)", [30, 100, 300, 500, 1000], index=2,
                                   key="bt_seldif_arriba")
        tipo_arriba_dif = c2.selectbox("Tipo aguas arriba", ["Instantáneo", "Selectivo (S)"],
                                        key="bt_seldif_tipo")
        idn_abajo = c3.selectbox("I∆n aguas abajo (mA)", [10, 30, 100, 300], index=1, key="bt_seldif_abajo")
        ratio_dif = idn_arriba / idn_abajo
        st.markdown(f"→ Relación I∆n arriba/abajo = **{ratio_dif:.1f}**")
        if tipo_arriba_dif == "Selectivo (S)":
            st.success("✅ El de aguas arriba es de tipo Selectivo (S): retardo intencionado de disparo → "
                       "selectividad cronométrica garantizada frente al de aguas abajo, "
                       f"{'y además cumple' if ratio_dif >= 3 else 'aunque conviene además'} la relación "
                       "amperimétrica ≥ 3.")
        elif ratio_dif >= 3:
            st.warning("⚠️ Ratio ≥ 3 pero el de aguas arriba NO es de tipo Selectivo (S): selectividad "
                       "amperimétrica probable, pero no garantizada por norma sin el retardo intencionado.")
        else:
            st.error("❌ Sin tipo Selectivo (S) arriba y con ratio < 3: selectividad no garantizada — un "
                     "defecto aguas abajo podría disparar también el diferencial de cabecera.")

    # ---------------------------------------------------------------- MOTORES
    with st.expander("⚙️ Motores: arranque y arrancador estrella-triángulo"):
        st.caption("Dimensionado orientativo del arranque de un motor asíncrono trifásico o monofásico. "
                   "Para el circuito completo (sección por Ib de servicio) usa la Calculadora principal con "
                   "'Es motor' activado.")
        c1, c2, c3, c4 = st.columns(4)
        p_mot = c1.number_input("Potencia motor (kW)", min_value=0.1, value=7.5, key="bt_mot_p")
        sistema_mot = c2.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI], index=1, key="bt_mot_sistema")
        v_mot = c3.number_input("Tensión (V)", min_value=1.0, value=400.0, key="bt_mot_v")
        cosphi_mot = c4.number_input("cos φ nominal", min_value=0.1, max_value=1.0, value=0.85, key="bt_mot_cosphi")
        c1, c2, c3 = st.columns(3)
        rend_mot = c1.number_input("Rendimiento η (%)", min_value=1.0, max_value=100.0, value=88.0,
                                    key="bt_mot_rend") / 100
        ratio_ia_in = c2.number_input("Relación Ia/In en arranque directo", min_value=1.0, value=7.0, step=0.5,
                                      key="bt_mot_ratio", help="Típica en motores asíncronos de jaula de "
                                      "ardilla: 5-8 veces In. Consulta la placa de características si la "
                                      "tienes.")
        tipo_arranque = c3.selectbox("Sistema de arranque", ["Directo", "Estrella-Triángulo"],
                                      key="bt_mot_tipo_arranque")
        c1, c2 = st.columns(2)
        l_mot = c1.number_input("Longitud del circuito (m)", min_value=0.1, value=25.0, key="bt_mot_l")
        s_mot = c2.selectbox("Sección adoptada (mm²)", SECCIONES_NORMALIZADAS, index=5, key="bt_mot_s")

        if sistema_mot == SISTEMA_TRI:
            in_mot = p_mot * 1000 / (math.sqrt(3) * v_mot * cosphi_mot * rend_mot)
        else:
            in_mot = p_mot * 1000 / (v_mot * cosphi_mot * rend_mot)
        ia_directo = in_mot * ratio_ia_in
        ia_efectiva = ia_directo / 3 if tipo_arranque == "Estrella-Triángulo" else ia_directo

        limite_ratio = next(r for pmin, pmax, r in RATIO_IA_IN_MAX_MOTOR if pmin <= p_mot < pmax)
        st.markdown(f"→ In nominal = **{in_mot:.2f} A** · Ia en arranque directo = **{ia_directo:.1f} A** "
                    f"({ratio_ia_in:g}×In)")
        if tipo_arranque == "Estrella-Triángulo":
            st.markdown(f"→ Con arrancador estrella-triángulo: Ia efectiva ≈ Ia/3 = **{ia_efectiva:.1f} A**, "
                        "par de arranque también se reduce a ~1/3 del par en arranque directo.")
        if ratio_ia_in > limite_ratio:
            st.warning(f"⚠️ La relación de arranque directo ({ratio_ia_in:g}) supera el máximo admitido por "
                       f"la ITC-BT-47 para {p_mot:g} kW ({limite_ratio:g}): "
                       f"{'correcto usar un arrancador que la reduzca, como el estrella-triángulo elegido' if tipo_arranque != 'Directo' else 'se requiere un sistema de arranque que la reduzca (estrella-triángulo u otro)'}.")
        else:
            st.success(f"✅ Relación de arranque ({ratio_ia_in:g}) dentro del límite ITC-BT-47 para {p_mot:g} "
                       f"kW ({limite_ratio:g}): el arranque directo estaría permitido.")

        kappa_mot = kappa_servicio("Cobre", "PVC")
        du_arranque_v = caida_tension_voltios(sistema_mot, ia_efectiva, l_mot, s_mot, 0.5, kappa_mot)
        tension_ref_mot = 230.0 if sistema_mot == SISTEMA_MONO else 400.0
        st.markdown(f"→ Caída de tensión durante el arranque (cos φ≈0,5 típico en arranque): "
                    f"**{du_arranque_v:.2f} V = {du_arranque_v/tension_ref_mot*100:.2f} %** con la sección "
                    f"adoptada de {s_mot:g} mm². Una caída elevada en el arranque puede impedir que el motor "
                    "llegue a girar o afectar a otros receptores del mismo cuadro (parpadeo de alumbrado).")

    # ---------------------------------------------------------------- ILUMINACION
    with st.expander("💡 Iluminación: luminotecnia y emergencia"):
        st.markdown("**Nº de luminarias por el método de los lúmenes**")
        c1, c2 = st.columns(2)
        local_lum = c1.selectbox("Tipo de local", list(ILUMINANCIA_POR_LOCAL.keys()), key="bt_lum_local")
        e_lum = c2.number_input("Iluminancia media exigida E (lux)",
                                 value=float(ILUMINANCIA_POR_LOCAL[local_lum]), min_value=1.0, key="bt_lum_e")
        c1, c2, c3, c4 = st.columns(4)
        s_lum = c1.number_input("Superficie (m²)", min_value=1.0, value=30.0, key="bt_lum_s")
        flujo_lum = c2.number_input("Flujo por luminaria (lm)", min_value=1.0, value=4000.0, key="bt_lum_flujo")
        cu_lum = c3.number_input("Coef. de utilización Cu", min_value=0.1, max_value=1.0, value=0.5,
                                  key="bt_lum_cu", help="Depende del índice del local (dimensiones/altura) y "
                                  "de las reflectancias de techo/paredes/suelo. 0,4-0,6 es un rango habitual "
                                  "en locales de tamaño medio con acabados claros.")
        cm_lum = c4.number_input("Coef. de mantenimiento Cm", min_value=0.1, max_value=1.0, value=0.8,
                                  key="bt_lum_cm", help="Depreciación del flujo por suciedad/envejecimiento; "
                                  "0,8 es un valor típico con mantenimiento periódico normal.")
        n_lum = math.ceil((e_lum * s_lum) / (flujo_lum * cu_lum * cm_lum)) if flujo_lum * cu_lum * cm_lum > 0 else 0
        st.markdown(f"→ N = (E·S)/(Φ·Cu·Cm) = **{n_lum} luminarias** de {flujo_lum:g} lm para "
                    f"{e_lum:g} lux en {s_lum:g} m².")
        st.caption("Cálculo simplificado por el método de los lúmenes (nivel medio en el plano de trabajo). "
                   "No sustituye a un estudio luminotécnico con software fotométrico para locales exigentes "
                   "o con geometría compleja.")

        st.divider()
        st.markdown("**Alumbrado de emergencia** (orden de magnitud — ITC-BT-28 / CTE DB-SI 4)")
        c1, c2 = st.columns(2)
        s_emerg = c1.number_input("Superficie o longitud de recorrido a cubrir (m²)", min_value=1.0,
                                   value=80.0, key="bt_emerg_s")
        flujo_emerg = c2.selectbox("Flujo de la luminaria autónoma (lm)", [90, 150, 240, 315, 400], index=2,
                                    key="bt_emerg_flujo")
        area_cobertura = flujo_emerg / 5.0  # aproximación orientativa: ~1 lux de servicio sobre el área cubierta
        n_emerg = math.ceil(s_emerg / area_cobertura) if area_cobertura > 0 else 0
        st.markdown(f"→ Área de cobertura orientativa por luminaria ≈ **{area_cobertura:.0f} m²** → "
                    f"**{n_emerg} luminarias** como orden de magnitud.")
        st.caption("⚠️ Muy simplificado: la disposición real NO depende solo de la superficie, sino de "
                   "colocar una luminaria en cada puerta de salida, cambio de dirección, escalera, cuadro "
                   "eléctrico y equipo de seguridad, con una separación máxima entre ellas (recorridos "
                   "≤ 25 m) — verifica siempre el trazado sobre plano. Autonomía mínima exigida: 1 hora. "
                   "Iluminancia mínima: 1 lux en el suelo del recorrido, 5 lux en cuadros y equipos de "
                   "seguridad.")

    # ---------------------------------------------------------------- VEHICULO ELECTRICO
    with st.expander("🚗 Punto de recarga de vehículo eléctrico (ITC-BT-52)"):
        c1, c2, c3 = st.columns(3)
        modo_carga = c1.selectbox("Modo de carga", ["Modo 2 — toma + cable con protección", "Modo 3 — wallbox dedicado"],
                                   key="bt_ve_modo")
        sistema_ve = c2.selectbox("Sistema", [SISTEMA_MONO, SISTEMA_TRI], key="bt_ve_sistema")
        p_ve = c3.number_input("Potencia del cargador (kW)", min_value=1.0, value=7.4 if sistema_ve == SISTEMA_MONO else 22.0,
                                key="bt_ve_p")
        c1, c2, c3 = st.columns(3)
        v_ve = c1.number_input("Tensión (V)", value=230.0 if sistema_ve == SISTEMA_MONO else 400.0, key="bt_ve_v")
        l_ve = c2.number_input("Longitud del circuito (m)", min_value=0.1, value=15.0, key="bt_ve_l")
        du_max_ve = c3.number_input("ΔU máx. recomendada (%)", min_value=0.5, value=3.0, key="bt_ve_du",
                                     help="El REBT admite hasta 5% en un circuito interior, pero para no "
                                     "perder potencia de carga se recomienda ser más exigente (≤3%).")

        in_ve = p_ve * 1000 / (v_ve * math.sqrt(3) if sistema_ve == SISTEMA_TRI else v_ve)  # cos phi ~ 1
        ib_calc_ve = 1.25 * in_ve  # ITC-BT-52: circuito dedicado, sin coeficiente de simultaneidad
        n_carg_ve = 2 if sistema_ve == SISTEMA_MONO else 3
        s_iz_ve, iz_ve, _, _ = seccion_por_criterio_termico(ib_calc_ve, METODO_B1, "PVC", "Cobre", n_carg_ve, 1.0)
        kappa_ve = kappa_servicio("Cobre", "PVC")
        s_du_ve, _, _ = seccion_por_caida_tension(sistema_ve, ib_calc_ve, l_ve, v_ve, 1.0, kappa_ve, du_max_ve, "Cobre")
        s_final_ve = max(s_iz_ve or 0, s_du_ve or 0) or None
        calibre_ve = calibre_magnetotermico_sugerido(ib_calc_ve)

        st.markdown(f"→ In del cargador ≈ **{in_ve:.2f} A** · Ib de cálculo (125%, ITC-BT-52) = "
                    f"**{ib_calc_ve:.2f} A**")
        cve1, cve2, cve3 = st.columns(3)
        cve1.metric("Sección mínima", f"{s_final_ve:g} mm²" if s_final_ve else "—")
        cve2.metric("Interruptor automático", f"{calibre_ve} A, curva C")
        cve3.metric("Diferencial mínimo", "Tipo A" if "Modo 2" in modo_carga else "Tipo A o B")
        st.caption("El circuito debe ser **dedicado** (sin compartir con otros usos) y llevar protección "
                   "diferencial de, como mínimo, tipo A; si el cargador no garantiza por diseño la ausencia "
                   "de componente continua de defecto, se exige tipo B — dato que debe indicar el fabricante "
                   "del cargador (igual criterio que en instalaciones fotovoltaicas, ITC-BT-52 / ITC-BT-40).")

    # ---------------------------------------------------------------- POTENCIA CONTRATADA Y RESPALDO
    with st.expander("🔋 Potencia contratada, grupo electrógeno y SAI"):
        st.markdown("**Potencia contratada óptima**")
        c1, c2, c3 = st.columns(3)
        p_max_obs = c1.number_input("Potencia máxima demandada observada/estimada (kW)", min_value=0.1,
                                     value=8.5, key="bt_pc_pmax")
        margen_pc = c2.number_input("Margen de seguridad (%)", min_value=0.0, value=10.0, key="bt_pc_margen")
        peaje_pc = c3.number_input("Peaje de potencia orientativo (€/kW y año)", min_value=0.0, value=38.0,
                                    key="bt_pc_peaje", help="Suma de peaje de acceso + cargos de potencia; "
                                    "varía según la comunidad autónoma y la comercializadora — ajusta este "
                                    "valor a tu caso real.")
        p_optima = p_max_obs * (1 + margen_pc / 100)
        st.markdown(f"→ Potencia a contratar recomendada ≈ **{p_optima:.2f} kW**")
        cp1, cp2, cp3 = st.columns(3)
        for col, factor, etiqueta in zip([cp1, cp2, cp3], [0.9, 1.0, 1.15], ["Ajustada (-10%)", "Recomendada", "Con holgura (+15%)"]):
            p_test = p_optima * factor
            coste_fijo = p_test * peaje_pc
            with col:
                st.metric(etiqueta, f"{p_test:.2f} kW", f"{_fmt_eur(coste_fijo)}/año término fijo")
        st.caption("Contratar por debajo de la potencia realmente demandada expone a penalización por "
                   "excesos de potencia en la factura; contratar de más solo encarece el término fijo sin "
                   "beneficio. El margen de seguridad amortigua picos puntuales no observados en la muestra.")

        st.divider()
        st.markdown("**Grupo electrógeno de respaldo**")
        c1, c2, c3 = st.columns(3)
        p_criticas_ge = c1.number_input("Suma de cargas críticas (kW)", min_value=0.1, value=15.0,
                                         key="bt_ge_p")
        cosphi_ge = c2.number_input("cos φ medio de las cargas", min_value=0.1, max_value=1.0, value=0.85,
                                     key="bt_ge_cosphi")
        factor_arranque_ge = c3.number_input("Factor por arranque de motores", min_value=1.0, value=1.25,
                                              step=0.05, key="bt_ge_factor", help="Sobredimensionado para "
                                              "cubrir la punta de corriente al arrancar motores u otras "
                                              "cargas con arranque brusco; 1,25-1,5 es habitual si hay "
                                              "motores en las cargas críticas.")
        s_ge = (p_criticas_ge / cosphi_ge) * factor_arranque_ge
        escalones_kva = [5, 8, 10, 15, 20, 25, 30, 40, 50, 65, 80, 100, 125, 150, 200, 250]
        ge_sugerido = next((e for e in escalones_kva if e >= s_ge), escalones_kva[-1])
        st.markdown(f"→ Potencia aparente necesaria ≈ **{s_ge:.1f} kVA** → grupo comercial recomendado: "
                    f"**{ge_sugerido:g} kVA** (siguiente escalón normalizado).")

        st.divider()
        st.markdown("**SAI / UPS de respaldo**")
        c1, c2, c3, c4 = st.columns(4)
        p_criticas_sai = c1.number_input("Carga crítica (kW)", min_value=0.05, value=2.0, key="bt_sai_p")
        autonomia_sai = c2.number_input("Autonomía deseada (min)", min_value=1.0, value=15.0, key="bt_sai_t")
        v_bateria_sai = c3.number_input("Tensión del banco de baterías (V)", min_value=1.0, value=48.0,
                                         key="bt_sai_v")
        eficiencia_sai = c4.number_input("Eficiencia del SAI (%)", min_value=50.0, max_value=100.0, value=90.0,
                                          key="bt_sai_eff") / 100
        energia_sai_wh = p_criticas_sai * 1000 * (autonomia_sai / 60) / eficiencia_sai
        capacidad_sai_ah = energia_sai_wh / v_bateria_sai
        st.markdown(f"→ Energía necesaria ≈ **{energia_sai_wh:.0f} Wh** → capacidad de batería ≈ "
                    f"**{capacidad_sai_ah:.1f} Ah** a {v_bateria_sai:g} V.")

    # ---------------------------------------------------------------- TARIFAS Y EFICIENCIA ENERGETICA
    with st.expander("💡 Tarifas eléctricas y eficiencia energética"):
        st.markdown("**Comparador de tarifas de acceso (2.0TD / 3.0TD)**")
        st.caption("⚠️ Modelo simplificado a 3 periodos (punta/llano/valle) para orientar la decisión — la "
                   "3.0TD real tiene 6 periodos estacionales, y los peajes/cargos son precios regulados que "
                   "cambian; ajusta los valores por defecto a la tarifa real de tu comercializadora.")
        ct1, ct2 = st.columns(2)
        with ct1:
            potencia_contratada_cmp = st.number_input("Potencia contratada (kW)", min_value=0.1, value=10.0,
                                                        step=0.5, key="bt_tar_potencia")
            tarifa_aplicable = "2.0TD (≤15 kW)" if potencia_contratada_cmp <= 15 else "3.0TD (>15 kW)"
            st.info(f"Tarifa de acceso aplicable: **{tarifa_aplicable}**")
        with ct2:
            consumo_anual_cmp = st.number_input("Consumo anual (kWh)", min_value=1.0, value=4000.0, step=100.0,
                                                 key="bt_tar_consumo")
        ct3, ct4, ct5 = st.columns(3)
        pct_punta = ct3.number_input("% consumo en punta", min_value=0.0, max_value=100.0, value=25.0,
                                      key="bt_tar_pct_punta")
        pct_llano = ct4.number_input("% consumo en llano", min_value=0.0, max_value=100.0, value=45.0,
                                      key="bt_tar_pct_llano")
        pct_valle = ct5.number_input("% consumo en valle", min_value=0.0, max_value=100.0, value=30.0,
                                      key="bt_tar_pct_valle")
        suma_pct = pct_punta + pct_llano + pct_valle
        if abs(suma_pct - 100) > 0.5:
            st.warning(f"⚠️ Los 3 porcentajes suman {suma_pct:.0f}%, deberían sumar 100%.")
        st.markdown("**Precios orientativos** (edítalos según tu tarifa real)")
        pt1, pt2, pt3, pt4 = st.columns(4)
        precio_potencia_ano = pt1.number_input("Precio potencia (€/kW·año)", min_value=0.0,
                                                value=38.0 if potencia_contratada_cmp <= 15 else 29.0,
                                                key="bt_tar_precio_pot")
        precio_punta = pt2.number_input("Precio energía punta (€/kWh)", min_value=0.0, value=0.22, step=0.01,
                                         key="bt_tar_precio_punta")
        precio_llano = pt3.number_input("Precio energía llano (€/kWh)", min_value=0.0, value=0.16, step=0.01,
                                         key="bt_tar_precio_llano")
        precio_valle = pt4.number_input("Precio energía valle (€/kWh)", min_value=0.0, value=0.10, step=0.01,
                                         key="bt_tar_precio_valle")
        incluir_impuestos = st.checkbox("Incluir impuesto eléctrico (5,11269632%) e IVA (21%)", value=True,
                                         key="bt_tar_impuestos")

        termino_potencia_cmp = potencia_contratada_cmp * precio_potencia_ano
        termino_energia_cmp = (consumo_anual_cmp * pct_punta / 100 * precio_punta +
                               consumo_anual_cmp * pct_llano / 100 * precio_llano +
                               consumo_anual_cmp * pct_valle / 100 * precio_valle)
        subtotal_cmp = termino_potencia_cmp + termino_energia_cmp
        if incluir_impuestos:
            subtotal_cmp *= 1.0511269632
            total_cmp = subtotal_cmp * 1.21
        else:
            total_cmp = subtotal_cmp

        tc1, tc2, tc3 = st.columns(3)
        tc1.metric("Término de potencia (año)", _fmt_eur(termino_potencia_cmp))
        tc2.metric("Término de energía (año)", _fmt_eur(termino_energia_cmp))
        tc3.metric("Estimación factura anual", _fmt_eur(total_cmp),
                   f"≈ {_fmt_eur(total_cmp/12)}/mes")
        st.caption("Estimación orientativa para comparar escenarios (p. ej. bajar la potencia contratada o "
                   "desplazar consumo a horario valle) — no sustituye a la factura real de tu "
                   "comercializadora, que puede incluir otros conceptos (alquiler de contador, descuentos, "
                   "tarifas planas...).")

        st.divider()
        st.markdown("**Amortización de eficiencia energética** (p. ej. cambio a LED)")
        ef1, ef2, ef3 = st.columns(3)
        potencia_actual_ef = ef1.number_input("Potencia actual instalada (W)", min_value=1.0, value=1000.0,
                                               key="bt_ef_p_actual")
        potencia_nueva_ef = ef2.number_input("Potencia nueva (LED) (W)", min_value=1.0, value=350.0,
                                              key="bt_ef_p_nueva")
        horas_dia_ef = ef3.number_input("Horas de uso al día", min_value=0.1, value=8.0, key="bt_ef_horas")
        ef4, ef5 = st.columns(2)
        precio_kwh_ef = ef4.number_input("Precio kWh (€)", min_value=0.01, value=0.18, key="bt_ef_precio")
        inversion_ef = ef5.number_input("Inversión en el cambio (€)", min_value=0.0, value=600.0,
                                         key="bt_ef_inversion")

        kwh_ahorrado_dia = (potencia_actual_ef - potencia_nueva_ef) / 1000 * horas_dia_ef
        ahorro_anual_ef = kwh_ahorrado_dia * 365 * precio_kwh_ef
        payback_ef = inversion_ef / ahorro_anual_ef if ahorro_anual_ef > 0 else None

        ee1, ee2, ee3 = st.columns(3)
        ee1.metric("Reducción de potencia", f"{potencia_actual_ef - potencia_nueva_ef:.0f} W",
                   f"{(1 - potencia_nueva_ef/potencia_actual_ef)*100:.0f}%" if potencia_actual_ef else None)
        ee2.metric("Ahorro anual estimado", _fmt_eur(ahorro_anual_ef))
        ee3.metric("Retorno de la inversión", f"{payback_ef:.1f} años" if payback_ef else "—")
        if payback_ef and payback_ef < 5:
            st.success(f"✅ Con estos datos, el cambio se amortiza en menos de 5 años "
                       f"({payback_ef:.1f} años).")

    # ---------------------------------------------------------------- NEUTRO CON ARMONICOS
    with st.expander("➰ Sección de neutro con cargas no lineales (armónicos)"):
        st.caption("Justificación según la guía de armónicos de tercer orden (triplen) de la norma "
                   "UNE-HD 60364-5-52 / IEC 60364-5-523, para cargas no lineales (electrónica, iluminación "
                   "LED/fluorescente con balasto electrónico, variadores de frecuencia, informática...).")
        thd3 = st.slider("Contenido de 3er armónico en la corriente de fase, THD₃ (%)", 0, 60, 20,
                          key="bt_arm_thd3")
        if thd3 < 15:
            st.success(f"✅ THD₃ = {thd3}% < 15%: el efecto de los armónicos es despreciable. El neutro se "
                       "dimensiona con el criterio general (igual sección que la fase, sin coeficientes "
                       "adicionales).")
        elif thd3 <= 33:
            st.warning(f"⚠️ THD₃ = {thd3}% (entre 15% y 33%): debe aplicarse un factor de reducción de "
                       "**0,86** a la intensidad admisible del conductor de fase (equivale a sobredimensionar "
                       "la fase respecto al cálculo sin armónicos). El neutro debe tener, como mínimo, la "
                       "misma sección que la fase — no es reducible aunque la fase sea de sección grande.")
        else:
            st.error(f"❌ THD₃ = {thd3}% (> 33%): la corriente por el neutro puede superar a la de fase. El "
                     "neutro pasa a ser el conductor determinante del dimensionado: calcúlalo con el mismo "
                     "criterio térmico que una fase (su propia Iz), y aplica igualmente el factor de "
                     "reducción de **0,86** a la fase.")
        st.caption("La corriente de tercer armónico (y sus múltiplos, 3º-9º-15º...) se suma en fase en el "
                   "neutro de un sistema trifásico con neutro, en vez de cancelarse como ocurre con la "
                   "componente fundamental equilibrada — por eso puede llegar a superar la propia corriente "
                   "de fase con cargas no lineales importantes.")



def _render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-brand">
            <div class="logo">⚡</div>
            <div><div class="name">REBT Suite</div><div class="sub">Instalaciones eléctricas</div></div>
        </div>
        """, unsafe_allow_html=True)

        def nav_button(icono, nombre, ayuda=None):
            activo = st.session_state["pagina_actual"] == nombre
            if st.button(f"{icono}  {nombre}", key=f"nav_{nombre}", help=ayuda,
                         type="primary" if activo else "secondary", width='stretch'):
                st.session_state["pagina_actual"] = nombre
                st.rerun()

        AYUDA_NAV = {
            "Inicio": "Panel principal: resumen del proyecto, accesos rápidos y plantillas.",
            "Presentación cliente": "Vista simplificada del proyecto, sin jerga técnica, para explicarlo a alguien no técnico.",
            "Proyectos": "Guardar, abrir, duplicar y comparar tus proyectos.",
            "Estadísticas": "Gráficos del presupuesto y de la producción fotovoltaica.",
            "Calculadora": "Punto de partida: calcula la sección de un cable de baja tensión.",
            "Fórmulas": "Justificación paso a paso del cálculo de la Calculadora.",
            "Fotovoltaica": "Dimensiona una instalación solar de autoconsumo, de forma independiente.",
            "Cálculos BT": "Calculadoras sueltas de referencia rápida (Ohm, tierras, cortocircuito...).",
            "Presupuesto": "Mediciones y precios por capítulos, a partir de lo que hayas calculado.",
            "Documentación": "Genera la MTD, el Anexo de Cálculos y el Pliego de Condiciones en PDF.",
            "Tablas normativas": "Consulta las tablas de intensidades y factores de la Guía-BT-19.",
            "Metodología": "Qué criterios y normativa aplica cada cálculo de la app.",
            "Autoevaluación": "Test rápido con preguntas generadas a partir de las tablas normativas de la app.",
            "Configuración": "Tu nombre, logo y firma (para los PDF) y el tema de la app.",
            "Acerca de": "Qué es esta aplicación y sus limitaciones conocidas.",
        }

        st.markdown('<p class="nav-group-label">Principal</p>', unsafe_allow_html=True)
        nav_button("🏠", "Inicio", AYUDA_NAV["Inicio"])
        nav_button("📁", "Proyectos", AYUDA_NAV["Proyectos"])
        nav_button("📊", "Estadísticas", AYUDA_NAV["Estadísticas"])
        nav_button("🗣️", "Presentación cliente", AYUDA_NAV["Presentación cliente"])

        st.markdown('<p class="nav-group-label">Herramientas</p>', unsafe_allow_html=True)
        iconos_herr = {"Calculadora": "🔌", "Fórmulas": "🧮", "Fotovoltaica": "☀️",
                       "Presupuesto": "💰", "Documentación": "📄"}
        for nombre in PAGINAS_HERRAMIENTAS:
            nav_button(iconos_herr[nombre], nombre, AYUDA_NAV.get(nombre))

        st.markdown('<p class="nav-group-label">Consulta rápida</p>', unsafe_allow_html=True)
        nav_button("📐", "Cálculos BT", AYUDA_NAV["Cálculos BT"])

        st.markdown('<p class="nav-group-label">Normativa</p>', unsafe_allow_html=True)
        nav_button("📚", "Tablas normativas", AYUDA_NAV["Tablas normativas"])
        nav_button("📖", "Metodología", AYUDA_NAV["Metodología"])
        nav_button("🎓", "Autoevaluación", AYUDA_NAV["Autoevaluación"])

        st.markdown('<p class="nav-group-label">Sistema</p>', unsafe_allow_html=True)
        nav_button("⚙️", "Configuración", AYUDA_NAV["Configuración"])
        nav_button("ℹ️", "Acerca de", AYUDA_NAV["Acerca de"])

        st.markdown("<div style='margin-top:1.4rem;'></div>", unsafe_allow_html=True)
        st.caption(f"📌 {st.session_state['nombre_proyecto_actual']}")

        st.markdown("""
        <div class="sidebar-credit">
            <div class="sidebar-credit-name">Younesse Tikent Tifaoui</div>
            <div class="sidebar-credit-links">
                <a href="https://www.instagram.com/younes.tik/?hl=es" target="_blank" title="Instagram">📷</a>
                <a href="https://www.linkedin.com/in/younesse-tikent-tifaoui-5b9aa3241/" target="_blank" title="LinkedIn">💼</a>
                <a href="https://mail.google.com/mail/?view=cm&amp;fs=1&amp;to=younessetikenttifaoui@gmail.com" target="_blank" title="Gmail">✉️</a>
            </div>
        </div>
        """, unsafe_allow_html=True)


def _tarjeta_kpi(col, icono: str, color: str, valor: str, etiqueta: str):
    with col:
        st.markdown(f'''<div class="kpi-card">
            <div class="kpi-icon {color}">{icono}</div>
            <div class="kpi-value">{valor}</div>
            <div class="kpi-label">{etiqueta}</div>
        </div>''', unsafe_allow_html=True)


def _tarjeta_acceso_rapido(col, icono: str, titulo: str, subtitulo: str, pagina_destino: str):
    with col:
        st.markdown(f'''<div class="quick-card">
            <div class="qc-icon">{icono}</div>
            <div class="qc-title">{titulo}</div>
            <div class="qc-sub">{subtitulo}</div>
        </div>''', unsafe_allow_html=True)
        if st.button("Abrir →", key=f"qa_{pagina_destino}", width='stretch'):
            st.session_state["pagina_actual"] = pagina_destino
            st.rerun()


PLANTILLAS_CIRCUITO = {
    "🏠 Vivienda — circuito de alumbrado": dict(
        tipo_circuito="Instalación interior — Alumbrado", sistema=SISTEMA_MONO, potencia_kw=1.5,
        cos_phi=0.95, metodo=METODO_B1, aislamiento="PVC", longitud=15.0),
    "🏠 Vivienda — tomas de uso general": dict(
        tipo_circuito="Instalación interior — Otros usos / fuerza", sistema=SISTEMA_MONO, potencia_kw=3.45,
        cos_phi=0.90, metodo=METODO_B1, aislamiento="PVC", longitud=20.0),
    "🏭 Nave industrial — línea de fuerza": dict(
        tipo_circuito="Instalación interior — Otros usos / fuerza", sistema=SISTEMA_TRI, potencia_kw=15.0,
        cos_phi=0.85, metodo=METODO_F, aislamiento="XLPE/EPR", longitud=40.0),
    "🔋 Derivación individual estándar": dict(
        tipo_circuito="Derivación Individual — usuario único (sin LGA)", sistema=SISTEMA_MONO, potencia_kw=5.75,
        cos_phi=0.90, metodo=METODO_B1, aislamiento="PVC", longitud=12.0),
}

PLANTILLAS_FV = {
    "🏠 Vivienda unifamiliar — 5 kWp": dict(
        modo_dimensionado="Por potencia pico deseada (kWp)", potencia_pico_deseada=5.0,
        tipo_autoconsumo="Con excedentes acogido a compensación", potencia_inversor_kw=5.0,
        sistema_ca=SISTEMA_MONO, con_bateria=False),
    "🏭 Nave industrial — 20 kWp trifásica": dict(
        modo_dimensionado="Por potencia pico deseada (kWp)", potencia_pico_deseada=20.0,
        tipo_autoconsumo="Sin excedentes", potencia_inversor_kw=20.0,
        sistema_ca=SISTEMA_TRI, con_bateria=False),
    "🔋 Instalación aislada con batería": dict(
        modo_dimensionado="Por potencia pico deseada (kWp)", potencia_pico_deseada=3.0,
        tipo_autoconsumo="Instalación aislada (con batería)", potencia_inversor_kw=3.0,
        sistema_ca=SISTEMA_MONO, con_bateria=True),
    "☀️ Autoconsumo comercial — 10 kWp": dict(
        modo_dimensionado="Por potencia pico deseada (kWp)", potencia_pico_deseada=10.0,
        tipo_autoconsumo="Con excedentes no acogido a compensación", potencia_inversor_kw=10.0,
        sistema_ca=SISTEMA_TRI, con_bateria=False),
}

PLANTILLAS_PRESUPUESTO = {
    "🏠 Vivienda unifamiliar": ["CAPÍTULO I: DERIVACIÓN INDIVIDUAL", "CAPÍTULO II: CUADRO DE PROTECCIÓN",
                                "CAPÍTULO III: CIRCUITO DE ALUMBRADO",
                                "CAPÍTULO IV: CIRCUITO DE TOMAS DE USO GENERAL",
                                "CAPÍTULO V: CIRCUITO DE COCINA Y HORNO", "CAPÍTULO VI: PUESTA A TIERRA"],
    "🏭 Nave industrial": ["CAPÍTULO I: ACOMETIDA Y CGP", "CAPÍTULO II: CUADRO GENERAL DE DISTRIBUCIÓN",
                           "CAPÍTULO III: FUERZA MOTRIZ", "CAPÍTULO IV: ILUMINACIÓN INDUSTRIAL",
                           "CAPÍTULO V: PUESTA A TIERRA"],
    "☀️ Instalación fotovoltaica": ["CAPÍTULO I: GENERADOR FOTOVOLTAICO",
                                    "CAPÍTULO II: PROTECCIONES E INTERCONEXIÓN"],
}


def _render_flujo_recomendado(hay_cable: bool, hay_fv: bool, n_capitulos: int, doc_generada: bool):
    """Indicador de progreso permanente: no es solo para la primera visita,
    orienta en cada visita sobre dónde va cada cosa y qué falta."""
    st.markdown('<p class="section-label">Flujo recomendado</p>', unsafe_allow_html=True)
    pasos = [
        ("1", "Calcular", "Cable o fotovoltaica", hay_cable or hay_fv, "Calculadora"),
        ("2", "Presupuestar", "Capítulos y mediciones", n_capitulos > 0, "Presupuesto"),
        ("3", "Documentar", "MTD, Anexo, Pliego (PDF)", doc_generada, "Documentación"),
    ]
    cols = st.columns(3)
    for col, (num, titulo, sub, hecho, destino) in zip(cols, pasos):
        with col:
            estado_icono = "✅" if hecho else "⬜"
            estado_txt = "Hecho" if hecho else "Pendiente"
            clase = "success" if hecho else "info"
            st.markdown(f'''<div class="result-card">
                <div style="display:flex; align-items:center; gap:0.6rem; margin-bottom:0.3rem;">
                    <span style="font-family:'JetBrains Mono',monospace; font-weight:700; font-size:1.1rem;
                        color:var(--accent-primary);">{num}</span>
                    <span style="font-weight:700;">{titulo}</span>
                </div>
                <div class="result-sub">{sub}</div>
                <span class="badge {clase}" style="margin-top:0.5rem;">{estado_icono} {estado_txt}</span>
            </div>''', unsafe_allow_html=True)
            if st.button("Ir →" if not hecho else "Revisar →", key=f"flujo_{destino}", width='stretch'):
                st.session_state["pagina_actual"] = destino
                st.rerun()


def _render_dashboard():
    nombre_usuario = st.session_state["config_profesional"].get("nombre") or "ingeniero/a"
    st.session_state.setdefault("guia_bienvenida_oculta", False)

    hay_cable = st.session_state["resultado_cable"].get("seccion_final") is not None
    hay_fv = bool(st.session_state["resultado_fv"]) and st.session_state["resultado_fv"].get("p_pico_kwp") is not None
    n_capitulos = len(st.session_state.get("presupuesto_capitulos", []))
    doc_generada = any("descargad" in a["texto"].lower() for a in st.session_state.get("actividad", []))
    es_usuario_nuevo = not hay_cable and not hay_fv and n_capitulos == 0 and not st.session_state["historial_proyectos"]

    st.markdown(f"### 👋 Hola, {nombre_usuario}")
    st.caption(f"Proyecto actual: **{st.session_state['nombre_proyecto_actual']}** · "
               f"{datetime.now().strftime('%A, %d de %B de %Y')}")

    # ---------------------------------------------------------------- Guía de bienvenida
    if es_usuario_nuevo and not st.session_state["guia_bienvenida_oculta"]:
        st.markdown('<div class="welcome-banner">', unsafe_allow_html=True)
        st.markdown("#### 🚀 ¿Primera vez por aquí? Así funciona en 3 pasos")
        st.markdown(
            "Esta aplicación calcula instalaciones eléctricas de baja tensión (REBT) y genera "
            "automáticamente el presupuesto y la documentación técnica en PDF. No hace falta que uses "
            "todo — con el paso 1 ya tienes un resultado útil."
        )
        b1, b2, b3 = st.columns(3)
        with b1:
            st.markdown("**① Calcula**  \nRellena un formulario simple: potencia, longitud, tipo de "
                        "instalación. O usa una plantilla ya rellena, más abajo.")
            if st.button("🔌 Empezar por aquí", key="bienvenida_ir_calc", type="primary", width='stretch'):
                st.session_state["pagina_actual"] = "Calculadora"
                st.rerun()
        with b2:
            st.markdown("**② Presupuesta**  \nCon un clic importas lo calculado a un capítulo de "
                        "presupuesto, con precios editables.")
        with b3:
            st.markdown("**③ Documenta**  \nDescarga la memoria técnica, el anexo de cálculos y el "
                        "pliego de condiciones, ya con portada y paginación.")
        if st.button("Entendido, no volver a mostrar esto", key="ocultar_bienvenida"):
            st.session_state["guia_bienvenida_oculta"] = True
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        _render_flujo_recomendado(hay_cable, hay_fv, n_capitulos, doc_generada)

    cfg_p = st.session_state.get("presupuesto_config", {"pct_beneficio": PORCENTAJE_BENEFICIO_DEFECTO,
                                                          "pct_amortizacion": PORCENTAJE_AMORTIZACION_DEFECTO})
    total_presu = sum(calcular_totales_capitulo(c["items"], cfg_p["pct_beneficio"], cfg_p["pct_amortizacion"])
                       for c in st.session_state.get("presupuesto_capitulos", []))

    st.markdown('<p class="section-label">Resumen del proyecto</p>', unsafe_allow_html=True)
    k1, k2, k3, k4 = st.columns(4)
    _tarjeta_kpi(k1, "🔌", "blue",
                 f"{st.session_state['resultado_cable'].get('seccion_final', '—')} mm²" if hay_cable else "Sin calcular",
                 "Sección de cable calculada")
    _tarjeta_kpi(k2, "☀️", "orange",
                 f"{st.session_state['resultado_fv'].get('p_pico_kwp', 0):.1f} kWp" if hay_fv else "Sin calcular",
                 "Potencia fotovoltaica")
    _tarjeta_kpi(k3, "💰", "green", _fmt_eur(total_presu) if n_capitulos else "Sin capítulos",
                 f"Presupuesto ({n_capitulos} capítulos)")
    _tarjeta_kpi(k4, "📁", "blue", f"{len(st.session_state.get('historial_proyectos', []))}",
                 "Proyectos en el historial")

    st.markdown('<p class="section-label">Accesos rápidos</p>', unsafe_allow_html=True)
    q1, q2, q3, q4 = st.columns(4)
    _tarjeta_acceso_rapido(q1, "🔌", "Calculadora de cables", "Sección, protecciones, ΔU", "Calculadora")
    _tarjeta_acceso_rapido(q2, "☀️", "Fotovoltaica", "Dimensionado completo FV", "Fotovoltaica")
    _tarjeta_acceso_rapido(q3, "💰", "Presupuesto", "Capítulos y mediciones", "Presupuesto")
    _tarjeta_acceso_rapido(q4, "📄", "Documentación", "MTD, Anexo, Condiciones", "Documentación")

    col_izq, col_der = st.columns([1.3, 1])
    with col_izq:
        st.markdown('<p class="section-label">Plantillas</p>', unsafe_allow_html=True)
        st.caption("Aplican valores de partida habituales — luego puedes ajustarlos en cada pestaña.")
        tab_pc, tab_pfv, tab_pp = st.tabs(["🔌 Cable", "☀️ Fotovoltaica", "💰 Presupuesto"])

        with tab_pc:
            for nombre_plantilla, valores in PLANTILLAS_CIRCUITO.items():
                pc1, pc2 = st.columns([4, 1])
                pc1.markdown(f"**{nombre_plantilla}**  \n"
                             f"<span style='color:var(--text-secondary); font-size:0.8rem;'>"
                             f"{valores['potencia_kw']:g} kW · {valores['sistema']} · {valores['metodo'][:28]}...</span>",
                             unsafe_allow_html=True)
                if pc2.button("Usar", key=f"plantilla_cable_{nombre_plantilla}"):
                    st.session_state["plantilla_activa"] = valores
                    st.session_state["pagina_actual"] = "Calculadora"
                    _registrar_actividad("📋", f"Plantilla de cable aplicada: {nombre_plantilla}")
                    st.rerun()

        with tab_pfv:
            for nombre_plantilla, valores in PLANTILLAS_FV.items():
                pc1, pc2 = st.columns([4, 1])
                pc1.markdown(f"**{nombre_plantilla}**  \n"
                             f"<span style='color:var(--text-secondary); font-size:0.8rem;'>"
                             f"{valores['potencia_pico_deseada']:g} kWp · {valores['sistema_ca']} · "
                             f"{valores['tipo_autoconsumo'][:28]}...</span>", unsafe_allow_html=True)
                if pc2.button("Usar", key=f"plantilla_fv_{nombre_plantilla}"):
                    st.session_state["plantilla_activa_fv"] = valores
                    st.session_state["pagina_actual"] = "Fotovoltaica"
                    _registrar_actividad("📋", f"Plantilla FV aplicada: {nombre_plantilla}")
                    st.rerun()

        with tab_pp:
            st.caption("Crea la estructura de capítulos típica — se añaden vacíos, listos para rellenar.")
            for nombre_plantilla, capitulos_nombres in PLANTILLAS_PRESUPUESTO.items():
                pc1, pc2 = st.columns([4, 1])
                pc1.markdown(f"**{nombre_plantilla}**  \n"
                             f"<span style='color:var(--text-secondary); font-size:0.8rem;'>"
                             f"{len(capitulos_nombres)} capítulos</span>", unsafe_allow_html=True)
                if pc2.button("Usar", key=f"plantilla_presu_{nombre_plantilla}"):
                    st.session_state.setdefault("presupuesto_capitulos", [])
                    for nombre_cap in capitulos_nombres:
                        st.session_state["presupuesto_capitulos"].append({"nombre": nombre_cap, "items": []})
                    st.session_state["pagina_actual"] = "Presupuesto"
                    _registrar_actividad("📋", f"Plantilla de presupuesto aplicada: {nombre_plantilla}")
                    st.rerun()

    with col_der:
        st.markdown('<p class="section-label">Actividad reciente</p>', unsafe_allow_html=True)
        actividad = st.session_state.get("actividad", [])
        if actividad:
            html_items = "".join(
                f'<div class="timeline-item"><div class="timeline-dot"></div>'
                f'<div><div class="timeline-time">{a["hora"]}</div>{a["icono"]} {a["texto"]}</div></div>'
                for a in actividad[:10]
            )
            st.markdown(f'<div class="result-card">{html_items}</div>', unsafe_allow_html=True)
        else:
            st.info("Todavía no hay actividad registrada en esta sesión.")


def _render_proyectos():
    st.markdown('<p class="section-label">Mis proyectos</p>', unsafe_allow_html=True)
    st.caption("Guardar descarga un archivo .json a tu ordenador — eso es lo único que persiste de verdad "
               "entre sesiones. El historial de abajo es solo de esta sesión del navegador (se pierde al "
               "cerrar o recargar la pestaña).")

    with st.container(border=True):
        st.markdown("**💾 Guardar proyecto actual**")
        c1, c2 = st.columns([3, 1])
        nombre_nuevo = c1.text_input("Nombre del proyecto", st.session_state["nombre_proyecto_actual"])
        c2.markdown("<br>", unsafe_allow_html=True)
        if c2.button("Guardar", type="primary", width='stretch'):
            st.session_state["nombre_proyecto_actual"] = nombre_nuevo
            datos = _serializar_proyecto(nombre_nuevo)
            st.session_state["historial_proyectos"].append(datos)
            st.session_state["historial_proyectos"] = st.session_state["historial_proyectos"][-MAX_HISTORIAL_PROYECTOS:]
            _registrar_actividad("💾", f"Proyecto guardado: {nombre_nuevo}")
            st.success(f"Proyecto '{nombre_nuevo}' añadido al historial de la sesión. Descárgalo abajo para "
                       "conservarlo de verdad.")
        datos_actuales = _serializar_proyecto(st.session_state["nombre_proyecto_actual"])
        st.download_button(
            f"⬇️ Descargar '{st.session_state['nombre_proyecto_actual']}' (.json, "
            f"{_tamano_proyecto_kb(datos_actuales):.1f} KB)",
            data=json.dumps(datos_actuales, ensure_ascii=False, indent=2, default=str),
            file_name=f"{st.session_state['nombre_proyecto_actual'].replace(' ', '_')}.json",
            mime="application/json", width='stretch')

    with st.container(border=True):
        st.markdown("**📂 Abrir proyecto**")
        archivo = st.file_uploader("Sube un .json guardado previamente", type=["json"])
        if archivo is not None:
            try:
                datos_cargados = json.loads(archivo.read().decode("utf-8"))
                if st.button("Cargar en la sesión actual", type="primary"):
                    _cargar_proyecto(datos_cargados)
                    _registrar_actividad("📂", f"Proyecto cargado: {datos_cargados.get('__nombre__', archivo.name)}")
                    st.success("Proyecto cargado correctamente.")
                    st.rerun()
            except (json.JSONDecodeError, UnicodeDecodeError):
                st.error("El archivo no es un proyecto válido de esta aplicación.")

    st.markdown('<p class="section-label">Historial de la sesión</p>', unsafe_allow_html=True)
    historial = st.session_state.get("historial_proyectos", [])
    if not historial:
        st.info("Todavía no has guardado ningún proyecto en esta sesión.")
        return

    for i, proy in enumerate(reversed(historial)):
        idx_real = len(historial) - 1 - i
        with st.container(border=True):
            hc1, hc2, hc3, hc4, hc5 = st.columns([3, 2, 1, 1, 1])
            hc1.markdown(f"**{proy.get('__nombre__', 'Sin nombre')}**")
            try:
                fecha_fmt = datetime.fromisoformat(proy["__fecha__"]).strftime("%d/%m/%Y %H:%M")
            except (KeyError, ValueError):
                fecha_fmt = "—"
            hc2.caption(fecha_fmt)
            if hc3.button("Abrir", key=f"hist_abrir_{idx_real}"):
                _cargar_proyecto(proy)
                _registrar_actividad("📂", f"Proyecto reabierto: {proy.get('__nombre__')}")
                st.rerun()
            if hc4.button("Duplicar", key=f"hist_dup_{idx_real}"):
                copia = dict(proy)
                copia["__nombre__"] = f"{proy.get('__nombre__', 'Proyecto')} (copia)"
                copia["__fecha__"] = datetime.now().isoformat()
                st.session_state["historial_proyectos"].append(copia)
                st.session_state["historial_proyectos"] = st.session_state["historial_proyectos"][-MAX_HISTORIAL_PROYECTOS:]
                _registrar_actividad("📑", f"Proyecto duplicado: {copia['__nombre__']}")
                st.rerun()
            if hc5.button("🗑️", key=f"hist_del_{idx_real}"):
                st.session_state["historial_proyectos"].pop(idx_real)
                st.rerun()

    if len(historial) >= 2:
        st.markdown('<p class="section-label">Comparador de proyectos</p>', unsafe_allow_html=True)
        nombres_hist = [f"{i}: {p.get('__nombre__', '-')}" for i, p in enumerate(historial)]
        c1, c2 = st.columns(2)
        sel_a = c1.selectbox("Proyecto A", nombres_hist, index=len(nombres_hist) - 2)
        sel_b = c2.selectbox("Proyecto B", nombres_hist, index=len(nombres_hist) - 1)
        proy_a = historial[int(sel_a.split(":")[0])]
        proy_b = historial[int(sel_b.split(":")[0])]

        def _resumen(p):
            rc = p.get("resultado_cable") or {}
            rf = p.get("resultado_fv") or {}
            caps = p.get("presupuesto_capitulos") or []
            cfg = p.get("presupuesto_config") or {"pct_beneficio": 15, "pct_amortizacion": 5}
            total = sum(calcular_totales_capitulo(c["items"], cfg["pct_beneficio"], cfg["pct_amortizacion"])
                        for c in caps)
            return {
                "Sección de cable": f"{rc.get('seccion_final', '—')} mm²" if rc.get("seccion_final") else "—",
                "Potencia FV": f"{rf.get('p_pico_kwp', 0):.2f} kWp" if rf.get("p_pico_kwp") else "—",
                "Capítulos de presupuesto": f"{len(caps)}",
                "Total presupuesto": _fmt_eur(total) if caps else "—",
            }

        res_a, res_b = _resumen(proy_a), _resumen(proy_b)
        df_comp = pd.DataFrame([{"Magnitud": k, "Proyecto A": res_a[k], "Proyecto B": res_b[k]} for k in res_a])
        st.dataframe(df_comp, width='stretch', hide_index=True)


def _render_estadisticas():
    st.markdown('<p class="section-label">Estadísticas</p>', unsafe_allow_html=True)

    capitulos = st.session_state.get("presupuesto_capitulos", [])
    cfg_p = st.session_state.get("presupuesto_config", {"pct_beneficio": PORCENTAJE_BENEFICIO_DEFECTO,
                                                          "pct_amortizacion": PORCENTAJE_AMORTIZACION_DEFECTO,
                                                          "pct_iva": IVA_DEFECTO_PCT})
    resultado_fv = st.session_state.get("resultado_fv", {})
    resultado_cable = st.session_state.get("resultado_cable", {})
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None
    hay_cable = resultado_cable.get("seccion_final") is not None
    actividad = st.session_state.get("actividad", [])

    if not capitulos and not hay_fv and not hay_cable:
        st.info("Todavía no hay datos suficientes: calcula un circuito, dimensiona una instalación "
                "fotovoltaica o añade capítulos al presupuesto para ver estadísticas.")
        return

    import plotly.graph_objects as go
    import plotly.express as px
    colores_plot = ["#3b82f6", "#e8a33d", "#22c55e", "#ef4444", "#a78bfa", "#06b6d4", "#f59e0b", "#ec4899"]
    layout_base = dict(margin=dict(t=10, b=10, l=10, r=10), paper_bgcolor="rgba(0,0,0,0)",
                        plot_bgcolor="rgba(0,0,0,0)", font=dict(color="#8b96a8"))

    n_partidas = sum(len(c["items"]) for c in capitulos)
    subtotal = sum(calcular_totales_capitulo(c["items"], cfg_p["pct_beneficio"], cfg_p["pct_amortizacion"])
                   for c in capitulos)
    total_con_iva = subtotal * (1 + cfg_p.get("pct_iva", IVA_DEFECTO_PCT) / 100)

    st.markdown('<p class="section-label">Resumen</p>', unsafe_allow_html=True)
    k1, k2, k3, k4 = st.columns(4)
    _tarjeta_kpi(k1, "💰", "green", _fmt_eur(total_con_iva) if capitulos else "—", "Presupuesto (IVA incl.)")
    _tarjeta_kpi(k2, "📦", "blue", str(n_partidas), "Partidas presupuestadas")
    _tarjeta_kpi(k3, "☀️", "orange", f"{resultado_fv.get('p_pico_kwp', 0):.1f} kWp" if hay_fv else "—",
                 "Potencia fotovoltaica")
    _tarjeta_kpi(k4, "📋", "blue", str(len(actividad)), "Acciones en la sesión")

    tab_presu, tab_fv, tab_actividad = st.tabs(["💰 Presupuesto", "☀️ Fotovoltaica", "📋 Actividad"])

    with tab_presu:
        if capitulos:
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Distribución por capítulo**")
                nombres = [c["nombre"][:35] for c in capitulos]
                valores = [calcular_totales_capitulo(c["items"], cfg_p["pct_beneficio"], cfg_p["pct_amortizacion"])
                           for c in capitulos]
                fig = go.Figure(data=[go.Pie(labels=nombres, values=valores, hole=0.55,
                                              marker=dict(colors=colores_plot))])
                fig.update_layout(**layout_base, height=320, legend=dict(orientation="h", y=-0.15))
                st.plotly_chart(fig, width='stretch')
            with c2:
                st.markdown("**Desglose PEM / Amortización / Beneficio / IVA**")
                desglose = {"PEM": subtotal,
                            "Amortización": subtotal * cfg_p["pct_amortizacion"] / 100,
                            "Beneficio": subtotal * cfg_p["pct_beneficio"] / 100}
                desglose["IVA"] = (desglose["PEM"] + desglose["Amortización"] + desglose["Beneficio"]) * \
                    cfg_p.get("pct_iva", IVA_DEFECTO_PCT) / 100
                fig2 = px.bar(x=list(desglose.keys()), y=list(desglose.values()), color=list(desglose.keys()),
                              color_discrete_sequence=colores_plot)
                fig2.update_layout(**layout_base, height=320, showlegend=False, xaxis_title="", yaxis_title="€")
                st.plotly_chart(fig2, width='stretch')

            st.markdown("**Materiales frente a mano de obra**")
            partidas_compuestas_stats = st.session_state.get("partidas_compuestas", {})

            def _categorias_de_item(designacion):
                """Devuelve {categoria: fraccion} para una partida. Si es una
                partida compuesta, reparte la fracción entre las categorías
                reales de sus componentes (así la mano de obra que lleva
                dentro de una partida compuesta también cuenta como mano de
                obra, no como 'Otros')."""
                for cat, items_cat in CATALOGO_MATERIALES.items():
                    if designacion in items_cat:
                        return {cat: 1.0}
                if designacion in partidas_compuestas_stats:
                    componentes = partidas_compuestas_stats[designacion]["componentes"]
                    precio_total = sum(c["cantidad"] * c["precio_unitario"] for c in componentes)
                    if precio_total > 0:
                        reparto = {}
                        for c in componentes:
                            cat_comp = "Otros / manual"
                            for cat, items_cat in CATALOGO_MATERIALES.items():
                                if c["designacion"] in items_cat:
                                    cat_comp = cat
                                    break
                            frac = (c["cantidad"] * c["precio_unitario"]) / precio_total
                            reparto[cat_comp] = reparto.get(cat_comp, 0) + frac
                        return reparto
                return {"Otros / manual": 1.0}

            coste_por_categoria = {}
            coste_mano_obra, coste_materiales = 0.0, 0.0
            for cap in capitulos:
                for it in cap["items"]:
                    pu = calcular_precio_venta(it["precio_base"], cfg_p["pct_beneficio"], cfg_p["pct_amortizacion"])
                    importe_it = it["cantidad"] * pu
                    for cat, frac in _categorias_de_item(it["designacion"]).items():
                        coste_por_categoria[cat] = coste_por_categoria.get(cat, 0) + importe_it * frac
                        if cat == "Mano de obra":
                            coste_mano_obra += importe_it * frac
                        else:
                            coste_materiales += importe_it * frac

            mo1, mo2 = st.columns(2)
            with mo1:
                fig_mo = go.Figure(data=[go.Pie(
                    labels=["Materiales y equipos", "Mano de obra"],
                    values=[coste_materiales, coste_mano_obra], hole=0.6,
                    marker=dict(colors=["#3b82f6", "#e8a33d"]))])
                fig_mo.update_layout(**layout_base, height=260, legend=dict(orientation="h", y=-0.15))
                st.plotly_chart(fig_mo, width='stretch')
            with mo2:
                st.metric("Materiales y equipos", _fmt_eur(coste_materiales),
                          f"{coste_materiales/subtotal*100:.0f}% del PEM" if subtotal else None)
                st.metric("Mano de obra", _fmt_eur(coste_mano_obra),
                          f"{coste_mano_obra/subtotal*100:.0f}% del PEM" if subtotal else None)
                if coste_mano_obra == 0:
                    st.caption("No hay partidas de mano de obra en este presupuesto todavía — añádelas "
                               "desde la categoría «Mano de obra» del catálogo, sueltas o dentro de una "
                               "partida compuesta.")

            st.markdown("**Distribución por categoría** (mano de obra resaltada aparte)")
            if coste_por_categoria:
                nombres_cat = list(coste_por_categoria.keys())
                valores_cat = list(coste_por_categoria.values())
                colores_cat = ["#e8a33d" if c == "Mano de obra" else "#3b82f6" for c in nombres_cat]
                fig_cat = go.Figure(go.Bar(x=valores_cat, y=nombres_cat, orientation="h",
                                           marker=dict(color=colores_cat)))
                fig_cat.update_layout(**layout_base, height=max(220, 42 * len(coste_por_categoria)),
                                       showlegend=False, xaxis_title="€", yaxis_title="")
                st.plotly_chart(fig_cat, width='stretch')
        else:
            st.info("Añade capítulos en Presupuesto para ver esta sección.")

    with tab_fv:
        if hay_fv:
            f1, f2, f3 = st.columns(3)
            f1.metric("Producción año 1", f"{resultado_fv['produccion_anual_kwh']:,.0f} kWh".replace(",", "."))
            f2.metric("Producción año 10", f"{resultado_fv.get('produccion_ano10', 0):,.0f} kWh".replace(",", "."))
            f3.metric("Producción año 25", f"{resultado_fv.get('produccion_ano25', 0):,.0f} kWh".replace(",", "."))

            st.markdown("**Producción mensual estimada (año 1)**")
            pesos_mes = [0.055, 0.065, 0.085, 0.095, 0.105, 0.11, 0.115, 0.105, 0.09, 0.075, 0.055, 0.045]
            meses = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
            produccion_total = resultado_fv["produccion_anual_kwh"]
            valores_mes = [produccion_total * p for p in pesos_mes]
            fig3 = px.bar(x=meses, y=valores_mes, color=valores_mes, color_continuous_scale="Blues")
            fig3.update_layout(**layout_base, height=280, showlegend=False, coloraxis_showscale=False,
                               xaxis_title="", yaxis_title="kWh")
            st.plotly_chart(fig3, width='stretch')
            st.caption("Distribución mensual orientativa según el patrón estacional típico de irradiación en "
                       "España (no son datos de PVGIS de tu ubicación exacta).")

            st.markdown("**Evolución de la producción anual (degradación del panel)**")
            anos = list(range(1, 26))
            degradacion = st.session_state.get("inputs_fv", {}).get("degradacion_anual", DEGRADACION_ANUAL_DEFECTO)
            produccion_anos = [produccion_total * (1 - degradacion / 100) ** (a - 1) for a in anos]
            fig4 = go.Figure(go.Scatter(x=anos, y=produccion_anos, mode="lines", fill="tozeroy",
                                        line=dict(color="#f59e0b")))
            fig4.update_layout(**layout_base, height=260, xaxis_title="Año", yaxis_title="kWh/año")
            st.plotly_chart(fig4, width='stretch')

            if resultado_fv.get("ahorro_anual"):
                e1, e2 = st.columns(2)
                e1.metric("Ahorro anual estimado", _fmt_eur(resultado_fv["ahorro_anual"]))
                if resultado_fv.get("payback_anos"):
                    e2.metric("Retorno de la inversión", f"{resultado_fv['payback_anos']:.1f} años")
        else:
            st.info("Dimensiona una instalación en Fotovoltaica para ver esta sección.")

    with tab_actividad:
        if actividad:
            conteo_tipo = {}
            for a in actividad:
                clave = a["texto"].split(":")[0].split(" descargad")[0].split(" guardad")[0][:28]
                conteo_tipo[clave] = conteo_tipo.get(clave, 0) + 1
            fig5 = px.bar(x=list(conteo_tipo.values()), y=list(conteo_tipo.keys()), orientation="h",
                         color=list(conteo_tipo.keys()), color_discrete_sequence=colores_plot)
            fig5.update_layout(**layout_base, height=max(200, 40 * len(conteo_tipo)), showlegend=False,
                               xaxis_title="Nº de veces", yaxis_title="")
            st.plotly_chart(fig5, width='stretch')
            st.dataframe(pd.DataFrame([{"Hora": a["hora"], "Acción": f"{a['icono']} {a['texto']}"}
                                        for a in actividad]), hide_index=True, width='stretch')
        else:
            st.info("Todavía no hay actividad registrada en esta sesión.")


def _render_configuracion():
    st.markdown('<p class="section-label">Configuración</p>', unsafe_allow_html=True)

    tab_perfil, tab_apariencia = st.tabs(["🖋️ Modo profesional", "🎨 Apariencia"])

    with tab_perfil:
        st.caption("Estos datos se incluyen en la cabecera y portada de todos los PDF generados "
                   "(memoria de cálculo, MTD, Anexo, Condiciones Generales).")
        cfg = st.session_state["config_profesional"]
        c1, c2 = st.columns(2)
        with c1:
            cfg["nombre"] = st.text_input("Nombre del técnico/instalador", cfg.get("nombre", ""))
            cfg["empresa"] = st.text_input("Empresa / razón social", cfg.get("empresa", ""))
        with c2:
            cfg["firma"] = st.text_input("Texto de firma (p. ej. nº de colegiado)", cfg.get("firma", ""))
            logo_file = st.file_uploader("Logotipo (PNG/JPG, opcional)", type=["png", "jpg", "jpeg"])
            if logo_file is not None:
                cfg["logo_b64"] = base64.b64encode(logo_file.read()).decode("ascii")
                cfg["logo_mime"] = logo_file.type
        if cfg.get("logo_b64"):
            lc1, lc2 = st.columns([1, 4])
            lc1.image(base64.b64decode(cfg["logo_b64"]), width=80)
            if lc2.button("Quitar logotipo"):
                cfg["logo_b64"] = ""
                st.rerun()
        st.session_state["config_profesional"] = cfg

    with tab_apariencia:
        st.caption("El tema se aplica a toda la aplicación (no afecta a los PDF, que mantienen un diseño "
                   "fijo pensado para imprimir).")
        tema_sel = st.radio("Tema", ["Oscuro", "Claro"],
                            index=0 if st.session_state["tema"] == "Oscuro" else 1, horizontal=True)
        if tema_sel != st.session_state["tema"]:
            st.session_state["tema"] = tema_sel
            st.rerun()
        if st.session_state["tema"] == "Claro":
            st.info(
                "⚠️ Limitación conocida: los campos y tablas nativos de Streamlit (números, desplegables, "
                "editor de precios) mantienen un fondo oscuro incluso en modo claro. Streamlit fija ese "
                "estilo a nivel de servidor (`.streamlit/config.toml`), no por sesión de usuario, así que "
                "no hay forma fiable de sincronizarlo con el interruptor de tema de la app. El modo oscuro "
                "es el que tiene el acabado completo."
            )


def _render_acerca_de():
    st.markdown('<p class="section-label">Acerca de</p>', unsafe_allow_html=True)
    st.markdown("""
### REBT Suite — Calculadora de Instalaciones Eléctricas

Herramienta de apoyo al diseño de instalaciones de baja tensión conforme al REBT (RD 842/2002) y sus
Instrucciones Técnicas Complementarias, con módulos independientes de cálculo de secciones de cable,
dimensionado fotovoltaico, calculadoras de referencia rápida, presupuesto por capítulos y generación de
documentación técnica (MTD, Anexo de Cálculos, Pliego de Condiciones).

**Construida con:** Python, Streamlit, pandas, reportlab (PDF), openpyxl (Excel), plotly (gráficos).

**Qué es y qué no es esta app:**
- ✅ Herramienta de apoyo al predimensionado, con las fórmulas y tablas normativas documentadas en cada
  pestaña (ver "Metodología").
- ✅ Los proyectos se guardan como archivo `.json` que descargas tú — no hay servidor guardando tus datos.
- ❌ No sustituye la verificación de un técnico competente antes de firmar un proyecto o memoria.
- ❌ No es una base de datos persistente: recarga la página y el historial de sesión se pierde (por eso
  existe la descarga de proyectos).

**Aviso legal:** las tablas y fórmulas se han contrastado contra la Guía-BT-19 y las ITC-BT
correspondientes, pero pueden existir erratas. Verifica los valores críticos antes de un uso profesional.
    """)
    st.caption("Versión 4.0 · Última actualización de este documento: sesión actual.")


def _render_tablas():
    st.markdown('<p class="section-label">Tabla A — Intensidades admisibles (A), cobre, no enterrado, '
                'aire a 40°C</p>', unsafe_allow_html=True)
    filas = []
    for s, datos in TABLA_A_COBRE.items():
        b1, ce, f = datos["B1"], datos["CE"], datos["F"]
        a1 = tuple(round(v * FACTOR_A1_SOBRE_B1, 1) for v in b1)
        b2 = tuple(round(v * FACTOR_B2_SOBRE_B1, 1) for v in b1)
        filas.append({
            "Sección (mm²)": s,
            "A1* 2×PVC": a1[IDX_2PVC], "A1* 2×XLPE": a1[IDX_2XLPE],
            "B1 3×PVC": b1[IDX_3PVC], "B1 2×PVC": b1[IDX_2PVC],
            "B1 3×XLPE": b1[IDX_3XLPE], "B1 2×XLPE": b1[IDX_2XLPE],
            "B2* 2×PVC": b2[IDX_2PVC], "B2* 2×XLPE": b2[IDX_2XLPE],
            "C/E 3×PVC": ce[IDX_3PVC], "C/E 2×PVC": ce[IDX_2PVC],
            "C/E 3×XLPE": ce[IDX_3XLPE], "C/E 2×XLPE": ce[IDX_2XLPE],
            "F (S≥25, XLPE)": f"{f:g}" if f is not None else "—",
        })
    st.dataframe(pd.DataFrame(filas), width='stretch', hide_index=True)
    st.caption("Columnas B1, C/E y F: directas de la Guía-BT-19, Tabla A. Columnas marcadas con * (A1, B2): "
               "estimadas aplicando un ratio a B1, anclado a un valor real verificado en cada caso "
               "(ver pestaña Metodología). Aluminio en instalación NO enterrada: ratio orientativo 0,78 "
               "sobre el valor de cobre mostrado aquí (no se tabula aparte).")

    st.markdown('<p class="section-label">Tabla D — Enterrado bajo tubo, XLPE (terreno 25°C, '
                'ρ=1,5 K·m/W, 0,70 m)</p>', unsafe_allow_html=True)
    filas_d = []
    for s, (cu3, al3, cu2, al2) in TABLA_D_ENTERRADO_XLPE.items():
        filas_d.append({"Sección (mm²)": s, "3× Cobre": cu3,
                         "3× Aluminio": f"{al3:g}" if al3 is not None else "—",
                         "2× Cobre": cu2,
                         "2× Aluminio": f"{al2:g}" if al2 is not None else "—"})
    st.dataframe(pd.DataFrame(filas_d), width='stretch', hide_index=True)

    st.markdown('<p class="section-label">Tabla E — Factores de corrección por agrupamiento de '
                'circuitos</p>', unsafe_allow_html=True)
    claves = sorted(set(k for d in TABLA_E_AGRUPAMIENTO.values() for k in d.keys()))
    filas_e = []
    for disp, tabla in TABLA_E_AGRUPAMIENTO.items():
        fila = {"Disposición": disp}
        for k in claves:
            fila[f"{k} circ."] = f"{tabla[k]:.2f}" if k in tabla else "—"
        filas_e.append(fila)
    st.dataframe(pd.DataFrame(filas_e), width='stretch', hide_index=True)

    col_izq, col_der = st.columns(2)
    with col_izq:
        st.markdown('<p class="section-label">Caídas de tensión máximas admisibles</p>', unsafe_allow_html=True)
        filas_u = [{"Tramo": k, "ΔU máx. (%)": f"{v:g}" if v is not None else "según proyecto"}
                   for k, v in CAIDA_TENSION_MAX.items()]
        st.dataframe(pd.DataFrame(filas_u), width='stretch', hide_index=True)
    with col_der:
        st.markdown('<p class="section-label">Constantes k — cortocircuito (S≤300 mm²)</p>',
                    unsafe_allow_html=True)
        filas_k = [{"Conductor": c, "Aislamiento": a, "k": v} for (c, a), v in K_CORTOCIRCUITO.items()]
        st.dataframe(pd.DataFrame(filas_k), width='stretch', hide_index=True)
        st.caption("S_mín = Icc·√t / k. Fuente: IEC 60364-5-54 / Libro Blanco de la Instalación (Prysmian).")


def _render_metodologia():
    st.markdown(
        f"""
### Criterios de dimensionado aplicados

Esta calculadora sigue el procedimiento de la **ITC-BT-19, artículo 19.2**, que exige que la sección de
un conductor cumpla simultáneamente:

**1. Criterio térmico — Ib ≤ In ≤ Iz.**
La intensidad admisible del cable (Iz), obtenida de la tabla de su método de instalación y corregida por
temperatura ambiente, agrupamiento de circuitos y, si procede, resistividad del terreno, debe ser igual o
superior a la corriente de empleo del circuito.

**2. Criterio de caída de tensión.**
La caída de tensión a lo largo del circuito —calculada con la conductividad del conductor **a su
temperatura de servicio** (70°C PVC / 90°C XLPE, criterio más conservador que usar la conductividad a
20°C)— no debe superar el porcentaje máximo admisible según el tramo (ITC-BT-14, 15, 19 y 40).

**3. Criterio térmico de cortocircuito** (verificación opcional).
Si se aportan la Icc en origen y el tiempo de actuación de la protección, se comprueba S ≥ Icc·√t / k.

La sección final adoptada es la mayor de las que exigen los criterios 1 y 2, redondeada a la sección
normalizada superior.

### Métodos de instalación — qué es dato directo y qué es estimación

| Método | Origen de los valores |
|---|---|
| B1 — Tubo empotrado en obra o superficie | **Directo** de la Tabla A, Guía-BT-19 |
| C/E — Bandeja no perforada/perforada, superficie | **Directo** de la Tabla A. La propia Guía-BT-19 indica que C y E comparten columna cuando sus valores son "prácticamente iguales" |
| F — Bandeja perforada, unipolares (S≥25mm², XLPE) | **Directo** de la Tabla A |
| D — Enterrado bajo tubo (XLPE) | **Directo** de la Tabla D, Guía-BT-19 |
| A1 — Tubo empotrado en pared aislante | **Estimado**: A1 = B1 × {FACTOR_A1_SOBRE_B1:g}, ratio anclado a un valor real verificado de forma independiente (A1, 2,5 mm², 2 cargados, PVC = 17,5 A) |
| B2 — Cable multiconductor en tubo | **Estimado**: B2 = B1 × {FACTOR_B2_SOBRE_B1:g}, ratio anclado a un ejercicio resuelto independiente (B2, 25 mm², 2 cargados, PVC = 77 A) |

### Casuística contemplada
- Sistemas monofásico (230 V) y trifásico (400 V).
- 6 métodos de instalación (ver tabla anterior).
- Motores según ITC-BT-47: 125 % para motor único; 125 % del mayor + 100 % del resto para varios motores;
  +30 % adicional en ascensores/grúas.
- Alumbrado de descarga sin corregir el factor de potencia (factor orientativo ×1,8).
- Cargas con armónicos significativos → neutro a sección plena.
- Cobre y aluminio, con la sección mínima de 16 mm² del aluminio en instalación fija (ITC-BT-07) y la
  mínima de la LGA (ITC-BT-14).
- Conductores en paralelo cuando la intensidad supera la capacidad de un único conductor de 300 mm².
- Sección del neutro (mitad de fase cuando aplica) y del conductor de protección (ITC-BT-18).
- Justificación de fórmulas con los valores del cálculo sustituidos (pestaña "Fórmulas").
- Memoria de cálculo en PDF con cajetín, y presupuesto orientativo exportable a Excel.

### Limitaciones conocidas — léelas antes de usar en un proyecto firmado
- **A1 y B2 son estimaciones** derivadas de B1 mediante un ratio anclado a un único valor real cada una;
  no son la columna oficial completa de la Guía-BT-19 para esos métodos. Verifica si tu proyecto depende
  críticamente de ellos.
- Los valores de **aluminio en instalación NO enterrada** son una estimación (ratio 0,78 respecto al
  cobre); la Guía-BT-19 remite a UNE-HD 60364-5-52 para esa combinación.
- El factor de corrección por **resistividad térmica del terreno** es orientativo.
- La caída de tensión usa una reactancia lineal orientativa fija (0,08 Ω/km).
- La verificación de cortocircuito comprueba el criterio **térmico** del conductor, no la coordinación de
  protecciones (curvas, poder de corte, selectividad).
- **Los precios del presupuesto son un punto de partida editable**, no precios de mercado en tiempo real;
  el precio del cobre fluctúa con su cotización.

**En definitiva:** esta herramienta acelera el predimensionado y documenta el criterio seguido, pero no
sustituye la verificación de un técnico competente con la edición vigente de la Guía-BT-19 y el resto de
normativa aplicable antes de firmar un proyecto.
        """
    )


# ==============================================================================
# 9. PUNTO DE ENTRADA
# ==============================================================================

def _generar_pregunta_iz():
    s = random.choice(list(TABLA_A_COBRE.keys()))
    metodo_col = random.choice(["B1", "CE"])
    idx_col = random.choice([IDX_3PVC, IDX_2PVC])
    valor_real = TABLA_A_COBRE[s][metodo_col][idx_col]
    nombre_metodo = "B1 (tubo empotrado)" if metodo_col == "B1" else "C/E (bandeja o directo)"
    n_cond = "3 cargados (trifásico)" if idx_col == IDX_3PVC else "2 cargados (monofásico)"
    distractores = {round(valor_real * f, 1) for f in (0.75, 0.85, 1.15, 1.3) if round(valor_real * f, 1) != valor_real}
    opciones = list(distractores)[:3] + [valor_real]
    random.shuffle(opciones)
    return {
        "pregunta": f"Según la Guía-BT-19 (Tabla A), ¿cuál es la intensidad admisible (Iz) de un cable de "
                    f"cobre de {s:g} mm², aislamiento PVC, método {nombre_metodo}, {n_cond}?",
        "opciones": [f"{v:g} A" for v in opciones],
        "correcta": opciones.index(valor_real),
        "explicacion": f"La Tabla A de la Guía-BT-19 da {valor_real:g} A para esa combinación exacta de "
                       f"sección, método y nº de conductores cargados — antes de aplicar los factores de "
                       "corrección por temperatura, agrupamiento, etc.",
    }


def _generar_pregunta_agrupamiento():
    disposicion = random.choice(list(TABLA_E_AGRUPAMIENTO.keys()))
    n_circ = random.choice(list(TABLA_E_AGRUPAMIENTO[disposicion].keys()))
    valor_real = TABLA_E_AGRUPAMIENTO[disposicion][n_circ]
    distractores = {round(valor_real + d, 2) for d in (-0.15, -0.08, 0.08, 0.15) if 0.2 < round(valor_real + d, 2) <= 1.0}
    opciones = list(distractores)[:3] + [valor_real]
    opciones = list(dict.fromkeys(opciones))[:4]
    if valor_real not in opciones:
        opciones[0] = valor_real
    random.shuffle(opciones)
    return {
        "pregunta": f"Con {n_circ} circuitos agrupados en disposición «{disposicion}», ¿qué factor de "
                    "corrección por agrupamiento (Tabla E) se aplica a la Iz de cada circuito?",
        "opciones": [f"{v:.2f}" for v in opciones],
        "correcta": opciones.index(valor_real),
        "explicacion": f"La Tabla E de la Guía-BT-19 fija {valor_real:.2f} para {n_circ} circuitos en esa "
                       "disposición — cuantos más circuitos agrupados, menor es el factor, porque se "
                       "calientan entre sí.",
    }


def _generar_pregunta_seccion_proteccion():
    s_fase = random.choice([4, 10, 16, 25, 35, 50, 70])
    if s_fase <= 16:
        correcta = s_fase
    elif s_fase <= 35:
        correcta = 16
    else:
        correcta = s_fase / 2
    opciones_posibles = sorted({s_fase, 16, s_fase / 2, s_fase * 2})
    opciones = [o for o in opciones_posibles if o > 0][:4]
    while len(opciones) < 4:
        opciones.append(max(opciones) * 1.5)
    random.shuffle(opciones)
    return {
        "pregunta": f"Según la tabla de la ITC-BT-18, para un conductor de fase de {s_fase:g} mm², ¿qué "
                    "sección mínima debe tener el conductor de protección?",
        "opciones": [f"{v:g} mm²" for v in opciones],
        "correcta": opciones.index(correcta),
        "explicacion": "ITC-BT-18: Sf≤16 → Sp=Sf; 16<Sf≤35 → Sp=16 mm²; Sf>35 → Sp=Sf/2.",
    }


def _generar_pregunta_curva_magnetotermico():
    curva = random.choice(list(CURVA_MAGNETOTERMICO_RANGOS.keys()))
    rango_real = CURVA_MAGNETOTERMICO_RANGOS[curva]
    otras_curvas = [c for c in CURVA_MAGNETOTERMICO_RANGOS if c != curva]
    opciones_rangos = [rango_real] + [CURVA_MAGNETOTERMICO_RANGOS[c] for c in otras_curvas]
    random.shuffle(opciones_rangos)
    return {
        "pregunta": f"¿Entre qué múltiplos de In dispara instantáneamente (zona magnética) un interruptor "
                    f"automático de curva {curva}?",
        "opciones": [f"{r[0]}-{r[1]}×In" for r in opciones_rangos],
        "correcta": opciones_rangos.index(rango_real),
        "explicacion": f"La curva {curva} dispara instantáneamente entre {rango_real[0]} y {rango_real[1]} "
                       "veces la intensidad nominal (UNE-EN 60898). B es la más sensible (circuitos "
                       "resistivos/largos), D la menos (motores, transformadores).",
    }


def _generar_pregunta_ratio_motor():
    pmin, pmax, ratio_real = random.choice(RATIO_IA_IN_MAX_MOTOR)
    otros_ratios = [r for _, _, r in RATIO_IA_IN_MAX_MOTOR if r != ratio_real]
    opciones = [ratio_real] + otros_ratios
    random.shuffle(opciones)
    rango_txt = f"{pmin:g}-{pmax:g} kW" if pmax != float("inf") else f"más de {pmin:g} kW"
    return {
        "pregunta": f"Según la ITC-BT-47, para un motor de potencia en el rango {rango_txt}, ¿cuál es la "
                    "relación máxima admisible entre la intensidad de arranque directo y la nominal (Ia/In)?",
        "opciones": [f"{v:g}" for v in opciones],
        "correcta": opciones.index(ratio_real),
        "explicacion": f"La ITC-BT-47 limita Ia/In a {ratio_real:g} para motores en ese rango de potencia; "
                       "por encima, se exige un sistema de arranque que reduzca la corriente (p. ej. "
                       "estrella-triángulo).",
    }


def _generar_pregunta_resistividad_terreno():
    terrenos_validos = {k: v for k, v in RESISTIVIDAD_TERRENOS_REF.items() if v is not None}
    terreno = random.choice(list(terrenos_validos.keys()))
    valor_real = terrenos_validos[terreno]
    valores_unicos = list({v for v in terrenos_validos.values() if v != valor_real})
    otros = random.sample(valores_unicos, min(3, len(valores_unicos)))
    opciones = [valor_real] + otros
    random.shuffle(opciones)
    return {
        "pregunta": f"¿Cuál es la resistividad orientativa del terreno tipo «{terreno}» (Ω·m), a efectos de "
                    "cálculo de puesta a tierra (ITC-BT-18)?",
        "opciones": [f"{v:g} Ω·m" for v in opciones],
        "correcta": opciones.index(valor_real),
        "explicacion": f"Valor de referencia: {valor_real:g} Ω·m. La resistividad real del terreno debe "
                       "medirse in situ siempre que sea posible; estos valores son solo orientativos para "
                       "un anteproyecto.",
    }


def _generar_pregunta_colores_resistencia():
    col1 = random.choice([c for c in COLORES_DIGITO if c not in ("Negro",)])
    col2 = random.choice(list(COLORES_DIGITO.keys()))
    col3 = random.choice(list(COLORES_MULTIPLICADOR.keys()))
    valor_real, _ = valor_resistencia_4_bandas(col1, col2, col3, "Oro")
    distractores = {valor_real * f for f in (10, 0.1, 2) if valor_real * f != valor_real}
    opciones = list(distractores)[:3] + [valor_real]
    opciones = list(dict.fromkeys(opciones))
    random.shuffle(opciones)
    return {
        "pregunta": f"Una resistencia con bandas {col1}-{col2}-{col3} (código de 4 bandas), ¿qué valor "
                    "representa?",
        "opciones": [f"{v:g} Ω" for v in opciones],
        "correcta": opciones.index(valor_real),
        "explicacion": f"Valor = (dígito1×10 + dígito2) × multiplicador = "
                       f"({COLORES_DIGITO[col1]}×10+{COLORES_DIGITO[col2]})×{COLORES_MULTIPLICADOR[col3]:g} = "
                       f"{valor_real:g} Ω.",
    }


def _generar_pregunta_du_maxima():
    caso = random.choice([
        ("alumbrado, en instalación interior", 3.0),
        ("otros usos, en instalación interior", 5.0),
        ("derivación individual, sin línea general de alimentación", 1.5),
    ])
    descripcion, valor_real = caso
    opciones = sorted({1.5, 3.0, 5.0, 6.5})
    return {
        "pregunta": f"¿Cuál es la caída de tensión máxima admisible (%) para un circuito de {descripcion}?",
        "opciones": [f"{v:g} %" for v in opciones],
        "correcta": opciones.index(valor_real),
        "explicacion": "ITC-BT-19/14/15: 3% alumbrado y 5% otros usos en instalación interior; 1,5% en "
                       "derivación individual cuando no existe línea general de alimentación.",
    }


GENERADORES_PREGUNTAS = [
    _generar_pregunta_iz, _generar_pregunta_agrupamiento, _generar_pregunta_seccion_proteccion,
    _generar_pregunta_curva_magnetotermico, _generar_pregunta_ratio_motor,
    _generar_pregunta_resistividad_terreno, _generar_pregunta_colores_resistencia,
    _generar_pregunta_du_maxima,
]


def _render_autoevaluacion():
    st.markdown('<p class="section-label">Autoevaluación</p>', unsafe_allow_html=True)
    st.caption("Preguntas generadas en el momento a partir de las mismas tablas normativas que usa la "
               "app (Tabla A, Tabla E, ITC-BT-18, ITC-BT-47...) — no son un banco fijo de preguntas, cada "
               "test es distinto. Pensado para repasar antes de un examen o una prueba de aptitud.")

    n_preguntas = st.slider("Número de preguntas", 4, len(GENERADORES_PREGUNTAS), 8, key="quiz_n")
    if st.button("🎲 Generar test nuevo", type="primary"):
        generadores_elegidos = random.sample(GENERADORES_PREGUNTAS, min(n_preguntas, len(GENERADORES_PREGUNTAS)))
        st.session_state["quiz_preguntas"] = [g() for g in generadores_elegidos]
        st.session_state["quiz_respuestas"] = [None] * len(generadores_elegidos)
        st.session_state["quiz_corregido"] = False
        st.rerun()

    preguntas = st.session_state.get("quiz_preguntas")
    if not preguntas:
        st.info("Pulsa «Generar test nuevo» para empezar.")
        return

    respuestas = st.session_state["quiz_respuestas"]
    corregido = st.session_state.get("quiz_corregido", False)

    for i_p, preg in enumerate(preguntas):
        st.markdown(f"**{i_p + 1}. {preg['pregunta']}**")
        idx_elegido = st.radio("Respuesta", preg["opciones"], index=respuestas[i_p], key=f"quiz_resp_{i_p}",
                                label_visibility="collapsed")
        respuestas[i_p] = preg["opciones"].index(idx_elegido) if idx_elegido is not None else None
        if corregido:
            if respuestas[i_p] == preg["correcta"]:
                st.success(f"✅ Correcto — {preg['explicacion']}")
            else:
                st.error(f"❌ La respuesta correcta era «{preg['opciones'][preg['correcta']]}» — "
                         f"{preg['explicacion']}")
        st.divider()

    if not corregido:
        if st.button("✔️ Corregir test", type="primary"):
            if any(r is None for r in respuestas):
                st.warning("Responde todas las preguntas antes de corregir.")
            else:
                st.session_state["quiz_corregido"] = True
                st.rerun()
    else:
        aciertos = sum(1 for i, preg in enumerate(preguntas) if respuestas[i] == preg["correcta"])
        pct = aciertos / len(preguntas) * 100
        st.markdown(f"### Resultado: {aciertos}/{len(preguntas)} ({pct:.0f}%)")
        st.progress(pct / 100)
        if pct >= 80:
            st.success("🎉 Muy buen resultado.")
        elif pct >= 50:
            st.warning("Vas por buen camino — repasa lo que hayas fallado.")
        else:
            st.error("Conviene repasar estos temas antes de un examen o prueba real.")


def _render_presentacion_cliente():
    st.markdown('<p class="section-label">Presentación cliente</p>', unsafe_allow_html=True)
    st.caption("La misma información del proyecto, sin siglas ni jerga técnica (REBT, ITC-BT, mm²...) — "
               "pensada para compartir pantalla o imprimir cuando se lo explicas a alguien que no es del "
               "sector.")

    inputs_cable = st.session_state.get("inputs_cable", {})
    resultado_cable = st.session_state.get("resultado_cable", {})
    inputs_fv = st.session_state.get("inputs_fv", {})
    resultado_fv = st.session_state.get("resultado_fv", {})
    capitulos = st.session_state.get("presupuesto_capitulos", [])
    cfg_presu = st.session_state.get("presupuesto_config", {
        "pct_beneficio": PORCENTAJE_BENEFICIO_DEFECTO, "pct_amortizacion": PORCENTAJE_AMORTIZACION_DEFECTO,
        "pct_iva": IVA_DEFECTO_PCT,
    })
    datos = st.session_state.get("datos_proyecto", {})

    hay_cable = resultado_cable.get("seccion_final") is not None
    hay_fv = bool(resultado_fv) and resultado_fv.get("p_pico_kwp") is not None
    subtotal = sum(calcular_totales_capitulo(c["items"], cfg_presu["pct_beneficio"], cfg_presu["pct_amortizacion"])
                   for c in capitulos)
    total_presupuesto = subtotal * (1 + cfg_presu["pct_iva"] / 100) if capitulos else 0.0

    if not hay_cable and not hay_fv and not capitulos:
        st.info("Todavía no hay nada calculado — completa la Calculadora, la Fotovoltaica o el Presupuesto "
                "para ver aquí el resumen listo para presentar.")
        return

    nombre_cliente = datos.get("titular") or "tu proyecto"
    st.markdown(f"## 👋 Esto es lo que hemos preparado para {nombre_cliente}")

    if hay_cable:
        with st.container(border=True):
            st.markdown("### 🔌 La instalación eléctrica")
            grosor_mm = resultado_cable["seccion_final"]
            comparacion = ("un cable fino, como el de un cargador de móvil" if grosor_mm <= 2.5 else
                          "un cable de grosor medio, como el de un electrodoméstico grande" if grosor_mm <= 10 else
                          "un cable considerablemente grueso, pensado para mover mucha potencia con seguridad")
            st.markdown(f"Para esta parte de la instalación hace falta un cable de **{grosor_mm:g} mm²** de "
                        f"grosor — a modo de referencia, es {comparacion}.")
            cumple_du = resultado_cable.get("e_final_pct", 0) <= inputs_cable.get("delta_u_max", 5)
            if cumple_du:
                st.success("✅ Con este cable, la corriente llega con la fuerza necesaria de principio a "
                           "fin sin perder potencia por el camino.")
            st.markdown(f"Además, se instala una protección automática de **{resultado_cable['calibre_magnetotermico']} A** "
                        "que corta la luz sola si algo va mal, antes de que pueda ser peligroso.")

    if hay_fv:
        with st.container(border=True):
            st.markdown("### ☀️ Los paneles solares")
            p_pico = resultado_fv["p_pico_kwp"]
            produccion = resultado_fv["produccion_anual_kwh"]
            n_paneles = resultado_fv["n_paneles_configurados"]
            st.markdown(f"Se instalan **{n_paneles} paneles solares**, con una potencia conjunta de "
                        f"**{p_pico:.1f} kW** — suficiente para producir aproximadamente "
                        f"**{_miles(produccion)} kWh de electricidad al año**.")
            equivalente_hogares_dias = produccion / 10  # aprox 10 kWh/dia consumo hogar medio
            st.caption(f"Para hacerte una idea: es más o menos lo que consumiría una vivienda media durante "
                       f"{equivalente_hogares_dias:.0f} días.")
            if resultado_fv.get("ahorro_anual"):
                st.markdown(f"Esto se traduce en un ahorro estimado de **{_fmt_eur(resultado_fv['ahorro_anual'])} "
                            "al año** en la factura de la luz.")
            if resultado_fv.get("payback_anos"):
                st.markdown(f"Con la inversión que nos has indicado, los paneles se pagan solos en unos "
                            f"**{resultado_fv['payback_anos']:.0f} años** — a partir de ahí, el ahorro es "
                            "beneficio neto.")
            co2 = resultado_fv.get("co2_evitado_kg_ano", 0) / 1000
            if co2:
                st.markdown(f"De paso, se evita la emisión de unas **{co2:.1f} toneladas de CO₂ al año** — "
                            "el equivalente a lo que absorben decenas de árboles.")

    if capitulos:
        with st.container(border=True):
            st.markdown("### 💰 El coste total")
            st.markdown(f"## {_fmt_eur(total_presupuesto)}")
            st.caption(f"Impuestos incluidos. Repartido en {len(capitulos)} bloques de trabajo "
                       f"({', '.join(c['nombre'] for c in capitulos[:4])}"
                       f"{'...' if len(capitulos) > 4 else ''}).")

    st.divider()
    st.caption("💡 Los números técnicos completos (secciones exactas, normativa aplicada, desglose de "
               "precios) siguen disponibles en la pestaña Documentación, por si en algún momento hace "
               "falta el detalle.")

    with st.spinner("Generando ficha para el cliente..."):
        pdf_cliente = generar_pdf_presentacion_cliente(datos, inputs_cable, resultado_cable, inputs_fv,
                                                         resultado_fv, total_presupuesto,
                                                         st.session_state.get("config_profesional", {}))
    if st.download_button("⬇️ Descargar esta misma vista en PDF", data=pdf_cliente,
                           file_name="presentacion_cliente.pdf", mime="application/pdf"):
        _registrar_actividad("🗣️", "Ficha de presentación cliente descargada")


def main():
    st.set_page_config(page_title="REBT Suite · Instalaciones Eléctricas", page_icon="⚡", layout="wide",
                       initial_sidebar_state="expanded")
    _inicializar_estado()
    st.markdown(generar_css(st.session_state["tema"]), unsafe_allow_html=True)
    _render_sidebar()

    pagina = st.session_state["pagina_actual"]

    if pagina == "Inicio":
        _render_dashboard()
        return
    if pagina == "Proyectos":
        _render_proyectos()
        return
    if pagina == "Estadísticas":
        _render_estadisticas()
        return
    if pagina == "Configuración":
        _render_configuracion()
        return
    if pagina == "Acerca de":
        _render_acerca_de()
        return
    if pagina == "Tablas normativas":
        _render_tablas()
        return
    if pagina == "Metodología":
        _render_metodologia()
        return
    if pagina == "Autoevaluación":
        _render_autoevaluacion()
        return
    if pagina == "Presentación cliente":
        _render_presentacion_cliente()
        return

    # --- Páginas de herramientas: mismo cajetín de cabecera que antes ---
    eyebrow = {"Calculadora": "Cálculo de secciones · Baja tensión", "Fórmulas": "Justificación de cálculo",
               "Fotovoltaica": "Dimensionado de instalaciones solares", "Cálculos BT": "Calculadoras de referencia rápida",
               "Presupuesto": "Mediciones y precios", "Documentación": "MTD · Anexo · Pliego"}.get(pagina, "")
    st.markdown(f"""
        <div class="titleblock">
            <div class="titleblock-main">
                <span class="titleblock-eyebrow">{eyebrow}</span>
                <h1>{pagina}</h1>
            </div>
            <div class="titleblock-meta">
                <div><span>Norma</span><strong>REBT · ITC-BT</strong></div>
                <div><span>Proyecto</span><strong>{st.session_state['nombre_proyecto_actual'][:16]}</strong></div>
                <div><span>Rev.</span><strong>4.0</strong></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    if pagina == "Calculadora":
        inputs_cable = _render_inputs()
        resultado_cable = calcular(inputs_cable)
        st.session_state["inputs_cable"] = inputs_cable
        st.session_state["resultado_cable"] = resultado_cable
        seccion_previa = st.session_state.get("_ultima_seccion_cable")
        if resultado_cable.get("seccion_final") and resultado_cable["seccion_final"] != seccion_previa:
            st.session_state["_ultima_seccion_cable"] = resultado_cable["seccion_final"]
            _registrar_actividad("🔌", f"Cable calculado: {resultado_cable['seccion_final']:g} mm²")
        _render_resultados(inputs_cable, resultado_cable)

    elif pagina == "Fórmulas":
        ic, rc = st.session_state.get("inputs_cable", {}), st.session_state.get("resultado_cable", {})
        if rc.get("seccion_final") is not None:
            _render_formulas(ic, rc)
        else:
            _estado_vacio("Aún no has calculado ningún circuito. Hazlo en Calculadora y aquí verás la "
                          "justificación completa, fórmula a fórmula.", "Calculadora",
                          "🔌 Ir a Calculadora", "🧮")

    elif pagina == "Fotovoltaica":
        inputs_fv = _render_inputs_fv()
        resultado_fv = calcular_fv(inputs_fv)
        st.session_state["inputs_fv"] = inputs_fv
        st.session_state["resultado_fv"] = resultado_fv
        potencia_previa = st.session_state.get("_ultima_potencia_fv")
        if resultado_fv.get("p_pico_kwp") and resultado_fv["p_pico_kwp"] != potencia_previa:
            st.session_state["_ultima_potencia_fv"] = resultado_fv["p_pico_kwp"]
            _registrar_actividad("☀️", f"FV calculada: {resultado_fv['p_pico_kwp']:.1f} kWp")
        _render_resultados_fv(inputs_fv, resultado_fv)
        _render_comparador_fv(inputs_fv, resultado_fv)

    elif pagina == "Cálculos BT":
        _render_calculos_bt()

    elif pagina == "Presupuesto":
        _render_presupuesto(st.session_state.get("inputs_cable", {}), st.session_state.get("resultado_cable", {}),
                            st.session_state.get("inputs_fv", {}), st.session_state.get("resultado_fv", {}))

    elif pagina == "Documentación":
        _render_documentacion(st.session_state.get("inputs_cable", {}), st.session_state.get("resultado_cable", {}),
                              st.session_state.get("inputs_fv", {}), st.session_state.get("resultado_fv", {}))


if __name__ == "__main__":
    main()
