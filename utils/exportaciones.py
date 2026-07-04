"""Exportacion de resultados del proyecto a distintos formatos.

La clase ExportadorProyecto agrupa la generacion de Excel, Word, PDF y
DXF, reproduciendo exactamente la misma logica que las funciones
originales 'exportar_excel', 'exportar_memoria_word',
'exportar_memoria_pdf' y 'generar_esquema_unifilar_dxf'.

Se deja preparado (metodo 'exportar_bc3') el hueco para anadir en el
futuro la exportacion a formato BC3, sin necesidad de reestructurar el
resto del codigo.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any, Dict, Optional

import pandas as pd

from datos.constantes import TIPOS_SUMINISTRO
from utils.logging_config import obtener_logger

logger = obtener_logger(__name__)


class ExportadorProyecto:
    """Genera los distintos formatos de exportacion del proyecto electrico."""

    @staticmethod
    def exportar_excel(
        df_bt_calc: Optional[pd.DataFrame],
        df_motores_calc: Optional[pd.DataFrame],
        df_mediciones: Optional[pd.DataFrame],
        resumen_presupuesto: Optional[Dict[str, float]],
    ) -> io.BytesIO:
        """Exporta los resultados de calculo, mediciones y presupuesto a un libro Excel."""
        buffer = io.BytesIO()
        hoja_escrita = False
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            if df_bt_calc is not None and not df_bt_calc.empty:
                df_bt_calc.to_excel(writer, sheet_name="Baja Tension", index=False)
                hoja_escrita = True
            if df_motores_calc is not None and not df_motores_calc.empty:
                df_motores_calc.to_excel(writer, sheet_name="Industrial-Motores", index=False)
                hoja_escrita = True
            if df_mediciones is not None and not df_mediciones.empty:
                df_mediciones.to_excel(writer, sheet_name="Mediciones", index=False)
                hoja_escrita = True
            if resumen_presupuesto:
                pd.DataFrame([resumen_presupuesto]).T.rename(columns={0: "Importe (EUR)"}).to_excel(writer, sheet_name="Presupuesto")
                hoja_escrita = True
            if not hoja_escrita:
                pd.DataFrame(
                    {"Aviso": ["Todavia no se han introducido datos en el proyecto. Anade circuitos, motores o mediciones antes de exportar."]}
                ).to_excel(writer, sheet_name="Aviso", index=False)
        buffer.seek(0)
        logger.info("Excel de resultados generado correctamente")
        return buffer

    @staticmethod
    def exportar_memoria_word(texto_memoria: str) -> io.BytesIO:
        """Exporta el texto de la memoria a un documento Word (.docx)."""
        from docx import Document

        doc = Document()
        for linea in texto_memoria.split("\n"):
            if linea.isupper() and linea.strip() and len(linea.strip()) > 3 and linea[0].isdigit() is False and linea.strip()[0].isalpha():
                doc.add_heading(linea, level=1)
            elif linea.strip().startswith(tuple(f"{n}." for n in range(1, 10))):
                doc.add_heading(linea, level=1)
            else:
                doc.add_paragraph(linea)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("Memoria en Word generada correctamente")
        return buffer

    @staticmethod
    def exportar_memoria_pdf(
        estado: Dict[str, Any],
        texto_memoria: str,
        df_bt_calc: Optional[pd.DataFrame],
        df_fv_calc: Optional[pd.DataFrame],
        df_motores_calc: Optional[pd.DataFrame],
        df_mediciones: Optional[pd.DataFrame],
        resumen_presupuesto: Optional[Dict[str, float]],
    ) -> io.BytesIO:
        """Exporta la memoria completa (portada, cuerpo y anexos) a PDF."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
        styles = getSampleStyleSheet()
        titulo = ParagraphStyle("TituloPortada", parent=styles["Title"], fontSize=22, spaceAfter=20)
        subt = ParagraphStyle("Subt", parent=styles["Normal"], fontSize=12, spaceAfter=8)
        h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=colors.HexColor("#0B3D91"))
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#0B3D91"))
        normal = styles["Normal"]

        datos = estado.get("datos_proyecto", {})
        story = []

        # ---- Portada ----
        story.append(Spacer(1, 4 * cm))
        story.append(Paragraph("MEMORIA TECNICA DE DISENO (MTD)", titulo))
        story.append(Paragraph("Instalacion electrica en Baja Tension conforme al REBT", subt))
        story.append(Spacer(1, 2 * cm))
        story.append(Paragraph(f"<b>Titular:</b> {datos.get('titular','-')}", subt))
        story.append(Paragraph(f"<b>Emplazamiento:</b> {datos.get('emplazamiento','-')}", subt))
        story.append(Paragraph(f"<b>Referencia:</b> {datos.get('referencia','-')}", subt))
        story.append(Paragraph(f"<b>Tipo de suministro:</b> {estado.get('tipo_suministro','-')}", subt))
        story.append(Paragraph(f"<b>Fecha:</b> {datos.get('fecha', str(date.today()))}", subt))
        story.append(PageBreak())

        # ---- Cuerpo de la memoria ----
        for bloque in texto_memoria.split("\n\n"):
            lineas = bloque.split("\n")
            for i, linea in enumerate(lineas):
                if linea.strip()[:2].rstrip(".").isdigit() or (linea.strip() and linea.strip()[0].isdigit() and "." in linea[:3]):
                    story.append(Paragraph(linea, h1))
                elif linea.strip().startswith(" - "):
                    story.append(Paragraph(linea.strip(), normal))
                else:
                    story.append(Paragraph(linea if linea.strip() else "&nbsp;", normal))
            story.append(Spacer(1, 6))
        story.append(PageBreak())

        def tabla_desde_df(df: Optional[pd.DataFrame], titulo_tabla: str) -> None:
            story.append(Paragraph(titulo_tabla, h2))
            if df is None or df.empty:
                story.append(Paragraph("Sin datos introducidos para este capitulo.", normal))
                story.append(Spacer(1, 12))
                return
            cols = list(df.columns)
            data = [cols] + df.astype(str).values.tolist()
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D91")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FF")]),
            ]))
            story.append(t)
            story.append(Spacer(1, 14))

        # ---- Anexo I: BT ----
        story.append(Paragraph("ANEXO I - CALCULOS DE CIRCUITOS DE BAJA TENSION", h1))
        tabla_desde_df(df_bt_calc, "Resultados de calculo de circuitos")
        story.append(PageBreak())

        # ---- Anexo II: FV ----
        story.append(Paragraph("ANEXO II - CALCULOS DE INSTALACION FOTOVOLTAICA", h1))
        tabla_desde_df(df_fv_calc, "Resultados de calculo fotovoltaico")
        story.append(PageBreak())

        # ---- Anexo III: Motores ----
        story.append(Paragraph("ANEXO III - CALCULOS DE MOTORES / INDUSTRIAL", h1))
        tabla_desde_df(df_motores_calc, "Resultados de calculo de motores")
        story.append(PageBreak())

        # ---- Anexo IV: Mediciones y presupuesto ----
        story.append(Paragraph("ANEXO IV - MEDICIONES Y PRESUPUESTO", h1))
        tabla_desde_df(df_mediciones, "Mediciones")
        if resumen_presupuesto:
            story.append(Paragraph("Resumen del presupuesto", h2))
            data = [["Concepto", "Importe (EUR)"]] + [[k, f"{v:,.2f}"] for k, v in resumen_presupuesto.items()]
            t = Table(data)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D91")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ]))
            story.append(t)

        doc.build(story)
        buffer.seek(0)
        logger.info("Memoria en PDF generada correctamente")
        return buffer

    @staticmethod
    def generar_esquema_unifilar_dxf(estado: Dict[str, Any], df_bt_calc: Optional[pd.DataFrame]) -> bytes:
        """Genera un esquema unifilar en formato DXF (abierto, compatible con
        AutoCAD mediante 'Guardar como .dwg'), con una distribucion inspirada en
        los configuradores profesionales tipo CIEBT: bloques en cascada para la
        acometida / proteccion / medida segun el tipo de suministro elegido, y
        derivaciones horizontales para cada circuito interior.

        NOTA IMPORTANTE: no es posible generar un binario .dwg nativo (formato
        propietario de Autodesk) sin las librerias oficiales de Autodesk/ODA. Se
        genera un .dxf, formato abierto totalmente compatible, que se puede abrir
        y volver a guardar como .dwg desde AutoCAD u otro programa CAD.
        """
        import ezdxf

        doc = ezdxf.new(setup=True)
        msp = doc.modelspace()

        capas = {
            "MARCO": 7, "ACOMETIDA": 5, "PROTECCION": 1, "MEDIDA": 3,
            "CUADRO": 6, "CIRCUITOS": 4, "TEXTOS": 7,
        }
        for nombre, color in capas.items():
            doc.layers.add(name=nombre, color=color)

        tipo_sum = estado.get("tipo_suministro", list(TIPOS_SUMINISTRO.keys())[0])
        elementos = TIPOS_SUMINISTRO.get(tipo_sum, {}).get("elementos", [])

        x0 = 0
        y = 0
        paso_y = 25
        ancho_caja = 90
        alto_caja = 14

        # ---- Marco / cajetin tipo plano profesional ----
        alto_total = paso_y * (len(elementos) + 2) + 40
        ancho_total = 260
        msp.add_lwpolyline(
            [(-20, -alto_total), (-20, 30), (ancho_total, 30), (ancho_total, -alto_total), (-20, -alto_total)],
            dxfattribs={"layer": "MARCO"},
        )
        msp.add_text("ESQUEMA UNIFILAR - INSTALACION ELECTRICA EN BAJA TENSION",
                     dxfattribs={"layer": "TEXTOS", "height": 4}).set_placement((-15, 15))
        msp.add_text(f"Tipo de suministro: {tipo_sum}",
                     dxfattribs={"layer": "TEXTOS", "height": 2.5}).set_placement((-15, 8))

        # ---- Cadena vertical de elementos del suministro (estilo CIEBT) ----
        puntos_conexion = []
        for elemento in elementos:
            cx, cy = x0, y
            msp.add_lwpolyline(
                [(cx, cy), (cx + ancho_caja, cy), (cx + ancho_caja, cy - alto_caja), (cx, cy - alto_caja), (cx, cy)],
                dxfattribs={"layer": "CUADRO"},
            )
            msp.add_text(elemento, dxfattribs={"layer": "TEXTOS", "height": 3}).set_placement(
                (cx + 4, cy - alto_caja / 2 - 1)
            )
            centro = (cx + ancho_caja / 2, cy - alto_caja)
            puntos_conexion.append(centro)
            if len(puntos_conexion) > 1:
                ant = puntos_conexion[-2]
                msp.add_line((ant[0], ant[1]), (centro[0], centro[1] - (paso_y - alto_caja) + alto_caja),
                             dxfattribs={"layer": "ACOMETIDA"})
            y -= paso_y

        ultimo_punto = puntos_conexion[-1] if puntos_conexion else (x0 + ancho_caja / 2, 0)

        # ---- Embarrado del cuadro general ----
        barra_y = ultimo_punto[1] - 10
        msp.add_line((ultimo_punto[0], ultimo_punto[1]), (ultimo_punto[0], barra_y), dxfattribs={"layer": "CUADRO"})
        msp.add_line((-10, barra_y), (ancho_total - 10, barra_y), dxfattribs={"layer": "CUADRO"})
        msp.add_text("EMBARRADO - CUADRO GENERAL DE MANDO Y PROTECCION",
                     dxfattribs={"layer": "TEXTOS", "height": 2.5}).set_placement((-10, barra_y + 2))

        # ---- Derivaciones a cada circuito interior (segun datos introducidos) ----
        if df_bt_calc is not None and not df_bt_calc.empty:
            n = len(df_bt_calc)
            espacio_x = (ancho_total - 10) / max(n, 1)
            for i, (_, fila) in enumerate(df_bt_calc.iterrows()):
                cx = -10 + espacio_x * i + espacio_x / 2
                msp.add_line((cx, barra_y), (cx, barra_y - 10), dxfattribs={"layer": "CIRCUITOS"})
                # simbolo de proteccion magnetotermica (rectangulo pequeno)
                msp.add_lwpolyline(
                    [(cx - 3, barra_y - 10), (cx + 3, barra_y - 10), (cx + 3, barra_y - 16), (cx - 3, barra_y - 16), (cx - 3, barra_y - 10)],
                    dxfattribs={"layer": "PROTECCION"},
                )
                msp.add_line((cx, barra_y - 16), (cx, barra_y - 26), dxfattribs={"layer": "CIRCUITOS"})
                nombre_circ = str(fila.get("Circuito", f"C{i+1}"))
                seccion = fila.get("Seccion (mm2)", "-")
                proteccion = fila.get("Proteccion (A)", "-")
                lineas_etiqueta = [str(nombre_circ), f"Cu {seccion} mm2 / {proteccion} A"]
                for j, linea_txt in enumerate(lineas_etiqueta):
                    msp.add_text(linea_txt, dxfattribs={"layer": "TEXTOS", "height": 1.8}).set_placement(
                        (cx - 8, barra_y - 30 - j * 3)
                    )
        else:
            msp.add_text("Sin circuitos interiores definidos todavia (ver pestana Baja Tension)",
                         dxfattribs={"layer": "TEXTOS", "height": 2.5}).set_placement((-10, barra_y - 12))

        # ---- Cajetin inferior (estilo profesional) ----
        cy_cajetin = -alto_total + 5
        msp.add_lwpolyline(
            [(-20, cy_cajetin), (ancho_total, cy_cajetin), (ancho_total, cy_cajetin - 20), (-20, cy_cajetin - 20), (-20, cy_cajetin)],
            dxfattribs={"layer": "MARCO"},
        )
        datos = estado.get("datos_proyecto", {})
        msp.add_text(f"Titular: {datos.get('titular','-')}", dxfattribs={"layer": "TEXTOS", "height": 2.2}).set_placement((-15, cy_cajetin - 6))
        msp.add_text(f"Emplazamiento: {datos.get('emplazamiento','-')}", dxfattribs={"layer": "TEXTOS", "height": 2.2}).set_placement((-15, cy_cajetin - 10))
        msp.add_text(f"Fecha: {datos.get('fecha', str(date.today()))} Escala: S/E", dxfattribs={"layer": "TEXTOS", "height": 2.2}).set_placement((-15, cy_cajetin - 14))

        buffer = io.StringIO()
        doc.write(buffer)
        contenido = buffer.getvalue()
        logger.info("Esquema unifilar DXF generado correctamente")
        return contenido.encode("utf-8")

    @staticmethod
    def exportar_bc3(*args: Any, **kwargs: Any) -> None:
        """Punto de extension reservado para la futura exportacion a formato BC3.

        Actualmente no implementado. Se deja definido para que anadir esta
        funcionalidad en el futuro no requiera cambios en el resto de la
        aplicacion (unicamente completar este metodo).
        """
        raise NotImplementedError("La exportacion a formato BC3 todavia no esta implementada.")


def exportar_excel(df_bt_calc, df_motores_calc, df_mediciones, resumen_presupuesto):
    return ExportadorProyecto.exportar_excel(df_bt_calc, df_motores_calc, df_mediciones, resumen_presupuesto)


def exportar_memoria_word(texto_memoria):
    return ExportadorProyecto.exportar_memoria_word(texto_memoria)


def exportar_memoria_pdf(estado, texto_memoria, df_bt_calc, df_fv_calc, df_motores_calc, df_mediciones, resumen_presupuesto):
    return ExportadorProyecto.exportar_memoria_pdf(estado, texto_memoria, df_bt_calc, df_fv_calc, df_motores_calc, df_mediciones, resumen_presupuesto)


def generar_esquema_unifilar_dxf(estado, df_bt_calc):
    return ExportadorProyecto.generar_esquema_unifilar_dxf(estado, df_bt_calc)
